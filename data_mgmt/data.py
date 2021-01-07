import datetime
import difflib
import os
import re
import typing
from collections import Counter
from typing import List, Dict, Union, Optional, Tuple

import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from datetime import timedelta, date

import numpy
from dateutil import parser
import numpy as np
from datamaps.api import project_data_from_master
import platform
from pathlib import Path

from dateutil.parser import ParserError
from docx import Document, table
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, Cm, RGBColor, Inches
from matplotlib import cm, pyplot as plt
from matplotlib.patches import Wedge, Rectangle, Circle
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.workbook import workbook
from textwrap import wrap


def _platform_docs_dir() -> Path:
    #  Cross plaform file path handling
    if platform.system() == "Linux":
        return Path.home() / "Documents" / "analysis_engine"
    if platform.system() == "Darwin":
        return Path.home() / "Documents" / "analysis_engine"
    else:
        return Path.home() / "Documents" / "analysis_engine"


root_path = _platform_docs_dir()


def get_master_data() -> List[
    Dict[str, Union[str, int, datetime.date, float]]
]:  # how specify a list of dictionaries?
    """Returns a list of dictionaries each containing quarter data"""
    master_data_list = [
        project_data_from_master(root_path / "core_data/master_3_2020_draft.xlsx", 3, 2020),
        project_data_from_master(root_path / "core_data/master_2_2020.xlsx", 2, 2020),
        project_data_from_master(root_path / "core_data/master_1_2020.xlsx", 1, 2020),
        project_data_from_master(root_path / "core_data/master_4_2019.xlsx", 4, 2019),
        project_data_from_master(root_path / "core_data/master_3_2019.xlsx", 3, 2019),
        project_data_from_master(root_path / "core_data/master_2_2019.xlsx", 2, 2019),
        project_data_from_master(root_path / "core_data/master_1_2019.xlsx", 1, 2019),
        project_data_from_master(root_path / "core_data/master_4_2018.xlsx", 4, 2018),
        project_data_from_master(root_path / "core_data/master_3_2018.xlsx", 3, 2018),
        project_data_from_master(root_path / "core_data/master_2_2018.xlsx", 2, 2018),
        project_data_from_master(root_path / "core_data/master_1_2018.xlsx", 1, 2018),
        project_data_from_master(root_path / "core_data/master_4_2017.xlsx", 4, 2017),
        project_data_from_master(root_path / "core_data/master_3_2017.xlsx", 3, 2017),
        project_data_from_master(root_path / "core_data/master_2_2017.xlsx", 2, 2017),
        project_data_from_master(root_path / "core_data/master_1_2017.xlsx", 1, 2017),
        project_data_from_master(root_path / "core_data/master_4_2016.xlsx", 4, 2016),
        project_data_from_master(root_path / "core_data/master_3_2016.xlsx", 3, 2016),
    ]

    return master_data_list


def get_master_data_file_paths() -> List[typing.TextIO]:
    file_list = [
        root_path / "core_data/master_2_2020.xlsx",
        root_path / "core_data/master_1_2020.xlsx",
        root_path / "core_data/master_4_2019.xlsx",
        root_path / "core_data/master_3_2019.xlsx",
        root_path / "core_data/master_2_2019.xlsx",
        root_path / "core_data/master_1_2019.xlsx",
        root_path / "core_data/master_4_2018.xlsx",
        root_path / "core_data/master_3_2018.xlsx",
        root_path / "core_data/master_2_2018.xlsx",
        root_path / "core_data/master_1_2018.xlsx",
        root_path / "core_data/master_4_2017.xlsx",
        root_path / "core_data/master_3_2017.xlsx",
        root_path / "core_data/master_2_2017.xlsx",
        root_path / "core_data/master_1_2017.xlsx",
        root_path / "core_data/master_4_2016.xlsx",
        root_path / "core_data/master_3_2016.xlsx",
    ]

    return file_list


def get_master_data_file_paths_fy_16_17() -> List[typing.TextIO]:
    file_list = [
        root_path / "core_data/master_4_2016.xlsx",
        root_path / "core_data/master_3_2016.xlsx",
    ]
    return file_list


def get_master_data_file_paths_fy_17_18() -> List[typing.TextIO]:
    file_list = [
        root_path / "core_data/master_4_2017.xlsx",
        root_path / "core_data/master_3_2017.xlsx",
        root_path / "core_data/master_2_2017.xlsx",
        root_path / "core_data/master_1_2017.xlsx",
    ]
    return file_list


def get_master_data_file_paths_fy_18_19() -> List[typing.TextIO]:
    file_list = [
        root_path / "core_data/master_4_2018.xlsx",
        root_path / "core_data/master_3_2018.xlsx",
        root_path / "core_data/master_2_2018.xlsx",
        root_path / "core_data/master_1_2018.xlsx",
    ]
    return file_list


def get_master_data_file_paths_fy_19_20() -> List[typing.TextIO]:
    file_list = [
        root_path / "core_data/master_4_2019.xlsx",
        root_path / "core_data/master_3_2019.xlsx",
        root_path / "core_data/master_2_2019.xlsx",
        root_path / "core_data/master_1_2019.xlsx",
    ]

    return file_list


def get_master_data_file_paths_fy_20_21() -> List[typing.TextIO]:
    file_list = [
        root_path / "core_data/master_2_2020.xlsx",
        root_path / "core_data/master_1_2020.xlsx",
    ]

    return file_list


def get_datamap_file_paths():
    pass


def get_project_information() -> Dict[str, Union[str, int]]:
    """Returns dictionary containing all project meta data"""
    return project_data_from_master(
        root_path / "core_data/data_mgmt/project_info.xlsx", 2, 2020
    )


def get_project_information_file_path() -> typing.TextIO:
    return root_path / "core_data/other/project_info.xlsx"


def get_gmpp_projects(project_info_dict) -> List[str]:
    # returns list of projects that are in gmpp
    project_list = list(project_info_dict.data.keys())
    output_list = []
    for p in project_list:
        if project_info_dict.data[p]["GMPP"]:
            output_list.append(p)

    return output_list


def get_error_list(seq: List[str]):
    seen = set()
    seen_add = seen.add
    return [x for x in seq if not (x in seen or seen_add(x))]


# for project summary pages
SRO_CONF_TABLE_LIST = [
    "SRO DCA",
    "Finance DCA",
    "Benefits DCA",
    "Resourcing DCA",
    "Schedule DCA",
]
SRO_CONF_KEY_LIST = [
    "Departmental DCA",
    "SRO Finance confidence",
    "SRO Benefits RAG",
    "Overall Resource DCA - Now",
    "SRO Schedule Confidence",
]

IPDC_DATE = datetime.date(
    2021, 2, 22
)  # ipdc date. Python date format is Year, Month, day

# abbreviations. Used in analysis instead of full projects names
ABBREVIATION = {
    "2nd Generation UK Search and Rescue Aviation": "SARH2",
    "A12 Chelmsford to A120 widening": "A12",
    "A14 Cambridge to Huntingdon Improvement Scheme": "A14",
    "A303 Amesbury to Berwick Down": "A303",
    "A358 Taunton to Southfields Dualling": "A358",
    "A417 Air Balloon": "A417",
    "A428 Black Cat to Caxton Gibbet": "A428",
    "A66 Northern Trans-Pennine": "A66",
    "Crossrail Programme": "Crossrail",
    "East Coast Digital Programme": "ECDP",
    "East Coast Mainline Programme": "ECMP",
    "East West Rail Programme (Central Section)": "EWR (Central)",
    "East West Rail Programme (Western Section)": "EWR (Western)",
    "East West Rail Configuration State 1": "EWR Config 1",
    "East West Rail Configuration State 2": "EWR Config 2",
    "East West Rail Configuration State 3": "EWR Config 3",
    "Future Theory Test Service (FTTS)": "FTTS",
    "Great Western Route Modernisation (GWRM) including electrification": "GWRM",
    "Heathrow Expansion": "HEP",
    "Hexagon": "Hexagon",
    "High Speed Rail Programme (HS2)": "HS2 Prog",
    "HS2 Phase 2b": "HS2 2b",
    "HS2 Phase1": "HS2 1",
    "HS2 Phase2a": "HS2 2a",
    "Integrated and Smart Ticketing - creating an account based back office": "IST",
    "Intercity Express Programme": "IEP",
    "Lower Thames Crossing": "LTC",
    "M4 Junctions 3 to 12 Smart Motorway": "M4",
    "Manchester North West Quadrant": "MNWQ",
    "Midland Main Line Programme": "MML Prog",
    "Midlands Rail Hub": "Mid Rail Hub",
    "North Western Electrification": "NWE",
    "Northern Powerhouse Rail": "NPR",
    "Oxford-Cambridge Expressway": "Ox-Cam Expressway",
    "Rail Franchising Programme": "Rail Franchising",
    "South West Route Capacity": "SWRC",
    "Thameslink Programme": "Thameslink",
    "Transpennine Route Upgrade (TRU)": "TRU",
    "Western Rail Link to Heathrow": "WRLtH",
}


class Projects:
    # project names as variables
    a12 = "A12 Chelmsford to A120 widening"
    a14 = "A14 Cambridge to Huntingdon Improvement Scheme"
    a303 = "A303 Amesbury to Berwick Down"
    a385 = "A358 Taunton to Southfields Dualling"
    a417 = "A417 Air Balloon"
    a428 = "A428 Black Cat to Caxton Gibbet"
    a66 = "A66 Northern Trans-Pennine"
    brighton_ml = "Brighton Mainline Upgrade Programme"
    cvs = "Commercial Vehicle Services (CVS)"
    east_coast_digital = "East Coast Digital Programme"
    east_coast_mainline = "East Coast Mainline Programme"
    em_franchise = "East Midlands Franchise"
    ewr_central = "East West Rail Programme (Central Section)"
    ewr_western = "East West Rail Programme (Western Section)"
    ewr_config1 = "East West Rail Configuration State 1"
    ewr_config2 = "East West Rail Configuration State 2"
    ewr_config3 = "East West Rail Configuration State 3"
    ftts = "Future Theory Test Service (FTTS)"
    heathrow_expansion = "Heathrow Expansion"
    hexagon = "Hexagon"
    hs2_programme = "High Speed Rail Programme (HS2)"
    hs2_2b = "HS2 Phase 2b"
    hs2_1 = "HS2 Phase1"
    hs2_2a = "HS2 Phase2a"
    ist = "Integrated and Smart Ticketing - creating an account based back office"
    lower_thames_crossing = "Lower Thames Crossing"
    m4 = "M4 Junctions 3 to 12 Smart Motorway"
    manchester_north_west_quad = "Manchester North West Quadrant"
    midland_mainline = "Midland Main Line Programme"
    midlands_rail_hub = "Midlands Rail Hub"
    north_of_england = "North of England Programme"
    northern_powerhouse = "Northern Powerhouse Rail"
    nwe = "North Western Electrification"
    ox_cam_expressway = "Oxford-Cambridge Expressway"
    rail_franchising = "Rail Franchising Programme"
    west_coast_partnership = "West Coast Partnership Franchise"
    crossrail = "Crossrail Programme"
    gwrm = "Great Western Route Modernisation (GWRM) including electrification"
    iep = "Intercity Express Programme"
    sarh2 = "2nd Generation UK Search and Rescue Aviation"
    south_west_route_capacity = "South West Route Capacity"
    thameslink = "Thameslink Programme"
    tru = "Transpennine Route Upgrade"
    wrlth = "Western Rail Link to Heathrow"

    # lists of projects names in groups
    he = [lower_thames_crossing, a303, a14, a66, a12, m4, a428, a417, a385]
    rail = [
        crossrail,
        thameslink,
        iep,
        east_coast_mainline,
        east_coast_digital,
        midland_mainline,
        nwe,
        south_west_route_capacity,
        brighton_ml,
        midlands_rail_hub,
        gwrm,
        tru,
        wrlth,
    ]
    hs2 = [hs2_1, hs2_2a, hs2_2b]
    hsmrpg = [
        hs2_1,
        hs2_2a,
        hs2_2b,
        ewr_config1,
        ewr_config2,
        ewr_config3,
        hexagon,
        northern_powerhouse,
    ]
    ewr = [ewr_config1, ewr_config2, ewr_config3]
    dvsa = [ftts, ist]
    all_not_hs2 = [
        "2nd Generation UK Search and Rescue Aviation",
        "A12 Chelmsford to A120 widening",
        "A14 Cambridge to Huntingdon Improvement Scheme",
        "A303 Amesbury to Berwick Down",
        "A358 Taunton to Southfields Dualling",
        "A417 Air Balloon",
        "A428 Black Cat to Caxton Gibbet",
        "A66 Northern Trans-Pennine",
        "Brighton Mainline Upgrade Programme",
        "Crossrail Programme",
        "East Coast Digital Programme",
        "East Coast Mainline Programme",
        "East West Rail Configuration State 1",
        "East West Rail Configuration State 2",
        "East West Rail Configuration State 3",
        "Future Theory Test Service (FTTS)",
        "Great Western Route Modernisation (GWRM) including electrification",
        "Hexagon",
        "Integrated and Smart Ticketing - creating an account based back office",
        "Intercity Express Programme",
        "Lower Thames Crossing",
        "M4 Junctions 3 to 12 Smart Motorway",
        "Midland Main Line Programme",
        "Midlands Rail Hub",
        "North Western Electrification",
        "Northern Powerhouse Rail",
        "Rail Franchising Programme",
        "South West Route Capacity",
        "Thameslink Programme",
        "Transpennine Route Upgrade",
        "Western Rail Link to Heathrow",
    ]
    fbc_stage = [
        hs2_1,
        crossrail,
        east_coast_mainline,
        iep,
        thameslink,
        south_west_route_capacity,
        hexagon,
        gwrm,
        nwe,
        midland_mainline,
        m4,
        a14,
        ewr_config1,
    ]
    obc_stage = [
        lower_thames_crossing,
        hs2_2a,
        tru,
        east_coast_digital,
        a303,
        a12,
        a428,
        a417,
        a385,
        ftts,
    ]
    sobc_stage = [
        hs2_2b,
        brighton_ml,
        ewr_config3,
        sarh2,
        midlands_rail_hub,
        wrlth,
        a66,
        ewr_config2,
    ]
    rail_infrastructure = [
        crossrail,
        iep,
        gwrm,
        midland_mainline,
        midlands_rail_hub,
        thameslink,
        east_coast_mainline,
        tru,
        wrlth,
        south_west_route_capacity,
        nwe,
        brighton_ml,
    ]


LIST_OF_GROUPS = [  # master.current_projects,
    Projects.he,
    Projects.rail,
    Projects.rail_franchising,
    Projects.hs2,
    Projects.hsmrpg,
    Projects.sarh2,
    Projects.all_not_hs2,
    Projects.fbc_stage,
    Projects.obc_stage,
    Projects.sobc_stage,
]
LIST_OF_TITLES = [
    "ALL",
    "HE",
    "RAIL INFRASTRUCTURE",
    "RAIL FRANCHISING",
    "HS2",
    "HSMRPG",
    "AMIS (SARH2)",
    "ALL, NOT HS2,",
    "FBC Projects",
    "OBC Projects",
    "SOBC Projects",
]

#  list of different baseline types. hold at global level?
BASELINE_TYPES = {
    "Re-baseline this quarter": "quarter",
    "Re-baseline ALB/Programme milestones": "programme_milestones",
    "Re-baseline ALB/Programme cost": "programme_costs",
    "Re-baseline ALB/Programme benefits": "programme_benefits",
    "Re-baseline IPDC milestones": "ipdc_milestones",
    "Re-baseline IPDC cost": "ipdc_costs",
    "Re-baseline IPDC benefits": "ipdc_benefits",
    "Re-baseline HMT milestones": "hmt_milestones",
    "Re-baseline HMT cost": "hmt_costs",
    "Re-baseline HMT benefits": "hmt_benefits",
}

YEAR_LIST = [
    "16-17",
    "17-18",
    "18-19",
    "19-20",
    "20-21",
    "21-22",
    "22-23",
    "23-24",
    "24-25",
    "25-26",
    "26-27",
    "27-28",
    "28-29",
    "29-30",
    "30-31",
    "31-32",
    "32-33",
    "33-34",
    "34-35",
    "35-36",
    "36-37",
    "37-38",
    "38-39",
    "39-40",
]

COST_KEY_LIST = [
    " RDEL Forecast Total",
    " CDEL Forecast one off new costs",
    " Forecast Non-Gov",
]
COST_TYPE_KEY_LIST = [
    (
        "Pre-profile RDEL",
        "Pre-profile CDEL Forecast one off new costs",
        "Pre-profile Forecast Non-Gov",
    ),
    (
        "Total RDEL Forecast Total",
        "Total CDEL Forecast one off new costs",
        "Non-Gov Total Forecast",
    ),
    (
        "Unprofiled RDEL Forecast Total",
        "Unprofiled CDEL Forecast one off new costs",
        "Unprofiled Forecast Non-Gov",
    ),
]
BEN_KEY_LIST = [
    "Pre-profile BEN Total",
    "Total BEN Forecast - Total Monetised Benefits",
    "Unprofiled Remainder BEN Forecast - Total Monetised Benefits",
]
BEN_TYPE_KEY_LIST = [
    (
        "Pre-profile BEN Forecast Gov Cashable",
        "Pre-profile BEN Forecast Gov Non-Cashable",
        "Pre-profile BEN Forecast - Economic (inc Private Partner)",
        "Pre-profile BEN Forecast - Disbenefit UK Economic",
    ),
    (
        "Total BEN Forecast - Gov. Cashable",
        "Total BEN Forecast - Gov. Non-Cashable",
        "Total BEN Forecast - Economic (inc Private Partner)",
        "Total BEN Forecast - Disbenefit UK Economic",
    ),
    (
        "Unprofiled Remainder BEN Forecast - Gov. Cashable",
        "Unprofiled Remainder BEN Forecast - Gov. Non-Cashable",
        "Unprofiled Remainder BEN Forecast - Economic (inc Private Partner)",
        "Unprofiled Remainder BEN Forecast - Disbenefit UK Economic",
    ),
]

# Matplotlib file formats
FILE_FORMATS = [
    "eps",
    "jpeg",
    "jpg",
    "pdf",
    "png",
    "ps",
    "raw",
    "rgba",
    "svg",
    "svgz",
    "tif",
    "tiff",
]

FIGURE_STYLE = {1: "half_horizontal", 2: "full_horizontal"}


def calculate_profiled(p: List[int], s: List[int], unpro: List[int]) -> list:
    """small helper function to calculate the proper profiled amount. This is necessary as
    other wise 'profiled' would actually be the total figure.
    p = profiled list
    s = spent list
    unpro = unprofiled list"""
    f_profiled = []
    for y, amount in enumerate(p):
        t = amount - (s[y] + unpro[y])
        f_profiled.append(t)
    return f_profiled


class Master:
    def __init__(
        self,
        master_data: List[Dict[str, Union[str, int, datetime.date, float]]],
        project_information: Dict[str, Union[str, int]],
    ) -> None:
        self.master_data = master_data
        self.project_information = project_information
        self.current_quarter = self.master_data[0].quarter
        self.abbreviations = {}
        self.current_projects = self.get_current_projects()
        self.bl_info = {}
        self.bl_index = {}
        self.dft_groups = {}
        self.project_stage = {}
        self.fbc_projects = []
        self.get_baseline_data()
        self.check_project_information()
        self.check_baselines()
        self.get_project_abbreviations()
        self.get_project_groups()

    """This is the entry point for all data. It converts a list of excel wbs (note at the moment)
    this is actually done prior to being passed into the Master class. The Master class does a number
    of things. 
    compiles and checks all baseline data for projects. These index reference points. 
    compiles lists of different project groups. e.g stage and DfT Group
    gets a list of projects currently in the portfolio. 
    checks data returned by projects is consistent with whats in project_information
    gets project abbreviations
    
    """
    # function to check abbreviations

    def get_baseline_data(self) -> dict:
        """
        Returns the two dictionaries baseline_info and baseline_index for all projects for all
        baseline types
        """

        baseline_info = {}
        baseline_index = {}

        for b_type in list(BASELINE_TYPES.keys()):
            project_baseline_info = {}
            project_baseline_index = {}
            for name in self.current_projects:
                bc_list = []
                lower_list = []
                for i, master in reversed(list(enumerate(self.master_data))):
                    if name in master.projects:
                        try:
                            approved_bc = master.data[name][b_type]
                            quarter = str(master.quarter)
                        # exception handling in here in case data keys across masters are not consistent.
                        except KeyError:
                            print(
                                str(b_type)
                                + " keys not present in "
                                + str(master.quarter)
                            )
                        if approved_bc == "Yes":
                            bc_list.append(approved_bc)
                            lower_list.append((approved_bc, quarter, i))
                    else:
                        pass
                for i in reversed(range(2)):
                    if name in self.master_data[i].projects:
                        approved_bc = self.master_data[i][name][b_type]
                        quarter = str(self.master_data[i].quarter)
                        lower_list.append((approved_bc, quarter, i))
                    else:
                        quarter = str(self.master_data[i].quarter)
                        lower_list.append((None, quarter, None))

                index_list = []
                for x in lower_list:
                    index_list.append(x[2])

                project_baseline_info[name] = list(reversed(lower_list))
                project_baseline_index[name] = list(reversed(index_list))

            baseline_info[BASELINE_TYPES[b_type]] = project_baseline_info
            baseline_index[BASELINE_TYPES[b_type]] = project_baseline_index

        self.bl_info = baseline_info
        self.bl_index = baseline_index

    def get_current_projects(self) -> list:
        """Returns a list of all the project names in the latest master"""
        return self.master_data[0].projects

    def check_project_information(self) -> None:
        """Checks that project names in master are present/the same as in project info.
        Stops the programme if not"""
        for p in self.current_projects:
            if p not in self.project_information.projects:
                print(
                    p
                    + " is not in the projects information document. Project names must be identical "
                    " in both documents. Programme stopping. Please amend."
                )
                break
            else:
                if p == self.current_projects[-1]:
                    print("The latest master and project information match")

    def check_baselines(self) -> None:
        """checks that projects have the correct baseline information. stops the
        programme if baselines are missing"""
        # work through best way to stop the programme.
        for v in BASELINE_TYPES.values():
            for p in self.current_projects:
                baselines = self.bl_index[v][p]
                if len(baselines) <= 2:
                    print(
                        p
                        + " does not have a baseline point for "
                        + v
                        + " this could cause the programme to "
                        "crash. Therefore the programme is stopping. "
                        "Please amend the data for " + p + " so that "
                        " it has at least one baseline point for " + v
                    )
            else:
                continue
            break

    def get_project_abbreviations(self) -> None:
        """gets the abbreviations for all current projects.
        held in the project info document"""
        for p in self.project_information.projects:
            self.abbreviations[p] = self.project_information[p]["Abbreviations"]

    def get_project_groups(self) -> None:
        """gets the groups that projects are part of e.g. business case
        stage or dft group"""

        raw_dict = {}
        raw_list = []
        group_list = []
        stage_list = []
        for i, master in enumerate(self.master_data):
            lower_dict = {}
            for p in master.projects:
                dft_group = master[p]["DfT Group"]
                stage = master[p]["IPDC approval point"]
                raw_list.append(("group", dft_group))
                raw_list.append(("stage", stage))
                lower_dict[p] = dict(raw_list)
                group_list.append(dft_group)
                stage_list.append(stage)
            raw_dict[str(master.quarter)] = lower_dict  # here

        group_list = list(set(group_list))
        stage_list = list(set(stage_list))

        group_dict = {}
        stage_dict = {}
        for quarter in raw_dict.keys():
            lower_g_dict = {}
            lower_s_dict = {}
            for group_type in group_list:
                g_list = []
                for p in raw_dict[quarter].keys():
                    p_group = raw_dict[quarter][p]["group"]
                    if p_group == group_type:
                        g_list.append(p)

                lower_g_dict[group_type] = g_list

            for stage_type in stage_list:
                s_list = []
                for p in raw_dict[quarter].keys():
                    p_stage = raw_dict[quarter][p]["stage"]
                    if p_stage == stage_type:
                        s_list.append(p)

                lower_s_dict[stage_type] = s_list

            group_dict[quarter] = lower_g_dict
            stage_dict[quarter] = lower_s_dict

        self.dft_groups = group_dict
        self.project_stage = stage_dict
        # self.fbc_projects = stage_dict[str(self.current_quarter)]["Full Business Case"]


#  check cdel cost profile
class CostData:
    def __init__(
        self,
        master: Master,
        project_group: List[str] or str,
        baseline_type: str = "ipdc_costs",
    ):
        self.master = master
        self.project_group = project_group
        self.baseline_type = baseline_type
        self.cat_spent = []
        self.cat_profiled = []
        self.cat_unprofiled = []
        self.spent = []
        self.profiled = []
        self.unprofiled = []
        self.current_profile = []
        self.last_profile = []
        self.baseline_profile_one = []
        self.baseline_profile_two = []
        self.baseline_profile_three = []
        self.rdel_profile = []
        self.cdel_profile = []
        self.ngov_profile = []
        self.y_scale_max = 0
        self.wlc_change = {}
        self.get_cost_totals()
        self.get_cost_profile()
        self.calculate_wlc_change()

    def get_cost_totals(self) -> None:
        """Returns lists containing the sum total of group (of projects) costs,
        sliced in different ways. Cumbersome for loop used at the moment, but
        is the least cumbersome loop I could design!"""
        spent = []
        profiled = []
        unprofiled = []
        group_rdel_spent = 0
        group_cdel_spent = 0
        group_ngov_spent = 0
        group_rdel_profiled = 0
        group_cdel_profiled = 0
        group_ngov_profiled = 0
        group_rdel_unprofiled = 0
        group_cdel_unprofiled = 0
        group_ngov_unprofiled = 0

        self.project_group = string_conversion(self.project_group)

        for i in range(3):
            for x, key in enumerate(COST_TYPE_KEY_LIST):
                group_total = 0
                for project_name in self.project_group:
                    cost_bl_index = self.master.bl_index[self.baseline_type][
                        project_name
                    ]
                    try:
                        rdel = self.master.master_data[cost_bl_index[i]].data[
                            project_name
                        ][key[0]]
                        if rdel is None:
                            rdel = 0

                        cdel = self.master.master_data[cost_bl_index[i]].data[
                            project_name
                        ][key[1]]
                        if cdel is None:
                            cdel = 0

                        ngov = self.master.master_data[cost_bl_index[i]].data[
                            project_name
                        ][key[2]]
                        if ngov is None:
                            ngov = 0

                        total = round(rdel + cdel + ngov)
                        group_total += total
                    except TypeError:  # handle None types, which are present if project not reporting last quarter.
                        rdel = 0
                        cdel = 0
                        ngov = 0
                        total = 0
                        group_total += total

                    if i == 0:  # current quarter
                        if x == 0:  # spent
                            try:  # handling for spend to date figures which are not present in all masters
                                rdel_std = self.master.master_data[
                                    cost_bl_index[i]
                                ].data[project_name]["20-21 RDEL STD one off new costs"]
                                if rdel_std is None:
                                    rdel_std = 0
                                cdel_std = self.master.master_data[
                                    cost_bl_index[i]
                                ].data[project_name]["20-21 CDEL STD one off new costs"]
                                if cdel_std is None:
                                    cdel_std = 0
                                ngov_std = self.master.master_data[
                                    cost_bl_index[i]
                                ].data[project_name]["20-21 CDEL STD Non Gov costs"]
                                if ngov_std is None:
                                    ngov_std = 0
                                group_rdel_spent += round(rdel + rdel_std)
                                group_cdel_spent += round(cdel + cdel_std)
                                group_ngov_spent += round(ngov + ngov_std)
                            except KeyError:
                                group_rdel_spent += rdel
                                group_cdel_spent += cdel
                                group_ngov_spent += ngov
                        if x == 1:  # profiled
                            group_rdel_profiled += rdel
                            group_cdel_profiled += cdel
                            group_ngov_profiled += ngov
                        if x == 2:  # unprofiled
                            group_rdel_unprofiled += rdel
                            group_cdel_unprofiled += cdel
                            group_ngov_unprofiled += ngov

                if x == 0:  # spent
                    try:  # handling for spend to date figures which are not present in all masters
                        rdel_std = self.master.master_data[cost_bl_index[i]].data[
                            project_name
                        ]["20-21 RDEL STD one off new costs"]
                        cdel_std = self.master.master_data[cost_bl_index[i]].data[
                            project_name
                        ]["20-21 CDEL STD one off new costs"]
                        ngov_std = self.master.master_data[cost_bl_index[i]].data[
                            project_name
                        ]["20-21 CDEL STD Non Gov costs"]
                        std_list = [
                            rdel_std,
                            cdel_std,
                            ngov_std,
                        ]  # converts none types to zero
                        for s, std in enumerate(std_list):
                            if std is None:
                                std_list[s] = 0
                        spent.append(round(group_total + sum(std_list)))
                    except (
                        KeyError,
                        TypeError,
                    ):  # Note. TypeError here as projects may have no baseline
                        spent.append(group_total)
                if x == 1:  # profiled
                    profiled.append(group_total)
                if x == 2:  # unprofiled
                    unprofiled.append(group_total)

        cat_spent = [group_rdel_spent, group_cdel_spent, group_ngov_spent]
        cat_profiled = [group_rdel_profiled, group_cdel_profiled, group_ngov_profiled]
        cat_unprofiled = [
            group_rdel_unprofiled,
            group_cdel_unprofiled,
            group_ngov_unprofiled,
        ]
        final_cat_profiled = calculate_profiled(cat_profiled, cat_spent, cat_unprofiled)

        all_profiled = calculate_profiled(profiled, spent, unprofiled)

        self.cat_spent = cat_spent
        self.cat_profiled = final_cat_profiled
        self.cat_unprofiled = cat_unprofiled
        self.spent = spent
        self.profiled = all_profiled
        self.unprofiled = unprofiled
        self.y_scale_max = max(profiled)

    def get_cost_profile(self) -> None:
        """Returns several lists which contain the sum of different cost profiles for the group of project
        contained with the master"""

        current_profile = []
        last_profile = []
        baseline_profile_one = []
        baseline_profile_two = []
        baseline_profile_three = []
        rdel_current_profile = []
        cdel_current_profile = []
        ngov_current_profile = []
        missing_projects = []

        self.project_group = string_conversion(self.project_group)

        for i in range(5):
            yearly_profile = []
            rdel_yearly_profile = []
            cdel_yearly_profile = []
            ngov_yearly_profile = []
            for year in YEAR_LIST:
                cost_total = 0
                rdel_total = 0
                cdel_total = 0
                ngov_total = 0
                for cost_type in COST_KEY_LIST:
                    for project_name in self.project_group:
                        project_bl_index = self.master.bl_index[self.baseline_type][
                            project_name
                        ]
                        try:
                            cost = self.master.master_data[project_bl_index[i]].data[
                                project_name
                            ][year + cost_type]
                            if cost is None:
                                cost = 0
                            cost_total += cost
                        except KeyError:  # to handle data across different financial years
                            # TODO come back and check this is working properly
                            try:
                                cost = self.master.project_information.data[project_name][year + cost_type]
                            except KeyError:
                                cost = 0
                            if cost is None:
                                cost = 0
                            cost_total += cost
                        except TypeError:  # Handles projects not present in the previous quarter
                            missing_projects.append(
                                str(project_name)
                            )  # projects added here. message is below.
                            cost = 0
                            cost_total += cost
                        except IndexError:  # Handles project baseline index
                            # TODO improve this loop
                            if i == 3:
                                try:
                                    cost = self.master.master_data[
                                        project_bl_index[2]
                                    ].data[project_name][year + cost_type]
                                    if cost is None:
                                        cost = 0
                                    cost_total += cost
                                except KeyError:  # to handle data across different financial years
                                    cost = 0
                                    cost_total += cost
                            if i == 4:
                                try:
                                    cost = self.master.master_data[
                                        project_bl_index[3]
                                    ].data[project_name][year + cost_type]
                                    if cost is None:
                                        cost = 0
                                    cost_total += cost
                                except KeyError:  # to handle data across different financial years
                                    cost = 0
                                    cost_total += cost
                                except IndexError:
                                    try:
                                        cost = self.master.master_data[
                                            project_bl_index[2]
                                        ].data[project_name][year + cost_type]
                                        if cost is None:
                                            cost = 0
                                        cost_total += cost
                                    except KeyError:  # to handle data across different financial years
                                        cost = 0
                                        cost_total += cost

                        if cost_type == COST_KEY_LIST[0]:  # rdel
                            rdel_total += cost
                        if cost_type == COST_KEY_LIST[1]:  # cdel
                            cdel_total += cost
                        if cost_type == COST_KEY_LIST[2]:  # ngov
                            ngov_total += cost

                yearly_profile.append(cost_total)
                rdel_yearly_profile.append(rdel_total)
                cdel_yearly_profile.append(cdel_total)
                ngov_yearly_profile.append(ngov_total)

            if i == 0:
                current_profile = yearly_profile
                rdel_current_profile = rdel_yearly_profile
                cdel_current_profile = cdel_yearly_profile
                ngov_current_profile = ngov_yearly_profile
            if i == 1:
                last_profile = yearly_profile
            if i == 2:
                baseline_profile_one = yearly_profile
            if i == 3:
                baseline_profile_two = yearly_profile
            if i == 4:
                baseline_profile_three = yearly_profile

        missing_projects = list(set(missing_projects))  # if TypeError raised above
        if len(missing_projects) is not 0:
            print(
                "NOTE: The following project(s) were not part of the portfolio last quarter "
                + str(missing_projects)
                + " this means current quarter and last quarter cost profiles are not like for like."
                " If you would like a like for like comparison between current and last quarter"
                " remove this project(s) from the master group."
            )

        self.current_profile = current_profile
        self.last_profile = last_profile
        self.baseline_profile_one = baseline_profile_one
        self.baseline_profile_two = baseline_profile_two
        self.baseline_profile_three = baseline_profile_three
        self.rdel_profile = rdel_current_profile
        self.cdel_profile = cdel_current_profile
        self.ngov_profile = ngov_current_profile

    def calculate_wlc_change(self) -> None:
        """calculates changes in whole life cost of project. Current against baselines"""

        wlc_change_dict = {}
        for project_name in self.project_group:
            wlc_list = []
            current_wlc = self.master.master_data[0].data[project_name][
                "Total Forecast"
            ]
            for i in range(1, 3):  # only taking last quarter and first baseline for now
                try:
                    cost_bl_index = self.master.bl_index[self.baseline_type][
                        project_name
                    ]
                    baseline_wlc = self.master.master_data[cost_bl_index[i]].data[
                        project_name
                    ]["Total Forecast"]
                    percentage_change = int(
                        ((current_wlc - baseline_wlc) / current_wlc) * 100
                    )
                    if i == 1:
                        wlc_list.append(("last quarter", percentage_change))
                    if i == 2:
                        wlc_list.append(("baseline one", percentage_change))
                except TypeError:  # handles NoneTypes
                    pass

            wlc_change_dict[project_name] = dict(wlc_list)

        self.wlc_change = wlc_change_dict


class BenefitsData:
    def __init__(
        self,
        master: Master,
        project_group: List[str] or str,
        baseline_type: str = "ipdc_benefits",
    ):
        self.master = master
        self.project_group = project_group
        self.baseline_type = baseline_type
        self.cat_delivered = []
        self.cat_profiled = []
        self.cat_unprofiled = []
        self.delivered = []
        self.profiled = []
        self.unprofiled = []
        self.y_scale_max = 0
        self.y_scale_min = 0
        self.economic_max = 0
        self.get_ben_totals()

    def get_ben_totals(self) -> None:
        """Returns lists containing the sum total of group (of projects) benefits,
        sliced in different ways. Cumbersome for loop used at the moment, but
        is the least cumbersome loop I could design!"""
        delivered = []
        profiled = []
        unprofiled = []
        group_cash_dev = 0
        group_uncash_dev = 0
        group_economic_dev = 0
        group_disben_dev = 0
        group_cash_profiled = 0
        group_uncash_profiled = 0
        group_economic_profiled = 0
        group_disben_profiled = 0
        group_cash_unprofiled = 0
        group_uncash_unprofiled = 0
        group_economic_unprofiled = 0
        group_disben_unprofiled = 0

        self.project_group = string_conversion(self.project_group)

        for i in range(3):
            for x, key in enumerate(BEN_TYPE_KEY_LIST):
                group_total = 0
                for project in self.project_group:
                    ben_bl_index = self.master.bl_index[self.baseline_type][project]
                    try:
                        cash = round(
                            self.master.master_data[ben_bl_index[i]].data[project][
                                key[0]
                            ]
                        )
                        uncash = round(
                            self.master.master_data[ben_bl_index[i]].data[project][
                                key[1]
                            ]
                        )
                        economic = round(
                            self.master.master_data[ben_bl_index[i]].data[project][
                                key[2]
                            ]
                        )
                        disben = round(
                            self.master.master_data[ben_bl_index[i]].data[project][
                                key[3]
                            ]
                        )

                        total = round(cash + uncash + economic + disben)
                        group_total += total
                    except TypeError:  # handle None types, which are present if project not reporting last quarter.
                        cash = 0
                        uncash = 0
                        economic = 0
                        disben = 0
                        total = 0
                        group_total += total

                    if i == 0:  # current quarter
                        if x == 0:  # spent
                            group_cash_dev += cash
                            group_uncash_dev += uncash
                            group_economic_dev += economic
                            group_disben_dev += disben
                        if x == 1:  # profiled
                            group_cash_profiled += cash
                            group_uncash_profiled += uncash
                            group_economic_profiled += economic
                            group_disben_profiled += disben
                        if x == 2:  # unprofiled
                            group_cash_unprofiled += cash
                            group_uncash_unprofiled += uncash
                            group_economic_unprofiled += economic
                            group_disben_unprofiled += disben

                if x == 0:  # spent
                    delivered.append(group_total)
                if x == 1:  # profiled
                    profiled.append(group_total)
                if x == 2:  # unprofiled
                    unprofiled.append(group_total)

        cat_spent = [
            group_cash_dev,
            group_uncash_dev,
            group_economic_dev,
            group_disben_dev,
        ]
        cat_profiled = [
            group_cash_profiled,
            group_uncash_profiled,
            group_economic_profiled,
            group_disben_profiled,
        ]
        cat_unprofiled = [
            group_cash_unprofiled,
            group_uncash_unprofiled,
            group_economic_unprofiled,
            group_disben_unprofiled,
        ]
        final_cat_profiled = calculate_profiled(cat_profiled, cat_spent, cat_unprofiled)
        all_profiled = calculate_profiled(profiled, delivered, unprofiled)

        self.cat_delivered = cat_spent
        self.cat_profiled = final_cat_profiled
        self.cat_unprofiled = cat_unprofiled
        self.delivered = delivered
        self.profiled = all_profiled
        self.unprofiled = unprofiled
        self.y_scale_max = max(profiled)
        self.y_scale_min = min(
            [group_disben_dev, group_disben_profiled, group_disben_unprofiled]
        )
        self.economic_max = max(
            [group_economic_dev, group_economic_unprofiled, group_economic_profiled]
        )


def milestone_info_handling(output_list: list, t_list: list) -> list:
    """helper function for handling and cleaning up milestone date generated
    via MilestoneDate class. Removes none type milestone names and non date
    string values"""
    if t_list[1][1] is not None:
        if isinstance(t_list[3][1], datetime.date):
            return output_list.append(t_list)
        else:
            try:
                d = parser.parse(t_list[3][1], dayfirst=True)
                t_list[3] = ("Date", d.date())
                return output_list.append(t_list)
            # ParserError for non-date string. TypeError for None types
            except (ParserError, TypeError):
                pass


def remove_project_name(project_name: str, milestone_key_list: List[str]) -> List[str]:
    """In this instance project_name is the abbreviation"""
    output_list = []
    for key in milestone_key_list:
        alter_key = key.replace(project_name + ", ", "")
        output_list.append(alter_key)
    return output_list


def remove_none_types(input_list):
    return [x for x in input_list if x is not None]


def get_milestone_date(
        project_name: str,
        milestone_dictionary: Dict[str, Union[datetime.date, str]],
        milestone_name: str,
) -> datetime:
    for k in milestone_dictionary.keys():
        if milestone_dictionary[k]["Project"] == project_name:
            if milestone_dictionary[k]["Milestone"] == milestone_name[1:]:
                return milestone_dictionary[k]["Date"]


class MilestoneData:
    def __init__(
        self,
        master: Master,
        project_group: List[str] or str,
        baseline_type: str = "ipdc_milestones",
    ):
        self.master = master
        self.project_group = project_group
        self.baseline_type = baseline_type
        self.current = {}
        self.last_quarter = {}
        self.baseline_dict = {}
        self.baseline_two = {}
        self.ordered_list_current = []
        self.ordered_list_last = []
        self.ordered_list_bl = []
        self.ordered_list_bl_two = []
        self.key_names = []
        self.key_names_last = []
        self.key_names_baseline = []
        self.type_list = []
        self.md_current = []
        self.md_last = []
        self.md_last_po = []
        self.md_baseline = []
        self.md_baseline_po = []
        self.md_baseline_two = []
        self.max_date = None
        self.min_date = None
        self.schedule_change = {}
        self.schedule_key_last = None
        self.schedule_key_baseline = None
        self.get_milestones()
        self.get_chart_info()
        # self.calculate_schedule_changes()

    def get_milestones(self) -> None:
        """
        Creates project milestone dictionaries for current, last_quarter, and
        baselines when provided with group and baseline type.
        """

        self.project_group = string_conversion(self.project_group)

        for bl in range(4):
            lower_dict = {}
            raw_list = []
            for project_name in self.project_group:
                project_list = []
                milestone_bl_index = self.master.bl_index[self.baseline_type][
                    project_name
                ]
                try:
                    p_data = self.master.master_data[milestone_bl_index[bl]].data[
                        project_name
                    ]
                # IndexError handles len of project bl index.
                # TypeError handles None Type present if project not reporting last quarter
                except (IndexError, TypeError):
                    continue

                # i loops below removes None Milestone names and rejects non-datetime date values.
                for i in range(1, 50):
                    try:
                        t = [
                            ("Project", self.master.abbreviations[project_name]),
                            ("Milestone", p_data["Approval MM" + str(i)]),
                            ("Type", "Approval"),
                            (
                                "Date",
                                p_data["Approval MM" + str(i) + " Forecast / Actual"],
                            ),
                            ("Notes", p_data["Approval MM" + str(i) + " Notes"]),
                        ]
                        milestone_info_handling(project_list, t)
                        t = [
                            ("Project", self.master.abbreviations[project_name]),
                            ("Milestone", p_data["Assurance MM" + str(i)]),
                            ("Type", "Assurance"),
                            (
                                "Date",
                                p_data["Assurance MM" + str(i) + " Forecast - Actual"],
                            ),
                            ("Notes", p_data["Assurance MM" + str(i) + " Notes"]),
                        ]
                        milestone_info_handling(project_list, t)
                    except KeyError:  # handles inconsistent keys naming for approval milestones.
                        try:
                            t = [
                                ("Project", self.master.abbreviations[project_name]),
                                ("Milestone", p_data["Approval MM" + str(i)]),
                                ("Type", "Approval"),
                                (
                                    "Date",
                                    p_data[
                                        "Approval MM" + str(i) + " Forecast - Actual"
                                    ],
                                ),
                                ("Notes", p_data["Approval MM" + str(i) + " Notes"]),
                            ]
                            milestone_info_handling(project_list, t)
                        except KeyError:
                            pass

                # handles inconsistent number of Milestone. could be incorporated above.
                for i in range(18, 67):
                    try:
                        t = [
                            ("Project", self.master.abbreviations[project_name]),
                            ("Milestone", p_data["Project MM" + str(i)]),
                            ("Type", "Delivery"),
                            (
                                "Date",
                                p_data["Project MM" + str(i) + " Forecast - Actual"],
                            ),
                            ("Notes", p_data["Project MM" + str(i) + " Notes"]),
                        ]
                        milestone_info_handling(project_list, t)
                    except KeyError:
                        pass

                # change in Q3. Some milestones collected via HMT approval section.
                # this loop picks them up
                # TODO check these are coming through in q3 data
                for i in range(1, 4):
                    try:
                        t = [
                            ("Project", self.master.abbreviations[project_name]),
                            ("Milestone", p_data["HMT Approval " + str(i)]),
                            ("Type", "Approval"),
                            (
                                "Date",
                                p_data["HMT Approval " + str(i) + " Forecast / Actual"],
                            ),
                            ("Notes", p_data["HMT Approval " + str(i) + " Notes"]),
                        ]
                        milestone_info_handling(project_list, t)
                    except KeyError:
                        pass

                # loop to stop keys names being the same. Done at project level.
                # not particularly concise code.
                upper_counter_list = []
                for entry in project_list:
                    upper_counter_list.append(entry[1][1])
                upper_count = Counter(upper_counter_list)
                lower_counter_list = []
                for entry in project_list:
                    if upper_count[entry[1][1]] > 1:
                        lower_counter_list.append(entry[1][1])
                        lower_count = Counter(lower_counter_list)
                        new_milestone_key = (
                            entry[1][1] + " (" + str(lower_count[entry[1][1]]) + ")"
                        )
                        entry[1] = ("Milestone", new_milestone_key)
                        raw_list.append(entry)
                    else:
                        raw_list.append(entry)

            # puts the list in chronological order
            sorted_list = sorted(raw_list, key=lambda k: (k[3][1] is None, k[3][1]))

            for r in range(len(sorted_list)):
                lower_dict["Milestone " + str(r)] = dict(sorted_list[r])

            if bl == 0:
                self.current = lower_dict
                self.ordered_list_current = sorted_list
            if bl == 1:
                self.last_quarter = lower_dict
                self.ordered_list_last = sorted_list
            if bl == 2:
                self.baseline_dict = lower_dict
                self.ordered_list_bl = sorted_list
            if bl == 3:
                self.baseline_two = lower_dict
                self.ordered_list_bl_two = sorted_list

    def get_chart_info(self) -> None:
        """returns data lists for matplotlib chart"""
        # Note this code could refactored so that it collects all milestones
        # reported across current, last and baseline. At the moment it only
        # uses milestones that are present in the current quarter.
        key_names = []
        key_names_last = []
        keys_names_baseline = []
        md_current = []
        md_last = []
        md_last_po = []  # po is for printout
        md_baseline = []
        md_baseline_po = []  #
        md_baseline_two = []
        type_list = []

        for m in self.current.values():
            m_project = m["Project"]
            m_name = m["Milestone"]
            m_date = m["Date"]
            m_type = m["Type"]
            key_names.append(m_project + ", " + m_name)
            md_current.append(m_date)
            type_list.append(m_type)

            # In two loops below NoneType has to be replaced with a datetime object
            # due to matplotlib being unable to handle NoneTypes when milestone_chart
            # is created. Haven't been able to find a solution to this.
            m_last_date = None
            for m_last in self.last_quarter.values():
                if m_last["Project"] == m_project:
                    if m_last["Milestone"] == m_name:
                        key_names_last.append(m_project + ", " + m_name)
                        m_last_date = m_last["Date"]
                        md_last.append(m_last_date)
                        md_last_po.append(m_last_date)
            if m_last_date is None:
                md_last.append(m_date)
                md_last_po.append(None)

            m_bl_date = None
            for m_bl in self.baseline_dict.values():
                if m_bl["Project"] == m_project:
                    if m_bl["Milestone"] == m_name:
                        keys_names_baseline.append(m_project + ", " + m_name)
                        m_bl_date = m_bl["Date"]
                        md_baseline.append(m_bl_date)
                        md_baseline_po.append(m_bl_date)
            if m_bl_date is None:
                md_baseline.append(m_date)
                md_baseline_po.append(None)

            m_bl_two_date = None
            for m_bl_two in self.baseline_two.values():
                if m_bl_two["Project"] == m_project:
                    if m_bl_two["Milestone"] == m_name:
                        m_bl_date = m_bl_two["Date"]
                        md_baseline_two.append(m_bl_date)
            if m_bl_two_date is None:
                md_baseline_two.append(m_date)

        if len(self.project_group) == 1:
            key_names = remove_project_name(
                self.master.abbreviations[self.project_group[0]], key_names
            )
        else:
            pass

        self.key_names = key_names
        self.key_names_last = key_names_last
        self.key_names_baseline = keys_names_baseline
        self.md_current = md_current
        self.md_last = md_last
        self.md_last_po = md_last_po
        self.md_baseline = md_baseline
        self.md_baseline_po = md_baseline_po
        self.md_baseline_two = md_baseline_two
        self.type_list = type_list
        self.max_date = max(
            remove_none_types(self.md_current)
            + remove_none_types(self.md_last)
            + remove_none_types(self.md_baseline)
        )
        self.min_date = min(
            remove_none_types(self.md_current)
            + remove_none_types(self.md_last)
            + remove_none_types(self.md_baseline)
        )

    def filter_chart_info(
        self,
        milestone_type: str or List[str] = "All",
        key_of_interest: str or List[str] = None,
        start_date: str = "1/1/2000",
        end_date: str = "1/1/2041",
    ):
        # bug handling required in the event that there are no milestones with the filter.
        # i.e. the filter returns no milestones.

        #  Filter milestone type
        milestone_type = string_conversion(milestone_type)
        if milestone_type != ["All"]:  # needs to be list as per string conversion
            for i, v in enumerate(self.type_list):
                if v not in milestone_type:
                    self.key_names[i] = "remove"
                    self.md_current[i] = "remove"
                    self.md_last[i] = "remove"
                    self.md_last_po[i] = "remove"
                    self.md_baseline[i] = "remove"
                    self.md_baseline_po[i] = "remove"
                    self.md_baseline_two[i] = "remove"
                    self.type_list[i] = "remove"
                else:
                    pass

            self.key_names = [x for x in self.key_names if x is not "remove"]
            self.md_current = [x for x in self.md_current if x is not "remove"]
            self.md_last = [x for x in self.md_last if x is not "remove"]
            self.md_last_po = [x for x in self.md_last_po if x is not "remove"]
            self.md_baseline = [x for x in self.md_baseline if x is not "remove"]
            self.md_baseline_po = [x for x in self.md_baseline_po if x is not "remove"]
            self.md_baseline_two = [
                x for x in self.md_baseline_two if x is not "remove"
            ]
            self.type_list = [x for x in self.type_list if x is not "remove"]
        else:
            pass

        #  Filter milestone names of interest
        key_of_interest = string_conversion(key_of_interest)
        filtered_list = []
        if key_of_interest is not None:
            # if developed further clearly good use regex
            for s in key_of_interest:  # s is string
                for v in self.key_names:  # v is value
                    if s in v:
                        filtered_list.append(v)
            for i, v in enumerate(self.key_names):  # fv is filtered value
                if v not in filtered_list:
                    self.key_names[i] = "remove"
                    self.md_current[i] = "remove"
                    self.md_last[i] = "remove"
                    self.md_last_po[i] = "remove"
                    self.md_baseline[i] = "remove"
                    self.md_baseline_po[i] = "remove"
                    self.md_baseline_two[i] = "remove"
                    self.type_list[i] = "remove"
                else:
                    pass
            self.key_names = [x for x in self.key_names if x is not "remove"]
            self.md_current = [x for x in self.md_current if x is not "remove"]
            self.md_last = [x for x in self.md_last if x is not "remove"]
            self.md_last_po = [x for x in self.md_last_po if x is not "remove"]
            self.md_baseline = [x for x in self.md_baseline if x is not "remove"]
            self.md_baseline_po = [x for x in self.md_baseline_po if x is not "remove"]
            self.md_baseline_two = [
                x for x in self.md_baseline_two if x is not "remove"
            ]
            self.type_list = [x for x in self.type_list if x is not "remove"]
        else:
            pass

        #  Filter milestones based on date.
        start = parser.parse(start_date, dayfirst=True)
        end = parser.parse(end_date, dayfirst=True)
        for i, d in enumerate(self.md_current):
            if start.date() <= d <= end.date():
                pass
            else:
                self.key_names[i] = "remove"
                self.md_current[i] = "remove"
                self.md_last[i] = "remove"
                self.md_last_po[i] = "remove"
                self.md_baseline[i] = "remove"
                self.md_baseline_po[i] = "remove"
                self.md_baseline_two[i] = "remove"
                self.type_list[i] = "remove"
        self.key_names = [x for x in self.key_names if x is not "remove"]
        self.md_current = [x for x in self.md_current if x is not "remove"]
        self.md_last = [x for x in self.md_last if x is not "remove"]
        self.md_last_po = [x for x in self.md_last_po if x is not "remove"]
        self.md_baseline = [x for x in self.md_baseline if x is not "remove"]
        self.md_baseline_po = [x for x in self.md_baseline_po if x is not "remove"]
        self.md_baseline_two = [x for x in self.md_baseline_two if x is not "remove"]
        self.type_list = [x for x in self.type_list if x is not "remove"]

        self.max_date = max(
            remove_none_types(self.md_current)
            + remove_none_types(self.md_last)
            + remove_none_types(self.md_baseline)
        )

        self.min_date = min(
            remove_none_types(self.md_current)
            + remove_none_types(self.md_last)
            + remove_none_types(self.md_baseline)
        )

    def calculate_schedule_changes(self) -> None:
        """calculates the changes in project schedules. If standard key for calculation
        not available it using the best next one available"""

        self.project_group = string_conversion(self.project_group)
        self.filter_chart_info(milestone_type=["Delivery", "Approval"])

        def schedule_info(
            project_name: str,
            other_key_list: List[str],
            c_key_list: List[str],
            other_dict: dict,
            current_dict: dict,
            dict_label: str,
        ):
            output_dict = {}
            schedule_info = []
            for key in reversed(other_key_list):
                if key in c_key_list:
                    sop = get_milestone_date(
                        project_name, other_dict, " Start of Project"
                    )
                    if sop is None:
                        sop = get_milestone_date(
                            project_name, current_dict, other_key_list[0]
                        )
                        schedule_info.append(("start key", other_key_list[0]))
                    else:
                        schedule_info.append(("start key", " Start of Project"))
                    schedule_info.append(("start", sop))
                    schedule_info.append(("end key", key))
                    date = get_milestone_date(project_name, current_dict, key)
                    schedule_info.append(("end current date", date))
                    other_date = get_milestone_date(project_name, other_dict, key)
                    schedule_info.append(("end other date", other_date))
                    project_length = (other_date - sop).days
                    schedule_info.append(("project length", project_length))
                    change = (date - other_date).days
                    schedule_info.append(("change", change))
                    p_change = int((change / project_length) * 100)
                    schedule_info.append(("percent change", p_change))
                    output_dict[dict_label] = dict(schedule_info)
                    break

            return output_dict

        output_dict = {}
        for project_name in self.project_group:
            project_name = self.master.abbreviations[project_name]
            current_key_list = []
            last_key_list = []
            baseline_key_list = []
            for key in self.key_names:
                try:
                    p = key.split(",")[0]
                    milestone_key = key.split(",")[1]
                    if project_name == p:
                        if milestone_key != " Project - Business Case End Date":
                            current_key_list.append(milestone_key)
                except IndexError:
                    # patch of single project group. In this instance the project name
                    # is removed from the key_name via remove_project_name function as
                    # part of get chart info.
                    if len(self.project_group) == 1:
                        current_key_list.append(key)
            for last_key in self.key_names_last:
                p = last_key.split(",")[0]
                milestone_key_last = last_key.split(",")[1]
                if project_name == p:
                    if milestone_key_last != " Project - Business Case End Date":
                        last_key_list.append(milestone_key_last)
            for baseline_key in self.key_names_baseline:
                p = baseline_key.split(",")[0]
                milestone_key_baseline = baseline_key.split(",")[1]
                if project_name == p:
                    if (
                        milestone_key_baseline
                        != " Project - Business Case End Date"
                        # and milestone_key_baseline != " Project End Date"
                    ):
                        baseline_key_list.append(milestone_key_baseline)

            b_dict = schedule_info(
                project_name,
                baseline_key_list,
                current_key_list,
                self.baseline_dict,
                self.current,
                "baseline",
            )
            l_dict = schedule_info(
                project_name,
                last_key_list,
                current_key_list,
                self.last_quarter,
                self.current,
                "last",
            )
            lower_dict = {**b_dict, **l_dict}

            # baseline_schedule_info = []
            # for b_key in reversed(baseline_key_list):
            #     if b_key in current_key_list:
            #         sop = get_milestone_date(project_name, self.baseline, " Start of Project")
            #         if sop is None:
            #             sop = get_milestone_date(project_name, self.current, baseline_key_list[0])
            #             baseline_schedule_info.append(("start key", baseline_key_list[0]))
            #         else:
            #             baseline_schedule_info.append(("start key", " Start of Project"))
            #         baseline_schedule_info.append(("start", sop))
            #         baseline_schedule_info.append(("end key", b_key))
            #         date = get_milestone_date(project_name, self.current, b_key)
            #         baseline_schedule_info.append(("end current date", date))
            #         b_date = get_milestone_date(project_name, self.baseline, b_key)
            #         baseline_schedule_info.append(("end baseline date", b_date))
            #         project_length = (b_date - sop).days
            #         baseline_schedule_info.append(("project length", project_length))
            #         change = (date - b_date).days
            #         baseline_schedule_info.append(("change", change))
            #         p_change = int((change / project_length) * 100)
            #         baseline_schedule_info.append(("percent change", p_change))
            #         lower_dict["baseline"] = dict(baseline_schedule_info)
            #         break

            # last_schedule_info = []
            # for l_key in reversed(last_key_list):
            #     last_schedule_info.append(("start key", sop))
            #     if l_key in current_key_list:
            #         date = get_milestone_date(project_name, self.current, l_key)
            #         l_date = get_milestone_date(project_name, self.last_quarter, l_key)
            #         lower_dict["last"] = l_key
            #         break

            output_dict[project_name] = lower_dict

        self.schedule_change = output_dict


# class CombinedData:
#     def __init__(self, wb, pfm_milestone_data):
#         self.wb = wb
#         self.pfm_milestone_data = pfm_milestone_data
#         # self.project_current = {}
#         # self.project_last = {}
#         # self.project_baseline = {}
#         # self.project_baseline_two = {}
#         self.group_current = {}
#         self.group_last = {}
#         self.group_baseline = {}
#         self.group_baseline_two = {}
#         self.combined_tuple_list_forecast = []
#         self.combined_tuple_list_baseline = []
#         self.combine_mi_pfm_data()
#
#     def combine_mi_pfm_data(self):
#         """
#         coverts data from MI system into usable format for graphical outputs
#         """
#         ws = self.wb.active
#
#         mi_milestone_name_list = []  # handles duplicates
#         mi_tuple_list_forecast = []
#         mi_tuple_list_baseline = []
#         for r in range(4, ws.max_row + 1):
#             mi_milestone_key_name_raw = ws.cell(row=r, column=3).value
#             mi_milestone_key_name = "MI, " + mi_milestone_key_name_raw
#             forecast_date = ws.cell(row=r, column=8).value
#             baseline_date = ws.cell(row=r, column=9).value
#             notes = ws.cell(row=r, column=10).value
#             if mi_milestone_key_name not in mi_milestone_name_list:
#                 mi_milestone_name_list.append(mi_milestone_key_name)
#                 mi_tuple_list_forecast.append(
#                     (mi_milestone_key_name, forecast_date.date(), notes)
#                 )
#                 mi_tuple_list_baseline.append(
#                     (mi_milestone_key_name, baseline_date.date(), notes)
#                 )
#             else:
#                 for i in range(
#                         2, 15
#                 ):  # alters duplicates by adding number to end of keys
#                     mi_altered_milestone_key_name = mi_milestone_key_name + " " + str(i)
#                     if mi_altered_milestone_key_name in mi_milestone_name_list:
#                         continue
#                     else:
#                         mi_tuple_list_forecast.append(
#                             (mi_altered_milestone_key_name, forecast_date.date(), notes)
#                         )
#                         mi_tuple_list_baseline.append(
#                             (mi_altered_milestone_key_name, baseline_date.date(), notes)
#                         )
#                         mi_milestone_name_list.append(mi_altered_milestone_key_name)
#                         break
#
#         mi_tuple_list_forecast = sorted(
#             mi_tuple_list_forecast, keys=lambda k: (k[1] is None, k[1])
#         )  # put the list in chronological order
#         mi_tuple_list_baseline = sorted(
#             mi_tuple_list_baseline, keys=lambda k: (k[1] is None, k[1])
#         )  # put the list in chronological order
#
#         pfm_tuple_list_forecast = []
#         pfm_tuple_list_baseline = []
#         for data in self.pfm_milestone_data.ordered_list_bl_two:
#             pfm_tuple_list_forecast.append(("PfM, " + data[0], data[1], data[2]))
#         for data in self.pfm_milestone_data.group_choronological_list_baseline:
#             pfm_tuple_list_baseline.append(("PfM, " + data[0], data[1], data[2]))
#
#         combined_tuple_list_forecast = mi_tuple_list_forecast + pfm_tuple_list_forecast
#         combined_tuple_list_baseline = mi_tuple_list_baseline + pfm_tuple_list_baseline
#
#         combined_tuple_list_forecast = sorted(
#             combined_tuple_list_forecast, keys=lambda k: (k[1] is None, k[1])
#         )  # put the list in chronological order
#         combined_tuple_list_baseline = sorted(
#             combined_tuple_list_baseline, keys=lambda k: (k[1] is None, k[1])
#         )  # put the list in chronological order
#
#         milestone_dict_forecast = {}
#         for series_one in combined_tuple_list_forecast:
#             if series_one[0] is not None:
#                 milestone_dict_forecast[series_one[0]] = {series_one[1]: series_one[2]}
#         milestone_dict_baseline = {}
#         for series_one in combined_tuple_list_baseline:
#             if series_one[0] is not None:
#                 milestone_dict_baseline[series_one[0]] = {series_one[1]: series_one[2]}
#
#         self.group_current = milestone_dict_forecast
#         self.group_last = {}
#         self.group_baseline = milestone_dict_baseline
#         self.group_baseline_two = {}
#         self.combined_tuple_list_forecast = combined_tuple_list_forecast
#         self.combined_tuple_list_baseline = combined_tuple_list_baseline
#
#
# class MilestoneCharts:
#     def __init__(
#             self,
#             latest_milestone_names,
#             latest_milestone_dates,
#             last_milestone_dates,
#             baseline_milestone_dates,
#             graph_title,
#             ipdc_date,
#     ):
#         self.latest_milestone_names = latest_milestone_names
#         self.latest_milestone_dates = latest_milestone_dates
#         self.last_milestone_dates = last_milestone_dates
#         self.baseline_milestone_dates = baseline_milestone_dates
#         self.graph_title = graph_title
#         self.ipdc_date = ipdc_date
#         # self.milestone_swimlane_charts()
#         self.build_charts()
#
#     def milestone_swimlane_charts(self):
#         # build scatter chart
#         fig, ax1 = plt.subplots()
#         fig.suptitle(self.graph_title, fontweight="bold")  # title
#         # set fig size
#         fig.set_figheight(4)
#         fig.set_figwidth(8)
#
#         ax1.scatter(
#             self.baseline_milestone_dates, self.latest_milestone_names, label="Baseline"
#         )
#         ax1.scatter(
#             self.last_milestone_dates, self.latest_milestone_names, label="Last Qrt"
#         )
#         ax1.scatter(
#             self.latest_milestone_dates, self.latest_milestone_names, label="Latest Qrt"
#         )
#
#         # format the series_one ticks
#         years = mdates.YearLocator()  # every year
#         months = mdates.MonthLocator()  # every month
#         years_fmt = mdates.DateFormatter("%Y")
#         months_fmt = mdates.DateFormatter("%b")
#
#         # calculate the length of the time period covered in chart. Not perfect as baseline dates can distort.
#         try:
#             td = (self.latest_milestone_dates[-1] - self.latest_milestone_dates[0]).days
#             if td <= 365 * 3:
#                 ax1.xaxis.set_major_locator(years)
#                 ax1.xaxis.set_minor_locator(months)
#                 ax1.xaxis.set_major_formatter(years_fmt)
#                 ax1.xaxis.set_minor_formatter(months_fmt)
#                 plt.setp(ax1.xaxis.get_minorticklabels(), rotation=45)
#                 plt.setp(
#                     ax1.xaxis.get_majorticklabels(), rotation=45, weight="bold"
#                 )  # milestone_swimlane_charts(key_name,
#                 #                           current_m_data,
#                 #                           last_m_data,
#                 #                           baseline_m_data,
#                 #                           'All Milestones')
#                 # scaling series_one axis
#                 # series_one axis value to no more than three months after last latest milestone date, or three months
#                 # before first latest milestone date. Hack, can be improved. Text highlights movements off chart.
#                 x_max = self.latest_milestone_dates[-1] + timedelta(days=90)
#                 x_min = self.latest_milestone_dates[0] - timedelta(days=90)
#                 for date in self.baseline_milestone_dates:
#                     if date > x_max:
#                         ax1.set_xlim(x_min, x_max)
#                         plt.figtext(
#                             0.98,
#                             0.03,
#                             "Check full schedule to see all milestone movements",
#                             horizontalalignment="right",
#                             fontsize=6,
#                             fontweight="bold",
#                         )
#                     if date < x_min:
#                         ax1.set_xlim(x_min, x_max)
#                         plt.figtext(
#                             0.98,
#                             0.03,
#                             "Check full schedule to see all milestone movements",
#                             horizontalalignment="right",
#                             fontsize=6,
#                             fontweight="bold",
#                         )
#             else:
#                 ax1.xaxis.set_major_locator(years)
#                 ax1.xaxis.set_minor_locator(months)
#                 ax1.xaxis.set_major_formatter(years_fmt)
#                 plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, weight="bold")
#         except IndexError:  # if milestone dates list is empty:
#             pass
#
#         ax1.legend()  # insert legend
#
#         # reverse series_two axis so order is earliest to oldest
#         ax1 = plt.gca()
#         ax1.set_ylim(ax1.get_ylim()[::-1])
#         ax1.tick_params(axis="series_two", which="major", labelsize=7)
#         ax1.yaxis.grid()  # horizontal lines
#         ax1.set_axisbelow(True)
#         # ax1.get_yaxis().set_visible(False)
#
#         # for i, txt in enumerate(latest_milestone_names):
#         #     ax1.annotate(txt, (i, latest_milestone_dates[i]))
#
#         # Add line of IPDC date, but only if in the time period
#         try:
#             if (
#                     self.latest_milestone_dates[0]
#                     <= self.ipdc_date
#                     <= self.latest_milestone_dates[-1]
#             ):
#                 plt.axvline(self.ipdc_date)
#                 plt.figtext(
#                     0.98,
#                     0.01,
#                     "Line represents when IPDC will discuss Q1 20_21 portfolio management report",
#                     horizontalalignment="right",
#                     fontsize=6,
#                     fontweight="bold",
#                 )
#         except IndexError:
#             pass
#
#         # size of chart and fit
#         fig.canvas.draw()
#         fig.tight_layout(rect=[0, 0.03, 1, 0.95])  # for title
#
#         fig.savefig(
#             root_path / "output/{}.png".format(self.graph_title), bbox_inches="tight"
#         )
#
#         # plt.close() #automatically closes figure so don't need to do manually.
#
#     def build_charts(self):
#
#         # add \n to series_two axis labels and cut down if two long
#         # labels = ['\n'.join(wrap(l, 40)) for l in latest_milestone_names]
#         labels = self.latest_milestone_names
#         final_labels = []
#         for l in labels:
#             if len(l) > 40:
#                 final_labels.append(l[:35])
#             else:
#                 final_labels.append(l)
#
#         # Chart
#         no_milestones = len(self.latest_milestone_names)
#
#         if no_milestones <= 30:
#             (
#                 np.array(final_labels),
#                 np.array(self.latest_milestone_dates),
#                 np.array(self.last_milestone_dates),
#                 np.array(self.baseline_milestone_dates),
#                 self.graph_title,
#                 self.ipdc_date,
#             )
#
#         if 31 <= no_milestones <= 60:
#             half = int(no_milestones / 2)
#             MilestoneCharts(
#                 np.array(final_labels[:half]),
#                 np.array(self.latest_milestone_dates[:half]),
#                 np.array(self.last_milestone_dates[:half]),
#                 np.array(self.baseline_milestone_dates[:half]),
#                 self.graph_title,
#                 self.ipdc_date,
#             )
#             title = self.graph_title + " cont."
#             MilestoneCharts(
#                 np.array(final_labels[half:no_milestones]),
#                 np.array(self.latest_milestone_dates[half:no_milestones]),
#                 np.array(self.last_milestone_dates[half:no_milestones]),
#                 np.array(self.baseline_milestone_dates[half:no_milestones]),
#                 title,
#                 self.ipdc_date,
#             )
#
#         if 61 <= no_milestones <= 90:
#             third = int(no_milestones / 3)
#             MilestoneCharts(
#                 np.array(final_labels[:third]),
#                 np.array(self.latest_milestone_dates[:third]),
#                 np.array(self.last_milestone_dates[:third]),
#                 np.array(self.baseline_milestone_dates[:third]),
#                 self.graph_title,
#                 self.ipdc_date,
#             )
#             title = self.graph_title + " cont. 1"
#             MilestoneCharts(
#                 np.array(final_labels[third: third * 2]),
#                 np.array(self.latest_milestone_dates[third: third * 2]),
#                 np.array(self.last_milestone_dates[third: third * 2]),
#                 np.array(self.baseline_milestone_dates[third: third * 2]),
#                 title,
#                 self.ipdc_date,
#             )
#             title = self.graph_title + " cont. 2"
#             MilestoneCharts(
#                 np.array(final_labels[third * 2: no_milestones]),
#                 np.array(self.latest_milestone_dates[third * 2: no_milestones]),
#                 np.array(self.last_milestone_dates[third * 2: no_milestones]),
#                 np.array(self.baseline_milestone_dates[third * 2: no_milestones]),
#                 title,
#                 self.ipdc_date,
#             )
#         pass
#


def vfm_matplotlib_graph(labels, current_qrt, last_qrt, title):
    #  Need to split this strings over two lines on series_one axis
    for n, i in enumerate(labels):
        if i == "Very High and Financially Positive":
            labels[n] = "Very High and \n Financially Positive"
        if i == "Economically Positive":
            labels[n] = "Economically \n Positive"

    x = np.arange(len(labels))  # the label locations
    width = 0.35  # the width of the bars

    fig, ax = plt.subplots()
    rects_one = ax.bar(x - width / 2, current_qrt, width, label="This quarter")
    rects_two = ax.bar(x + width / 2, last_qrt, width, label="Last quarter")

    # Add some text for labels, title and custom series_one-axis tick labels, etc.
    # ax.set_ylabel('Number')
    ax.set_title(title)
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    # Rotate the tick labels and set their alignment.
    # plt.setp(ax.get_xticklabels(), alignment=)
    ax.legend()

    def autolabel(rects):
        """Attach a text label above each bar in *rects*, displaying its height."""
        for rect in rects:
            height = rect.get_height()
            ax.annotate(
                "{}".format(height),
                xy=(rect.get_x() + rect.get_width() / 2, height),
                xytext=(0, 3),  # 3 points vertical offset
                textcoords="offset points",
                ha="center",
                va="bottom",
            )

    autolabel(rects_one)
    autolabel(rects_two)

    fig.tight_layout()

    fig.savefig(root_path / "output/{}.png".format(title), bbox_inches="tight")

    plt.show()


def set_figure_size(graph_type: str) -> Tuple[int, int]:
    if graph_type == "half_horizontal":
        return 11.69, 4.10
    if graph_type == "full_horizontal":
        return 11.69, 8.20


def cost_profile_graph(cost_master: CostData, **kwargs) -> plt.figure:
    """Compiles a matplotlib line chart for costs of GROUP of projects contained within cost_master class"""

    fig, (ax1) = plt.subplots(1)  # two subplots for this chart

    # fig size
    try:
        fig_size = kwargs["fig_size"]
        fig.set_size_inches(set_figure_size(fig_size))
    except KeyError:
        fig.set_size_inches(set_figure_size(FIGURE_STYLE[2]))
        pass

    # title
    if len(cost_master.project_group) == 1:
        title = cost_master.master.abbreviations[cost_master.project_group[0]] + " cost profile change"
    else:
        try:
            title = kwargs["title"] + " cost profile change"
        except KeyError:
            pass
            title = ""
            print("You need to provide a title for this chart")

    plt.suptitle(title, fontweight="bold", fontsize=25)

    # Overall cost profile chart
    if (
        sum(cost_master.baseline_profile_one) != 0 or cost_master.baseline_profile_one == []
    ):  # handling in the event that group of projects have no baseline profile.
        ax1.plot(
            YEAR_LIST,
            np.array(cost_master.baseline_profile_one),  # baseline profile
            label="Baseline",
            linewidth=5.0,
            marker="o",
        )
    else:
        pass
    if (
        sum(cost_master.last_profile) != 0 or cost_master.last_profile == [] or cost_master.last_profile != cost_master.baseline_profile_one
    ):  # handling for no cost profile, project not present last quarter and last/baseline profiles being the same.
        ax1.plot(
            YEAR_LIST,
            np.array(cost_master.last_profile),  # last quarter profile
            label="Last quarter",
            linewidth=5.0,
            marker="o",
        )
    else:
        pass
    ax1.plot(
        YEAR_LIST,
        np.array(cost_master.current_profile),  # current profile
        label="Latest",
        linewidth=5.0,
        marker="o",
    )

    # Chart styling
    plt.xticks(rotation=45, size=14)
    plt.yticks(size=14)
    ax1.tick_params(axis="series_one", which="major")
    ax1.set_ylabel("Cost (£m)")
    ylab1 = ax1.yaxis.get_label()
    ylab1.set_style("italic")
    ylab1.set_size(16)
    ax1.grid(color="grey", linestyle="-", linewidth=0.2)
    ax1.legend(prop={"size": 16})
    # ax1.set_title(
    #     "Change in project cost profile",
    #     loc="left",
    #     fontsize=12,
    #     fontweight="bold",
    # )

    # # plot rdel, cdel, non-gov chart data
    # if (
    #         sum(cost_master.ngov_profile) != 0
    # ):  # if statement as most projects don't have ngov cost.
    #     ax2.plot(
    #         YEAR_LIST,
    #         np.array(cost_master.ngov_profile),
    #         label="Non-Gov",
    #         linewidth=3.0,
    #         marker="o",
    #     )
    # ax2.plot(
    #     YEAR_LIST,
    #     np.array(cost_master.cdel_profile),
    #     label="CDEL",
    #     linewidth=3.0,
    #     marker="o",
    # )
    # ax2.plot(
    #     YEAR_LIST,
    #     np.array(cost_master.rdel_profile),
    #     label="RDEL",
    #     linewidth=3.0,
    #     marker="o",
    # )
    #
    # # rdel/cdel profile chart styling
    # ax2.tick_params(axis="series_one", which="major", labelsize=6, rotation=45)
    # ax2.set_xlabel("Financial Years")
    # ax2.set_ylabel("Cost (£m)")
    # xlab2 = ax2.xaxis.get_label()
    # ylab2 = ax2.yaxis.get_label()
    # xlab2.set_style("italic")
    # xlab2.set_size(8)
    # ylab2.set_style("italic")
    # ylab2.set_size(8)
    # ax2.grid(color="grey", linestyle="-", linewidth=0.2)
    # ax2.legend(prop={"size": 6})
    # ax2.set_title(
    #     "Fig 2 - current cost type profile", loc="left", fontsize=8, fontweight="bold"
    # )

    fig.tight_layout(rect=[0, 0.03, 1, 0.95])  # size/fit of chart

    try:
        kwargs["show"] == "No"
    except KeyError:
        plt.show()

    return fig


def cost_profile_baseline_graph(
    cost_master: CostData, *title: Tuple[Optional[str]]
) -> plt.figure:
    """Compiles a matplotlib line chart for costs of GROUP of projects contained within cost_master class.
    As as default last quarters profile is not included. It creates two plots. First plot shows overall
    profile in current, last quarters anb baseline form. Second plot shows rdel, cdel, and 'non-gov' cost profile"""

    fig, (ax1, ax2) = plt.subplots(2)  # two subplots for this chart

    """cost profile charts"""
    if len(cost_master.entity) == 1:
        fig.suptitle(cost_master.entity[0] + " Cost Profile", fontweight="bold")
    else:
        fig.suptitle(title[0] + " Cost Profile", fontweight="bold")  # title

    # Overall cost profile chart
    if (
        sum(cost_master.baseline_profile_three) != 0
    ):  # handling in the event that group of projects have no baseline profile.
        ax1.plot(
            YEAR_LIST,
            np.array(cost_master.baseline_profile_three),  # baseline profile
            label="Baseline 3",
            linewidth=3.0,
            marker="o",
        )
    else:
        pass
    if (
        sum(cost_master.baseline_profile_two) != 0
    ):  # handling in the event that group of projects have no baseline profile.
        ax1.plot(
            YEAR_LIST,
            np.array(cost_master.baseline_profile_two),  # baseline profile
            label="Baseline 2",
            linewidth=3.0,
            marker="o",
        )
    else:
        pass
    if (
        sum(cost_master.baseline_profile_one) != 0
    ):  # handling in the event that group of projects have no last quarter profile
        ax1.plot(
            YEAR_LIST,
            np.array(cost_master.baseline_profile_one),  # last quarter profile
            label="Baseline 1",
            linewidth=3.0,
            marker="o",
        )
    else:
        pass
    ax1.plot(
        YEAR_LIST,
        np.array(cost_master.current_profile),  # current profile
        label="Latest",
        linewidth=3.0,
        marker="o",
    )

    # Chart styling
    ax1.tick_params(axis="series_one", which="major", labelsize=6, rotation=45)
    ax1.set_ylabel("Cost (£m)")
    ylab1 = ax1.yaxis.get_label()
    ylab1.set_style("italic")
    ylab1.set_size(8)
    ax1.grid(color="grey", linestyle="-", linewidth=0.2)
    ax1.legend(prop={"size": 6})
    ax1.set_title(
        "Fig 1 - cost profile changes", loc="left", fontsize=8, fontweight="bold"
    )

    # plot rdel, cdel, non-gov chart data
    if (
        sum(cost_master.ngov_profile) != 0
    ):  # if statement as most projects don't have ngov cost.
        ax2.plot(
            YEAR_LIST,
            np.array(cost_master.ngov_profile),
            label="Non-Gov",
            linewidth=3.0,
            marker="o",
        )
    ax2.plot(
        YEAR_LIST,
        np.array(cost_master.cdel_profile),
        label="CDEL",
        linewidth=3.0,
        marker="o",
    )
    ax2.plot(
        YEAR_LIST,
        np.array(cost_master.rdel_profile),
        label="RDEL",
        linewidth=3.0,
        marker="o",
    )

    # rdel/cdel profile chart styling
    ax2.tick_params(axis="series_one", which="major", labelsize=6, rotation=45)
    ax2.set_xlabel("Financial Years")
    ax2.set_ylabel("Cost (£m)")
    xlab2 = ax2.xaxis.get_label()
    ylab2 = ax2.yaxis.get_label()
    xlab2.set_style("italic")
    xlab2.set_size(8)
    ylab2.set_style("italic")
    ylab2.set_size(8)
    ax2.grid(color="grey", linestyle="-", linewidth=0.2)
    ax2.legend(prop={"size": 6})
    ax2.set_title(
        "Fig 2 - current cost type profile", loc="left", fontsize=8, fontweight="bold"
    )

    plt.show()

    return fig


def spent_calculation(
    master: Dict[str, Union[str, datetime.date, int, float]], project: str
) -> int:
    keys = [
        "Pre-profile RDEL",
        "20-21 RDEL STD Total",
        "Pre-profile CDEL Forecast one off new costs",
        "20-21 CDEL STD Total",
        "Pre-profile Forecast Non-Gov",
        "20-21 CDEL STD Non Gov costs",
    ]

    total = 0
    for k in keys:
        try:
            total += master.data[project][k]
        except TypeError:  # None types
            pass

    return total


def open_word_doc(wd_path: str) -> Document:
    """Function stores an empty word doc as a variable"""
    return Document(wd_path)


def get_word_doc() -> Document():
    """returns the summary temp doc"""
    wd_path = root_path / "input/summary_temp.docx"
    return open_word_doc(wd_path)


def wd_heading(
    doc: Document, project_info: Dict[str, Union[str, int]], project_name: str
) -> None:
    """Function adds header to word doc"""
    font = doc.styles["Normal"].font
    font.name = "Arial"
    font.size = Pt(12)

    heading = str(
        project_info.data[project_name]["Abbreviations"]
    )  # integrate into master
    intro = doc.add_heading(str(heading), 0)
    intro.alignment = 1
    intro.bold = True


def key_contacts(doc: Document, master: Master, project_name: str) -> None:
    """Function adds keys contact details"""
    sro_name = master.master_data[0].data[project_name][
        "Senior Responsible Owner (SRO)"
    ]
    if sro_name is None:
        sro_name = "tbc"

    sro_email = master.master_data[0].data[project_name][
        "Senior Responsible Owner (SRO) - Email"
    ]
    if sro_email is None:
        sro_email = "email: tbc"

    sro_phone = master.master_data[0].data[project_name]["SRO Phone No."]
    if sro_phone == None:
        sro_phone = "phone number: tbc"

    doc.add_paragraph(
        "SRO: " + str(sro_name) + ", " + str(sro_email) + ", " + str(sro_phone)
    )

    pd_name = master.master_data[0].data[project_name]["Project Director (PD)"]
    if pd_name is None:
        pd_name = "TBC"

    pd_email = master.master_data[0].data[project_name]["Project Director (PD) - Email"]
    if pd_email is None:
        pd_email = "email: tbc"

    pd_phone = master.master_data[0].data[project_name]["PD Phone No."]
    if pd_phone is None:
        pd_phone = "TBC"

    doc.add_paragraph(
        "PD: " + str(pd_name) + ", " + str(pd_email) + ", " + str(pd_phone)
    )

    contact_name = master.master_data[0].data[project_name]["Working Contact Name"]
    if contact_name is None:
        contact_name = "TBC"

    contact_email = master.master_data[0].data[project_name]["Working Contact Email"]
    if contact_email is None:
        contact_email = "email: tbc"

    contact_phone = master.master_data[0].data[project_name][
        "Working Contact Telephone"
    ]
    if contact_phone is None:
        contact_phone = "TBC"

    doc.add_paragraph(
        "PfM reporting lead: "
        + str(contact_name)
        + ", "
        + str(contact_email)
        + ", "
        + str(contact_phone)
    )


def dca_table(doc: Document, master: Master, project_name: str) -> None:
    """Creates SRO confidence table"""
    w_table = doc.add_table(rows=1, cols=5)
    hdr_cells = w_table.rows[0].cells
    hdr_cells[0].text = "Delivery confidence"
    hdr_cells[1].text = "This quarter"
    hdr_cells[2].text = str(master.master_data[1].quarter)
    hdr_cells[3].text = str(master.master_data[2].quarter)
    hdr_cells[4].text = str(master.master_data[3].quarter)

    for x, dca_key in enumerate(SRO_CONF_KEY_LIST):
        row_cells = w_table.add_row().cells
        row_cells[0].text = dca_key
        for i, m in enumerate(master.master_data[:4]):  # last four masters taken
            try:
                rating = convert_rag_text(m.data[project_name][dca_key])
                row_cells[i + 1].text = rating
                cell_colouring(row_cells[i + 1], rating)
            except (KeyError, TypeError):
                row_cells[i + 1].text = "N/A"

    w_table.style = "Table Grid"
    make_rows_bold([w_table.rows[0]])  # makes top of table bold.
    # make_columns_bold([table.columns[0]]) #right cells in table bold
    column_widths = (Cm(3.9), Cm(2.9), Cm(2.9), Cm(2.9), Cm(2.9))
    set_col_widths(w_table, column_widths)


def dca_narratives(doc: Document, master: Master, project_name: str) -> None:
    """Places all narratives into document and checks for differences between
    current and last quarter"""

    doc.add_paragraph()
    p = doc.add_paragraph()
    text = "*Red text highlights changes in narratives from last quarter"
    p.add_run(text).font.color.rgb = RGBColor(255, 0, 0)

    headings_list = [
        "SRO delivery confidence narrative",
        "Financial cost narrative",
        "Financial comparison with last quarter",
        "Financial comparison with baseline",
        "Benefits Narrative",
        "Benefits comparison with last quarter",
        "Benefits comparison with baseline",
        "Milestone narrative",
    ]

    narrative_keys_list = [
        "Departmental DCA Narrative",
        "Project Costs Narrative",
        "Cost comparison with last quarters cost narrative",
        "Cost comparison within this quarters cost narrative",
        "Benefits Narrative",
        "Ben comparison with last quarters cost - narrative",
        "Ben comparison within this quarters cost - narrative",
        "Milestone Commentary",
    ]

    for x in range(len(headings_list)):
        doc.add_paragraph().add_run(str(headings_list[x])).bold = True
        text_one = str(master.master_data[0].data[project_name][narrative_keys_list[x]])
        try:
            text_two = str(
                master.master_data[1].data[project_name][narrative_keys_list[x]]
            )
        except KeyError:
            text_two = text_one

        # There are two options here for comparing text. Have left this for now.
        # compare_text_showall(dca_a, dca_b, doc)
        compare_text_new_and_old(text_one, text_two, doc)


def change_word_doc_landscape(doc: Document) -> Document:
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)  # new page
    new_width, new_height = new_section.page_height, new_section.page_width
    new_section.orientation = WD_ORIENTATION.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    return doc


def change_word_doc_portrait(doc: Document) -> Document:
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_width, new_height = new_section.page_height, new_section.page_width
    new_section.orientation = WD_ORIENTATION.PORTRAIT
    new_section.page_width = new_width
    new_section.page_height = new_height
    return doc


def put_matplotlib_fig_into_word(doc: Document, fig) -> None:
    """Places line graph cost profile into word document"""
    # Place fig in word doc.
    fig.savefig("cost_profile.png")
    doc.add_picture("cost_profile.png", width=Inches(8))  # to place nicely in doc
    os.remove("cost_profile.png")
    plt.close()  # automatically closes figure so don't need to do manually.


def convert_rag_text(dca_rating: str) -> None:
    """Converts RAG name into a acronym"""

    if dca_rating == "Green":
        return "G"
    elif dca_rating == "Amber/Green":
        return "A/G"
    elif dca_rating == "Amber":
        return "A"
    elif dca_rating == "Amber/Red":
        return "A/R"
    elif dca_rating == "Red":
        return "R"
    else:
        return ""


def cell_colouring(word_table_cell: table.Table.cell, colour: str) -> None:
    """Function that handles cell colouring for word documents"""

    try:
        if colour == "R":
            colour = parse_xml(r'<w:shd {} w:fill="cb1f00"/>'.format(nsdecls("w")))
        elif colour == "A/R":
            colour = parse_xml(r'<w:shd {} w:fill="f97b31"/>'.format(nsdecls("w")))
        elif colour == "A":
            colour = parse_xml(r'<w:shd {} w:fill="fce553"/>'.format(nsdecls("w")))
        elif colour == "A/G":
            colour = parse_xml(r'<w:shd {} w:fill="a5b700"/>'.format(nsdecls("w")))
        elif colour == "G":
            colour = parse_xml(r'<w:shd {} w:fill="17960c"/>'.format(nsdecls("w")))

        word_table_cell._tc.get_or_add_tcPr().append(colour)

    except TypeError:
        pass


def make_rows_bold(rows: list) -> None:
    """This function makes text bold in a list of row numbers for a word document"""
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def set_col_widths(word_table: table, widths: list) -> None:
    """This function sets the width of table in a word document"""
    for row in word_table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def compare_text_new_and_old(text_1: str, text_2: str, doc: Document) -> None:
    """compares two sets of text and highlights differences in red text."""

    comp = difflib.Differ()
    diff = list(comp.compare(text_2.split(), text_1.split()))
    new_text = diff
    y = doc.add_paragraph()

    for i in range(0, len(diff)):
        f = len(diff) - 1
        if i < f:
            a = i - 1
        else:
            a = i

        if diff[i][0:3] == "  |":
            j = i + 1
            if diff[i][0:3] and diff[a][0:3] == "  |":
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == "+ |":
            if diff[i][0:3] and diff[a][0:3] == "+ |":
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == "- |":
            pass
        elif diff[i][0:3] == "  -":
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0:3] == "  •":
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0] == "+":
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.color.rgb = RGBColor(255, 0, 0)
        elif diff[i][0] == "-":
            pass
        elif diff[i][0] == "?":
            pass
        else:
            if diff[i] != "+ |":
                y.add_run(diff[i][1:])


def make_file_friendly(quarter_str: str) -> str:
    """Converts datamaps.api project_data_from_master quarter data into a string to use when
    saving output files. Courtesy of M Lemon."""
    regex = r"Q(\d) (\d+)\/(\d+)"
    return re.sub(regex, r"Q\1_\2_\3", quarter_str)


def total_costs_benefits_bar_chart(
    cost_master: CostData, ben_master: BenefitsData, **kwargs
) -> plt.figure:
    """compiles a matplotlib bar chart which shows total project costs"""
    fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2)  # four sub plots

    try:
        fig_size = kwargs["fig_size"]
        fig.set_size_inches(set_figure_size(fig_size))
    except KeyError:
        fig.set_size_inches(set_figure_size(FIGURE_STYLE[2]))
        pass

    # cost profile charts.
    if len(cost_master.project_group) == 1:
        title = cost_master.master.abbreviations[cost_master.project_group[0]] + " cost and benefit totals"
    else:
        try:
            title = kwargs["title"] + " cost and benefit totals"
        except KeyError:
            title = ""
            print("You need to provide a title for this chart")

    plt.suptitle(title, fontweight="bold", fontsize=25)
    plt.xticks(size=12)
    plt.yticks(size=10)

    # Y AXIS SCALE MAX
    highest_int = max([cost_master.y_scale_max, ben_master.y_scale_max, ben_master.economic_max])  # check in refactor
    y_max = highest_int + percentage(5, highest_int)
    ax1.set_ylim(0, y_max)

    # COST SPENT, PROFILED AND UNPROFILED
    labels = ["Latest", "Last quarter", "Baseline"]
    width = 0.5
    ax1.bar(labels, np.array(cost_master.spent), width, label="Spent")
    ax1.bar(
        labels,
        np.array(cost_master.profiled),
        width,
        bottom=np.array(cost_master.spent),
        label="Profiled",
    )
    ax1.bar(
        labels,
        np.array(cost_master.unprofiled),
        width,
        bottom=np.array(cost_master.spent) + np.array(cost_master.profiled),
        label="Unprofiled",
    )
    ax1.legend(prop={"size": 10})
    ax1.xaxis.set_tick_params(labelsize=12)
    ax1.yaxis.set_tick_params(labelsize=12)
    ax1.set_ylabel("Cost (£m)")
    ylab1 = ax1.yaxis.get_label()
    ylab1.set_style("italic")
    ylab1.set_size(12)
    # ax1.tick_params(axis="series_one", which="major")
    # ax1.tick_params(axis="series_two", which="major")
    ax1.set_title(
        "Fig 1. Change in total cost",
        loc="left",
        fontsize=12,
        fontweight="bold",
    )

    # RDEL, CDEL AND NON-GOV TOTALS BAR CHART
    labels = ["RDEL", "CDEL", "Non Gov"]
    width = 0.5
    ax2.bar(labels, np.array(cost_master.cat_spent), width, label="Spent")
    ax2.bar(
        labels,
        np.array(cost_master.cat_profiled),
        width,
        bottom=np.array(cost_master.cat_spent),
        label="Profiled",
    )
    ax2.bar(
        labels,
        np.array(cost_master.cat_unprofiled),
        width,
        bottom=np.array(cost_master.cat_spent) + np.array(cost_master.cat_profiled),
        label="Unprofiled",
    )
    ax2.legend(prop={"size": 10})
    ax2.xaxis.set_tick_params(labelsize=12)
    ax2.yaxis.set_tick_params(labelsize=12)
    ax2.set_ylabel("Costs (£m)")
    ylab3 = ax2.yaxis.get_label()
    ylab3.set_style("italic")
    ylab3.set_size(12)
    # ax2.tick_params(axis="series_one", which="major", labelsize=6)
    # ax2.tick_params(axis="series_two", which="major", labelsize=6)
    ax2.set_title(
        "Fig 2. Current cost breakdown",
        loc="left",
        fontsize=12,
        fontweight="bold",
    )

    ax2.set_ylim(0, y_max)  # scale series_two axis max

    # BENEFITS SPENT, PROFILED AND UNPROFILED
    labels = ["Latest", "Last Quarter", "Baseline"]
    width = 0.5
    ax3.bar(labels, np.array(ben_master.delivered), width, label="Delivered")
    ax3.bar(
        labels,
        np.array(ben_master.profiled),
        width,
        bottom=np.array(ben_master.delivered),
        label="Profiled",
    )
    ax3.bar(
        labels,
        np.array(ben_master.unprofiled),
        width,
        bottom=np.array(ben_master.delivered) + np.array(ben_master.profiled),
        label="Unprofiled",
    )
    ax3.legend(prop={"size": 10})
    ax3.xaxis.set_tick_params(labelsize=12)
    ax3.yaxis.set_tick_params(labelsize=12)
    ax3.set_ylabel("Benefits (£m)")
    ylab3 = ax3.yaxis.get_label()
    ylab3.set_style("italic")
    ylab3.set_size(12)
    # ax3.tick_params(axis="series_one", which="major", labelsize=6)
    # ax3.tick_params(axis="series_two", which="major", labelsize=6)
    ax3.set_title(
        "Fig 3. Change in total benefits",
        loc="left",
        fontsize=12,
        fontweight="bold",
    )

    ax3.set_ylim(0, y_max)

    # BENEFITS CASHABLE, NON-CASHABLE, ECONOMIC, DISBENEFIT BAR CHART
    labels = ["Cashable", "Non-Cashable", "Economic", "Disbenefit"]
    width = 0.5
    ax4.bar(labels, np.array(ben_master.cat_delivered), width, label="Delivered")
    ax4.bar(
        labels,
        np.array(ben_master.cat_profiled),
        width,
        bottom=np.array(ben_master.cat_delivered),
        label="Profiled",
    )
    ax4.bar(
        labels,
        np.array(ben_master.cat_unprofiled),
        width,
        bottom=np.array(ben_master.cat_delivered) + np.array(ben_master.cat_profiled),
        label="Unprofiled",
    )
    ax4.legend(prop={"size": 10})
    ax4.xaxis.set_tick_params(labelsize=12)
    ax4.yaxis.set_tick_params(labelsize=12)
    ax4.set_ylabel("Benefits (£m)")
    ylab4 = ax4.yaxis.get_label()
    ylab4.set_style("italic")
    ylab4.set_size(12)
    # ax4.tick_params(axis="series_one", which="major", labelsize=6)
    # ax4.tick_params(axis="series_two", which="major", labelsize=6)
    ax4.set_title(
        "Fig 4. Current benefit breakdown",
        loc="left",
        fontsize=12,
        fontweight="bold",
    )

    if ben_master.y_scale_min == 0:
        ax4.set_ylim(0, y_max)
    else:  #  for negative benefits
        y_min = ben_master.y_scale_min + percentage(40, ben_master.y_scale_min)  # arbitrary 40 percent
        ax4.set_ylim(y_min, y_max)

    fig.tight_layout(rect=[0, 0.03, 1, 0.95])  # size/fit of chart

    try:
        kwargs["show"] == "No"
    except KeyError:
        plt.show()

    return fig


def check_baselines(master: Master) -> None:
    """checks that projects have the correct baseline information. stops the
    programme if baselines are missing"""

    for v in BASELINE_TYPES.values():
        for p in master.current_projects:
            baselines = master.bl_index[v][p]
            if len(baselines) <= 2:
                print(
                    p
                    + " does not have a baseline point for "
                    + v
                    + " this could cause the programme to"
                    "crash. Therefore the programme is stopping. "
                    "Please amend the data for " + p + " so that "
                    " it has at least one baseline point for " + v
                )
                break
        else:
            continue
        break


def percentage(percent: int, whole: float) -> int:
    return round((percent * whole) / 100.0)


def get_old_fy_cost_data(
    master_file: typing.TextIO, project_id_wb: typing.TextIO
) -> None:
    """
    Gets all old financial data from a specified master and places into project id document.
    """
    master = project_data_from_master(
        master_file, 1, 2010
    )  # random year specified as not in use
    wb = load_workbook(project_id_wb)
    ws = wb.active

    for i in range(1, ws.max_column + 1):
        project_name = ws.cell(row=1, column=1 + i).value
        for row_num in range(2, ws.max_row + 1):
            key = ws.cell(row=row_num, column=1).value
            try:
                if key in master.data[project_name].keys():
                    ws.cell(row=row_num, column=1 + i).value = master.data[
                        project_name
                    ][key]
            except KeyError:  # project might not be present in quarter
                pass

    wb.save(project_id_wb)


def run_get_old_fy_data(master_files_list: list, project_id_wb: typing.TextIO) -> None:
    for f in reversed(
        master_files_list
    ):  # reversed so it gets the latest data in masters
        get_old_fy_cost_data(f, project_id_wb)


def place_old_fy_data_into_master_wb(
    master_file: typing.TextIO, project_id_wb: typing.TextIO
) -> None:
    """
    places all old financial year data into master files.
    """
    id_master = project_data_from_master(
        project_id_wb, 2, 2020
    )  # random year specify as not used
    wb = load_workbook(master_file)
    ws = wb.active

    for i in range(1, ws.max_column + 1):
        project_name = ws.cell(row=1, column=1 + i).value
        for row_num in range(2, ws.max_row + 1):
            key = ws.cell(row=row_num, column=1).value
            try:
                if key in id_master.data[project_name].keys():
                    ws.cell(row=row_num, column=1 + i).value = id_master.data[
                        project_name
                    ][key]
            except KeyError:  # project might not be present in quarter
                pass

    wb.save(master_file)


def run_place_old_fy_data_into_masters(
    master_files_list: list, project_id_wb: typing.TextIO
) -> None:
    for f in master_files_list:
        place_old_fy_data_into_master_wb(f, project_id_wb)


def put_key_change_master_into_dict(key_change_file: typing.TextIO) -> Dict[str, str]:
    """
    places keys information i.e. keys old and new names from wb into a python dictionary
    """
    wb = load_workbook(key_change_file)
    ws = wb.active

    output_dict = {}
    for x in range(1, ws.max_row + 1):
        key = ws.cell(row=x, column=1).value
        codename = ws.cell(row=x, column=2).value
        output_dict[key] = codename

    return output_dict


def alter_wb_master_file_key_names(
    master_file: typing.TextIO, key_change_dict: Dict[str, str]
) -> workbook:
    """
    places altered keys names, from the keys change master dictionary, into master wb(s).
    """
    wb = load_workbook(master_file)
    ws = wb.active

    for row_num in range(2, ws.max_row + 1):
        for (
            key
        ) in key_change_dict.keys():  # changes stored in the altered keys change log wb
            if ws.cell(row=row_num, column=1).value == key:
                ws.cell(row=row_num, column=1).value = key_change_dict[key]
        for year in YEAR_LIST:  # changes to yearly profile keys
            if ws.cell(row=row_num, column=1).value == year + " CDEL Forecast Total":
                ws.cell(row=row_num, column=1).value = (
                    year + " CDEL Forecast one off new costs"
                )

    return wb.save(master_file)


def run_change_keys(master_files_list: list, key_dict: Dict[str, str]) -> None:
    """
    runs code which replaces old keys names with new names in master excel workbooks.
    """
    for f in master_files_list:
        alter_wb_master_file_key_names(f, key_dict)


def string_conversion(name):
    if isinstance(name, str):
        return [name]
    else:
        return name


def compare_masters(files: List[typing.TextIO], projects: List[str] or str) -> workbook:
    """Takes two masters and compares if there have been any changes to project data values,
    as well as any new data keys.
    files = list of file paths to masters. Only last two masters are used
    projects = list of those projects that require data checking"""

    projects = string_conversion(projects)

    last_master = project_data_from_master(files[1], 4, 2090)

    wb = load_workbook(files[0])
    ws = wb.active

    project_count = []
    change_count = 0
    new_key_count = 0

    for row_num in range(2, ws.max_row + 1):
        key = ws.cell(row=row_num, column=1).value
        for column_num in range(2, ws.max_column + 1):
            project_name = ws.cell(row=1, column=column_num).value
            if project_name not in projects:
                pass
            else:
                wb_value = ws.cell(row=row_num, column=column_num).value
                try:
                    dict_value = last_master.data[project_name][key]
                    if wb_value == dict_value:
                        pass
                    else:
                        ws.cell(row=row_num, column=column_num).fill = PatternFill(
                            start_color="ffba00", end_color="ffba00", fill_type="solid"
                        )
                        project_count.append(project_name)
                        change_count += 1
                except KeyError:
                    if (
                        project_name in last_master.projects
                    ):  # keys error due to keys not being present.
                        ws.cell(row=row_num, column=1).fill = PatternFill(
                            start_color="ffba00", end_color="ffba00", fill_type="solid"
                        )
                    else:  # keys error due to project not being present.
                        ws.cell(row=1, column=column_num).fill = PatternFill(
                            start_color="ffba00", end_color="ffba00", fill_type="solid"
                        )
    # separate lop to calculate this
    for row_num in range(2, ws.max_row + 1):
        key = ws.cell(row=row_num, column=1).value
        if key not in list(last_master.data[projects[0]].keys()):
            new_key_count += 1

    count_ws = wb.create_sheet("Count", 1)

    count_ws.cell(row=1, column=1).value = "No of changes"
    count_ws.cell(row=1, column=2).value = change_count
    count_ws.cell(row=2, column=1).value = "No of new keys"
    count_ws.cell(row=2, column=2).value = new_key_count

    project_count = Counter(project_count)
    for i, p in enumerate(project_count.keys()):
        count_ws.cell(row=i + 3, column=1).value = p
        count_ws.cell(row=i + 3, column=2).value = project_count[p]

    return wb


def totals_chart(costs: CostData, benefits: BenefitsData, **kwargs) -> None:
    """Small function to hold together code to create and save a total_costs_benefits_bar_chart"""
    if kwargs == {}:
        f = total_costs_benefits_bar_chart(costs, benefits)
        f.savefig(root_path / "output/{}_profile.png".format(costs.project_group[0]))
    else:
        if list(kwargs.keys()) == ["title", "fig_size"]:
            f = total_costs_benefits_bar_chart(
                costs, benefits, fig_size=kwargs["fig_size"], title=kwargs["title"]
            )
            f.savefig(root_path / "output/{}_profile.png".format(str(kwargs["title"])))
        if list(kwargs.keys()) == ["fig_size", "title"]:
            f = total_costs_benefits_bar_chart(
                costs, benefits, fig_size=kwargs["fig_size"], title=kwargs["title"]
            )
            f.savefig(root_path / "output/{}_profile.png".format(str(kwargs["title"])))
        if list(kwargs.keys()) == ["title"]:
            f = total_costs_benefits_bar_chart(costs, benefits, title=kwargs["title"])
            f.savefig(root_path / "output/{}_profile.png".format(str(kwargs["title"])))
        if list(kwargs.keys()) == ["fig_size"]:
            f = total_costs_benefits_bar_chart(
                costs, benefits, fig_size=kwargs["fig_size"]
            )
            f.savefig(
                root_path / "output/{}_profile.png".format(costs.project_group[0])
            )


def standard_profile(costs: CostData, **kwargs):
    """Small function to hold together code to create and save a cost_profile_graph"""
    if kwargs == {}:
        f = cost_profile_graph(costs)
        f.savefig(root_path / "output/{}_profile.png".format(costs.project_group[0]))
    else:
        if list(kwargs.keys()) == ["title", "fig_size"]:
            f = cost_profile_graph(
                costs, fig_size=kwargs["fig_size"], title=kwargs["title"]
            )
            f.savefig(root_path / "output/{}_profile.png".format(str(kwargs["title"])))
        if list(kwargs.keys()) == ["fig_size", "title"]:
            f = cost_profile_graph(
                costs, fig_size=kwargs["fig_size"], title=kwargs["title"]
            )
            f.savefig(root_path / "output/{}_profile.png".format(str(kwargs["title"])))
        if list(kwargs.keys()) == ["title"]:
            f = cost_profile_graph(costs, title=kwargs["title"])
            f.savefig(root_path / "output/{}_profile.png".format(str(kwargs["title"])))
        if list(kwargs.keys()) == ["fig_size"]:
            f = cost_profile_graph(costs, fig_size=kwargs["fig_size"])
            f.savefig(
                root_path / "output/{}_profile.png".format(costs.project_group[0])
            )


def save_graph(fig: plt.figure, file_name: str, **kwargs) -> None:
    """Generic function for saving matplotlib figure into a word document"""
    if "orientation" in list(kwargs.keys()):
        if kwargs["orientation"] == "landscape":
            fig.savefig("temp_file.png")
            doc = open_word_doc(root_path / "input/summary_temp_landscape.docx")
            doc.add_picture("temp_file.png", width=Inches(8))
            doc.save(root_path / "output/{}.docx".format(file_name))
            os.remove("temp_file.png")
        if kwargs["orientation"] == "portrait":
            fig.savefig("temp_file.png")
            doc = open_word_doc(root_path / "input/summary_temp.docx")
            doc.add_picture("temp_file.png", width=Inches(8))
            doc.save(root_path / "output/{}.docx".format(file_name))
            os.remove("temp_file.png")
    else:
        fig.savefig("temp_file.png")
        doc = open_word_doc(root_path / "input/summary_temp_landscape.docx")
        doc.add_picture("temp_file.png", width=Inches(8))
        doc.save(root_path / "output/{}.docx".format(file_name))
        os.remove("temp_file.png")


# from stackoverflow.
def do_mask(x: List[datetime.date], y: List[datetime.date]):
    """
    helper function for putting series of datetime.date values with NoneType into
    matplotlib.
    """
    mask = None
    mask = ~(x == None)
    return np.array(x)[mask], np.array(y)[mask]


def milestone_chart(
    milestone_data: MilestoneData,
    **kwargs,
) -> plt.figure:
    # build scatter chart
    fig, ax1 = plt.subplots()

    # set figure size
    try:
        fig_size = kwargs["fig_size"]
        fig.set_size_inches(set_figure_size(fig_size))
    except KeyError:
        fig.set_size_inches(set_figure_size(FIGURE_STYLE[2]))
        pass

    # title
    if len(milestone_data.project_group) == 1:
        try:
            title = kwargs["title"]
        except KeyError:
            title = milestone_data.master.abbreviations[milestone_data.project_group[0]] + " Schedule"
    else:
        try:
            title = kwargs["title"]
        except KeyError:
            pass
            print("You need to provide a title for this chart")

    fig.suptitle(title, fontweight="bold", fontsize=25)

    def handle_long_keys(key_names: List[str]) -> List[str]:
        labels = ['\n'.join(wrap(l, 40)) for l in key_names]
        final_labels = []
        for l in labels:
            if len(l) > 70:
                final_labels.append(l[:70])
            else:
                final_labels.append(l)
        return final_labels

    m_key_names = handle_long_keys(milestone_data.key_names)

    # convert lists into numpy arrays.
    # milestone_data.md_baseline = np.array(milestone_data.md_baseline)
    # milestone_data.md_last = np.array(milestone_data.md_last)
    # milestone_data.md_current = np.array(milestone_data.md_current)

    # fom stackoverflow, Since plotting series as a scatter plot, the order
    # is not crucial. index nparrays for non-zero elements. not using at moment.
    # idx_three = milestone_data.md_current.nonzero()[0].tolist()
    # ax1.scatter(
    #     milestone_data.md_current[idx_three],
    #     np.array(milestone_data.key_names)[idx_three],
    #     label="Current",
    #     zorder=10,
    #     c='g'
    # )
    # idx_two = milestone_data.md_last.nonzero()[0].tolist()
    # ax1.scatter(
    #     milestone_data.md_last[idx_two],
    #     np.array(milestone_data.key_names)[idx_two],
    #     label="Last quarter",
    #     zorder=5,
    #     c='orange'
    # )
    # idx = milestone_data.md_baseline.nonzero()[0].tolist()
    # ax1.scatter(
    #     milestone_data.md_baseline[idx],
    #     np.array(milestone_data.key_names)[idx],
    #     label="Baseline",
    #     zorder=1,
    #     c='b'
    # )

    # this method does not handle NoneTypes. Therefore get_chart_info returns md_current
    # instead of NoneTypes. Works fine, but underlying data is incorrect. Although this is
    # hidden from the user, preference for not making data wrong. but using at the moment.
    ax1.scatter(milestone_data.md_baseline, m_key_names, label="Baseline", s=200)
    ax1.scatter(milestone_data.md_last, m_key_names, label="Last quarter", s=200)
    ax1.scatter(milestone_data.md_current, m_key_names, label="Current", s=200)

    # ax1.scatter(*do_mask(milestone_data.md_current, milestone_data.key_names), label="Current", zorder=10, c='g')
    # ax1.scatter(*do_mask(milestone_data.md_last, milestone_data.key_names), label="Last quarter", zorder=5, c='orange')
    # ax1.scatter(*do_mask(milestone_data.md_baseline, milestone_data.key_names), label="Baseline", zorder=1, c='b')

    # format the series_one ticks
    years = mdates.YearLocator()  # every year
    months = mdates.MonthLocator()  # every month
    years_fmt = mdates.DateFormatter("%Y")
    months_fmt = mdates.DateFormatter("%b")
    # ax1.xaxis.set_major_locator(years)
    # ax1.xaxis.set_minor_locator(months)
    # ax1.xaxis.set_major_formatter(years_fmt)
    # ax1.xaxis.set_minor_formatter(months_fmt)
    # plt.setp(ax1.xaxis.get_minorticklabels(), rotation=45)
    # plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, weight="bold")

    """calculate the length of the time period covered in chart.
    Not perfect as baseline dates can distort."""
    td = (milestone_data.max_date - milestone_data.min_date).days
    if td >= 365 * 3:
        ax1.xaxis.set_major_locator(years)
        ax1.xaxis.set_minor_locator(months)
        ax1.xaxis.set_major_formatter(years_fmt)
        # ax1.xaxis.set_minor_formatter(months_fmt)
        plt.setp(ax1.xaxis.get_minorticklabels(), rotation=45, size=14)
        plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, weight="bold", size=16)

        # scaling series_one axis. Keeping for now in case useful.
        # series_one axis value to no more than three months after last latest milestone date, or three months
        # before first latest milestone date. Hack, can be improved. Text highlights movements off chart.
        # x_max = milestone_data.md_current[-1] + timedelta(days=90)
        # x_min = milestone_data.md_current[0] - timedelta(days=90)
        # for date in milestone_data.md_baseline:
        #     if date > x_max:
        #         ax1.set_xlim(x_min, x_max)
        #         plt.figtext(
        #             0.98,
        #             0.03,
        #             "Check full schedule to see all milestone movements",
        #             horizontalalignment="right",
        #             fontsize=6,
        #             fontweight="bold",
        #         )
        #     if date < x_min:
        #         ax1.set_xlim(x_min, x_max)
        #         plt.figtext(
        #             0.98,
        #             0.03,
        #             "Check full schedule to see all milestone movements",
        #             horizontalalignment="right",
        #             fontsize=6,
        #             fontweight="bold",
        #         )
    else:
        ax1.xaxis.set_major_locator(years)
        ax1.xaxis.set_minor_locator(months)
        ax1.xaxis.set_major_formatter(years_fmt)
        ax1.xaxis.set_minor_formatter(months_fmt)
        plt.setp(ax1.xaxis.get_minorticklabels(), rotation=45, size=14)
        plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, weight="bold", size=16)

    ax1.legend(prop={"size": 14})  # insert legend

    # plt.xticks(rotation=45, size=14)
    plt.yticks(size=12)

    # reverse series_two axis so order is earliest to oldest
    ax1 = plt.gca()
    ax1.set_ylim(ax1.get_ylim()[::-1])
    ax1.tick_params(axis="series_two", which="major", labelsize=7)
    ax1.yaxis.grid()  # horizontal lines
    ax1.set_axisbelow(True)

    try:
        if kwargs["show_keys"] == "no":
            ax1.get_yaxis().set_visible(False)
    except KeyError:
        pass

    # Add line of analysis date, but only if in the time period
    try:
        blue_line = kwargs["blue_line"]
        if blue_line == "Today":
            if (
                milestone_data.min_date
                <= datetime.date.today()
                <= milestone_data.max_date
            ):
                plt.axvline(datetime.date.today())
                plt.figtext(
                    0.98,
                    0.01,
                    "Line represents date analysis compiled",
                    horizontalalignment="right",
                    fontsize=10,
                    fontweight="bold",
                )
        if blue_line == "ipdc_date":
            if milestone_data.min_date <= IPDC_DATE <= milestone_data.max_date:
                plt.axvline(IPDC_DATE)
                plt.figtext(
                    0.98,
                    0.01,
                    "Line represents PfM report at IPDC",
                    horizontalalignment="right",
                    fontsize=10,
                    fontweight="bold",
                )
    except KeyError:
        pass

    # size of chart and fit
    # fig.canvas.draw()
    fig.tight_layout(rect=[0, 0.03, 1, 0.95])  # for title

    try:
        kwargs["show"] == "No"
    except KeyError:
        plt.show()

    return fig

    # fig.savefig(root_path / 'output/{}.png'.format(graph_title), bbox_inches='tight')

    # plt.close() #automatically closes figure so don't need to do manually.


# def compile_all_profiles():
#     report_doc = open_word_doc(wd_path)
#     for i, p in enumerate(LIST_OF_GROUPS):
#         costs.get_cost_profile(p, 'ipdc_costs')
#         graph = cost_profile_graph(costs, LIST_OF_TITLES[i])
#         put_matplotlib_fig_into_word(report_doc, graph)
#         report_doc.save(root_path / "output/different_cost_profiles.docx")
#
#
# def compile_all_totals():
#     report_doc = open_word_doc(wd_path)
#     for i, p in enumerate(LIST_OF_GROUPS):
#         costs.get_cost_totals(p, 'ipdc_costs')
#         benefits.get_ben_totals(p, 'ipdc_benefits')
#         graph = total_costs_benefits_bar_chart(costs, benefits, LIST_OF_TITLES[i])
#         put_matplotlib_fig_into_word(report_doc, graph)
#         report_doc.save(root_path / "output/different_total_cost_profiles.docx")

# def build_charts(latest_milestone_names,
#                  latest_milestone_dates,
#                  last_milestone_dates,
#                  baseline_milestone_dates,
#                  baseline_milestone_dates_two,
#                  graph_title,
#                  ipdc_date,
#                  no_of_labels):
#     """
#     calculates how many graphical outputs to produced
#     based on number of milestones. Milestone keys names,
#     dates, graph title, date for blue line to represent,
#     and number of labels to have on each graph
#     are passed in.
#     """
#
#     # axis labels are reduced if two long
#     labels = latest_milestone_names
#     final_labels = []
#     for l in labels:
#         if len(l) >= 40:
#             final_labels.append(l[:40])
#         else:
#             final_labels.append(l)
#
#     #  Charts are built
#     no_milestones = len(latest_milestone_names)
#
#     if no_milestones <= no_of_labels:
#         milestone_swimlane_charts(np.array(final_labels), np.array(latest_milestone_dates),
#                                   np.array(last_milestone_dates),
#                                   np.array(baseline_milestone_dates),
#                                   np.array(baseline_milestone_dates_two),
#                                   graph_title, ipdc_date)
#
#     if no_of_labels + 1 <= no_milestones <= no_of_labels*2:
#         half = int(no_milestones / 2)
#         milestone_swimlane_charts(np.array(final_labels[:half]),
#                                   np.array(latest_milestone_dates[:half]),
#                                   np.array(last_milestone_dates[:half]),
#                                   np.array(baseline_milestone_dates[:half]),
#                                   np.array(baseline_milestone_dates_two[:half]),
#                                   graph_title, ipdc_date)
#         title = graph_title + ' cont.'
#         milestone_swimlane_charts(np.array(final_labels[half:no_milestones]),
#                                   np.array(latest_milestone_dates[half:no_milestones]),
#                                   np.array(last_milestone_dates[half:no_milestones]),
#                                   np.array(baseline_milestone_dates[half:no_milestones]),
#                                   np.array(baseline_milestone_dates_two[half:no_milestones]),
#                                   title, ipdc_date)
#
#     if (no_of_labels*2) + 1 <= no_milestones <= no_of_labels*3:
#         third = int(no_milestones / 3)
#         milestone_swimlane_charts(np.array(final_labels[:third]),
#                                   np.array(latest_milestone_dates[:third]),
#                                   np.array(last_milestone_dates[:third]),
#                                   np.array(baseline_milestone_dates[:third]),
#                                   np.array(baseline_milestone_dates_two[:third]),
#                                   graph_title, ipdc_date)
#         title = graph_title + ' cont. 1'
#         milestone_swimlane_charts(np.array(final_labels[third:third * 2]),
#                                   np.array(latest_milestone_dates[third:third * 2]),
#                                   np.array(last_milestone_dates[third:third * 2]),
#                                   np.array(baseline_milestone_dates[third:third * 2]),
#                                   np.array(baseline_milestone_dates_two[third:third * 2]),
#                                   title, ipdc_date)
#         title = graph_title + ' cont. 2'
#         milestone_swimlane_charts(np.array(final_labels[third * 2:no_milestones]),
#                                   np.array(latest_milestone_dates[third * 2:no_milestones]),
#                                   np.array(last_milestone_dates[third * 2:no_milestones]),
#                                   np.array(baseline_milestone_dates[third * 2:no_milestones]),
#                                   np.array(baseline_milestone_dates_two[third * 2:no_milestones]),
#                                   title, ipdc_date)
#     pass

DCA_KEYS = {
    "SRO": "Departmental DCA",
    "FINANCE": "SRO Finance confidence",
    "BENEFITS": "SRO Benefits RAG",
    "SCHEDULE": "SRO Schedule Confidence",
}

DCA_RATING_SCORES = {
    "Green": 5,
    "Amber/Green": 4,
    "Amber": 3,
    "Amber/Red": 2,
    "Red": 1,
    None: None,
}


class DcaData:
    def __init__(self, master: Master):
        self.master = master
        self.dca_dictionary = {}
        self.dca_changes = {}
        self.dca_count = {}
        self.get_dictionary()
        self.get_count()

    def get_dictionary(self) -> None:
        """
        collects all data for calculating changes in dca ratings.
        """
        quarter_dict = {}
        for i in range(len(self.master.master_data)):
            try:
                type_dict = {}
                for dca_type in list(DCA_KEYS.values()):
                    dca_dict = {}
                    for project_name in self.master.master_data[i].projects:
                        colour = self.master.master_data[i].data[project_name][dca_type]
                        score = DCA_RATING_SCORES[
                            self.master.master_data[i].data[project_name][dca_type]
                        ]
                        costs = self.master.master_data[i].data[project_name][
                            "Total Forecast"
                        ]
                        dca_colour = [("DCA", colour)]
                        dca_score = [("DCA score", score)]
                        t = [("Type", dca_type)]
                        cost_amount = [("Costs", costs)]
                        quarter = [("Quarter", str(self.master.master_data[i].quarter))]
                        dca_dict[self.master.abbreviations[project_name]] = dict(
                            dca_colour + t + cost_amount + quarter + dca_score
                        )
                    type_dict[dca_type] = dca_dict
            except KeyError:  # handles dca_type e.g. schedule confidence key not present
                pass

            quarter_dict[str(self.master.master_data[i].quarter)] = type_dict

        self.dca_dictionary = quarter_dict

    def get_changes(self, quarter_one: str, quarter_two: str) -> None:
        """compiles dictionary of changes in dca ratings when provided with two quarter arguments"""
        self.q_one = quarter_one
        self.q_two = quarter_two

        c_dict = {}
        for dca_type in list(self.dca_dictionary[self.q_one].keys()):
            lower_dict = {}
            for project_name in list(self.dca_dictionary[self.q_one][dca_type].keys()):
                t = [("Type", dca_type)]
                try:
                    dca_one_colour = self.dca_dictionary[quarter_one][dca_type][
                        project_name
                    ]["DCA"]
                    dca_two_colour = self.dca_dictionary[quarter_two][dca_type][
                        project_name
                    ]["DCA"]
                    dca_one_score = self.dca_dictionary[quarter_one][dca_type][
                        project_name
                    ]["DCA score"]
                    dca_two_score = self.dca_dictionary[quarter_two][dca_type][
                        project_name
                    ]["DCA score"]
                    if dca_one_score == dca_two_score:
                        status = [("Status", "Same")]
                        change = [("Change", "Unchanged")]
                    if dca_one_score > dca_two_score:
                        status = [
                            (
                                "Status",
                                "Improved from "
                                + dca_two_colour
                                + " to "
                                + dca_one_colour,
                            )
                        ]
                        change = [("Change", "Up")]
                    if dca_one_score < dca_two_score:
                        status = [
                            (
                                "Status",
                                "Worsened from "
                                + dca_two_colour
                                + " to "
                                + dca_one_colour,
                            )
                        ]
                        change = [("Change", "Down")]
                except TypeError:  # This picks up None types
                    status = [("Status", "Missing")]
                    change = [("Change", "Unknown")]
                except KeyError:  # This picks up projects not being present in the quarters being analysed.
                    status = [("Status", "New entry")]
                    change = [("Change", "New entry")]

                lower_dict[project_name] = dict(t + status + change)

            c_dict[dca_type] = lower_dict
        self.dca_changes = c_dict

    def get_count(self) -> None:
        """Returns dictionary containing a count of dcas"""
        output_dict = {}
        error_list = []
        for quarter in self.dca_dictionary.keys():
            dca_dict = {}
            for i, dca_type in enumerate(list(self.dca_dictionary[quarter].keys())):
                colour_count = []
                total_count = []
                for x, colour in enumerate(list(DCA_RATING_SCORES.keys())):
                    count = 0
                    cost = 0
                    total = 0
                    cost_total = 0
                    for y, project in enumerate(
                        list(self.dca_dictionary[quarter][dca_type].keys())
                    ):
                        total += 1
                        try:
                            cost_total += self.dca_dictionary[quarter][dca_type][
                                project
                            ]["Costs"]
                        except TypeError:
                            error_list.append(
                                project
                                + " total costs for "
                                + str(quarter)
                                + " are in an incorrect data type and need changing"
                            )
                            pass
                        if (
                            self.dca_dictionary[quarter][dca_type][project]["DCA"]
                            == colour
                        ):
                            count += 1
                            try:
                                cost += self.dca_dictionary[quarter][dca_type][project][
                                    "Costs"
                                ]
                            except TypeError:  # error message above doesn't need repeating
                                pass
                    colour_count.append((colour, (count, cost, cost / cost_total)))
                    total_count.append(
                        ("Total", (total, cost_total, cost_total / cost_total))
                    )

                dca_dict[dca_type] = dict(colour_count + total_count)
            output_dict[quarter] = dca_dict

        error_list = get_error_list(error_list)
        for x in error_list:
            print(x)

        self.dca_count = output_dict


def dca_changes_into_word(dca_data: DcaData, doc: Document) -> Document:
    for i, dca_type in enumerate(list(dca_data.dca_changes.keys())):
        if i != 0:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        else:
            pass
        title = dca_type + " " + "Confidence changes"
        top = doc.add_paragraph()
        top.add_run(title).bold = True

        doc.add_paragraph()
        sub_head = "Improvements"
        sub = doc.add_paragraph()
        sub.add_run(sub_head).bold = True
        count = 0
        for project_name in list(dca_data.dca_changes[dca_type].keys()):
            if dca_data.dca_changes[dca_type][project_name]["Change"] == "Up":
                doc.add_paragraph(
                    project_name
                    + " "
                    + dca_data.dca_changes[dca_type][project_name]["Status"]
                )
                count += 1
        total_line = str(count) + " project(s) in total improved"
        doc.add_paragraph(total_line)

        doc.add_paragraph()
        sub_head = "Decreases"
        sub = doc.add_paragraph()
        sub.add_run(sub_head).bold = True
        count = 0
        for project_name in list(dca_data.dca_changes[dca_type].keys()):
            if dca_data.dca_changes[dca_type][project_name]["Change"] == "Down":
                doc.add_paragraph(
                    project_name
                    + " "
                    + dca_data.dca_changes[dca_type][project_name]["Status"]
                )
                count += 1
        total_line = str(count) + " project(s) in total have decreased"
        doc.add_paragraph(total_line)

        doc.add_paragraph()
        sub_head = "Missing ratings"
        sub = doc.add_paragraph()
        sub.add_run(sub_head).bold = True
        count = 0
        for project_name in list(dca_data.dca_changes[dca_type].keys()):
            if dca_data.dca_changes[dca_type][project_name]["Change"] == "Unknown":
                doc.add_paragraph(
                    project_name
                    + " "
                    + dca_data.dca_changes[dca_type][project_name]["Status"]
                )
                count += 1
        total_line = str(count) + " project(s) in total are missing a rating"
        doc.add_paragraph(total_line)

    return doc


def dca_changes_into_excel(dca_data: DcaData, quarter: List[str] or str) -> workbook:
    wb = Workbook()

    quarter = string_conversion(quarter)

    for q in quarter:
        start_row = 3
        ws = wb.create_sheet(
            make_file_friendly(q)
        )  # creating worksheets. names restricted to 30 characters.
        ws.title = make_file_friendly(q)  # title of worksheet
        for i, dca_type in enumerate(list(dca_data.dca_count[q].keys())):
            ws.cell(row=start_row + i, column=2).value = dca_type
            ws.cell(row=start_row + i, column=3).value = "Count"
            ws.cell(row=start_row + i, column=4).value = "Costs"
            ws.cell(row=start_row + i, column=5).value = "Proportion costs"
            for x, colour in enumerate(list(dca_data.dca_count[q][dca_type].keys())):
                ws.cell(row=start_row + i + x + 1, column=2).value = colour
                ws.cell(row=start_row + i + x + 1, column=3).value = (
                    dca_data.dca_count[q][dca_type][colour]
                )[0]
                ws.cell(row=start_row + i + x + 1, column=4).value = (
                    dca_data.dca_count[q][dca_type][colour]
                )[1]
                ws.cell(row=start_row + i + x + 1, column=5).value = (
                    dca_data.dca_count[q][dca_type][colour]
                )[2]
                if colour is None:
                    ws.cell(row=start_row + i + x + 1, column=2).value = "None"

            start_row += 9
    wb.remove(wb["Sheet"])
    return wb


RISK_LIST = [
    "Brief Risk Description ",
    "BRD Risk Category",
    "BRD Primary Risk to",
    "BRD Internal Control",
    "BRD Mitigation - Actions taken (brief description)",
    "BRD Residual Impact",
    "BRD Residual Likelihood",
    "Severity Score Risk Category",
    "BRD Has this Risk turned into an Issue?",
]

RISK_SCORES = {"Very Low": 0, "Low": 1, "Medium": 2, "High": 3, "Very High": 4}


def risk_score(risk_impact: str, risk_likelihood: str) -> str:
    impact_score = RISK_SCORES[risk_impact]
    likelihood_score = RISK_SCORES[risk_likelihood]
    score = impact_score + likelihood_score
    if score <= 4:
        if risk_impact == "Medium" and risk_likelihood == "Medium":
            return "Medium"
        else:
            return "Low"
    if 5 <= score <= 6:
        if risk_impact == "High" and risk_likelihood == "High":
            return "High"
        if risk_impact == "Low" and risk_likelihood == "Very High":
            return "Low"
        if risk_impact == "Very High" and risk_likelihood == "Low":
            return "Low"
        else:
            return "Medium"
    if score > 6:
        return "High"


class RiskData:
    def __init__(self, master: Master):
        self.master = master
        self.risk_dictionary = {}
        self.risk_count = {}
        self.risk_impact_count = {}
        self.get_dictionary()
        self.get_count()

    def get_dictionary(self):
        quarter_dict = {}
        for i in range(len(self.master.master_data)):
            project_dict = {}
            try:
                for project_name in self.master.master_data[i].projects:
                    number_dict = {}
                    for x in range(1, 11):  # currently 10 risks
                        risk_list = []
                        for risk_type in RISK_LIST:
                            try:
                                amended_risk_type = risk_type + str(x)
                                risk = (
                                    risk_type,
                                    self.master.master_data[i].data[project_name][
                                        amended_risk_type
                                    ],
                                )
                                risk_list.append(risk)
                            except KeyError:
                                try:
                                    amended_risk_type = (
                                        risk_type[:4] + str(x) + risk_type[3:]
                                    )
                                    risk = (
                                        risk_type,
                                        self.master.master_data[i].data[project_name][
                                            amended_risk_type
                                        ],
                                    )
                                    risk_list.append(risk)
                                except KeyError:
                                    try:
                                        if risk_type == "Severity Score Risk Category":
                                            impact = (
                                                "BRD Residual Impact"[:4]
                                                + str(x)
                                                + "BRD Residual Impact"[3:]
                                            )
                                            likelihoood = (
                                                "BRD Residual Likelihood"[:4]
                                                + str(x)
                                                + "BRD Residual Likelihood"[3:]
                                            )
                                            score = risk_score(
                                                self.master.master_data[i].data[
                                                    project_name
                                                ][impact],
                                                self.master.master_data[i].data[
                                                    project_name
                                                ][likelihoood],
                                            )
                                            risk = (
                                                "Severity Score Risk Category",
                                                score,
                                            )
                                            risk_list.append(risk)
                                    except KeyError:
                                        if risk_type == "Severity Score Risk Category":
                                            pass
                                        else:
                                            print(
                                                "check "
                                                + project_name
                                                + " "
                                                + str(x)
                                                + " "
                                                + risk_type
                                            )

                            number_dict[x] = dict(risk_list)

                    project_dict[self.master.abbreviations[project_name]] = number_dict
            except KeyError:  # handles dca_type e.g. schedule confidence key not present
                pass
            quarter_dict[str(self.master.master_data[i].quarter)] = project_dict

        self.risk_dictionary = quarter_dict

    def get_count(self):
        count_output_dict = {}
        impact_output_dict = {}
        for quarter in self.risk_dictionary.keys():
            count_lower_dict = {}
            impact_lower_dict = {}
            for i in range(len(RISK_LIST)):
                count_list = []
                impact_list = []
                for y, project_name in enumerate(
                    list(self.risk_dictionary[quarter].keys())
                ):
                    for x, number in enumerate(
                        list(self.risk_dictionary[quarter][project_name].keys())
                    ):
                        try:
                            risk_value = self.risk_dictionary[quarter][project_name][
                                number
                            ][RISK_LIST[i]]
                            impact = self.risk_dictionary[quarter][project_name][
                                number
                            ]["Severity Score Risk Category"]
                            count_list.append(risk_value)
                            impact_list.append((risk_value, impact))
                        except KeyError:
                            pass

                count_lower_dict[RISK_LIST[i]] = Counter(count_list)
                impact_lower_dict[RISK_LIST[i]] = Counter(impact_list)

            count_output_dict[quarter] = count_lower_dict
            impact_output_dict[quarter] = impact_lower_dict

        self.risk_count = count_output_dict
        self.risk_impact_count = impact_output_dict


def risks_into_excel(risk_data: RiskData, quarter: List[str] or str) -> workbook:
    wb = Workbook()

    quarter = string_conversion(quarter)

    for q in quarter:
        start_row = 3
        ws = wb.create_sheet(
            make_file_friendly(q + " all data")
        )  # creating worksheets. names restricted to 30 characters.
        ws.title = make_file_friendly(q + " all data")  # title of worksheet

        for y, project_name in enumerate(list(risk_data.risk_dictionary[q].keys())):
            for x, number in enumerate(
                list(risk_data.risk_dictionary[q][project_name].keys())
            ):
                if (
                    risk_data.risk_dictionary[q][project_name][number][
                        "Brief Risk Description "
                    ]
                    is None
                ):
                    break
                else:
                    ws.cell(row=start_row + 1 + x, column=2).value = project_name
                    ws.cell(row=start_row + 1 + x, column=3).value = str(number)
                    for i in range(len(RISK_LIST)):
                        try:
                            ws.cell(
                                row=start_row + 1 + x, column=4 + i
                            ).value = risk_data.risk_dictionary[q][project_name][
                                number
                            ][
                                RISK_LIST[i]
                            ]
                        except KeyError:
                            pass

            start_row += x

        for i in range(len(RISK_LIST)):
            ws.cell(row=3, column=4 + i).value = RISK_LIST[i]
        ws.cell(row=3, column=2).value = "Project Name"
        ws.cell(row=3, column=3).value = "Risk Number"

        ws = wb.create_sheet(
            make_file_friendly(q + " Count")
        )  # creating worksheets. names restricted to 30 characters.
        ws.title = make_file_friendly(q + " Count")  # title of worksheet

        start_row = 3
        for v, risk_cat in enumerate(list(risk_data.risk_count[q].keys())):
            if (
                risk_cat == "Brief Risk Description "
                or risk_cat == "BRD Mitigation - Actions taken (brief description)"
            ):
                pass
            else:
                ws.cell(row=start_row, column=2).value = risk_cat
                ws.cell(row=start_row, column=3).value = "Low"
                ws.cell(row=start_row, column=4).value = "Medium"
                ws.cell(row=start_row, column=5).value = "High"
                ws.cell(row=start_row, column=6).value = "Total"
                for b, cat in enumerate(list(risk_data.risk_count[q][risk_cat].keys())):
                    ws.cell(row=start_row + b + 1, column=2).value = cat
                    ws.cell(
                        row=start_row + b + 1, column=3
                    ).value = risk_data.risk_impact_count[q][risk_cat][(cat, "Low")]
                    ws.cell(
                        row=start_row + b + 1, column=4
                    ).value = risk_data.risk_impact_count[q][risk_cat][(cat, "Medium")]
                    ws.cell(
                        row=start_row + b + 1, column=5
                    ).value = risk_data.risk_impact_count[q][risk_cat][(cat, "High")]
                    ws.cell(
                        row=start_row + b + 1, column=6
                    ).value = risk_data.risk_count[q][risk_cat][cat]

                start_row += b + 4

    wb.remove(wb["Sheet"])

    return wb


VFM_LIST = [
    "NPV for all projects and NPV for programmes if available",
    "Adjusted Benefits Cost Ratio (BCR)",
    "Initial Benefits Cost Ratio (BCR)",
    "VfM Category single entry",
    "VfM Category lower range",
    "VfM Category upper range",
    "Present Value Cost (PVC)",
    "Present Value Benefit (PVB)",
    "Benefits Narrative",
]

VFM_CAT = [
    "Poor",
    "Low",
    "Medium",
    "High",
    "Very High",
    "Very High and Financially Positive",
    "Economically Positive",
    "Total",
]


class VfMData:
    def __init__(self, master: Master):
        self.master = master
        self.vfm_dictionary = {}
        self.vfm_cat_count = {}
        self.vfm_cat_pvc = {}
        self.get_dictionary()
        self.get_count()

    def get_dictionary(self) -> None:
        quarter_dict = {}
        for i in range(len(self.master.master_data)):
            project_dict = {}
            for project_name in self.master.master_data[i].projects:
                vfm_list = []
                for vfm_type in VFM_LIST:
                    try:
                        vfm = (
                            vfm_type,
                            self.master.master_data[i].data[project_name][vfm_type],
                        )
                        vfm_list.append(vfm)
                    except KeyError:  # vfm range keys not in all masters
                        pass

                project_dict[self.master.abbreviations[project_name]] = dict(vfm_list)
            quarter_dict[str(self.master.master_data[i].quarter)] = project_dict

            self.vfm_dictionary = quarter_dict

    def get_count(self) -> None:
        """Returns dictionary containing a count of vfm categories and pvc totals"""
        count_output_dict = {}
        pvc_output_dict = {}
        error_list = []
        for i, quarter in enumerate(self.vfm_dictionary.keys()):
            pvc_list = []
            cat_list = []
            for cat in VFM_CAT:
                cat_pvc_count = 0
                total_pvc_count = 0
                cat_count = 0
                total_count = 0
                for y, project in enumerate(list(self.vfm_dictionary[quarter].keys())):
                    try:
                        project_pvc = self.vfm_dictionary[quarter][project][
                            "Present Value Cost (PVC)"
                        ]
                        total_pvc_count += project_pvc
                        proj_cat = self.vfm_dictionary[quarter][project][
                            "VfM Category single entry"
                        ]
                        if proj_cat == cat:
                            cat_pvc_count += project_pvc
                    except TypeError:
                        if project_pvc is not None:
                            error_list.append(
                                quarter + " " + project + " PVC data needs checking"
                            )
                            pass
                    proj_cat = self.vfm_dictionary[quarter][project][
                        "VfM Category single entry"
                    ]
                    if proj_cat is not None:
                        total_count += 1
                        if proj_cat == cat:
                            cat_count += 1
                    if proj_cat is None:
                        if i == 0:
                            error_list.append(
                                quarter + " " + project + " VfM Category is None"
                            )

                pvc_list.append((cat, cat_pvc_count))
                cat_list.append((cat, cat_count))
            pvc_list.append(("Total", total_pvc_count))
            cat_list.append(("Total", total_count))
            pvc_output_dict[quarter] = dict(pvc_list)
            count_output_dict[quarter] = dict(cat_list)

        error_list = get_error_list(error_list)
        for x in error_list:
            print(x)

        self.vfm_cat_pvc = pvc_output_dict
        self.vfm_cat_count = count_output_dict


def vfm_into_excel(vfm_data: VfMData, quarter: List[str] or str) -> workbook:
    wb = Workbook()

    quarter = string_conversion(quarter)

    for q in quarter:
        start_row = 3
        ws = wb.create_sheet(
            make_file_friendly(q)
        )  # creating worksheets. names restricted to 30 characters.
        ws.title = make_file_friendly(q)  # title of worksheet
        for i, project_name in enumerate(list(vfm_data.vfm_dictionary[q].keys())):
            ws.cell(row=start_row + i, column=2).value = project_name
            for x, key in enumerate(
                list(vfm_data.vfm_dictionary[q][project_name].keys())
            ):
                ws.cell(row=2, column=3 + x).value = key
                ws.cell(
                    row=start_row + i, column=3 + x
                ).value = vfm_data.vfm_dictionary[q][project_name][key]

        ws.cell(row=2, column=2).value = "Project/Programme"

    start_row = 4
    ws = wb.create_sheet("Count")
    ws.title = "Count"
    for x, q in enumerate(quarter):
        ws.cell(row=3, column=3 + x).value = q
        ws.cell(row=3 + 12, column=3 + x).value = q
        for i, cat in enumerate(VFM_CAT):
            ws.cell(row=start_row + i, column=2).value = cat
            ws.cell(row=start_row + i + 12, column=2).value = cat
            try:
                ws.cell(row=start_row + i, column=3 + x).value = vfm_data.vfm_cat_pvc[
                    q
                ][cat]
                ws.cell(
                    row=start_row + i + 12, column=3 + x
                ).value = vfm_data.vfm_cat_count[q][cat]
            except KeyError:
                pass

    ws.cell(row=2, column=2).value = "PVC total per category"
    ws.cell(row=3, column=2).value = "Category"
    ws.cell(row=2 + 12, column=2).value = "Category count"
    ws.cell(row=3 + 12, column=2).value = "Category"

    wb.remove(wb["Sheet"])
    return wb


# for speed_dial analysis
# degree_range, rot_text, and gauge all in early development. Code taken from
# http://nicolasfauchereau.github.io/climatecode/posts/drawing-a-gauge-with-matplotlib/
def degree_range(n):
    start = np.linspace(0, 180, n + 1, endpoint=True)[0:-1]
    end = np.linspace(0, 180, n + 1, endpoint=True)[1::]
    mid_points = start + ((end - start) / 2.0)
    return np.c_[start, end], mid_points


def rot_text(ang):
    rotation = np.degrees(np.radians(ang) * np.pi / np.pi - np.radians(90))
    return rotation


def gauge(
    labels=["LOW", "MEDIUM", "HIGH", "VERY HIGH", "EXTREME"],
    colors="jet_r",
    arrow=1,
    arrow_two=2,
    title="",
    fname=False,
):
    """
    some sanity checks first

    """

    N = len(labels)

    if arrow > N:
        raise Exception(
            "\n\nThe category ({}) is greated than \
        the length\nof the labels ({})".format(
                arrow, N
            )
        )

    """
    if colors is a string, we assume it's a matplotlib colormap
    and we discretize in N discrete colors 
    """

    if isinstance(colors, str):
        cmap = cm.get_cmap(colors, N)
        cmap = cmap(np.arange(N))
        colors = cmap[::-1, :].tolist()
    if isinstance(colors, list):
        if len(colors) == N:
            colors = colors[::-1]
        else:
            raise Exception(
                "\n\nnumber of colors {} not equal \
            to number of categories{}\n".format(
                    len(colors), N
                )
            )

    """
    begins the plotting
    """

    fig, ax = plt.subplots()

    ang_range, mid_points = degree_range(N)
    print(ang_range)
    print(mid_points)

    labels = labels[::-1]

    """
    plots the sectors and the arcs
    """
    patches = []
    for ang, c in zip(ang_range, colors):
        # sectors
        patches.append(Wedge((0.0, 0.0), 0.4, *ang, facecolor="w", lw=2))
        # arcs
        patches.append(
            Wedge((0.0, 0.0), 0.4, *ang, width=0.10, facecolor=c, lw=2, alpha=0.5)
        )

    [ax.add_patch(p) for p in patches]

    """
    set the labels (e.g. 'LOW','MEDIUM',...)
    """

    for mid, lab in zip(mid_points, labels):
        ax.text(
            0.35 * np.cos(np.radians(mid)),
            0.35 * np.sin(np.radians(mid)),
            lab,
            horizontalalignment="center",
            verticalalignment="center",
            fontsize=14,
            fontweight="bold",
            rotation=rot_text(mid),
        )

    """
    set the bottom banner and the title
    """
    r = Rectangle((-0.4, -0.1), 0.8, 0.1, facecolor="w", lw=2)
    ax.add_patch(r)

    ax.text(
        0,
        -0.05,
        title,
        horizontalalignment="center",
        verticalalignment="center",
        fontsize=22,
        fontweight="bold",
    )

    """
    plots the arrow now
    """

    # pos = abs(arrow - N)
    pos = mid_points[abs(arrow - N)]
    print(pos)

    ax.arrow(
        0,
        0,
        0.225 * np.cos(np.radians(pos)),
        0.225 * np.sin(np.radians(pos)),
        width=0.04,
        head_width=0.09,
        head_length=0.1,
        fc="k",
        ec="k",
    )

    pos_two = mid_points[abs(arrow_two - N)]

    ax.arrow(
        0,
        0,
        0.225 * np.cos(np.radians(pos_two)),
        0.225 * np.sin(np.radians(pos_two)),
        width=0.04,
        head_width=0.09,
        head_length=0.1,
        fill=False,
        ec="k",
    )

    ax.add_patch(Circle((0, 0), radius=0.02, facecolor="k"))
    ax.add_patch(Circle((0, 0), radius=0.01, facecolor="w", zorder=11))

    """
    removes frame and ticks, and makes axis equal and tight
    """

    ax.set_frame_on(False)
    ax.axes.set_xticks([])
    ax.axes.set_yticks([])
    ax.axis("equal")
    plt.tight_layout()
    if fname:
        fig.savefig(root_path / "output/speed_dial_graph.png", dpi=200)
        # doc = open_word_doc(root_path / "input/summary_temp.docx")
        # doc.add_picture("temp_file.png", width=Inches(8))
        # doc.save(root_path / "output/speed_dial.docx")
        # os.remove("temp_file.png")


def sort_projects_by_dca(
    master_data: List[Dict[str, Union[str, int, datetime.date, float]]],
    projects: List[str] or str,
) -> List[str]:
    # returns a list of projects sorted by dca rag rating
    rag_list = []
    for project_name in projects:
        rag = master_data.data[project_name]["Departmental DCA"]
        rag_list.append((project_name, rag))

    rag_list_sorted = sorted(rag_list, key=lambda x: x[1])

    return rag_list_sorted


def cost_v_schedule_chart(milestones: MilestoneData, costs: CostData):

    rags = []
    for project_name in milestones.project_group:
        rag = milestones.master.master_data[0].data[project_name]["Departmental DCA"]
        rags.append((project_name, rag))

    rags = sorted(rags, key=lambda x: x[1])

    rag_count = Counter(x[1] for x in rags)

    wb = Workbook()
    ws = wb.active

    ws.cell(row=2, column=2).value = "Project Name"
    ws.cell(row=2, column=3).value = "Schedule change"
    ws.cell(row=2, column=4).value = "WLC Change"
    ws.cell(row=2, column=5).value = "WLC"
    ws.cell(row=2, column=6).value = "DCA"
    ws.cell(row=2, column=7).value = "Start key"
    ws.cell(row=2, column=8).value = "End key"

    for x, project_name in enumerate(rags):
        ab = milestones.master.abbreviations[project_name[0]]
        ws.cell(row=x + 3, column=2).value = ab
        ws.cell(row=x + 3, column=3).value = milestones.schedule_change[ab]["baseline"][
            "percent change"
        ]
        ws.cell(row=x + 3, column=4).value = costs.wlc_change[project_name[0]][
            "baseline one"
        ]
        ws.cell(row=x + 3, column=5).value = costs.master.master_data[0].data[
            project_name[0]
        ]["Total Forecast"]
        ws.cell(row=x + 3, column=6).value = costs.master.master_data[0].data[
            project_name[0]
        ]["Departmental DCA"]
        ws.cell(row=x + 3, column=7).value = milestones.schedule_change[ab]["baseline"][
            "start key"
        ]
        ws.cell(row=x + 3, column=8).value = milestones.schedule_change[ab]["baseline"][
            "end key"
        ]

    # bubble_chart(ws, rag_occurrence)

    return wb


def make_columns_bold(columns: list) -> None:
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def change_text_size(columns: list, size: int) -> None:
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(size)


def convert_bc_stage_text(bc_stage: str) -> str:
    '''
    function that converts bc stage.
    :param bc_stage: the string name for business cases that it kept in the master
    :return: standard/shorter string name
    '''

    if bc_stage == 'Strategic Outline Case':
        return 'SOBC'
    elif bc_stage == 'Outline Business Case':
        return 'OBC'
    elif bc_stage == 'Full Business Case':
        return 'FBC'
    elif bc_stage == 'pre-Strategic Outline Case':
        return 'pre-SOBC'
    else:
        return bc_stage


def make_text_red(columns: list) -> None:
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text == 'Not reported':
                        run.font.color.rgb = RGBColor(255, 0, 0)


def project_report_meta_data(doc: Document,
                             costs: CostData,
                             milestones: MilestoneData,
                             benefits: BenefitsData,
                             project_name: str):
    '''Meta data table'''
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run('Annex A. High level MI data and analysis').bold = True

    '''Costs meta data'''
    run = doc.add_paragraph().add_run('Costs')
    font = run.font
    font.bold = True
    font.underline = True
    t = doc.add_table(rows=1, cols=4)
    hdr_cells = t.rows[0].cells
    hdr_cells[0].text = 'WLC:'
    hdr_cells[1].text = '£' + str(round(costs.master.master_data[0].data[project_name]["Total Forecast"])) + 'm'
    hdr_cells[2].text = 'Spent:'
    # spent = spent_calculation(costs.master.master_data[0], project_name)
    hdr_cells[3].text = '£' + str(round(costs.spent[0])) + 'm'
    row_cells = t.add_row().cells
    row_cells[0].text = 'RDEL Total:'
    row_cells[1].text = '£' + str(round(sum(costs.rdel_profile))) + 'm'
    row_cells[2].text = 'Profiled:'
    row_cells[3].text = '£' + str(round(costs.profiled[0])) + 'm'  #  first in list is current
    row_cells = t.add_row().cells
    row_cells[0].text = 'CDEL Total:'
    row_cells[1].text = '£' + str(round(sum(costs.cdel_profile))) + 'm'
    row_cells[2].text = 'Unprofiled:'
    row_cells[3].text = '£' + str(round(costs.unprofiled[0])) + 'm'

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(t, column_widths)
    # make column keys bold
    make_columns_bold([t.columns[0], t.columns[2]])
    change_text_size([t.columns[0], t.columns[1], t.columns[2], t.columns[3]], 10)

    '''Financial data'''
    doc.add_paragraph()
    run = doc.add_paragraph().add_run('Financial')
    font = run.font
    font.bold = True
    font.underline = True
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Type of funding:'
    hdr_cells[1].text = str(costs.master.master_data[0].data[project_name]['Source of Finance'])
    hdr_cells[2].text = 'Contingency:'
    contingency = costs.master.master_data[0].data[project_name]['Overall contingency (£m)']
    if contingency is None: #  can this be refactored?
        hdr_cells[3].text = "None"
    else:
        hdr_cells[3].text = '£' + str(round(contingency)) + 'm'
    row_cells = table.add_row().cells
    row_cells[0].text = 'Optimism Bias (OB):'
    ob = costs.master.master_data[0].data[project_name]['Overall figure for Optimism Bias (£m)']
    if ob is None:
        row_cells[1].text = str(ob)
    else:
        try:
            row_cells[1].text = '£' + str(round(ob)) + 'm'
        except TypeError:
            row_cells[1].text = ob
    row_cells[2].text = 'Contingency in costs:'
    con_included_wlc = costs.master.master_data[0].data[project_name] \
        ['Is this Continency amount included within the WLC?']
    if con_included_wlc is None:
        row_cells[3].text = 'Not reported'
    else:
        row_cells[3].text = con_included_wlc
    row_cells = table.add_row().cells
    row_cells[0].text = 'OB in costs:'
    ob_included_wlc = costs.master.master_data[0].data[project_name] \
        ['Is this Optimism Bias included within the WLC?']
    if ob_included_wlc is None:
        row_cells[1].text = 'Not reported'
    else:
        row_cells[1].text = ob_included_wlc
    row_cells[2].text = ''
    row_cells[3].text = ''

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(table, column_widths)
    # make column keys bold
    make_columns_bold([table.columns[0], table.columns[2]])
    change_text_size([table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10)

    '''Project Stage data'''
    doc.add_paragraph()
    run = doc.add_paragraph().add_run('Stage')
    font = run.font
    font.bold = True
    font.underline = True
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Business case stage :'
    hdr_cells[1].text = convert_bc_stage_text(
            costs.master.master_data[0].data[project_name]['IPDC approval point'])
    hdr_cells[2].text = 'Delivery stage:'
    delivery_stage = costs.master.master_data[0].data[project_name]['Project stage']
    if delivery_stage is None:
        hdr_cells[3].text = "Not reported"
    else:
        hdr_cells[3].text = delivery_stage

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(table, column_widths)
    # make column keys bold
    make_columns_bold([table.columns[0], table.columns[2]])
    change_text_size([table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10)
    make_text_red([table.columns[1], table.columns[3]])  # make 'not reported red'

    '''Milestone/Stage meta data'''
    abb = milestones.master.abbreviations[project_name]
    doc.add_paragraph()
    run = doc.add_paragraph().add_run('Schedule/Milestones')
    font = run.font
    font.bold = True
    font.underline = True
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Start date:'
    try:
        start_project = get_milestone_date(abb, milestones.current, " Start of Project")
        hdr_cells[1].text = start_project.strftime("%d/%m/%Y")
    except KeyError:
        hdr_cells[1].text = "Not reported"
    except AttributeError:
        hdr_cells[1].text = "Not reported"
    hdr_cells[2].text = 'Start of operations:'
    try:
        start_ops = get_milestone_date(abb, milestones.current, " Start of Operation")
        hdr_cells[3].text = start_ops.strftime("%d/%m/%Y")
    except KeyError:
        hdr_cells[3].text = "Not reported"
    except AttributeError:
        hdr_cells[3].text = "Not reported"
    row_cells = table.add_row().cells
    row_cells[0].text = 'Start of construction:'
    try:
        start_con = get_milestone_date(abb, milestones.current, " Start of Construction/build")
        row_cells[1].text = start_con.strftime("%d/%m/%Y")
    except KeyError:
        row_cells[1].text = "Not reported"
    except AttributeError:
        row_cells[1].text = "Not reported"
    row_cells[2].text = 'Full Operations:'  # check
    try:
        full_ops = get_milestone_date(abb, milestones.current, " Full Operations")
        row_cells[3].text = full_ops.strftime("%d/%m/%Y")
    except KeyError:
        row_cells[3].text = "Not reported"
    except AttributeError:
        row_cells[3].text = "Not reported"

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(table, column_widths)
    # make column keys bold
    make_columns_bold([table.columns[0], table.columns[2]])
    change_text_size([table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10)
    make_text_red([table.columns[1], table.columns[3]])  # make 'not reported red'

    '''vfm meta data'''
    doc.add_paragraph()
    run = doc.add_paragraph().add_run('VfM')
    font = run.font
    font.bold = True
    font.underline = True
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'VfM category:'
    vfm_cat = costs.master.master_data[0].data[project_name]['VfM Category single entry']
    if vfm_cat is None:
        hdr_cells[1].text = "Not reported"
    else:
        hdr_cells[1].text = vfm_cat
    hdr_cells[2].text = 'BCR:'
    bcr = costs.master.master_data[0].data[project_name]['Adjusted Benefits Cost Ratio (BCR)']
    hdr_cells[3].text = str(bcr)

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(table, column_widths)
    # make column keys bold
    make_columns_bold([table.columns[0], table.columns[2]])
    change_text_size([table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10)
    make_text_red([table.columns[1], table.columns[3]])  # make 'not reported red'

    '''benefits meta data'''
    doc.add_paragraph()
    run = doc.add_paragraph().add_run('Benefits')
    font = run.font
    font.bold = True
    font.underline = True
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Total Benefits:'
    hdr_cells[1].text = '£' + str(round(benefits.master.master_data[0].data[project_name]["BEN Totals Forecast"])) + 'm'
    hdr_cells[2].text = 'Benefits delivered:'
    hdr_cells[3].text = '£' + str(round(benefits.delivered[0])) + 'm' #  first in list is current
    row_cells = table.add_row().cells
    row_cells[0].text = 'Benefits profiled:'
    row_cells[1].text = '£' + str(round(benefits.profiled[0])) + 'm'
    row_cells[2].text = 'Benefits unprofiled:'
    row_cells[3].text = '£' + str(round(benefits.unprofiled[0])) + 'm'

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(table, column_widths)
    # make column keys bold
    make_columns_bold([table.columns[0], table.columns[2]])
    change_text_size([table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10)
    return doc


def print_out_project_milestones(doc: Document, milestones: MilestoneData, project_name: str) -> Document:
    def plus_minus_days(change_value):
        '''mini function to place plus or minus sign before time delta
        value in milestone_table function. Only need + signs to be added
        as negative numbers have minus already'''
        try:
            if change_value > 0:
                text = '+ ' + str(change_value)
            else:
                text = str(change_value)
        except TypeError:
            text = change_value

        return text

    def get_milestone_notes(
            project_name: str,
            milestone_dictionary: Dict[str, Union[datetime.date, str]],
            milestone_name: str,
    ) -> datetime:
        for k in milestone_dictionary.keys():
            if milestone_dictionary[k]["Project"] == project_name:
                if milestone_dictionary[k]["Milestone"] == milestone_name:
                    return milestone_dictionary[k]["Notes"]

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    # table heading
    ab = milestones.master.abbreviations[project_name]
    doc.add_paragraph().add_run(str(ab + ' milestone table (2021 - 22)')).bold = True
    # some_text = 'The below table presents all project reported remaining high-level milestones, with six months grace ' \
                # 'from close of the current quarter. Milestones are sorted in chronological order. Changes in milestones' \
                # ' dates in comparison to last quarter and baseline have been calculated and are provided.'
    # doc.add_paragraph().add_run(str(some_text)).italic = True

    ab = milestones.master.abbreviations[project_name]

    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Milestone'
    hdr_cells[1].text = 'Date'
    hdr_cells[2].text = 'Change from last quarter'
    hdr_cells[3].text = 'Change from baseline'
    hdr_cells[4].text = 'Notes'

    for i, m in enumerate(milestones.key_names):
        row_cells = table.add_row().cells
        row_cells[0].text = m
        row_cells[1].text = milestones.md_current[i].strftime("%d/%m/%Y")
        try:
            row_cells[2].text = plus_minus_days((milestones.md_current[i] - milestones.md_last_po[i]).days)
        except TypeError:
            row_cells[2].text = "Not reported"
        try:
            row_cells[3].text = plus_minus_days((milestones.md_current[i] - milestones.md_baseline_po[i]).days)
        except TypeError:
            row_cells[3].text = "Not reported"
        try:
            row_cells[4].text = get_milestone_notes(ab, milestones.current, m)
            paragraph = row_cells[4].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size = Pt(8)  # font size = 8
        except TypeError:
            pass

        # try:
        #     if milestone_filter_start_date <= milestone_date:  # filter based on date
        #         row_cells = table.add_row().cells
        #         row_cells[0].text = milestone
        #         if milestone_date is None:
        #             row_cells[1].text = 'No date'
        #         else:
        #             row_cells[1].text = milestone_date.strftime("%d/%m/%Y")
        #         b_one_value = first_diff_data[project_name][milestone]
        #         row_cells[2].text = plus_minus_days(b_one_value)
        #         b_two_value = second_diff_data[project_name][milestone]
        #         row_cells[3].text = plus_minus_days(b_two_value)
        #
        #         notes = p_current_milestones[project_name][milestone][milestone_date]
        #         row_cells[4].text = str(notes)
        #         # trying to high changes to narratuve in red text
        #         # if milestone in p_last_milestones[project_name].keys():
        #         #     last_milestone_date = p_last_milestones[project_name][milestone]
        #         #     last_note = p_last_milestones[project_name][milestone][last_milestone_date]
        #         #     row_cells[4] = compare_text_newandold(notes, last_note, doc)
        #         # elif milestone not in p_last_milestones[project_name].keys():
        #         #     row_cells[4].text = str(notes)
        #
        #         paragraph = row_cells[4].paragraphs[0]
        #         run = paragraph.runs
        #         font = run[0].font
        #         font.size = Pt(8)  # font size = 8
        #
        #
        # except TypeError:  # this is to deal with none types which are still placed in output
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = milestone
        #     if milestone_date is None:
        #         row_cells[1].text = 'No date'
        #     else:
        #         row_cells[1].text = milestone_date.strftime("%d/%m/%Y")
        #     b_one_value = first_diff_data[project_name][milestone]
        #     row_cells[2].text = plus_minus_days(b_one_value)
        #     b_two_value = second_diff_data[project_name][milestone]
        #     row_cells[3].text = plus_minus_days(b_two_value)
        #     notes = p_current_milestones[project_name][milestone][milestone_date]
        #     row_cells[4].text = str(notes)
        #     paragraph = row_cells[4].paragraphs[0]
        #     run = paragraph.runs
        #     font = run[0].font
        #     font.size = Pt(8)  # font size = 8

    table.style = 'Table Grid'

    # column widths
    column_widths = (Cm(6), Cm(2.6), Cm(2), Cm(2), Cm(8.95))
    set_col_widths(table, column_widths)
    # make_columns_bold([table.columns[0], table.columns[3]])  # make keys bold
    # make_text_red([table.columns[1], table.columns[4]])  # make 'not reported red'

    make_rows_bold([table.rows[0]])  # makes top of table bold. Found function on stack overflow.
    return doc


def project_scope_text(doc: Document, master: Master, project_name: str) -> Document:
    doc.add_paragraph().add_run('Project Scope').bold = True
    text_one = str(master.master_data[0].data[project_name]['Project Scope'])
    try:
        text_two = str(master.master_data[1].data[project_name]['Project Scope'])
    except KeyError:
        text_two = text_one
    # different options for comparing costs
    # compare_text_showall(dca_a, dca_b, doc)
    compare_text_new_and_old(text_one, text_two, doc)
    return doc