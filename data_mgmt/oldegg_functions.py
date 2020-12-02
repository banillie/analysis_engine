"""place for storing all master templates"""
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

'''

Store of common functions found in all analysis engine

This is also where the date of BICC is set. This is the date from which much of the analysis is set.
NOTE. Python date format is (YYYY,MM,DD)

'''

import random
from datetime import datetime
from collections import Counter
from openpyxl.styles import Font, PatternFill
from openpyxl.styles.differential import DifferentialStyle
from openpyxl.formatting import Rule
import difflib
from docx.shared import RGBColor


def compare_text_newandold(text_1: str, text_2: str, doc):
    """
    This function places text into word document highlighting in red
    new text (against what was reported the previous quarter. It is old
    and could probably benefit from some refactoring.
    text_1: latest text. string.
    text_2: last text. string
    doc: word doc
    """

    comp = difflib.Differ()
    diff = list(comp.compare(text_2.split(), text_1.split()))
    new_text = diff
    y = doc.add_paragraph()
    y.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for i in range(0, len(diff)):
        f = len(diff) - 1
        if i < f:
            a = i - 1
        else:
            a = i

        if diff[i][0:3] == '  |':
            j = i + 1
            if diff[i][0:3] and diff[a][0:3] == '  |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '+ |':
            if diff[i][0:3] and diff[a][0:3] == '+ |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '- |':
            pass
        elif diff[i][0:3] == '  -':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0:3] == '  •':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0] == '+':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.color.rgb = RGBColor(255, 0, 0)
        elif diff[i][0] == '-':
            pass
        elif diff[i][0] == '?':
            pass
        else:
            if diff[i] != '+ |':
                y.add_run(diff[i][1:])

    return doc


def cell_colouring(cell, colour):
    """
    Function that handles cell colouring for word documents.
    cell: cell reference
    color: colour reference
    """

    try:
        if colour == 'R':
            colour = parse_xml(r'<w:shd {} w:fill="cb1f00"/>'.format(nsdecls('w')))
        elif colour == 'A/R':
            colour = parse_xml(r'<w:shd {} w:fill="f97b31"/>'.format(nsdecls('w')))
        elif colour == 'A':
            colour = parse_xml(r'<w:shd {} w:fill="fce553"/>'.format(nsdecls('w')))
        elif colour == 'A/G':
            colour = parse_xml(r'<w:shd {} w:fill="a5b700"/>'.format(nsdecls('w')))
        elif colour == 'G':
            colour = parse_xml(r'<w:shd {} w:fill="17960c"/>'.format(nsdecls('w')))

        cell._tc.get_or_add_tcPr().append(colour)

    except TypeError:
        pass


def make_rows_bold(rows: list):
    """This function makes text bold in a list of row numbers for a word document"""
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def set_col_widths(table, widths):
    """This function sets the width of table in a word document"""
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def project_all_milestones_dict(project_names,
                                master_data,
                                baseline_index,
                                data_to_return=int):
    '''
    Function that puts project milestone data in dictionary in order of newest date first.

    Project_names: list of project names of interest / in range
    Master_data: quarter master data set

    Dictionary is structured as {'project name': {'milestone name': datetime.date: 'notes'}}

    '''

    upper_dict = {}

    for name in project_names:
        lower_dict = {}
        raw_list = []
        try:
            p_data = master_data[baseline_index[name][data_to_return]].data[name]
            for i in range(1, 50):
                try:
                    m_date = p_data['Approval MM' + str(i) + ' Forecast / Actual']
                    if type(m_date) == str:
                        m_date = datetime.strptime(m_date, "%d/%m/%Y").date()
                    t = (p_data['Approval MM' + str(i)],
                         m_date,
                         p_data['Approval MM' + str(i) + ' Notes'])
                    raw_list.append(t)
                except KeyError:
                    try:
                        m_date = p_data['Approval MM' + str(i) + ' Forecast - Actual']
                        if type(m_date) == str:
                            m_date = datetime.strptime(m_date, "%d/%m/%Y").date()
                        t = (p_data['Approval MM' + str(i)],
                             m_date,
                             p_data['Approval MM' + str(i) + ' Notes'])
                        raw_list.append(t)
                    except KeyError:
                        pass

            for i in range(1, 50):
                try:
                    m_date = p_data['Assurance MM' + str(i) + ' Forecast - Actual']
                    if type(m_date) == str:
                        m_date = datetime.strptime(m_date, "%d/%m/%Y").date()
                    t = (p_data['Assurance MM' + str(i)],
                         m_date,
                         p_data['Assurance MM' + str(i) + ' Notes'])
                    raw_list.append(t)
                except KeyError:
                    pass

            for i in range(18, 67):
                try:
                    m_date = p_data['Project MM' + str(i) + ' Forecast - Actual']
                    if type(m_date) == str:
                        m_date = datetime.strptime(m_date, "%d/%m/%Y").date()
                    t = (p_data['Project MM' + str(i)],
                         m_date,
                         p_data['Project MM' + str(i) + ' Notes'])
                    raw_list.append(t)
                except KeyError:
                    pass

        except (KeyError, TypeError, ValueError):  # further testing required to understand this exception handling.
            pass

        except IndexError:
            print('warning ' + name + ' does not have complete baseline index list')

        # put the list in chronological order
        sorted_list = sorted(raw_list, key=lambda k: (k[1] is None, k[1]))

        # loop to stop keys names being the same. Not ideal as doesn't handle keys that may already have numbers as
        # strings at end of names. But still useful.
        for x in sorted_list:
            if x[0] is not None:
                if x[0] in lower_dict:
                    for i in range(2, 15):
                        key_name = x[0] + ' ' + str(i)
                        if key_name in lower_dict:
                            continue
                        else:
                            lower_dict[key_name] = {x[1]: x[2]}
                            break
                else:
                    lower_dict[x[0]] = {x[1]: x[2]}
            else:
                pass

        upper_dict[name] = lower_dict

    return upper_dict


def project_time_difference(proj_m_data_1, proj_m_data_2):
    """Function that calculates time different between milestone dates"""
    upper_dictionary = {}

    for proj_name in proj_m_data_1:
        td_dict = {}
        for milestone in proj_m_data_1[proj_name]:
            milestone_date = tuple(proj_m_data_1[proj_name][milestone])[0]
            if milestone_date is None:
                td_dict[milestone] = 'No date'
            else:
                try:
                    old_milestone_date = tuple(proj_m_data_2[proj_name][milestone])[0]
                    time_delta = (milestone_date - old_milestone_date).days  # time_delta calculated here
                    if time_delta == 0:
                        td_dict[milestone] = 0
                    else:
                        td_dict[milestone] = time_delta
                except (KeyError, TypeError):
                    td_dict[milestone] = 'Not reported'  # not reported that quarter

        upper_dictionary[proj_name] = td_dict

    return upper_dictionary


def filter_gmpp(master):
    project_list = []
    for project_name in master.projects:
        if master.data[project_name]['GMPP - IPA ID Number'] is not None:
            project_list.append(project_name)

    return project_list


def convert_rag_text(dca_rating):
    if dca_rating == 'Green':
        return 'G'
    elif dca_rating == 'Amber/Green':
        return 'A/G'
    elif dca_rating == 'Amber':
        return 'A'
    elif dca_rating == 'Amber/Red':
        return 'A/R'
    elif dca_rating == 'Red':
        return 'R'
    else:
        return ''


def filter_project_group(master_data, group):
    '''
    function for return a list of projects according to their group
    :param master_data: one quarters master data
    :param group: group name of interest. this is a string.
    options are 'Rail Group', 'HSMRPG', 'International Security and Environment', 'Roads Devolution & Motoring'.
    Note this list should be kept up to date as group names change.
    :return: list of projects in specified group
    '''

    project_name_list = master_data.projects

    output_list = []

    for project_name in project_name_list:
        if master_data.data[project_name]['DfT Group'] == group:
            output_list.append(project_name)
        else:
            pass

    return output_list


def get_all_project_names(masters_list):
    '''
    function returns list of all projects across multiple dictionaries

    useful if you need project names across multiple quarters

    masters_list: list of masters containing quarter information
    '''

    output_list = []
    for master in masters_list:
        for name in master.projects:
            if name not in output_list:
                output_list.append(name)

    return output_list


def get_quarter_stamp(masters_list):
    '''
    Function used to specify the quarter being reported.

    masters_list: list of masters containing quarter information
    '''

    output_list = []
    for master in masters_list:
        project_name = random.choice(master.projects)
        quarter_stamp = master.data[project_name]['Reporting period (GMPP - Snapshot Date)']
        output_list.append(quarter_stamp)

    return output_list


def concatenate_dates(date, bicc_date):
    '''
    function for converting dates into concatenated written time periods
    :param date: datetime.date
    :return: concatenated date
    '''
    if date != None:
        a = (date - bicc_date).days
        year = 365
        month = 30
        fortnight = 14
        week = 7
        if a >= 365:
            yrs = int(a / year)
            holding_days_years = a % year
            months = int(holding_days_years / month)
            holding_days_months = a % month
            fortnights = int(holding_days_months / fortnight)
            weeks = int(holding_days_months / week)
        elif 0 <= a <= 365:
            yrs = 0
            months = int(a / month)
            holding_days_months = a % month
            fortnights = int(holding_days_months / fortnight)
            weeks = int(holding_days_months / week)
            # if 0 <= a <=60:
        elif a <= -365:
            yrs = int(a / year)
            holding_days = a % -year
            months = int(holding_days / month)
            holding_days_months = a % -month
            fortnights = int(holding_days_months / fortnight)
            weeks = int(holding_days_months / week)
        elif -365 <= a <= 0:
            yrs = 0
            months = int(a / month)
            holding_days_months = a % -month
            fortnights = int(holding_days_months / fortnight)
            weeks = int(holding_days_months / week)
            # if -60 <= a <= 0:
        else:
            print('something is wrong and needs checking')

        if yrs == 1:
            if months == 1:
                return ('{} yr, {} mth'.format(yrs, months))
            if months > 1:
                return ('{} yr, {} mths'.format(yrs, months))
            else:
                return ('{} yr'.format(yrs))
        elif yrs > 1:
            if months == 1:
                return ('{} yrs, {} mth'.format(yrs, months))
            if months > 1:
                return ('{} yrs, {} mths'.format(yrs, months))
            else:
                return ('{} yrs'.format(yrs))
        elif yrs == 0:
            if a == 0:
                return ('Today')
            elif 1 <= a <= 6:
                return ('This week')
            elif 7 <= a <= 13:
                return ('Next week')
            elif -7 <= a <= -1:
                return ('Last week')
            elif -14 <= a <= -8:
                return ('-2 weeks')
            elif 14 <= a <= 20:
                return ('2 weeks')
            elif 20 <= a <= 60:
                if bicc_date.month == date.month:
                    return ('Later this mth')
                elif (date.month - bicc_date.month) == 1:
                    return ('Next mth')
                else:
                    return ('2 mths')
            elif -60 <= a <= -15:
                if bicc_date.month == date.month:
                    return ('Earlier this mth')
                elif (date.month - bicc_date.month) == -1:
                    return ('Last mth')
                else:
                    return ('-2 mths')
            elif months == 12:
                return ('1 yr')
            else:
                return ('{} mths'.format(months))

        elif yrs == -1:
            if months == -1:
                return ('{} yr, {} mth'.format(yrs, -(months)))
            if months < -1:
                return ('{} yr, {} mths'.format(yrs, -(months)))
            else:
                return ('{} yr'.format(yrs))
        elif yrs < -1:
            if months == -1:
                return ('{} yrs, {} mth'.format(yrs, -(months)))
            if months < -1:
                return ('{} yrs, {} mths'.format(yrs, -(months)))
            else:
                return ('{} yrs'.format(yrs))
    else:
        return ('None')


def up_or_down(latest_dca, last_dca):
    '''
    function that calculates if confidence has increased or decreased
    :param latest_dca:
    :param last_dca:
    :return:
    '''

    if latest_dca == last_dca:
        return (int(0))
    elif latest_dca != last_dca:
        if last_dca == 'Green':
            if latest_dca != 'Amber/Green':
                return (int(-1))
        elif last_dca == 'Amber/Green':
            if latest_dca == 'Green':
                return (int(1))
            else:
                return (int(-1))
        elif last_dca == 'Amber':
            if latest_dca == 'Green':
                return (int(1))
            elif latest_dca == 'Amber/Green':
                return (int(1))
            else:
                return (int(-1))
        elif last_dca == 'Amber/Red':
            if latest_dca == 'Red':
                return (int(-1))
            else:
                return (int(1))
        else:
            return (int(1))


def convert_bc_stage_text(bc_stage):
    '''
    function that converts bc stage.
    :param bc_stage: the string name for business cases that it kept in the master
    :return: standard/shorter string name
    '''

    if bc_stage == 'Strategic Outline Case':
        return 'SOBC'
    elif bc_stage == 'Outline Business Case':
        return 'OBC'
    elif bc_stage == 'Full Business Case':
        return 'FBC'
    elif bc_stage == 'pre-Strategic Outline Case':
        return 'pre-SOBC'
    else:
        return bc_stage


def combine_narrtives(project_name, master_data, key_list):
    '''
    Function that combines narratives across keys
    :param project_name: project name
    :param master_data: master data from one quarter
    :param key_list: list of keys that contain the narrative (values) to be combined.
    :return: combined narrative
    '''
    output = ''
    for key in key_list:
        output = output + str(master_data[project_name][key])

    return output


def baseline_information_bc(project_list, masters_list):
    '''
    Function that calculates when project business case has changed. Only returns where there have been changes.
    :param project_list: list of project names
    :param masters_list: list of masters with quarter information
    :return: python dictionary in format 'project name':('BC', 'Quarter Stamp', index position in masters_list)
    '''
    output = {}

    for project_name in project_list:
        lower_list = []
        for i, master in enumerate(masters_list):
            if project_name in master.projects:
                approved_bc = master.data[project_name]['IPDC approval point']
                quarter = master.data[project_name]['Reporting period (GMPP - Snapshot Date)']
                try:
                    previous_approved_bc = masters_list[i + 1].data[project_name]['IPDC approval point']
                    if approved_bc != previous_approved_bc:
                        lower_list.append((approved_bc, quarter, i))
                except IndexError:
                    # this captures the last available quarter data if project was in portfolio from beginning
                    lower_list.append((approved_bc, quarter, i))
                except KeyError:
                    # this captures the first quarter the project reported if not in portfolio from beginning
                    lower_list.append((approved_bc, quarter, i))

        output[project_name] = lower_list

    return output


def baseline_information(project_list, masters_list, data_baseline):
    '''
    Function that calculates in information within masters has been baselined
    :param project_list: list of project names
    :param masters_list: list of quarter masters.
    :param data_baseline: type of information to check for baselines. options are: 'ALB milestones' etc
    :return: python dictionary structured as 'project name': ('yes (if reported that quarter), 'quarter stamp',
    index position of master in master quarter list)
    '''

    output = {}

    for project_name in project_list:
        lower_list = []
        for i, master in enumerate(masters_list):
            if project_name in master.projects:
                try:
                    approved_bc = master.data[project_name][data_baseline]
                    quarter = str(master.quarter)
                    if approved_bc == 'Yes':
                        lower_list.append((approved_bc, quarter, i))
                except KeyError:
                    pass

        output[project_name] = lower_list

    return output


def baseline_index(baseline_data, master_list):
    '''
    Function that calculates the index list for baseline data
    :param baseline_data: output created by either baseline_information or baseline_information_bc functions
    :return: python dictionary in format 'project name':[index list]
    '''

    left_lst_qrt = [x for x in master_list[1].projects if x not in master_list[0].projects]
    new_this_qrt = [x for x in master_list[0].projects if x not in master_list[1].projects]

    output = {}

    for project_name in baseline_data:
        if project_name in master_list[0].projects:
            if project_name in new_this_qrt:
                lower_list = [0, None, 0]  # this is to handle new projects that have only one quarters reporting
            else:
                lower_list = [0, 1]
                for tuple_info in baseline_data[project_name]:
                    lower_list.append(tuple_info[2])
        else:
            if project_name in left_lst_qrt:
                lower_list = [None, 1]
                for tuple_info in baseline_data[project_name]:
                    lower_list.append(tuple_info[2])
            else:
                lower_list = [None, None]
                for tuple_info in baseline_data[project_name]:
                    lower_list.append(tuple_info[2])

        output[project_name] = lower_list

    return output


def get_project_cost_profile(project_name_list, q_masters_data_list, cost_list, year_list, bc_index, index):
    '''
    Function that gets projects project cost information and returns it in a python dictionary format.
    :param project_name_list: list of project names
    :param q_masters_data_list: list of master python dictionaries containing quarter information
    :param cost_list: list of cost keys names. this is necessary due to the total cost having be calculated across
    rdel, cdel and non-gov breakdown.
    :param year_list: list of year keys e.g. '19-20', '20-21'
    :param index: index value for which master to use from the q_master_data_list . 0 is for latest, 1 last and
    2 baseline. The actual index list q_master_list is set at a global level in this programme.
    :return: a dictionary structured 'project_name': 'year rdel' : value, 'year cdel' : value, 'year Non-Gov' : value,
    'year total' : value
    '''

    upper_dictionary = {}

    for project_name in project_name_list:
        lower_dictionary = {}
        for year in year_list:
            try:
                project_data = q_masters_data_list[bc_index[project_name][index]].data[project_name]
            except (IndexError, TypeError):  # TypeError deals with None Types.
                project_data = q_masters_data_list[bc_index[project_name][-1]].data[project_name]

            try:
                total = 0
                for type in cost_list:

                    try:
                        lower_dictionary[year + type] = project_data[year + type]
                    except KeyError:
                        lower_dictionary[year + type] = None

                    if year + type in project_data.keys():
                        cost = project_data[year + type]
                        try:
                            total = total + cost
                        except TypeError:
                            pass

                lower_dictionary[year + ' total'] = total

            except TypeError:
                lower_dictionary[year + ' total'] = None

        upper_dictionary[project_name] = lower_dictionary

    return upper_dictionary


def get_project_income_profile(project_name_list, q_masters_data_list, income_list, year_list, bc_index, index):
    '''
    Function that gets projects project income information and returns it in a python dictionary format.
    :param project_name_list: list of project names
    :param q_masters_data_list: list of master python dictionaries containing quarter information
    :param income_list: list of income keys names.
    :param year_list: list of year keys e.g. '19-20', '20-21'
    :param index: index value for which master to use from the q_master_data_list . 0 is for latest, 1 last and
    2 baseline. The actual index list q_master_list is set at a global level in this programme.
    :return: a dictionary structured 'project_name' : 'year income' : value
    '''

    upper_dictionary = {}

    for project_name in project_name_list:
        lower_dictionary = {}
        for year in year_list:
            try:
                project_data = q_masters_data_list[bc_index[project_name][index]].data[project_name]
            except (IndexError, TypeError):  # TypeError deals with None Types
                project_data = q_masters_data_list[bc_index[project_name][-1]].data[project_name]

            try:
                for type in income_list:
                    try:
                        lower_dictionary[year + type] = project_data[year + type]
                    except KeyError:
                        lower_dictionary[year + type] = 0
            except TypeError:
                for type in income_list:
                    lower_dictionary[year + type] = 0

        upper_dictionary[project_name] = lower_dictionary

    return upper_dictionary


def calculate_group_project_total(project_name_list, master_data, project_name_no_count_list, type_list, year_list):
    '''
    calculates the total cost figure for each year and type of spend e.g. RDEL 19-20, for all projects of interest.
    :param project_name_list: list of project names
    :param master_data: master data set as created by the get_project_cost_profile
    :param project_name_no_count_list: list of project names to remove from total figures, to ensure no double counting
    e.g. if there are separate schemes as well as overall programme reporting.
    :param type_list: the type of financial figure list being counted. e.g. costs or income
    :return: python dictionary in format 'year + spend type': total
    '''

    output = {}

    project_list = [x for x in project_name_list if x not in project_name_no_count_list]

    for cost in type_list:
        for year in year_list:
            total = 0
            for project_name in project_list:
                try:
                    total = total + master_data[project_name][year + cost]
                except (KeyError, TypeError):
                    total = total + 0

            output[year + cost] = total

    return output


def grey_conditional_formatting(ws):
    '''
    function applies grey conditional formatting for 'Not Reporting'.
    :param worksheet: ws
    :return: cf of sheet
    '''

    grey_text = Font(color="f0f0f0")
    grey_fill = PatternFill(bgColor="f0f0f0")
    dxf = DifferentialStyle(font=grey_text, fill=grey_fill)
    rule = Rule(type="containsText", operator="containsText", text="Not reporting", dxf=dxf)
    rule.formula = ['NOT(ISERROR(SEARCH("Not reporting",A1)))']
    ws.conditional_formatting.add('A1:X80', rule)

    grey_text = Font(color="cfcfea")
    grey_fill = PatternFill(bgColor="cfcfea")
    dxf = DifferentialStyle(font=grey_text, fill=grey_fill)
    rule = Rule(type="containsText", operator="containsText", text="Data not collected", dxf=dxf)
    rule.formula = ['NOT(ISERROR(SEARCH("Data not collected",A1)))']
    ws.conditional_formatting.add('A1:X80', rule)

    return ws


def conditional_formatting(ws, list_columns, list_conditional_text, list_text_colours, list_background_colours,
                           row_start, row_end):
    for column in list_columns:
        for i, txt in enumerate(list_conditional_text):
            text = list_text_colours[i]
            fill = list_background_colours[i]
            dxf = DifferentialStyle(font=text, fill=fill)
            rule = Rule(type="containsText", operator="containsText", text=txt, dxf=dxf)
            for_rule_formula = 'NOT(ISERROR(SEARCH("' + txt + '",' + column + '1)))'
            rule.formula = [for_rule_formula]
            ws.conditional_formatting.add(column + row_start + ':' + column + row_end, rule)

    return ws


def compare_text_newandold(text_1, text_2, doc):
    '''
    function that places text into doc highlighting new and old text
    text_1: latest text. string.
    text_2: last text. string
    doc: word doc
    '''

    comp = difflib.Differ()
    diff = list(comp.compare(text_2.split(), text_1.split()))
    new_text = diff
    y = doc.add_paragraph()

    for i in range(0, len(diff)):
        f = len(diff) - 1
        if i < f:
            a = i - 1
        else:
            a = i

        if diff[i][0:3] == '  |':
            j = i + 1
            if diff[i][0:3] and diff[a][0:3] == '  |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '+ |':
            if diff[i][0:3] and diff[a][0:3] == '+ |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '- |':
            pass
        elif diff[i][0:3] == '  -':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0:3] == '  •':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0] == '+':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.color.rgb = RGBColor(255, 0, 0)
        elif diff[i][0] == '-':
            pass
        elif diff[i][0] == '?':
            pass
        else:
            if diff[i] != '+ |':
                y.add_run(diff[i][1:])

    return doc


def get_ben_totals(project_name, benefits_bl_index, list_of_masters_all):
    '''gets benefits data to place into the bar chart element in the financial analysis graphs'''

    ben_change_key_list = ['Pre-profile BEN Total',
                           "Total BEN Forecast - Total Monetised Benefits",
                           'Unprofiled Remainder BEN Forecast - Total Monetised Benefits']

    ben_type_key_list = ['Pre-profile BEN Forecast Gov Cashable',
                         'Pre-profile BEN Forecast Gov Non-Cashable',
                         'Pre-profile BEN Forecast - Economic (inc Private Partner)',
                         'Pre-profile BEN Forecast - Disbenefit UK Economic',
                         'Unprofiled Remainder BEN Forecast - Gov. Cashable',
                         'Unprofiled Remainder BEN Forecast - Gov. Non-Cashable',
                         'Unprofiled Remainder BEN Forecast - Economic (inc Private Partner)',
                         'Unprofiled Remainder BEN Forecast - Disbenefit UK Economic',
                         'Total BEN Forecast - Gov. Cashable',
                         'Total BEN Forecast - Gov. Non-Cashable',
                         'Total BEN Forecast - Economic (inc Private Partner)',
                         'Total BEN Forecast - Disbenefit UK Economic']

    ben_list = []
    index_1 = benefits_bl_index[project_name]
    index_2 = index_1[0:3]
    index_2.reverse()
    for x in index_2:
        if x is not None:
            for y in ben_change_key_list:
                ben = list_of_masters_all[x].data[project_name][y]
                ben_list.append(ben)
        else:
            for i in range(len(ben_change_key_list)):
                ben = 0
                ben_list.append(ben)

    ben_type_list = []
    for y in ben_type_key_list:
        ben = list_of_masters_all[0].data[project_name][y]
        if ben is not None:
            ben_type_list.append(ben)
        else:
            ben_type_list.append(0)

    return ben_list, ben_type_list


def spent_calculation(master: dict, project: str) -> int:
    keys = ['Pre-profile RDEL', '20-21 RDEL STD Total', 'Pre-profile CDEL', '20-21 CDEL STD Total']

    total = 0
    for k in keys:
        total += master.data[project][k]

    return total
