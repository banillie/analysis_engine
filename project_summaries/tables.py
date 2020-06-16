from analysis.engine_functions import all_milestones_dict
from analysis.data import list_of_masters_all, bc_index, baseline_bc_stamp, root_path, milestone_analysis_date
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT

def milestone_table(function):

    doc = Document()
    sections = doc.sections
    section = sections[0]

    font = doc.styles['Normal'].font
    font.name = 'Arial'
    font.size = Pt(10)

    heading = str('flip out')
    intro = doc.add_heading(str(heading), 0)
    intro.alignment = 1
    intro.bold = True

    section_2 = sections[0]
    new_width, new_height = section_2.page_height, section.page_width
    section_2.orientation = WD_ORIENT.LANDSCAPE
    section_2.page_width = new_width
    section_2.page_height = new_height

    #Paragraph spacing
    #doc.paragraph_format.space_before = Pt(18)
    #doc.paragraph_format.space_after = Pt(12)

    for project_name in list_of_masters_all[0].projects:

        print(project_name)
        y = doc.add_paragraph()
        heading = project_name
        y.add_run(str(heading)).bold = True

        p_oldest_milestones = function([project_name], list_of_masters_all[bc_index[project_name][2]])
        second_diff_data = project_time_difference(p_current_milestones, p_oldest_milestones)

        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Milestone'
        hdr_cells[1].text = 'Date'
        hdr_cells[2].text = 'Change from Lst Qrt'
        hdr_cells[3].text = 'Change from BL'
        hdr_cells[4].text = 'Notes'

        # TODO specify column widths

        for milestone in p_current_milestones[project_name].keys():

            milestone_date = tuple(p_current_milestones[project_name][milestone])[0]

            try:
                if milestone_analysis_date <= milestone_date: # filter based on date
                    row_cells = table.add_row().cells
                    row_cells[0].text = milestone
                    if milestone_date is None:
                        row_cells[1].text = 'No date'
                    else:
                        row_cells[1].text = milestone_date.strftime("%d/%m/%Y")
                    b_one_value = first_diff_data[project_name][milestone]
                    row_cells[2].text = str(b_one_value)
                    b_two_value = second_diff_data[project_name][milestone]
                    row_cells[3].text = str(b_two_value)
                    notes = p_current_milestones[project_name][milestone][milestone_date]
                    if notes is not None:
                        row_cells[4].text = notes

            except TypeError:  # this is to deal with none types which are still placed in output
                row_cells = table.add_row().cells
                row_cells[0].text = milestone
                if milestone_date is None:
                    row_cells[1].text = 'No date'
                else:
                    row_cells[1].text = milestone_date.strftime("%d/%m/%Y")
                b_one_value = first_diff_data[project_name][milestone]
                row_cells[2].text = str(b_one_value)
                b_two_value = second_diff_data[project_name][milestone]
                row_cells[3].text = str(b_two_value)
                notes = p_current_milestones[project_name][milestone][milestone_date]
                if notes is not None:
                    row_cells[4].text = notes

        table.style = 'Table Grid'

        make_rows_bold(table.rows[0]) # makes top of table bold. Found function on stack overflow.

    doc.save(root_path / 'output/table.docx')

def make_rows_bold(*rows):
    '''Makes text bold in specified row'''
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def project_time_difference(proj_m_data_1, proj_m_data_2):
    """Function that calculates time different between milestone dates"""
    upper_dictionary = {}

    for proj_name in proj_m_data_1:
        td_dict = {}
        for milestone in proj_m_data_1[proj_name]:
            milestone_date = tuple(proj_m_data_1[proj_name][milestone])[0]
            if milestone_date is None:
                td_dict[milestone] = 'No date'
            else:
                try:
                    old_milestone_date = tuple(proj_m_data_2[proj_name][milestone])[0]
                    time_delta = (milestone_date - old_milestone_date).days  # time_delta calculated here
                    if time_delta == 0:
                        td_dict[milestone] = 0
                    else:
                        td_dict[milestone] = time_delta
                except (KeyError, TypeError):
                    td_dict[milestone] = 'Not reported' # not reported that quarter

        upper_dictionary[proj_name] = td_dict

    return upper_dictionary


'''get all milestone data'''
p_current_milestones = all_milestones_dict(list_of_masters_all[0].projects, list_of_masters_all[0])
p_last_milestones = all_milestones_dict(list_of_masters_all[1].projects, list_of_masters_all[1])

'''calculate time current and last quarter'''
first_diff_data = project_time_difference(p_current_milestones, p_last_milestones)

table(all_milestones_dict)



