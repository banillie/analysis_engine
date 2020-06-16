import matplotlib.pyplot as plt
from analysis.data import latest_cost_profiles, last_cost_profiles, baseline_1_cost_profiles, \
    year_list, root_path, a66, a303, crossrail, thameslink
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

doc = Document()
amended_year_list = year_list[:-1]
#plt.style.use('fivethirtyeight') #maybe? or sort graph style?

def get_financial_data(project_name, cost_type):
    '''gets project financial data'''
    latest = []
    last = []
    baseline = []
    for year in amended_year_list:
        baseline.append(baseline_1_cost_profiles[project_name][year + cost_type])
        last.append(last_cost_profiles[project_name][year + cost_type])
        latest.append(latest_cost_profiles[project_name][year + cost_type])

    return latest, last, baseline

def single_chart(project_name_list):
    '''creates matplotlob chart and places it in word doc'''
    for project_name in project_name_list:
        y = doc.add_paragraph()
        heading = 'Financial Analysis - Cost Profile'
        y.add_run(str(heading)).bold = True

        '''financial chart'''

        profile_data = get_financial_data(project_name, ' total')

        year = ['19/20', '20/21', '21/22', '22/23', '23/24', '24/25', '25/26', '26/27', '27/28', '28/29']
        baseline_profile = profile_data[2]
        last_profile = profile_data[1]
        latest_profile = profile_data[0]


        fig, ax = plt.subplots(figsize=(20,10))
        #plt.style.use('fivethirtyeight')
        #plt.style.use('seaborn-muted')

        plt.subplots_adjust(left=0.05, bottom=0.15, right=0.92, top=0.92, wspace=0.2, hspace=0.2)

        ax.plot(year, baseline_profile, color='blue', label='baseline', linewidth=4.0, marker="o")
        ax.plot(year, last_profile, color='yellow', label='last quarter', linewidth=4.0, marker="o")
        ax.plot(year, latest_profile, color='green', label='latest', linewidth=4.0, marker="o")

        ax.set_xlabel('Financial Years')
        ax.set_ylabel('Cost (£m)')
        xlab = ax.xaxis.get_label()
        ylab = ax.yaxis.get_label()
        xlab.set_style('italic')
        xlab.set_size(20)
        ylab.set_style('italic')
        ylab.set_size(20)

        ax.set_title(str(project_name) + ' Cost Profile Changes')  # is this needed?

        ax.grid(color='grey', linestyle='-', linewidth=0.2)

        ax.legend(borderpad=2)

        ax.figure.savefig('cost_profile.png')

        doc.add_picture('cost_profile.png', width=Inches(5.8)) #to place nicely in doc

        doc.save(root_path/'output/graph.docx')

def two_graph_chart(project_list):
    for project_name in project_list:
        profile_data_total = get_financial_data(project_name, ' total')
        profile_data_rdel = get_financial_data(project_name, ' RDEL Forecast Total')
        profile_data_cdel = get_financial_data(project_name, ' CDEL Forecast Total')

        year = ['19/20', '20/21', '21/22', '22/23', '23/24', '24/25', '25/26', '26/27', '27/28', '28/29']
        baseline_profile_total = profile_data_total[2]
        last_profile_total = profile_data_total[1]
        latest_profile_total = profile_data_total[0]

        latest_profile_rdel = profile_data_rdel[0]
        latest_profile_cdel = profile_data_cdel[0]

        fig, (ax1, ax2) = plt.subplots(2, figsize=(20,10))
        fig.suptitle(str(project_name) + ' Cost Analysis')  # title

        ax1.plot(year, baseline_profile_total, color='blue', label='baseline', linewidth=4.0, marker="o")
        ax1.plot(year, last_profile_total, color='yellow', label='last quarter', linewidth=4.0, marker="o")
        ax1.plot(year, latest_profile_total, color='green', label='latest', linewidth=4.0, marker="o")

        ax2.plot(year, latest_profile_cdel, color='red', label='CDEL', linewidth=4.0, marker="o")
        ax2.plot(year, latest_profile_rdel, color='blue', label='RDEL', linewidth=4.0, marker="o")

        ax1.set_xlabel('Financial Years')
        ax1.set_ylabel('Cost (£m)')
        xlab1 = ax1.xaxis.get_label()
        ylab1 = ax1.yaxis.get_label()
        xlab1.set_style('italic')
        xlab1.set_size(10)
        ylab1.set_style('italic')
        ylab1.set_size(10)
        ax1.grid(color='grey', linestyle='-', linewidth=0.2)
        ax1.legend(borderpad=2)
        #TODO give legends a header

        ax2.set_xlabel('Financial Years')
        ax2.set_ylabel('Cost (£m)')
        xlab2 = ax2.xaxis.get_label()
        ylab2 = ax2.yaxis.get_label()
        xlab2.set_style('italic')
        xlab2.set_size(10)
        ylab2.set_style('italic')
        ylab2.set_size(10)
        ax2.grid(color='grey', linestyle='-', linewidth=0.2)
        ax2.legend(borderpad=2)

        fig.savefig('cost_profile.png')

        doc.add_picture('cost_profile.png', width=Inches(12))  # to place nicely in doc



        doc.save(root_path / 'output/graph_2.docx')

def three_graph_chart(project_list):
    for project_name in project_list:
        profile_data_total = get_financial_data(project_name, ' total')
        profile_data_rdel = get_financial_data(project_name, ' RDEL Forecast Total')
        profile_data_cdel = get_financial_data(project_name, ' CDEL Forecast Total')

        year = ['19/20', '20/21', '21/22', '22/23', '23/24', '24/25', '25/26', '26/27', '27/28', '28/29']
        baseline_profile_total = profile_data_total[2]
        last_profile_total = profile_data_total[1]
        latest_profile_total = profile_data_total[0]

        latest_profile_rdel = profile_data_rdel[0]
        latest_profile_cdel = profile_data_cdel[0]

        plt.subplot(2, 1, 1)
        plt.plot(year, baseline_profile_total, color='blue', label='baseline', linewidth=4.0, marker="o")
        plt.plot(year, last_profile_total, color='yellow', label='last quarter', linewidth=4.0, marker="o")
        plt.plot(year, latest_profile_total, color='green', label='latest', linewidth=4.0, marker="o")

        plt.subplot(2, 2, 3)
        plt.plot(year, latest_profile_cdel, color='red', label='CDEL', linewidth=4.0, marker="o")
        plt.plot(year, latest_profile_rdel, color='blue', label='RDEL', linewidth=4.0, marker="o")

        plt.subplot(2, 2, 4)

        # fig, (ax1, ax2, ax3) = plt.subplots(3, figsize=(20,10))
        # fig.suptitle(str(project_name) + ' Cost Analysis')  # title

        # ax1.plot(year, baseline_profile_total, color='blue', label='baseline', linewidth=4.0, marker="o")
        # ax1.plot(year, last_profile_total, color='yellow', label='last quarter', linewidth=4.0, marker="o")
        # ax1.plot(year, latest_profile_total, color='green', label='latest', linewidth=4.0, marker="o")
        #
        # ax2.plot(year, latest_profile_cdel, color='red', label='CDEL', linewidth=4.0, marker="o")
        # ax2.plot(year, latest_profile_rdel, color='blue', label='RDEL', linewidth=4.0, marker="o")
        #
        # ax1.set_xlabel('Financial Years')
        # ax1.set_ylabel('Cost (£m)')
        # xlab1 = ax1.xaxis.get_label()
        # ylab1 = ax1.yaxis.get_label()
        # xlab1.set_style('italic')
        # xlab1.set_size(10)
        # ylab1.set_style('italic')
        # ylab1.set_size(10)
        # ax1.grid(color='grey', linestyle='-', linewidth=0.2)
        # ax1.legend(borderpad=2)
        # #TODO give legends a header
        #
        # ax2.set_xlabel('Financial Years')
        # ax2.set_ylabel('Cost (£m)')
        # xlab2 = ax2.xaxis.get_label()
        # ylab2 = ax2.yaxis.get_label()
        # xlab2.set_style('italic')
        # xlab2.set_size(10)
        # ylab2.set_style('italic')
        # ylab2.set_size(10)
        # ax2.grid(color='grey', linestyle='-', linewidth=0.2)
        # ax2.legend(borderpad=2)
        #
        # fig.savefig('cost_profile.png')
        #
        # #put into word doc
        # y = doc.add_paragraph()
        # heading = 'Annex - Financial Analysis'
        # y.add_run(str(heading)).bold = True
        #
        # sections = doc.sections
        # section_2 = sections[0]
        # new_width, new_height = section_2.page_height, section_2.page_width
        # section_2.orientation = WD_ORIENT.LANDSCAPE
        # section_2.page_width = new_width
        # section_2.page_height = new_height
        #
        # doc.add_picture('cost_profile.png', width=Inches(5.8))  # to place nicely in doc
        #
        # doc.save(root_path / 'output/graph_2.docx')



project_name_list = [a66, crossrail, a303, thameslink]

#single_chart(project_list)

two_graph_chart(project_name_list)

#three_graph_chart([a66])