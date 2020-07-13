'''
Programme that compiles project dashboards/summary sheets.
Input:
1) Nothing for the user to input manually. when the programme run its default is to take the last four quarters
worth of data.

Output:
1) Word document with all information e.g. narratives, tables and charts, built into it.
'''

from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, RGBColor, Inches, Pt
from docx.enum.section import WD_SECTION_START, WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import difflib
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import numpy as np
import datetime
from datetime import timedelta
from textwrap import wrap


from analysis.engine_functions import convert_rag_text, project_time_difference, convert_bc_stage_text
from analysis.data import list_of_masters_all, root_path, latest_cost_profiles, last_cost_profiles, \
    baseline_1_cost_profiles, year_list, SRO_conf_key_list, p_current_milestones, \
    p_last_milestones, p_baseline_milestones, first_diff_data, ipdc_date, abbreviations, \
    benefits_bl_index, costs_bl_index


import os

milestone_filter_start_date = ipdc_date - timedelta(days=30*6)
milestone_filter_end_date = ipdc_date + timedelta(days=545)

def cell_colouring(cell, colour):
    '''
    Function that handles cell colouring
    cell: cell reference
    color: colour reference
    '''

    try:
        if colour == 'R':
            colour = parse_xml(r'<w:shd {} w:fill="cb1f00"/>'.format(nsdecls('w')))
        elif colour == 'A/R':
            colour = parse_xml(r'<w:shd {} w:fill="f97b31"/>'.format(nsdecls('w')))
        elif colour == 'A':
            colour = parse_xml(r'<w:shd {} w:fill="fce553"/>'.format(nsdecls('w')))
        elif colour == 'A/G':
            colour = parse_xml(r'<w:shd {} w:fill="a5b700"/>'.format(nsdecls('w')))
        elif colour == 'G':
            colour = parse_xml(r'<w:shd {} w:fill="17960c"/>'.format(nsdecls('w')))

        cell._tc.get_or_add_tcPr().append(colour)

    except TypeError:
        pass

def compare_text_showall(text_1, text_2, doc):
    '''
    Function places text into doc highlighting all changes.
    text_1: latest text. string.
    text_2: last text. string
    doc: word doc
    '''

    comp = difflib.Differ()
    diff = list(comp.compare(text_2.split(), text_1.split()))
    y = doc.add_paragraph()

    for i in range(0, len(diff)):
        f = len(diff) - 1
        if i < f:
            a = i - 1
        else:
            a = i

        if diff[i][0:3] == '  |':
            j = i + 1
            if diff[i][0:3] and diff[a][0:3] == '  |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '+ |':
            if diff[i][0:3] and diff[a][0:3] == '+ |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '- |':
            pass
        elif diff[i][0:3] == '  -':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0:3] == '  •':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0] == '+':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.color.rgb = RGBColor(255, 0, 0)
        elif diff[i][0] == '-':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.strike = True
        elif diff[i][0] == '?':
            pass
        else:
            if diff[i] != '+ |':
                y.add_run(diff[i])

    return doc

def compare_text_newandold(text_1, text_2, doc):
    '''
    function that places text into doc highlighting new and old text
    text_1: latest text. string.
    text_2: last text. string
    doc: word doc
    '''

    comp = difflib.Differ()
    diff = list(comp.compare(text_2.split(), text_1.split()))
    new_text = diff
    y = doc.add_paragraph()
    y.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for i in range(0, len(diff)):
        f = len(diff) - 1
        if i < f:
            a = i - 1
        else:
            a = i

        if diff[i][0:3] == '  |':
            j = i + 1
            if diff[i][0:3] and diff[a][0:3] == '  |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '+ |':
            if diff[i][0:3] and diff[a][0:3] == '+ |':
                y = doc.add_paragraph()
            else:
                pass
        elif diff[i][0:3] == '- |':
            pass
        elif diff[i][0:3] == '  -':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0:3] == '  •':
            y = doc.add_paragraph()
            g = diff[i][2]
            y.add_run(g)
        elif diff[i][0] == '+':
            w = len(diff[i])
            g = diff[i][1:w]
            y.add_run(g).font.color.rgb = RGBColor(255, 0, 0)
        elif diff[i][0] == '-':
            pass
        elif diff[i][0] == '?':
            pass
        else:
            if diff[i] != '+ |':
                y.add_run(diff[i][1:])

    return doc

def set_col_widths(t, widths):
    for row in t.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def produce_word_doc(projects):
    '''Function that compiles each summary sheet'''

    masters = list_of_masters_all[0:4]

    for project_name in projects:
        doc = Document(root_path/'input/summary_temp.docx')
        print(project_name)

        font = doc.styles['Normal'].font
        font.name = 'Arial'
        font.size = Pt(12)

        heading = str(abbreviations[project_name])
        intro = doc.add_heading(str(heading), 0)
        intro.alignment = 1
        intro.bold = True

        para_1 = doc.add_paragraph()
        #para_1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sro_name = list_of_masters_all[0].data[project_name]['Senior Responsible Owner (SRO)']
        if sro_name is None:
            sro_name = 'tbc'

        sro_email = list_of_masters_all[0].data[project_name]['Senior Responsible Owner (SRO) - Email']
        if sro_email is None:
            sro_email = 'email: tbc'

        sro_phone = list_of_masters_all[0].data[project_name]['SRO Phone No.']
        if sro_phone == None:
            sro_phone = 'phone number: tbc'

        para_1.add_run('SRO: ' + str(sro_name) + ', ' + str(sro_email) + ', ' + str(sro_phone))

        para_2 = doc.add_paragraph()
        #para_2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pd_name = list_of_masters_all[0].data[project_name]['Project Director (PD)']
        if pd_name is None:
            pd_name = 'TBC'

        pd_email = list_of_masters_all[0].data[project_name]['Project Director (PD) - Email']
        if pd_email is None:
            pd_email = 'email: tbc'

        pd_phone = list_of_masters_all[0].data[project_name]['PD Phone No.']
        if pd_phone is None:
            pd_phone = 'TBC'

        para_2.add_run('PD: ' + str(pd_name) + ', ' + str(pd_email) + ', ' + str(pd_phone))

        para_3 = doc.add_paragraph()
        #para_3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        contact_name = list_of_masters_all[0].data[project_name]['Working Contact Name']
        if contact_name is None:
            contact_name = 'TBC'

        contact_email = list_of_masters_all[0].data[project_name]['Working Contact Email']
        if contact_email is None:
            contact_email = 'email: tbc'

        contact_phone = list_of_masters_all[0].data[project_name]['Working Contact Telephone']
        if contact_phone is None:
            contact_phone = 'TBC'

        para_3.add_run('PfM reporting lead: ' + str(contact_name) + ', ' + str(contact_email)
                       + ', ' + str(contact_phone))

        '''Start of table with DCA confidence ratings'''
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Delivery confidence'
        hdr_cells[1].text = 'This quarter'
        hdr_cells[2].text = str(list_of_masters_all[1].quarter)
        hdr_cells[3].text = str(list_of_masters_all[2].quarter)
        hdr_cells[4].text = str(list_of_masters_all[3].quarter)

        #RAG ratings
        for x, dca_key in enumerate(SRO_conf_key_list):
            row_cells = table.add_row().cells
            row_cells[0].text = dca_key
            for i, master in enumerate(masters):
                try:
                    rating = convert_rag_text(master.data[project_name][dca_key])
                    row_cells[i+1].text = rating
                    cell_colouring(row_cells[i+1], rating)
                except (KeyError, TypeError):
                    row_cells[i+1].text = 'N/A'

        table.style = 'Table Grid'
        make_rows_bold([table.rows[0]])  # makes top of table bold.
        #make_columns_bold([table.columns[0]]) #right cells in table bold
        column_widths = (Cm(3.9), Cm(2.9), Cm(2.9), Cm(2.9), Cm(2.9))
        set_col_widths(table, column_widths)

        doc.add_paragraph()
        p = doc.add_paragraph()
        text = '*Red text highlights changes in narratives from last quarter'
        p.add_run(text).font.color.rgb = RGBColor(255, 0, 0)

        '''DCA narrative'''
        #doc.add_paragraph()

        headings_list = ['SRO delivery confidence narrative',
                         'Financial cost narrative',
                         'Financial comparison with last quarter',
                         'Financial comparison with baseline',
                         'Benefits Narrative',
                         'Benefits comparison with last quarter',
                         'Benefits comparison with baseline',
                         'Milestone narrative']

        narrative_keys_list = ['Departmental DCA Narrative',
                               'Project Costs Narrative',
                               'Cost comparison with last quarters cost narrative',
                               'Cost comparison within this quarters cost narrative',
                               'Benefits Narrative',
                               'Ben comparison with last quarters cost - narrative',
                               'Ben comparison within this quarters cost - narrative',
                               'Milestone Commentary']

        for x in range(len(headings_list)):
            doc.add_paragraph().add_run(str(headings_list[x])).bold = True
            text_one = str(list_of_masters_all[0].data[project_name][narrative_keys_list[x]])
            try:
                text_two = str(list_of_masters_all[1].data[project_name][narrative_keys_list[x]])
            except KeyError:
                text_two = text_one

            #different options for comparing costs
            # compare_text_showall(dca_a, dca_b, doc)
            compare_text_newandold(text_one, text_two, doc)

        '''Financial data'''
        project_profile_data_total = get_financial_profile(project_name, ' total')
        project_profile_data_rdel = get_financial_profile(project_name, ' RDEL Forecast Total')
        project_profile_data_cdel = get_financial_profile(project_name, ' CDEL Forecast Total')

        fin_all = get_financial_totals(project_name) # all totals
        #print(fin_all)

        #totals by spent, profiled and unprofiled
        total_fin = fin_all[0]
        total_profiled_bl = total_fin[1] - (total_fin[0] + total_fin[2])
        total_profiled_lst = total_fin[4] - (total_fin[3] + total_fin[5])
        total_profiled_now = total_fin[7] - (total_fin[6] + total_fin[8])
        t_spent = np.array([total_fin[0], total_fin[3], total_fin[6]])
        t_profiled = np.array([total_profiled_bl, total_profiled_lst, total_profiled_now])
        t_unprofiled = np.array([total_fin[2], total_fin[5], total_fin[8]])

        #totals for rdel and cdel
        rdel_fin = fin_all[1]
        cdel_fin = fin_all[2]
        rdel_profiled = rdel_fin[7] - (rdel_fin[6] + rdel_fin[8])
        cdel_profiled = cdel_fin[7] - (cdel_fin[6] + cdel_fin[8])
        rc_spent = np.array([rdel_fin[6], cdel_fin[6]])
        rc_profiled = np.array([rdel_profiled, cdel_profiled])
        rc_unprofiled = np.array([rdel_fin[8], cdel_fin[8]])

        source_of_finance = list_of_masters_all[0].data[project_name]['Source of Finance']
        contingency = list_of_masters_all[0].data[project_name]['Overall contingency (£m)']
        con_included_wlc = list_of_masters_all[0].data[project_name]\
                           ['Is this Continency amount included within the WLC?']

        ob = list_of_masters_all[0].data[project_name]['Overall figure for Optimism Bias (£m)']
        #print(ob)

        ob_included_wlc = list_of_masters_all[0].data[project_name]['Is this Optimism Bias included within the WLC?']
        '''vfm category now'''
        if list_of_masters_all[0].data[project_name]['VfM Category single entry'] is None:
            vfm_cat = str(list_of_masters_all[0].data[project_name]['VfM Category lower range']) + ' - ' + \
                      str(list_of_masters_all[0].data[project_name]['VfM Category upper range'])
        else:
            vfm_cat = list_of_masters_all[0].data[project_name]['VfM Category single entry']
        bcr = list_of_masters_all[0].data[project_name]['Adjusted Benefits Cost Ratio (BCR)']

        '''milestone data'''
        ipdc_business_case_stage = convert_bc_stage_text(
            list_of_masters_all[0].data[project_name]['IPDC approval point'])
        delivery_stage = list_of_masters_all[0].data[project_name]['Project stage']
        if delivery_stage is None:
            delivery_stage = 'Not reported'
        try:
            start_project = tuple(p_current_milestones[project_name]['Start of Project'])[0]
            start_project_text = start_project.strftime("%d/%m/%Y")
        except (KeyError, AttributeError):
            start_project_text = 'Not reported'

        try:
            start_con_build = tuple(p_current_milestones[project_name]['Start of Construction/build'])[0]
            start_con_build_text = start_con_build.strftime("%d/%m/%Y")
        except (KeyError, AttributeError):
            start_con_build_text = 'Not reported'
        try:
            start_ops = tuple(p_current_milestones[project_name]['Start of Operation'])[0]
            start_ops_text = start_ops.strftime("%d/%m/%Y")
        except (KeyError, AttributeError):
            start_ops_text = 'Not reported'
        try:
            full_ops = tuple(p_current_milestones[project_name]['Full Operations'])[0]
            full_ops_text = full_ops.strftime("%d/%m/%Y")
        except (KeyError, AttributeError):
            full_ops_text = 'Not reported'

        '''ben data'''
        all_ben = get_ben_totals(project_name)  # all totals
        total_ben = all_ben[0]
        #print(total_ben)

        # totals by spent, profiled and unprofiled
        total_ben_profiled_bl = total_ben[1] - (total_ben[0] + total_ben[2])
        total_ben_profiled_lst = total_ben[4] - (total_ben[3] + total_ben[5])
        total_ben_profiled_now = total_ben[7] - (total_ben[6] + total_ben[8])
        b_achieved = np.array([total_ben[0], total_ben[3], total_ben[6]])
        b_profiled = np.array([total_ben_profiled_bl, total_ben_profiled_lst, total_ben_profiled_now])
        b_unprofiled = np.array([total_ben[2], total_ben[5], total_ben[8]])

        ben_type_all = all_ben[1]
        #print(ben_type_all)
        b_cashable_profiled = ben_type_all[8] - (ben_type_all[0] + ben_type_all[4])
        b_non_cashable_profiled = ben_type_all[9] - (ben_type_all[1] + ben_type_all[5])
        b_economic_profiled = ben_type_all[10] - (ben_type_all[2] + ben_type_all[6])
        b_disbenefit_profiled = ben_type_all[11] - (ben_type_all[3] + ben_type_all[7])
        b_type_achieved = np.array([ben_type_all[0], ben_type_all[1], ben_type_all[2], ben_type_all[3]])
        b_type_profiled = np.array([b_cashable_profiled,
                                   b_non_cashable_profiled,
                                   b_economic_profiled,
                                   b_disbenefit_profiled])
        b_type_unprofiled = np.array([ben_type_all[4], ben_type_all[5], ben_type_all[6], ben_type_all[7]])
        b_type_disbenefit = [ben_type_all[3], ben_type_all[7], ben_type_all[1]]

        '''Meta data table'''
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        '''Costs meta data'''
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.add_run('Annex A. High level MI data and analysis').bold = True

        run = doc.add_paragraph().add_run('Costs')
        font = run.font
        font.bold = True
        font.underline = True
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'WLC:'
        hdr_cells[1].text = '£' + str(round(total_fin[7])) + 'm'
        hdr_cells[2].text = 'Spent:'
        hdr_cells[3].text = '£' + str(round(total_fin[6])) + 'm'
        row_cells = table.add_row().cells
        row_cells[0].text = 'RDEL Total:'
        row_cells[1].text = '£' + str(round(rdel_fin[7])) + 'm'
        row_cells[2].text = 'Profiled:'
        row_cells[3].text = '£' + str(round(total_profiled_now)) + 'm'
        row_cells = table.add_row().cells
        row_cells[0].text = 'CDEL Total:'
        row_cells[1].text = '£' + str(round(cdel_fin[7])) + 'm'
        row_cells[2].text = 'Unprofiled:'
        row_cells[3].text = '£' + str(round(total_fin[8])) + 'm'

        #set column width
        column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
        set_col_widths(table, column_widths)
        #make column keys bold
        make_columns_bold([table.columns[0], table.columns[2]])
        change_text_size([table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10)

        '''Financial data'''
        doc.add_paragraph()
        run = doc.add_paragraph().add_run('Financial')
        font = run.font
        font.bold = True
        font.underline = True
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Type of funding:'
        hdr_cells[1].text = str(source_of_finance)
        hdr_cells[2].text = 'Contingency:'
        if contingency is None:
            hdr_cells[3].text = str(contingency)
        else:
            hdr_cells[3].text = '£' + str(round(contingency)) + 'm'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Optimism Bias (OB):'
        if ob is None:
            row_cells[1].text = str(ob)
        else:
            try:
                row_cells[1].text = '£' + str(round(ob)) + 'm'
            except TypeError:
                row_cells[1].text = ob
        row_cells[2].text = 'Contingency in costs:'
        if con_included_wlc is None:
            row_cells[3].text = 'Not reported'
        else:
            row_cells[3].text = con_included_wlc
        row_cells = table.add_row().cells
        row_cells[0].text = 'OB in costs:'
        if ob_included_wlc is None:
            row_cells[1].text = 'Not reported'
        else:
            row_cells[1].text = ob_included_wlc
        row_cells[2].text = ''
        row_cells[3].text = ''

        # set column width
        column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
        set_col_widths(table, column_widths)
        # make column keys bold
        make_columns_bold([table.columns[0], table.columns[2]])
        change_text_size([table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10)

        '''Project Stage data'''
        doc.add_paragraph()
        run = doc.add_paragraph().add_run('Stage')
        font = run.font
        font.bold = True
        font.underline = True
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Business case stage (IPDC approved):'
        hdr_cells[1].text = ipdc_business_case_stage
        hdr_cells[2].text = 'Delivery stage:'
        hdr_cells[3].text = delivery_stage

        # set column width
        column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
        set_col_widths(table, column_widths)
        # make column keys bold
        make_columns_bold([table.columns[0], table.columns[2]])
        change_text_size([table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10)
        make_text_red([table.columns[1], table.columns[3]])  # make 'not reported red'

        '''Milestone/Stage meta data'''
        doc.add_paragraph()
        run = doc.add_paragraph().add_run('Schedule/Milestones')
        font = run.font
        font.bold = True
        font.underline = True
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Start date:'
        hdr_cells[1].text = start_project_text
        hdr_cells[2].text = 'Start of operations:'
        hdr_cells[3].text = start_ops_text
        row_cells = table.add_row().cells
        row_cells[0].text = 'Start of construction:'
        row_cells[1].text = start_con_build_text
        row_cells[2].text = 'Full Operations:'
        row_cells[3].text = full_ops_text

        # set column width
        column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
        set_col_widths(table, column_widths)
        # make column keys bold
        make_columns_bold([table.columns[0], table.columns[2]])
        change_text_size([table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10)
        make_text_red([table.columns[1], table.columns[3]]) #make 'not reported red'

        '''vfm meta data'''
        doc.add_paragraph()
        run = doc.add_paragraph().add_run('VfM')
        font = run.font
        font.bold = True
        font.underline = True
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'VfM category:'
        hdr_cells[1].text = vfm_cat
        hdr_cells[2].text = 'BCR:'
        hdr_cells[3].text = str(bcr)

        # set column width
        column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
        set_col_widths(table, column_widths)
        # make column keys bold
        make_columns_bold([table.columns[0], table.columns[2]])
        change_text_size([table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10)
        make_text_red([table.columns[1], table.columns[3]])  # make 'not reported red'

        '''benefits meta data'''
        doc.add_paragraph()
        run = doc.add_paragraph().add_run('Benefits')
        font = run.font
        font.bold = True
        font.underline = True
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Total Benefits:'
        hdr_cells[1].text = '£' + str(round(total_ben[7])) + 'm'
        hdr_cells[2].text = 'Benefits delivered:'
        hdr_cells[3].text = '£' + str(round(total_ben[6])) + 'm'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Benefits profiled:'
        row_cells[1].text = '£' + str(round(total_ben_profiled_now)) + 'm'
        row_cells[2].text = 'Benefits unprofiled:'
        row_cells[3].text = '£' + str(round(total_ben[8])) + 'm'

        # set column width
        column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
        set_col_widths(table, column_widths)
        # make column keys bold
        make_columns_bold([table.columns[0], table.columns[2]])
        change_text_size([table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10)

        '''start of analysis'''
        new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        new_width, new_height = new_section.page_height, new_section.page_width
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = new_width
        new_section.page_height = new_height

        '''Financial charts'''
        fin_profile_graph(doc,
                          project_name,
                          project_profile_data_total,
                          project_profile_data_rdel,
                          project_profile_data_cdel)
        fin_ben_total_charts(doc,
                             project_name,
                             t_spent,
                             t_profiled,
                             t_unprofiled,
                             total_fin,
                             rc_spent,
                             rc_profiled,
                             rc_unprofiled,
                             b_achieved,
                             b_profiled,
                             b_unprofiled,
                             b_type_achieved,
                             b_type_profiled,
                             b_type_unprofiled,
                             b_type_disbenefit,
                             total_ben)

        '''milestone swimlane charts'''
        # #chart of with milestone over the next two years
        # m_data = milestone_schedule_data(p_current_milestones,
        #                                  p_last_milestones,
        #                                  p_baseline_milestones,
        #                                  project_name,
        #                                  milestone_filter_start_date,
        #                                  milestone_filter_end_date)
        #
        # # add \n to y axis labels and cut down if two long
        # labels = ['\n'.join(wrap(l, 40)) for l in m_data[0]]
        # final_labels = []
        # for l in labels:
        #     if len(l) > 70:
        #         final_labels.append(l[:70])
        #     else:
        #         final_labels.append(l)
        #
        # no_milestones = len(m_data[0])
        #
        # title = 'schedule two year window'
        # if no_milestones <= 15:
        #     milestone_swimlane_charts(doc, project_name, np.array(final_labels), np.array(m_data[1]), np.array(m_data[2]), \
        #                           np.array(m_data[3]), title)
        #
        # if 16 <= no_milestones <= 35:
        #     half = int(no_milestones/2)
        #     milestone_swimlane_charts(doc, project_name, np.array(final_labels[:half]), np.array(m_data[1][:half]),
        #                               np.array(m_data[2][:half]), np.array(m_data[3][:half]), title)
        #     title = title + ' cont.'
        #     milestone_swimlane_charts(doc, project_name, np.array(final_labels[half:no_milestones]),
        #                               np.array(m_data[1][half:no_milestones]),
        #                               np.array(m_data[2][half:no_milestones]),
        #                               np.array(m_data[3][half:no_milestones]), title)

        #chart with all milestones
        m_data = milestone_schedule_data(p_current_milestones,
                                         p_last_milestones,
                                         p_baseline_milestones,
                                         project_name)
        #print(m_data)
        # add \n to y axis labels and cut down if two long
        labels = ['\n'.join(wrap(l, 40)) for l in m_data[0]]
        final_labels = []
        for l in labels:
            if len(l) > 70:
                final_labels.append(l[:70])
            else:
                final_labels.append(l)

        # Chart
        no_milestones = len(m_data[0])

        title = 'total schedule'
        if no_milestones <= 20:
            milestone_swimlane_charts(doc, project_name, np.array(final_labels), np.array(m_data[1]),
                                      np.array(m_data[2]), \
                                      np.array(m_data[3]), title)

        if 21 <= no_milestones <= 40:
            half = int(no_milestones/2)
            milestone_swimlane_charts(doc, project_name, np.array(final_labels[:half]),
                                      np.array(m_data[1][:half]),
                                      np.array(m_data[2][:half]), np.array(m_data[3][:half]), title)
            title = title + ' cont.'
            milestone_swimlane_charts(doc, project_name, np.array(final_labels[half:no_milestones]),
                                      np.array(m_data[1][half:no_milestones]),
                                      np.array(m_data[2][half:no_milestones]),
                                      np.array(m_data[3][half:no_milestones]), title)

        if 41 <= no_milestones <= 70:
            third = int(no_milestones/3)
            milestone_swimlane_charts(doc, project_name, np.array(final_labels[:third]),
                                      np.array(m_data[1][:third]),
                                      np.array(m_data[2][:third]), np.array(m_data[3][:third]), title)
            title = title + ' cont.'
            milestone_swimlane_charts(doc, project_name, np.array(final_labels[third:third*2]),
                                      np.array(m_data[1][third:third*2]),
                                      np.array(m_data[2][third:third*2]),
                                      np.array(m_data[3][third:third*2]), title)
            #title = title + ' schedule all cont.'
            milestone_swimlane_charts(doc, project_name, np.array(final_labels[third*2:no_milestones]),
                                      np.array(m_data[1][third*2:no_milestones]),
                                      np.array(m_data[2][third*2:no_milestones]),
                                      np.array(m_data[3][third*2:no_milestones]), title)


        #milestone table
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        #table heading
        doc.add_paragraph().add_run(str('Project high-level milestones')).bold = True
        some_text = 'The below table presents all project reported remaining high-level milestones, with six months grace ' \
                    'from close of the current quarter. Milestones are sorted in chronological order. Changes in milestones' \
                    ' dates in comparison to last quarter and baseline have been calculated and are provided.'
        doc.add_paragraph().add_run(str(some_text)).italic = True

        milestone_table(doc, p_baseline_milestones, project_name)

        '''Back to portrait'''
        new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        new_width, new_height = new_section.page_height, new_section.page_width
        new_section.orientation = WD_ORIENT.PORTRAIT
        new_section.page_width = new_width
        new_section.page_height = new_height

        '''Project Scope'''
        doc.add_paragraph().add_run('Project Scope').bold = True
        text_one = str(list_of_masters_all[0].data[project_name]['Project Scope'])
        try:
            text_two = str(list_of_masters_all[1].data[project_name]['Project Scope'])
        except KeyError:
            text_two = text_one

        # different options for comparing costs
        # compare_text_showall(dca_a, dca_b, doc)
        compare_text_newandold(text_one, text_two, doc)

        #TODO add quarter info in title
        doc.save(root_path/'output/{}_summary.docx'.format(project_name))

def combine_narrtives(project_name, master, key_list):
    '''function that combines text across different keys'''
    output = ''
    for key in key_list:
        output = output + str(master.data[project_name][key])

    return output

def get_financial_profile(project_name, cost_type):
    '''gets project financial data'''
    latest = []
    last = []
    baseline = []
    amended_year_list = year_list[:-1] #to not use 'unprofiled'
    for year in amended_year_list:
        baseline.append(baseline_1_cost_profiles[project_name][year + cost_type])
        last.append(last_cost_profiles[project_name][year + cost_type])
        latest.append(latest_cost_profiles[project_name][year + cost_type])

    return latest, last, baseline

def fin_ben_total_charts(doc,
                         project_name,
                         total_spent,
                         total_profiled,
                         total_unprofiled,
                         total_fin,
                         rc_spent,
                         rc_profiled,
                         rc_unprofiled,
                         delivered_ben,
                         profiled_ben,
                         unprofiled_ben,
                         type_delivered_ben,
                         type_profiled_ben,
                         type_unprofiled_ben,
                         type_disbenefit_ben,
                         total_ben):

    fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2) #four sub plotsprint

    fig.suptitle(abbreviations[project_name] + ' costs and benefits analysis', fontweight='bold') # title

    #Spent, Profiled and Unprofile chart
    labels = ['Baseline', 'Last Quarter', 'Latest']
    width = 0.5
    ax1.bar(labels, total_spent, width, label='Spent')
    ax1.bar(labels, total_profiled, width, bottom=total_spent, label='Profiled')
    ax1.bar(labels, total_unprofiled, width, bottom=total_spent + total_profiled, label='Unprofiled')
    ax1.legend(prop={'size': 6})
    ax1.set_ylabel('Cost (£m)')
    ylab1 = ax1.yaxis.get_label()
    ylab1.set_style('italic')
    ylab1.set_size(8)
    ax1.tick_params(axis='x', which='major', labelsize=6)
    ax1.tick_params(axis='y', which='major', labelsize=6)
    ax1.set_title('Fig 1 - cost total change over time', loc='left', fontsize=8, fontweight='bold')

    #scaling y axis
    #y axis value setting so it takes either highest ben or cost figure
    cost_max = max(total_fin) + max(total_fin)/5
    ben_max = max(total_ben) + max(total_ben)/5
    y_max = max([cost_max, ben_max])
    ax1.set_ylim(0, y_max)

    #rdel/cdel bar chart

    labels = ['RDEL', 'CDEL']
    width = 0.5
    ax3.bar(labels, rc_spent, width, label='Spent')
    ax3.bar(labels, rc_profiled, width, bottom=rc_spent, label='Profiled')
    ax3.bar(labels, rc_unprofiled, width, bottom=rc_spent + rc_profiled, label='Unprofiled')
    ax3.legend(prop={'size': 6})
    ax3.set_ylabel('Costs (£m)')
    ylab3 = ax3.yaxis.get_label()
    ylab3.set_style('italic')
    ylab3.set_size(8)
    ax3.tick_params(axis='x', which='major', labelsize=6)
    ax3.tick_params(axis='y', which='major', labelsize=6)
    ax3.set_title('Fig 2 - wlc cost type break down', loc='left', fontsize=8, fontweight='bold')

    #y_max = max(total_fin) + max(total_fin) * 1 / 5
    ax3.set_ylim(0, y_max) #scale y axis max

    # benefits change
    labels = ['Baseline', 'Last Quarter', 'Latest']
    width = 0.5
    ax2.bar(labels, delivered_ben, width, label='Delivered')
    ax2.bar(labels, profiled_ben, width, bottom=delivered_ben, label='Profiled')
    ax2.bar(labels, unprofiled_ben, width, bottom=delivered_ben + profiled_ben, label='Unprofiled')
    ax2.legend(prop={'size': 6})
    ax2.set_ylabel('Benefits (£m)')
    ylab2 = ax2.yaxis.get_label()
    ylab2.set_style('italic')
    ylab2.set_size(8)
    ax2.tick_params(axis='x', which='major', labelsize=6)
    ax2.tick_params(axis='y', which='major', labelsize=6)
    ax2.set_title('Fig 3 - ben total change over time', loc='left', fontsize=8, fontweight='bold')

    ax2.set_ylim(0, y_max)

    # benefits break down
    labels = ['Cashable', 'Non-Cashable', 'Economic', 'Disbenefit']
    width = 0.5
    ax4.bar(labels, type_delivered_ben, width, label='Delivered')
    ax4.bar(labels, type_profiled_ben, width, bottom=type_delivered_ben, label='Profiled')
    ax4.bar(labels, type_unprofiled_ben, width, bottom=type_delivered_ben + type_profiled_ben, label='Unprofiled')
    ax4.legend(prop={'size': 6})
    ax4.set_ylabel('Benefits (£m)')
    ylab4 = ax4.yaxis.get_label()
    ylab4.set_style('italic')
    ylab4.set_size(8)
    ax4.tick_params(axis='x', which='major', labelsize=6)
    ax4.tick_params(axis='y', which='major', labelsize=6)
    ax4.set_title('Fig 4 - benefits profile type', loc='left', fontsize=8, fontweight='bold')

    y_min = min(type_disbenefit_ben)
    ax4.set_ylim(y_min, y_max)

    # size of chart and fit
    fig.canvas.draw()
    fig.tight_layout(rect=[0, 0.03, 1, 0.95])  # for title

    fig.savefig('cost_bens_overview.png')
    plt.close()  # automatically closes figure so don't need to do manually.

    doc.add_picture('cost_bens_overview.png', width=Inches(8))  # to place nicely in doc
    os.remove('cost_bens_overview.png')

    return doc

def fin_profile_graph(doc,
                      project_name,
                      profile_data_total,
                      profile_data_rdel,
                      profile_data_cdel):

    fig, (ax1, ax2) = plt.subplots(2) #four sub plotsprint

    '''cost profile charts'''
    year = year_list[:-1]
    baseline_profile_total = profile_data_total[2]
    last_profile_total = profile_data_total[1]
    latest_profile_total = profile_data_total[0]

    latest_profile_rdel = profile_data_rdel[0]
    latest_profile_cdel = profile_data_cdel[0]

    fig.suptitle(abbreviations[project_name] + ' financial analysis', fontweight='bold') # title

    #plot cost change profile chart
    ax1.plot(year, baseline_profile_total, label='Baseline', linewidth=3.0, marker="o")
    ax1.plot(year, last_profile_total, label='Last quarter', linewidth=3.0, marker="o")
    ax1.plot(year, latest_profile_total, label='Latest', linewidth=3.0, marker="o")

    #cost profile change chart styling
    ax1.tick_params(axis='x', which='major', labelsize=6, rotation=45)
    ax1.set_ylabel('Cost (£m)')
    ylab1 = ax1.yaxis.get_label()
    ylab1.set_style('italic')
    ylab1.set_size(8)
    ax1.grid(color='grey', linestyle='-', linewidth=0.2)
    ax1.legend(prop={'size': 6})
    ax1.set_title('Fig 1 - cost profile changes', loc='left', fontsize=8, fontweight='bold')

    # scaling y axis
    # y axis value setting so it takes highest cost profile yeah
    all = profile_data_total[0] + profile_data_total[1] + profile_data_total[2]
    y_max = max(all) + max(all) * 1 / 5
    ax1.set_ylim(0, y_max)

    # plot rdel/cdel chart data
    ax2.plot(year, latest_profile_cdel, label='CDEL', linewidth=3.0, marker="o")
    ax2.plot(year, latest_profile_rdel, label='RDEL', linewidth=3.0, marker="o")

    #rdel/cdel profile chart styling
    ax2.tick_params(axis='x', which='major', labelsize=6, rotation=45)
    ax2.set_xlabel('Financial Years')
    ax2.set_ylabel('Cost (£m)')
    xlab2 = ax2.xaxis.get_label()
    ylab2 = ax2.yaxis.get_label()
    xlab2.set_style('italic')
    xlab2.set_size(8)
    ylab2.set_style('italic')
    ylab2.set_size(8)
    ax2.grid(color='grey', linestyle='-', linewidth=0.2)
    ax2.legend(prop={'size': 6})

    ax2.set_ylim(0, y_max)

    ax2.set_title('Fig 2 - cost profile spend type', loc='left', fontsize=8, fontweight='bold')

    # size of chart and fit
    fig.canvas.draw()
    fig.tight_layout(rect=[0, 0.03, 1, 0.95])  # for title

    fig.savefig('cost_profile.png')
    plt.close()  # automatically closes figure so don't need to do manually.

    doc.add_picture('cost_profile.png', width=Inches(8))  # to place nicely in doc
    os.remove('cost_profile.png')

    return doc


def milestone_swimlane_charts(doc, project_name, latest_milestone_names, latest_milestone_dates, \
                              last_milestone_dates, baseline_milestone_dates, graph_title):
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    #build scatter chart
    fig, ax1 = plt.subplots()
    fig.suptitle(abbreviations[project_name] + ' ' + graph_title, fontweight='bold')  # title
    # set fig size
    # fig.set_figheight(4)
    # fig.set_figwidth(8)

    ax1.scatter(baseline_milestone_dates, latest_milestone_names, label='Baseline')
    ax1.scatter(last_milestone_dates, latest_milestone_names, label='Last Qrt')
    ax1.scatter(latest_milestone_dates, latest_milestone_names, label='Latest Qrt')

    # format the x ticks
    years = mdates.YearLocator()  # every year
    months = mdates.MonthLocator()  # every month
    years_fmt = mdates.DateFormatter('%Y')
    months_fmt = mdates.DateFormatter('%b')

    # calculate the length of the time period covered in chart. Not perfect as baseline dates can distort.
    try:
        td = (latest_milestone_dates[-1] - latest_milestone_dates[0]).days
        if td <= 365*2:
            ax1.xaxis.set_major_locator(years)
            ax1.xaxis.set_minor_locator(months)
            ax1.xaxis.set_major_formatter(years_fmt)
            ax1.xaxis.set_minor_formatter(months_fmt)
            plt.setp(ax1.xaxis.get_minorticklabels(), rotation=45, fontsize=6)
            plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, weight='bold', fontsize=8)
            # scaling x axis
            # x axis value to no more than three months after last latest milestone date, or three months
            # before first latest milestone date. Hack, can be improved. Text highlights movements off chart.
            x_max = last_milestone_dates[-1] + timedelta(days=90)
            x_min = last_milestone_dates[0] - timedelta(days=90)
            for date in baseline_milestone_dates:
                if date > x_max:
                    ax1.set_xlim(x_min, x_max)
                    plt.figtext(0.98, 0.03,
                                'Check full schedule to see all milestone movements',
                                horizontalalignment='right', fontsize=6, fontweight='bold')
                if date < x_min:
                    ax1.set_xlim(x_min, x_max)
                    plt.figtext(0.98, 0.03,
                                'Check full schedule to see all milestone movements',
                                horizontalalignment='right', fontsize=6, fontweight='bold')
        else:
            ax1.xaxis.set_major_locator(years)
            ax1.xaxis.set_minor_locator(months)
            ax1.xaxis.set_major_formatter(years_fmt)
            plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, weight='bold')
    except IndexError: #if milestone dates list is empty:
        pass

    ax1.legend() #insert legend

    #reverse y axis so order is earliest to oldest
    ax1 = plt.gca()
    ax1.set_ylim(ax1.get_ylim()[::-1])
    ax1.tick_params(axis='y', which='major', labelsize=7)
    ax1.yaxis.grid()  # horizontal lines
    ax1.set_axisbelow(True)

    #Add line of IPDC date, but only if in the time period
    try:
        if latest_milestone_dates[0] <= ipdc_date <= latest_milestone_dates[-1]:
            plt.axvline(ipdc_date)
            plt.figtext(0.98, 0.01, 'Line represents when IPDC will discuss Q1 20_21 portfolio management report',
                        horizontalalignment='right', fontsize=6, fontweight='bold')
    except IndexError:
        pass

    #size of chart and fit
    fig.canvas.draw()
    fig.tight_layout(rect=[0, 0.03, 1, 0.95]) #for title

    fig.savefig('schedule.png', bbox_inches='tight')
    plt.close() #automatically closes figure so don't need to do manually.

    doc.add_picture('schedule.png', width=Inches(8))  # to place nicely in doc
    os.remove('schedule.png')

    return doc

def milestone_table(doc, p_baseline_milestones, project_name):

    second_diff_data = project_time_difference(p_current_milestones, p_baseline_milestones)

    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Milestone'
    hdr_cells[1].text = 'Date'
    hdr_cells[2].text = 'Change from Lst Qrt'
    hdr_cells[3].text = 'Change from BL'
    hdr_cells[4].text = 'Notes'

    for milestone in p_current_milestones[project_name].keys():

        milestone_date = tuple(p_current_milestones[project_name][milestone])[0]

        try:
            if milestone_filter_start_date <= milestone_date: # filter based on date
                row_cells = table.add_row().cells
                row_cells[0].text = milestone
                if milestone_date is None:
                    row_cells[1].text = 'No date'
                else:
                    row_cells[1].text = milestone_date.strftime("%d/%m/%Y")
                b_one_value = first_diff_data[project_name][milestone]
                row_cells[2].text = plus_minus_days(b_one_value)
                b_two_value = second_diff_data[project_name][milestone]
                row_cells[3].text = plus_minus_days(b_two_value)

                notes = p_current_milestones[project_name][milestone][milestone_date]
                row_cells[4].text = str(notes)
                # trying to high changes to narratuve in red text
                # if milestone in p_last_milestones[project_name].keys():
                #     last_milestone_date = p_last_milestones[project_name][milestone]
                #     last_note = p_last_milestones[project_name][milestone][last_milestone_date]
                #     row_cells[4] = compare_text_newandold(notes, last_note, doc)
                # elif milestone not in p_last_milestones[project_name].keys():
                #     row_cells[4].text = str(notes)

                paragraph = row_cells[4].paragraphs[0]
                run = paragraph.runs
                font = run[0].font
                font.size = Pt(8)  # font size = 8


        except TypeError:  # this is to deal with none types which are still placed in output
            row_cells = table.add_row().cells
            row_cells[0].text = milestone
            if milestone_date is None:
                row_cells[1].text = 'No date'
            else:
                row_cells[1].text = milestone_date.strftime("%d/%m/%Y")
            b_one_value = first_diff_data[project_name][milestone]
            row_cells[2].text = plus_minus_days(b_one_value)
            b_two_value = second_diff_data[project_name][milestone]
            row_cells[3].text = plus_minus_days(b_two_value)
            notes = p_current_milestones[project_name][milestone][milestone_date]
            row_cells[4].text = str(notes)
            paragraph = row_cells[4].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size = Pt(8)  # font size = 8

    table.style = 'Table Grid'

    # column widths
    column_widths = (Cm(6), Cm(2.6), Cm(2), Cm(2), Cm(8.95))
    set_col_widths(table, column_widths)
    # make_columns_bold([table.columns[0], table.columns[3]])  # make keys bold
    # make_text_red([table.columns[1], table.columns[4]])  # make 'not reported red'

    make_rows_bold([table.rows[0]]) # makes top of table bold. Found function on stack overflow.

    return doc

def make_rows_bold(rows=list):
    '''Makes text bold in specified row'''
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def make_columns_bold(columns=list):
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def make_text_red(columns=list):
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text == 'Not reported':
                        run.font.color.rgb = RGBColor(255, 0, 0)

def change_text_size(columns=list, size=int):
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(size)

def plus_minus_days(change_value):
    '''mini function to place plus or minus sign before time delta
    value in milestone_table function. Only need + signs to be added
    as negative numbers have minus already'''
    try:
        if change_value > 0:
            text = '+ ' + str(change_value)
        else:
            text = str(change_value)
    except TypeError:
        text = change_value

    return text

def get_financial_totals(project_name):
    '''gets financial data to place into the bar chart element in the financial analysis graphs'''
    key_list = [('Pre-profile RDEL',
                'Pre-profile CDEL'),
                ('Total RDEL Forecast Total',
                 'Total CDEL Forecast Total WLC'),
                ('Unprofiled RDEL Forecast Total',
                 'Unprofiled CDEL Forecast Total WLC')]

    total_cost_list = []
    rdel_cost_list = []
    cdel_cost_list = []

    index_1 = costs_bl_index[project_name]
    index_2 = index_1[0:3]
    index_2.reverse()
    for x in index_2:
        if x is not None:
            for y in key_list:
                rdel = list_of_masters_all[x].data[project_name][y[0]]
                cdel = list_of_masters_all[x].data[project_name][y[1]]
                total = rdel + cdel
                total_cost_list.append(total)
                rdel_cost_list.append(rdel)
                cdel_cost_list.append(cdel)

        else:
            for i in range(len(key_list)):
                rdel = 0
                cdel = 0
                total = 0
                total_cost_list.append(total)
                rdel_cost_list.append(rdel)
                cdel_cost_list.append(cdel)

    return total_cost_list, rdel_cost_list, cdel_cost_list

def get_ben_totals(project_name):
    '''gets benefits data to place into the bar chart element in the financial analysis graphs'''

    ben_change_key_list = ['Pre-profile BEN Total',
                           'Total BEN Forecast - Total Monetised Benefits',
                           'Unprofiled Remainder BEN Forecast - Total Monetised Benefits']

    ben_type_key_list = ['Pre-profile BEN Forecast Gov Cashable',
                  'Pre-profile BEN Forecast Gov Non-Cashable',
                  'Pre-profile BEN Forecast - Economic (inc Private Partner)',
                  'Pre-profile BEN Forecast - Disbenefit UK Economic',
                  'Unprofiled Remainder BEN Forecast - Gov. Cashable',
                  'Unprofiled Remainder BEN Forecast - Gov. Non-Cashable',
                  'Unprofiled Remainder BEN Forecast - Economic (inc Private Partner)',
                  'Unprofiled Remainder BEN Forecast - Disbenefit UK Economic',
                  'Total BEN Forecast - Gov. Cashable',
                  'Total BEN Forecast - Gov. Non-Cashable',
                  'Total BEN Forecast - Economic (inc Private Partner)',
                  'Total BEN Forecast - Disbenefit UK Economic']


    ben_list = []
    index_1 = benefits_bl_index[project_name]
    index_2 = index_1[0:3]
    index_2.reverse()
    for x in index_2:
        if x is not None:
            for y in ben_change_key_list:
                ben = list_of_masters_all[x].data[project_name][y]
                ben_list.append(ben)
        else:
            for i in range(len(ben_change_key_list)):
                ben = 0
                ben_list.append(ben)

    ben_type_list = []
    for y in ben_type_key_list:
        ben = list_of_masters_all[0].data[project_name][y]
        if ben is not None:
            ben_type_list.append(ben)
        else:
            ben_type_list.append(0)

    return ben_list, ben_type_list

def milestone_schedule_data(latest_m_dict, last_m_dict, baseline_m_dict, project_name,
                                     filter_start_date = datetime.date(2000, 1, 1),
                                     filter_end_date = datetime.date(2050, 1, 1)):
    milestone_names = []
    mile_d_l_lst = []
    mile_d_last_lst = []
    mile_d_bl_lst = []

    #lengthy for loop designed so that all milestones and dates are stored and shown in output chart, even if they
    #were not present in last and baseline data reporting
    for m in list(latest_m_dict[project_name].keys()):
        if m == 'Project - Business Case End Date':  # filter out as to far in future
            pass
        else:
            if m is not None:
                m_d = tuple(latest_m_dict[project_name][m])[0]

            try:
                if m in list(last_m_dict[project_name].keys()):
                    m_d_lst = tuple(last_m_dict[project_name][m])[0]
                else:
                    m_d_lst = tuple(latest_m_dict[project_name][m])[0]
            except KeyError:
                m_d_lst = tuple(latest_m_dict[project_name][m])[0]

            if m in list(baseline_m_dict[project_name].keys()):
                m_d_bl = tuple(baseline_m_dict[project_name][m])[0]
            else:
                m_d_bl = tuple(latest_m_dict[project_name][m])[0]

            if m_d is not None:
                try:
                    if filter_start_date <= m_d <= filter_end_date:
                        milestone_names.append(m)
                        mile_d_l_lst.append(m_d)
                        if m_d_lst is not None:
                            mile_d_last_lst.append(m_d_lst)
                        else:
                            mile_d_last_lst.append(m_d)
                        if m_d_bl is not None:
                            mile_d_bl_lst.append(m_d_bl)
                        else:
                            if m_d_lst is not None:
                                mile_d_bl_lst.append(m_d_lst)
                            else:
                                mile_d_bl_lst.append(m_d)
                except TypeError:
                    print('check ' + m + ' as date give is a str. It is ' + m_d)

    return milestone_names, mile_d_l_lst, mile_d_last_lst, mile_d_bl_lst


'''RUNNING PROGRAMME'''

'''enter into the printing function the quarter details for the output files e.g. _q4_1920 (note put underscore at
front)'''
produce_word_doc(list_of_masters_all[0].projects)
