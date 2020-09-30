# from analysis.data import root_path
# from docx import Document
# from docx.shared import Pt, Cm, Inches

# def set_col_widths(t):
#     widths = (Inches(1), Inches(1), Inches(0.2), Inches(1), Inches(1))
#     for row in t.rows:
#         for idx, width in enumerate(widths):
#             row.cells[idx].width = width
#
# def table():
#
#     doc = Document()
#
#     font = doc.styles['Normal'].font
#     font.name = 'Arial'
#     font.size = Pt(10)
#
#     heading = str('How do you change column widths?')
#     doc.add_paragraph().add_run(str(heading)).bold = True
#
#     table = doc.add_table(rows=1, cols=5)
#     hdr_cells = table.rows[0].cells
#     hdr_cells[0].text = 'Milestone'
#     hdr_cells[1].text = 'Date'
#     hdr_cells[2].text = 'Change from Lst Qrt'
#     hdr_cells[3].text = 'Change from BL'
#     hdr_cells[4].text = 'Notes'
#
#     # TODO specify column widths
#
#     for i in range(0,10):
#         row_cells = table.add_row().cells
#         row_cells[0].text = 'h'
#         row_cells[1].text = 'e'
#         row_cells[2].text = 'l'
#         row_cells[3].text = 'l'
#         row_cells[4].text = 'oooo'
#
#     #change table column width
#     # for cell in table.columns[2].cells:
#     #      cell.width = Cm(4)  #not working
#
#     set_col_widths(table) #not working
#
#     #make_rows_bold(table.rows[0]) # makes top of table bold. Found function on stack overflow.
#
#     doc.save(root_path / 'output/table.docx')
#
# def make_rows_bold(*rows):
#     '''Makes text bold in specified row'''
#     for row in rows:
#         for cell in row.cells:
#             for paragraph in cell.paragraphs:
#                 for run in paragraph.runs:
#                     run.font.bold = True

from docx import Document
from docx.shared import Pt, Inches


def set_col_widths(t):
    widths = (Inches(1), Inches(1), Inches(0.2), Inches(1), Inches(1))
    for row in t.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def table():
    doc = Document()

    font = doc.styles["Normal"].font
    font.name = "Arial"
    font.size = Pt(10)

    heading = "How do you change widths?"
    doc.add_paragraph().add_run(heading).bold = True

    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Milestone'
    hdr_cells[1].text = 'Date'
    hdr_cells[2].text = 'Change from Lst Qrt'
    hdr_cells[3].text = 'Change from BL'
    hdr_cells[4].text = 'Notes'

    for i in range(0, 10):
        row_cells = table.add_row().cells
        row_cells[0].text = 'h'
        row_cells[1].text = 'e'
        row_cells[2].text = 'l'
        row_cells[3].text = 'l'
        row_cells[4].text = 'oooo'

    set_col_widths(table)

    doc.save("/home/will/Documents/analysis_engine/output/table.docx")


table()

table()



