import os
import re
import difflib

from typing import TextIO, Union, Tuple, List
from matplotlib import pyplot as plt
from docx import Document, table
from docx.shared import Pt, Cm, RGBColor, Inches
from pdf2image import convert_from_path
from openpyxl import Workbook, load_workbook

from analysis_engine.error_msgs import logger


def open_word_doc(wd_path: str) -> Document:
    """Function stores an empty word doc as a variable"""
    return Document(wd_path)


def get_input_doc(file_path: TextIO) -> Union[Workbook, Document, None]:
    """
    Returns blank documents in analysis_engine/input file used for saving outputs.
    Raises error and user message if files are not present
    """
    try:
        if str(file_path).endswith(".docx"):
            return open_word_doc(file_path)
        if str(file_path).endswith(".xlsx"):
            return load_workbook(file_path)
    except FileNotFoundError:
        base = os.path.basename(file_path)
        raise FileNotFoundError(
            str(base) + " document not present in input file. Stopping."
        )


def put_matplotlib_fig_into_word(doc, fig: plt.figure or plt, **kwargs) -> None:
    """Does rendering of matplotlib graph into word. Best method I could find for
    maintain high quality render output it to firstly save as pdf and then convert
    to jpeg!
    kwargs can be width=Inches(int) or transparent=False
    """
    fig.savefig("fig.pdf")
    page = convert_from_path("fig.pdf", 500)
    page[0].save("fig.jpeg", "JPEG")
    doc.add_picture("fig.jpeg", **kwargs)  # to place nicely in doc
    # doc.add_picture("fig.jpeg")
    os.remove("fig.jpeg")
    os.remove("fig.pdf")
    plt.close()  # automatically closes figure so don't need to do manually.


def make_file_friendly(quarter_str: str) -> str:
    """Converts datamaps.api project_data_from_master quarter data into a string to use when
    saving output files. Courtesy of M Lemon."""
    regex = r"Q(\d) (\d+)\/(\d+)"
    return re.sub(regex, r"Q\1_\2_\3", quarter_str)


FIGURE_STYLE = {1: "half_horizontal", 2: "full_horizontal"}


def set_figure_size(graph_type: str) -> Tuple[int, int]:
    if graph_type == "half_horizontal":
        return 11.69, 5.10
    if graph_type == "full_horizontal":
        return 11.69, 8.20


def set_fig_size(kwargs, fig: plt.figure) -> plt.figure:
    if "fig_size" in kwargs:
        fig.set_size_inches(set_figure_size(kwargs["fig_size"]))
    else:
        fig.set_size_inches(set_figure_size(FIGURE_STYLE[2]))

    return fig


def get_chart_title(
    **c_kwargs,  # chart kwargs
) -> str:
    if "title" in c_kwargs:
        return c_kwargs["title"]
    else:
        logger.info("Please note you can provide a title for this chart using --title.")
        return None


# helper function for milestone chart
def handle_long_keys(key_name: str, **kwargs) -> str:
    if "output_type" in kwargs:
        if kwargs["output_type"] == "milestones":
            output_list = []
            for name in key_name:
                if len(name) >= 30:
                    l = name.split()
                    l.insert(3, "\n")
                    new_str = " ".join(l)
                    output_list.append(re.sub("\s\\n\s", "\n", new_str))
                else:
                    output_list.append(name)

            return output_list

    else:  # this is for dandelion
        if len(key_name) >= 25:
            l = key_name.split()
            word_count = len(l)
            if word_count == 1:
                return key_name
            if word_count == 2:
                l.insert(1, "\n")
            if word_count >= 4:
                l.insert(3, "\n")
                l.insert(6, "\n")
            new_str = " ".join(l)
            return re.sub("\s\\n\s", "\n", new_str)
        else:
            return key_name


def plus_minus_days(change_value):
    """mini function to place plus or minus sign before time delta
    value in milestone_table function. Only need + signs to be added
    as negative numbers have minus already"""
    try:
        if change_value > 0:
            text = "+ " + str(change_value)
        else:
            text = str(change_value)
    except TypeError:
        text = change_value

    return text


def set_col_widths(word_table: table, widths: list) -> None:
    """This function sets the width of table in a word document"""
    for row in word_table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def compare_text_new_and_old(text_1: str, text_2: str, doc: Document) -> None:
    """
    Compares two sets of text and highlights differences in red text. In word.
    """

    comp = difflib.Differ()
    diff = list(comp.compare(text_2.split(), text_1.split()))
    diff = [x for x in diff if x[0] != "-"]  # remove all deleted text
    diff = [x for x in diff if x[0] != "?"]  # remove ?. not sure what these represent.
    y = doc.add_paragraph()
    for i, text in enumerate(diff):
        # f = len(diff) - 1
        # if i < f:
        #     a = i - 1
        # else:
        #     a = i

        if text[0:3] == "  |" or text[0:3] == "+ |":
            # j = i + 1
            # if diff[i][:3] == "  |"
            # if diff[i][0:3] and diff[a][0:3] == "  |":
            y = doc.add_paragraph()
            # else:
            #     pass
        # elif text[0:3] == "+ |":
        #     if diff[i][0:3] and diff[a][0:3] == "+ |":
        #         y = doc.add_paragraph()
        #     else:
        #         pass
        # if text[0:3] == "- |":
        #     pass
        # if text[0:3] == "  -":
        #     y = doc.add_paragraph()
        #     g = diff[i][2]
        #     y.add_run(g)
        # elif text[0:3] == "  •":
        #     y = doc.add_paragraph()
        #     g = text[2]
        #     y.add_run(g)
        if text[0] == " ":
            # total_nc += 1
            if i == 0:
                y.add_run(text[2:])
            # if total_nc == 1:
            #     y.add_run(text[2:])
            else:
                y.add_run(text[1:])
        if text[0] == "+":
            # w = len(diff[i])
            # g = diff[i][1:w]
            # total_plus += 1  # new text might not be first
            if i == 0:
                y.add_run(text[2:]).font.color.rgb = RGBColor(255, 0, 0)
            else:
                y.add_run(text[1:]).font.color.rgb = RGBColor(255, 0, 0)
        # if diff[i][0] == "-":
        #     pass
        # if diff[i][0] == "?":
        #     pass
        # else:
        #     if diff[i] != "+ |":
        #         y.add_run(diff[i][1:])


def make_columns_bold(columns: list) -> None:
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def change_text_size(columns: list, size: int) -> None:
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(size)


def make_text_red(columns: list) -> None:
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text == "Not reported":
                        run.font.color.rgb = RGBColor(255, 0, 0)
