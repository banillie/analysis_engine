from docx import table
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, RGBColor, Inches
from docx import Document

from analysis_engine.benefits import BenefitsData
from analysis_engine.dictionaries import (
    DCA_KEYS,
    SUMMARY_DCA_TEXT,
    CONVERT_RAG,
    SUMMARY_NARRATIVES,
)
from analysis_engine.segmentation import get_group
from analysis_engine.render_utils import (
    make_file_friendly,
    set_col_widths,
    compare_text_new_and_old,
    make_columns_bold,
    change_text_size,
    make_text_red,
    get_input_doc,
    put_matplotlib_fig_into_word,
)
from analysis_engine.costs import CostData, cost_profile_graph_new
from analysis_engine.milestones import MilestoneData, get_milestone_date
from analysis_engine.risks import RiskData
from analysis_engine.cleaning import convert_none_types


def run_p_reports(master, **kwargs) -> None:
    group = get_group(master, master["current_quarter"], **kwargs)
    for p in group:
        print("Compiling summary for " + p)
        kwargs["full_name"] = p
        kwargs["group"] = [master["project_information"][p]["Abbreviations"]]
        # report_doc = ""
        report_doc = get_input_doc(kwargs["root_path"] + kwargs["word_portrait"])
        # qrt = make_file_friendly(str(master["current_quarter"])
        out = compile_p_report_new(report_doc, master, **kwargs)
        out.save(
            str(kwargs["root_path"]) + kwargs["word_save_path"].format(f"{p}_summary")
        )
        # if kwargs["type"] == "long":
        #     output.save(root_path / "output/{}_long_report_{}.docx".format(p, qrt))
        # if kwargs["type"] == "short":
        #     output.save(root_path / "output/{}_short_report_{}.docx".format(p, qrt))


def compile_p_report_new(
    doc: Document,
    master,
    **kwargs,
) -> Document:
    wd_heading(doc, **kwargs)
    key_contacts(doc, master, **kwargs)
    dca_table(doc, master, **kwargs)
    dca_narratives(doc, master, **kwargs)
    costs = CostData(master, **kwargs)
    costs.get_forecast_cost_profile()
    # benefits = BenefitsData(master, **kwargs)  # not used
    milestones = MilestoneData(master, **kwargs)
    project_report_meta_data(doc, master, costs, milestones, **kwargs)
    change_word_doc_landscape(doc)
    kwargs["title"] = "Cost Profile (for each of the last four quarters)"
    ms_graph = cost_profile_graph_new(costs, **kwargs)
    put_matplotlib_fig_into_word(doc, ms_graph, width=Inches(8))
    # risks = RiskData(master, **kwargs)
    # print_out_project_risks(doc, risks, **kwargs)
    return doc


def delete_paragraph(paragraph):
    """helper function to remove empyt para at top of summary_temp doc.
    only used here."""
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def wd_heading(
    doc: Document,
    **kwargs,
) -> None:
    heading = str(kwargs["group"][0])
    intro = doc.add_heading(heading, 0)
    intro.alignment = 1
    intro.bold = True


def key_contacts(doc: Document, master, **kwargs) -> None:
    data = master["master_data"][0]["data"][kwargs["full_name"]]
    """Function adds keys contact details"""
    sro_name = data["Senior Responsible Owner (SRO)"]
    if sro_name is None:
        sro_name = "tbc"

    sro_email = data["Senior Responsible Owner (SRO) - Email"]
    if sro_email is None:
        sro_email = "email: tbc"

    sro_phone = data["SRO Phone No."]
    if sro_phone == None:
        sro_phone = "phone number: tbc"

    doc.add_paragraph(
        "SRO: " + str(sro_name) + ", " + str(sro_email) + ", " + str(sro_phone)
    )

    pd_name = data["Project Director (PD)"]
    if pd_name is None:
        pd_name = "TBC"

    pd_email = data["Project Director (PD) - Email"]
    if pd_email is None:
        pd_email = "email: tbc"

    pd_phone = data["PD Phone No."]
    if pd_phone is None:
        pd_phone = "phone: tbc"

    doc.add_paragraph(
        "PD: " + str(pd_name) + ", " + str(pd_email) + ", " + str(pd_phone)
    )


def dca_table(doc: Document, master, **kwargs) -> None:
    """Creates SRO confidence table"""

    # doc.add_paragraph()
    p = doc.add_paragraph()
    text = "* Note in Q2 2021/22 DCA ratings moved to a three point scale."
    p.add_run(text)
    # .font.color.rgb = RGBColor(255, 0, 0)

    w_table = doc.add_table(rows=1, cols=5)
    hdr_cells = w_table.rows[0].cells
    hdr_cells[0].text = "Delivery confidence"
    hdr_cells[1].text = "This quarter"
    hdr_cells[2].text = str(master["master_data"][1]["quarter"])
    hdr_cells[3].text = str(master["master_data"][2]["quarter"])
    hdr_cells[4].text = str(master["master_data"][3]["quarter"])

    SRO_CONF_KEY_LIST = list(DCA_KEYS["ipdc"].values())
    for x, dca_key in enumerate(SRO_CONF_KEY_LIST):
        row_cells = w_table.add_row().cells
        row_cells[0].text = SUMMARY_DCA_TEXT[dca_key]
        for i, m in enumerate(master["master_data"][:4]):  # last four masters taken
            try:
                rating = CONVERT_RAG[m["data"][kwargs["full_name"]][dca_key]]
                row_cells[i + 1].text = rating
                cell_colouring(row_cells[i + 1], rating)
            except (KeyError, TypeError):
                row_cells[i + 1].text = "N/A"

    w_table.style = "Table Grid"
    make_rows_bold([w_table.rows[0]])  # makes top of table bold.
    # make_columns_bold([table.columns[0]]) #right cells in table bold
    column_widths = (Cm(4.4), Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.8))
    set_col_widths(w_table, column_widths)


def dca_narratives(doc: Document, master, **kwargs) -> None:
    """Places all narratives into document and checks for differences between
    current and last quarter"""

    doc.add_paragraph()
    p = doc.add_paragraph()
    text = "*Red text highlights changes in narratives from last quarter"
    p.add_run(text).font.color.rgb = RGBColor(255, 0, 0)

    for k in list(SUMMARY_NARRATIVES.keys()):
        # keeping this statement in as don't want to crash output if a key name has changed.
        try:
            doc.add_paragraph().add_run(k).bold = True
            text_one = str(
                master["master_data"][0]["data"][kwargs["full_name"]][
                    SUMMARY_NARRATIVES[k]
                ]
            )
            try:
                text_two = str(
                    master["master_data"][1]["data"][kwargs["full_name"]][
                        SUMMARY_NARRATIVES[k]
                    ]
                )
            except (KeyError, IndexError):  # index error relates to data_bridge
                text_two = text_one
        except KeyError:
            break

        compare_text_new_and_old(text_one, text_two, doc)


def project_report_meta_data(
    doc: Document,
    master,
    costs: CostData,
    milestones: MilestoneData,
    **kwargs,
):
    """Meta data table"""
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Annex A. High level MI data and analysis").bold = True

    """Costs meta data"""
    # this chuck is pretty messy because the data is messy
    run = doc.add_paragraph().add_run("Costs - Forecast")
    font = run.font
    font.bold = True
    font.underline = True
    # master_data = costs.master['master_data'][0]["data"]
    t = doc.add_table(rows=1, cols=4)
    hdr_cells = t.rows[0].cells

    hdr_cells[0].text = "Total:"
    hdr_cells[1].text = (
        "£" + str(round(costs.totals[master["current_quarter"]]["total"])) + "m"
    )

    ## NOT USER REQUIREMENT AT MOMENT
    # hdr_cells[2].text = "CDEL:"
    # hdr_cells[3].text = (
    #     "£" + str(round(costs.totals[master['current_quarter']]['cdel'])) + "m"
    # )
    #
    # row_cells = t.add_row().cells
    # row_cells[0].text = "RDEL:"
    # row_cells[1].text = (
    #     "£" + str(round(costs.totals[master['current_quarter']]['rdel'])) + "m"
    # )
    # row_cells[2].text = "Non-Gov:"
    # row_cells[3].text = (
    #     "£" + str(round(costs.totals[master['current_quarter']]['n_gov'])) + "m"
    # )

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(t, column_widths)
    # make column keys bold
    make_columns_bold([t.columns[0], t.columns[2]])
    change_text_size([t.columns[0], t.columns[1], t.columns[2], t.columns[3]], 10)

    """Financial data"""
    doc.add_paragraph()
    run = doc.add_paragraph().add_run("Costing data")
    font = run.font
    font.bold = True
    font.underline = True
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Type of funding:"
    hdr_cells[1].text = str(
        master["master_data"][0]["data"][kwargs["full_name"]]["Source of Finance"]
    )
    hdr_cells[2].text = "Contingency:"
    contingency = convert_none_types(
        master["master_data"][0]["data"][kwargs["full_name"]][
            "Overall contingency (£m)"
        ]
    )
    if contingency is None:  # can this be refactored?
        hdr_cells[3].text = "None"
    else:
        hdr_cells[3].text = "£" + str(round(contingency)) + "m"
    row_cells = table.add_row().cells
    row_cells[0].text = "Optimism Bias (OB):"
    ob = convert_none_types(
        master["master_data"][0]["data"][kwargs["full_name"]][
            "Overall figure for Optimism Bias (£m)"
        ]
    )
    if ob is None:
        row_cells[1].text = str(ob)
    else:
        try:
            row_cells[1].text = "£" + str(round(ob)) + "m"
        except TypeError:
            row_cells[1].text = ob
    row_cells[2].text = "Contingency in costs:"
    con_included_wlc = master["master_data"][0]["data"][kwargs["full_name"]][
        "Is this Continency amount included within the WLC?"
    ]
    if con_included_wlc is None:
        row_cells[3].text = "Not reported"
    else:
        row_cells[3].text = con_included_wlc
    row_cells = table.add_row().cells
    row_cells[0].text = "OB in costs:"
    ob_included_wlc = master["master_data"][0]["data"][kwargs["full_name"]][
        "Is this Optimism Bias included within the WLC?"
    ]
    if ob_included_wlc is None:
        row_cells[1].text = "Not reported"
    else:
        row_cells[1].text = str(ob_included_wlc)
    row_cells[2].text = ""
    row_cells[3].text = ""

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(table, column_widths)
    # make column keys bold
    make_columns_bold([table.columns[0], table.columns[2]])
    change_text_size(
        [table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10
    )

    """Project Stage data"""
    doc.add_paragraph()
    run = doc.add_paragraph().add_run("Stage data")
    font = run.font
    font.bold = True
    font.underline = True
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Business case stage:"
    hdr_cells[1].text = master["master_data"][0]["data"][kwargs["full_name"]][
        "IPDC approval point"
    ]
    hdr_cells[2].text = "Delivery stage:"
    delivery_stage = str(
        convert_none_types(
            master["master_data"][0]["data"][kwargs["full_name"]]["Project stage"]
        )
    )
    if delivery_stage is None:
        hdr_cells[3].text = "Not reported"
    else:
        hdr_cells[3].text = delivery_stage

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(table, column_widths)
    # make column keys bold
    make_columns_bold([table.columns[0], table.columns[2]])
    change_text_size(
        [table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10
    )
    make_text_red([table.columns[1], table.columns[3]])  # make 'not reported red'

    """Milestone/Stage meta data"""
    abb = kwargs["group"][0]
    doc.add_paragraph()
    run = doc.add_paragraph().add_run("Schedule - Forecast")
    font = run.font
    font.bold = True
    font.underline = True
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Start date:"
    try:
        start_project = get_milestone_date(
            milestones.milestone_dict,
            "Project - Start Date",
            str(master["current_quarter"]),
            abb,
        )
        hdr_cells[1].text = start_project.strftime("%d/%m/%Y")
    except (KeyError, AttributeError):
        try:  # the team has two keys !!!!
            start_project = get_milestone_date(
                milestones.milestone_dict,
                "Start of Project",
                str(master["current_quarter"]),
                abb,
            )
            hdr_cells[1].text = start_project.strftime("%d/%m/%Y")
        except (KeyError, AttributeError):
            hdr_cells[1].text = "Not reported"

    hdr_cells[2].text = "Start of operations:"
    try:
        start_ops = get_milestone_date(
            milestones.milestone_dict,
            "Start of Operation",
            str(master["current_quarter"]),
            abb,
        )
        hdr_cells[3].text = start_ops.strftime("%d/%m/%Y")
    except (KeyError, AttributeError):
        hdr_cells[3].text = "Not reported"

    row_cells = table.add_row().cells
    row_cells[0].text = "Start of construction:"
    try:
        start_con = get_milestone_date(
            milestones.milestone_dict,
            "Start of Construction/build",
            str(master["current_quarter"]),
            abb,
        )
        row_cells[1].text = start_con.strftime("%d/%m/%Y")
    except (KeyError, AttributeError):
        row_cells[1].text = "Not reported"

    row_cells[2].text = "Full Operations:"  # check
    try:
        full_ops = get_milestone_date(
            milestones.milestone_dict,
            "Full Operations",
            str(master["current_quarter"]),
            abb,
        )
        row_cells[3].text = full_ops.strftime("%d/%m/%Y")
    except (KeyError, AttributeError):
        row_cells[3].text = "Not reported"

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(table, column_widths)
    # make column keys bold
    make_columns_bold([table.columns[0], table.columns[2]])
    change_text_size(
        [table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10
    )
    make_text_red([table.columns[1], table.columns[3]])  # make 'not reported red'

    """vfm meta data"""
    doc.add_paragraph()
    run = doc.add_paragraph().add_run("VfM data")
    font = run.font
    font.bold = True
    font.underline = True
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "VfM category:"
    vfm_cat = master["master_data"][0]["data"][kwargs["full_name"]][
        "VfM Category single entry"
    ]
    hdr_cells[1].text = str(vfm_cat)
    hdr_cells[2].text = "BCR:"
    bcr = master["master_data"][0]["data"][kwargs["full_name"]][
        "Adjusted Benefits Cost Ratio (BCR)"
    ]
    hdr_cells[3].text = str(bcr)

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(table, column_widths)
    # make column keys bold
    make_columns_bold([table.columns[0], table.columns[2]])
    change_text_size(
        [table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10
    )
    make_text_red([table.columns[1], table.columns[3]])  # make 'not reported red'

    """benefits meta data"""
    doc.add_paragraph()
    run = doc.add_paragraph().add_run("Benefits - Forecast")
    font = run.font
    font.bold = True
    font.underline = True
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Total:"
    hdr_cells[1].text = (
        "£"
        + str(
            round(
                convert_none_types(
                    master["master_data"][0]["data"][kwargs["full_name"]][
                        "Total BEN Forecast - Total Monetised Benefits"
                    ]
                )
            )
        )
        + "m"
    )
    hdr_cells[2].text = "Economic:"
    hdr_cells[3].text = (
        "£"
        + str(
            round(
                convert_none_types(
                    master["master_data"][0]["data"][kwargs["full_name"]][
                        "Total BEN Forecast - Economic (inc Private Partner)"
                    ]
                )
            )
        )
        + "m"
    )

    row_cells = table.add_row().cells
    row_cells[0].text = "Cashable:"
    row_cells[1].text = (
        "£"
        + str(
            round(
                convert_none_types(
                    master["master_data"][0]["data"][kwargs["full_name"]][
                        "Total BEN Forecast - Gov. Cashable"
                    ]
                )
            )
        )
        + "m"
    )
    row_cells[2].text = "Disbenefits:"
    row_cells[3].text = (
        "£"
        + str(
            round(
                convert_none_types(
                    master["master_data"][0]["data"][kwargs["full_name"]][
                        "Total BEN Forecast - Disbenefit UK Economic"
                    ]
                )
            )
        )
        + "m"
    )

    row_cells = table.add_row().cells
    row_cells[0].text = "Non-Cashable:"
    row_cells[1].text = (
        "£"
        + str(
            round(
                convert_none_types(
                    master["master_data"][0]["data"][kwargs["full_name"]][
                        "Total BEN Forecast - Gov. Non-Cashable"
                    ]
                )
            )
        )
        + "m"
    )

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(table, column_widths)
    # make column keys bold
    make_columns_bold([table.columns[0], table.columns[2]])
    change_text_size(
        [table.columns[0], table.columns[1], table.columns[2], table.columns[3]], 10
    )
    return doc


def change_word_doc_landscape(doc: Document) -> Document:
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)  # new page
    new_width, new_height = new_section.page_height, new_section.page_width
    new_section.orientation = WD_ORIENTATION.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    return doc


def change_word_doc_portrait(doc: Document) -> Document:
    new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    new_width, new_height = new_section.page_height, new_section.page_width
    new_section.orientation = WD_ORIENTATION.PORTRAIT
    new_section.page_width = new_width
    new_section.page_height = new_height
    return doc


def print_out_project_risks(doc: Document, risks: RiskData, **kwargs) -> Document:
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    # table heading
    ab = kwargs["group"][0]
    doc.add_paragraph().add_run(str(ab + " RISKS")).bold = True

    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Description "
    hdr_cells[1].text = "Internal Control"
    hdr_cells[2].text = "Mitigation"
    hdr_cells[3].text = "Impact"
    hdr_cells[4].text = "Likelihood"

    p_risks = risks.risk_dictionary[kwargs["quarter"][0]][ab]

    for i in p_risks:
        row_cells = table.add_row().cells
        row_cells[0].text = p_risks[i]["Brief Risk Description "]
        row_cells[1].text = p_risks[i]["BRD Internal Control"]
        row_cells[2].text = p_risks[i][
            "BRD Mitigation - Actions taken (brief description)"
        ]
        row_cells[3].text = p_risks[i]["BRD Residual Impact"]
        row_cells[4].text = p_risks[i]["BRD Residual Likelihood"]

    table.style = "Table Grid"
    # column widths
    column_widths = (Cm(5), Cm(1.5), Cm(11), Cm(1.5), Cm(1.5))
    set_col_widths(table, column_widths)
    # make_columns_bold([table.columns[0], table.columns[3]])  # make keys bold
    # make_text_red([table.columns[1], table.columns[4]])  # make 'not reported red'

    make_rows_bold(
        [table.rows[0]]
    )  # makes top of table bold. Found function on stack overflow.
    return doc


def make_rows_bold(rows: list) -> None:
    """This function makes text bold in a list of row numbers for a word document"""
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def cell_colouring(word_table_cell: table.Table.cell, colour: str) -> None:
    """Function that handles cell colouring for word documents"""

    try:
        if colour == "R":
            colour = parse_xml(r'<w:shd {} w:fill="cb1f00"/>'.format(nsdecls("w")))
        elif colour == "A/R":
            colour = parse_xml(r'<w:shd {} w:fill="f97b31"/>'.format(nsdecls("w")))
        elif colour == "A":
            colour = parse_xml(r'<w:shd {} w:fill="fce553"/>'.format(nsdecls("w")))
        elif colour == "A/G":
            colour = parse_xml(r'<w:shd {} w:fill="a5b700"/>'.format(nsdecls("w")))
        elif colour == "G":
            colour = parse_xml(r'<w:shd {} w:fill="17960c"/>'.format(nsdecls("w")))

        word_table_cell._tc.get_or_add_tcPr().append(colour)

    except TypeError:
        pass
