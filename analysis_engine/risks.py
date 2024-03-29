from collections import Counter

from openpyxl import workbook, Workbook
from docx import Document

from analysis_engine.dictionaries import (
    RISK_SCORES,
    PORTFOLIO_RISK_SCORES,
    RISK_LIST,
    PORTFOLIO_RISK_LIST,
    PORTFOLIO_RISKS_WORD,
    RISK_NO_DICTIONARY,
    PORTFOLIO_RISK_IMPACT_ASSESSMENT,
    PORTFOLIO_RISK_COUNT_FILTER_OUTED_KEYS
)
from analysis_engine.segmentation import get_group, get_correct_p_data
from analysis_engine.render_utils import (
    make_file_friendly,
    get_input_doc,
    compare_text_new_and_old,
)
from analysis_engine.error_msgs import logger


def risk_score(risk_impact: str, risk_likelihood: str) -> str:
    impact_score = RISK_SCORES[risk_impact]
    try:
        likelihood_score = RISK_SCORES[risk_likelihood]
    except KeyError:
        likelihood_score = PORTFOLIO_RISK_SCORES[risk_likelihood]
    try:
        score = impact_score + likelihood_score
    except TypeError:
        if risk_impact == "N/A" and risk_likelihood == "N/A":
            return "N/A"
        return None
    if score <= 4:
        if risk_impact == "Medium":
            if risk_likelihood == "Medium" or risk_likelihood == "Possible":
                return "Medium"
            else:
                return "Low"
        else:
            return "Low"
    if 5 <= score <= 6:
        if risk_impact == "High":
            if risk_likelihood == "High" or risk_likelihood == "Likely":
                return "High"
        if risk_impact == "Low":
            if risk_likelihood == "Very High" and risk_likelihood == "Very Likely":
                return "Low"
        if risk_impact == "Very High":
            if risk_likelihood == "Low" and risk_likelihood == "Unlikely":
                return "Low"
        else:
            return "Medium"
    if score > 6:
        return "High"


class RiskData:
    def __init__(self, master, **kwargs):
        self.master = master
        self.kwargs = kwargs
        self.report = kwargs["report"]
        self.quarters = self.master["quarter_list"]
        self.risk_dictionary = {}
        self.portfolio_risk_dictionary = {}
        self.risk_count = {}
        self.portfolio_risk_count = {}
        self.risk_impact_count = {}
        self.portfolio_risk_impact_count = {}
        self.portfolio_type_impact_count = {}
        self.get_project_dictionary()
        self.get_portfolio_dictionary()
        self.get_count()
        self.get_portfolio_count()

    def get_project_dictionary(self):
        quarter_dict = {}
        for tp in self.quarters:
            project_dict = {}
            group = get_group(self.master, tp, **self.kwargs)
            for p in group:
                p_data = get_correct_p_data(self.master, p, tp)
                if p_data is None:
                    continue
                try:
                    project_number_dict = {}
                    for x in range(1, 11):  # currently 10 risks
                        project_risk_list = []
                        group = (
                            "Group",
                            self.master["project_information"][p]["Group"],
                        )
                        stage = ("Stage", p_data["IPDC approval point"])
                        project_risk_list.append(group)
                        project_risk_list.append(stage)
                        for risk_type in RISK_LIST:
                            try:
                                amended_risk_type = risk_type + str(x)
                                risk = (
                                    risk_type,
                                    p_data[amended_risk_type],
                                )
                                project_risk_list.append(risk)
                            except KeyError:
                                try:
                                    amended_risk_type = (
                                        risk_type[:4] + str(x) + risk_type[3:]
                                    )
                                    risk = (
                                        risk_type,
                                        p_data[amended_risk_type],
                                    )
                                    project_risk_list.append(risk)
                                except KeyError:
                                    try:
                                        if risk_type == PORTFOLIO_RISK_IMPACT_ASSESSMENT:
                                            impact = (
                                                "BRD Residual Impact"[:4]
                                                + str(x)
                                                + "BRD Residual Impact"[3:]
                                            )
                                            likelihoood = (
                                                "BRD Residual Likelihood"[:4]
                                                + str(x)
                                                + "BRD Residual Likelihood"[3:]
                                            )
                                            score = risk_score(
                                                p_data[impact],
                                                p_data[likelihoood],
                                            )
                                            risk = (
                                                PORTFOLIO_RISK_IMPACT_ASSESSMENT,
                                                score,
                                            )
                                            project_risk_list.append(risk)
                                    except KeyError:
                                        if risk_type == PORTFOLIO_RISK_IMPACT_ASSESSMENT:
                                            pass
                                        else:
                                            print(
                                                "check "
                                                + p
                                                + " "
                                                + str(x)
                                                + " "
                                                + risk_type
                                            )
                            # description_key.append()
                            if risk[1] is None:
                                break
                            project_number_dict[x] = dict(project_risk_list)
                    project_dict[
                        self.master["project_information"][p]["Abbreviations"]
                    ] = project_number_dict
                except KeyError:
                    pass
                quarter_dict[tp] = project_dict

        self.risk_dictionary = quarter_dict

    def get_portfolio_dictionary(self):
        quarter_dict = {}
        missing_key_list = []
        for tp in self.quarters:
            portfolio_dict = {}
            group = get_group(self.master, tp, **self.kwargs)
            for p in group:
                p_data = get_correct_p_data(self.master, p, tp)
                if p_data is None:
                    continue
                portfolio_number_dict = {}
                for x in range(1, 9):  # currently 8 risks.
                    portfolio_risk_list = []
                    group = (
                        "Group",
                        self.master["project_information"][p]["Group"],
                    )
                    stage = ("Stage", p_data["IPDC approval point"])
                    portfolio_risk_list.append(group)
                    portfolio_risk_list.append(stage)
                    for risk_type in PORTFOLIO_RISK_LIST:
                        try:
                            amended_risk_type = risk_type + " " + str(x)
                            risk = (
                                risk_type,
                                p_data[amended_risk_type],
                            )
                            portfolio_risk_list.append(risk)
                        except KeyError:
                            if risk_type == PORTFOLIO_RISK_IMPACT_ASSESSMENT:
                                try:
                                    score = risk_score(
                                        p_data[
                                            "Portfolio Risk Impact Assessment " + str(x)
                                        ],
                                        p_data["Portfolio Risk Likelihood Assessment " + str(x)],
                                    )
                                    risk = (
                                        PORTFOLIO_RISK_IMPACT_ASSESSMENT,
                                        score,
                                    )
                                    portfolio_risk_list.append(risk)
                                except KeyError:
                                    pass
                            else:
                                msg = (
                                    str(tp)
                                    + " master does not have key: "
                                    + amended_risk_type
                                )
                                if msg not in missing_key_list:
                                    missing_key_list.append(msg)
                                pass
                        portfolio_number_dict[x] = dict(portfolio_risk_list)

                    portfolio_dict[
                        self.master["project_information"][p]["Abbreviations"]
                    ] = portfolio_number_dict

                quarter_dict[tp] = portfolio_dict

        if missing_key_list:
            for p in missing_key_list:
                logger.info(f"{p} key is missing")

        self.portfolio_risk_dictionary = quarter_dict

    def get_count(self):
        count_output_dict = {}
        impact_output_dict = {}
        for quarter in self.risk_dictionary.keys():
            count_lower_dict = {}
            impact_lower_dict = {}
            for i in range(len(RISK_LIST)):
                count_list = []
                impact_list = []
                for y, project_name in enumerate(
                    list(self.risk_dictionary[quarter].keys())
                ):
                    for x, number in enumerate(
                        list(self.risk_dictionary[quarter][project_name].keys())
                    ):
                        try:
                            risk_value = self.risk_dictionary[quarter][project_name][
                                number
                            ][RISK_LIST[i]]
                            impact = self.risk_dictionary[quarter][project_name][
                                number
                            ][PORTFOLIO_RISK_IMPACT_ASSESSMENT]
                            count_list.append(risk_value)
                            impact_list.append((risk_value, impact))
                        except KeyError:
                            pass

                count_lower_dict[RISK_LIST[i]] = Counter(count_list)
                impact_lower_dict[RISK_LIST[i]] = Counter(impact_list)

            count_output_dict[quarter] = count_lower_dict
            impact_output_dict[quarter] = impact_lower_dict

        self.risk_count = count_output_dict
        self.risk_impact_count = impact_output_dict

    def get_portfolio_count(self):
        count_output_dict = {}
        impact_output_dict = {}
        type_output_dict = {}
        for quarter in self.portfolio_risk_dictionary.keys():
            count_lower_dict = {}
            impact_lower_dict = {}
            type_lower_dict = {}
            for i in range(len(PORTFOLIO_RISK_LIST)):
                count_list = []
                impact_list = []
                for y, project_name in enumerate(
                    list(self.portfolio_risk_dictionary[quarter].keys())
                ):
                    for number in list(
                        self.portfolio_risk_dictionary[quarter][project_name].keys()
                    ):
                        try:
                            risk_value = self.portfolio_risk_dictionary[quarter][
                                project_name
                            ][number][PORTFOLIO_RISK_LIST[i]]
                            # impact = 'High'
                            impact = self.portfolio_risk_dictionary[quarter][
                                project_name
                            ][number][PORTFOLIO_RISK_IMPACT_ASSESSMENT]
                            count_list.append(risk_value)
                            # impact_list.append((number, impact))
                            impact_list.append((risk_value, impact))
                        except KeyError:
                            pass

                count_lower_dict[PORTFOLIO_RISK_LIST[i]] = Counter(count_list)
                impact_lower_dict[PORTFOLIO_RISK_LIST[i]] = Counter(impact_list)

            for i in range(1, 9):  # currently 6 risks. Changed from 5 to 6 in Q4 2021
                type_list = []
                for project_name in list(
                    self.portfolio_risk_dictionary[quarter].keys()
                ):
                    try:
                        risk_type = i
                        impact = self.portfolio_risk_dictionary[quarter][project_name][
                            i
                        ][PORTFOLIO_RISK_IMPACT_ASSESSMENT]
                        type_list.append(impact)
                    except KeyError:
                        pass
                type_lower_dict[i] = Counter(type_list)

            count_output_dict[quarter] = count_lower_dict
            impact_output_dict[quarter] = impact_lower_dict
            type_output_dict[quarter] = type_lower_dict

        self.portfolio_risk_count = count_output_dict
        self.portfolio_risk_impact_count = impact_output_dict
        self.portfolio_type_impact_count = type_output_dict


def risks_into_excel(risk_data: RiskData) -> workbook:
    wb = Workbook()

    for q in risk_data.risk_dictionary.keys():
        start_row = 1
        ws = wb.create_sheet(
            make_file_friendly(str(q) + " all data")
        )  # creating worksheets. names restricted to 30 characters.
        ws.title = make_file_friendly(q + " all data")  # title of worksheet

        for y, project_name in enumerate(list(risk_data.risk_dictionary[q].keys())):
            for x, number in enumerate(
                list(risk_data.risk_dictionary[q][project_name].keys())
            ):
                ws.cell(
                    row=start_row + number, column=1
                ).value = risk_data.risk_dictionary[q][project_name][number]["Group"]
                ws.cell(row=start_row + number, column=2).value = project_name
                ws.cell(
                    row=start_row + number, column=3
                ).value = risk_data.risk_dictionary[q][project_name][number]["Stage"]
                ws.cell(row=start_row + number, column=4).value = str(number)
                for i in range(len(RISK_LIST)):
                    try:
                        ws.cell(
                            row=start_row + number, column=5 + i
                        ).value = risk_data.risk_dictionary[q][project_name][number][
                            RISK_LIST[i]
                        ]
                    except KeyError:
                        # print(project_name)
                        pass

            start_row += number

        for i in range(len(RISK_LIST)):
            ws.cell(row=1, column=5 + i).value = RISK_LIST[i]
        ws.cell(row=1, column=1).value = "DfT Group"
        ws.cell(row=1, column=2).value = "Project Name"
        ws.cell(row=1, column=3).value = "Stage"
        ws.cell(row=1, column=4).value = "Risk Number"

        ws = wb.create_sheet(
            make_file_friendly(q + " Count")
        )  # creating worksheets. names restricted to 30 characters.
        ws.title = make_file_friendly(q + " Count")  # title of worksheet

        start_row = 3
        for v, risk_cat in enumerate(list(risk_data.risk_count[q].keys())):
            if (
                risk_cat == "Brief Risk Description "
                or risk_cat == "BRD Mitigation - Actions taken (brief description)"
            ):
                pass
            else:
                ws.cell(row=start_row, column=2).value = risk_cat
                ws.cell(row=start_row, column=3).value = "Low"
                ws.cell(row=start_row, column=4).value = "Medium"
                ws.cell(row=start_row, column=5).value = "High"
                ws.cell(row=start_row, column=6).value = "Total"
                for b, cat in enumerate(list(risk_data.risk_count[q][risk_cat].keys())):
                    ws.cell(row=start_row + b + 1, column=2).value = str(cat)
                    ws.cell(
                        row=start_row + b + 1, column=3
                    ).value = risk_data.risk_impact_count[q][risk_cat][(cat, "Low")]
                    ws.cell(
                        row=start_row + b + 1, column=4
                    ).value = risk_data.risk_impact_count[q][risk_cat][(cat, "Medium")]
                    ws.cell(
                        row=start_row + b + 1, column=5
                    ).value = risk_data.risk_impact_count[q][risk_cat][(cat, "High")]
                    ws.cell(
                        row=start_row + b + 1, column=6
                    ).value = risk_data.risk_count[q][risk_cat][cat]

                start_row += b + 4

    wb.remove(wb["Sheet"])

    return wb


def portfolio_risks_into_excel(risk_data: RiskData) -> workbook:
    wb = Workbook()

    for q in risk_data.portfolio_risk_dictionary.keys():
        start_row = 1
        ws = wb.create_sheet(
            make_file_friendly(str(q) + " all data")
        )  # creating worksheets. names restricted to 30 characters.
        ws.title = make_file_friendly(q + " all data")  # title of worksheet

        for y, project_name in enumerate(
            list(risk_data.portfolio_risk_dictionary[q].keys())
        ):
            # print(project_name)
            for x, number in enumerate(
                list(risk_data.portfolio_risk_dictionary[q][project_name].keys())
            ):
                ws.cell(
                    row=start_row + number, column=1
                ).value = risk_data.portfolio_risk_dictionary[q][project_name][number][
                    "Group"
                ]
                ws.cell(row=start_row + number, column=2).value = project_name
                ws.cell(
                    row=start_row + number, column=3
                ).value = risk_data.portfolio_risk_dictionary[q][project_name][number][
                    "Stage"
                ]
                ws.cell(row=start_row + number, column=4).value = number
                # print(project_name)
                for i in range(len(PORTFOLIO_RISK_LIST)):
                    try:
                        ws.cell(
                            row=start_row + number, column=5 + i
                        ).value = risk_data.portfolio_risk_dictionary[q][project_name][
                            number
                        ][
                            PORTFOLIO_RISK_LIST[i]
                        ]
                    except KeyError:
                        pass

            start_row += number

        for i in range(len(PORTFOLIO_RISK_LIST)):
            ws.cell(row=1, column=5 + i).value = PORTFOLIO_RISK_LIST[i]
        ws.cell(row=1, column=1).value = "DfT Group"
        ws.cell(row=1, column=2).value = "Project Name"
        ws.cell(row=1, column=3).value = "Stage"
        ws.cell(row=1, column=4).value = "Risk Number"

        ws = wb.create_sheet(
            make_file_friendly(q + " Count")
        )  # creating worksheets. names restricted to 30 characters.
        ws.title = make_file_friendly(q + " Count")  # title of worksheet

        start_row = 3
        for v, risk_cat in enumerate(list(risk_data.portfolio_risk_count[q].keys())):
            if risk_cat in PORTFOLIO_RISK_COUNT_FILTER_OUTED_KEYS:
                pass
            else:
                ws.cell(row=start_row, column=2).value = risk_cat
                ws.cell(row=start_row, column=3).value = "Low"
                ws.cell(row=start_row, column=4).value = "Medium"
                ws.cell(row=start_row, column=5).value = "High"
                ws.cell(row=start_row, column=6).value = "Total"
                for b, cat in enumerate(
                    list(risk_data.portfolio_risk_count[q][risk_cat].keys())
                ):
                    # print(b, cat)
                    ws.cell(row=start_row + b + 1, column=2).value = str(cat)
                    ws.cell(
                        row=start_row + b + 1, column=3
                    ).value = risk_data.portfolio_risk_impact_count[q][risk_cat][
                        (cat, "Low")
                    ]
                    ws.cell(
                        row=start_row + b + 1, column=4
                    ).value = risk_data.portfolio_risk_impact_count[q][risk_cat][
                        (cat, "Medium")
                    ]
                    ws.cell(
                        row=start_row + b + 1, column=5
                    ).value = risk_data.portfolio_risk_impact_count[q][risk_cat][
                        (cat, "High")
                    ]
                    ws.cell(
                        row=start_row + b + 1, column=6
                    ).value = risk_data.portfolio_risk_count[q][risk_cat][cat]

                try:
                    start_row += b + 4
                except UnboundLocalError:
                    # couldn't debug error, but exception working.
                    pass

        ws.cell(row=start_row, column=2).value = "Risk Type"
        ws.cell(row=start_row, column=3).value = "Low"
        ws.cell(row=start_row, column=4).value = "Medium"
        ws.cell(row=start_row, column=5).value = "High"
        ws.cell(row=start_row, column=6).value = "N/A"
        ws.cell(row=start_row, column=7).value = "None"
        ws.cell(row=start_row, column=8).value = "Total"
        for i, no in enumerate(risk_data.portfolio_type_impact_count[q].keys()):
            ws.cell(row=start_row + i + 1, column=2).value = str(no)
            ws.cell(
                row=start_row + i + 1, column=3
            ).value = risk_data.portfolio_type_impact_count[q][no][("Low")]
            ws.cell(
                row=start_row + i + 1, column=4
            ).value = risk_data.portfolio_type_impact_count[q][no][("Medium")]
            ws.cell(
                row=start_row + i + 1, column=5
            ).value = risk_data.portfolio_type_impact_count[q][no][("High")]
            ws.cell(
                row=start_row + i + 1, column=6
            ).value = risk_data.portfolio_type_impact_count[q][no][("N/A")]
            ws.cell(
                row=start_row + i + 1, column=7
            ).value = risk_data.portfolio_type_impact_count[q][no][(None)]
            ws.cell(row=start_row + i + 1, column=8).value = sum(
                risk_data.portfolio_type_impact_count[q][no].values()
            )

    wb.remove(wb["Sheet"])

    return wb


def portfolio_risks_into_word_by_project(risk_data: RiskData) -> Document:
    doc = get_input_doc(
        risk_data.kwargs["root_path"] + risk_data.kwargs["word_portrait"]
    )
    latest_q = risk_data.quarters[0]
    other_q = risk_data.quarters[1]
    for p in risk_data.risk_dictionary[latest_q].keys():
        heading = str(p)
        intro = doc.add_heading(heading, 0)
        intro.alignment = 1
        intro.bold = True
        for port_risk_no in range(1, 9):
            doc.add_paragraph().add_run(RISK_NO_DICTIONARY[port_risk_no]).bold = True
            for k in PORTFOLIO_RISKS_WORD:
                try:
                    doc.add_paragraph().add_run(k).italic = True
                    text_one = str(
                        risk_data.portfolio_risk_dictionary[latest_q][p][port_risk_no][
                            k
                        ]
                    )
                    try:
                        text_two = str(
                            risk_data.portfolio_risk_dictionary[other_q][p][
                                port_risk_no
                            ][k]
                        )
                    except (KeyError, IndexError):  # index error relates to data_bridge
                        text_two = text_one
                except KeyError:
                    break

                compare_text_new_and_old(text_one, text_two, doc)

    return doc


def portfolio_risks_into_word_by_risk(risk_data: RiskData) -> Document:
    doc = get_input_doc(
        risk_data.kwargs["root_path"] + risk_data.kwargs["word_portrait"]
    )
    latest_q = risk_data.quarters[0]
    other_q = risk_data.quarters[1]
    for port_risk_no in range(1, 9):
        heading = str(RISK_NO_DICTIONARY[port_risk_no])
        intro = doc.add_heading(heading, 0)
        intro.alignment = 1
        intro.bold = True
        for p in risk_data.risk_dictionary[latest_q].keys():
            doc.add_paragraph().add_run(p).bold = True
            for k in PORTFOLIO_RISKS_WORD:
                try:
                    doc.add_paragraph().add_run(k).italic = True
                    text_one = str(
                        risk_data.portfolio_risk_dictionary[latest_q][p][port_risk_no][
                            k
                        ]
                    )
                    try:
                        text_two = str(
                            risk_data.portfolio_risk_dictionary[other_q][p][
                                port_risk_no
                            ][k]
                        )
                    except (KeyError, IndexError):  # index error relates to data_bridge
                        text_two = text_one
                except KeyError:
                    break

                compare_text_new_and_old(text_one, text_two, doc)

    return doc
