import datetime

from typing import List, Dict, Union

from datetime import date

from datamaps.api import project_data_from_master
import platform
from pathlib import Path

from docx import Document, table

from docx.shared import Pt, Cm, RGBColor, Inches

from openpyxl import load_workbook, Workbook

from openpyxl.formatting import Rule
from openpyxl.styles import Font, PatternFill, Border
from openpyxl.styles.differential import DifferentialStyle

from analysis_engine.data import (
    convert_bc_stage_text,
    plus_minus_days,
    concatenate_dates,
    convert_rag_text,
    rag_txt_list,
    black_text,
    fill_colour_list,
    make_file_friendly,
    wd_heading,
    key_contacts,
    dca_table,
    dca_narratives,
    open_word_doc,
    set_col_widths,
    make_columns_bold,
    change_text_size,
    Master,
)


def _cdg_platform_docs_dir() -> Path:
    #  Cross plaform file path handling
    if platform.system() == "Linux":
        return Path.home() / "Documents" / "data_bridge"
    if platform.system() == "Darwin":
        return Path.home() / "Documents" / "data_bridge"
    else:
        return Path.home() / "Documents" / "data_bridge"


cdg_root_path = _cdg_platform_docs_dir()


def cdg_get_master_data() -> List[
    Dict[str, Union[str, int, datetime.date, float]]
]:  # how specify a list of dictionaries?
    """Returns a list of dictionaries each containing quarter data"""
    master_data_list = [
        project_data_from_master(
            cdg_root_path / "core_data/cdg_master_4_2020.xlsx", 4, 2020
        ),
        project_data_from_master(
            cdg_root_path / "core_data/cdg_master_3_2020.xlsx", 3, 2020
        ),
    ]
    return master_data_list


def cdg_get_project_information() -> Dict[str, Union[str, int]]:
    """Returns dictionary containing all project meta data"""
    return project_data_from_master(
        cdg_root_path / "core_data/cdg_project_info.xlsx", 2, 2020
    )


def place_data_into_new_master_format(master_data: Dict):  # throw away
    wb = load_workbook(cdg_root_path / "core_data/CDG_portfolio_report.xlsx")
    ws = wb.active

    for i, p in enumerate(master_data.projects):
        ws.cell(row=3, column=i + 5).value = p
        for row_num in range(2, ws.max_row + 1):
            key = ws.cell(row=row_num, column=2).value
            try:
                ws.cell(row=row_num, column=i + 5).value = master_data.data[p][key]
            except KeyError:
                pass

    return wb


CDG_BASELINE_TYPES = {
    "Re-baseline this quarter": "quarter",
}
# CDG_GROUP_DICT = {"Corporate Finance": "CF", "Group Finance": "GF"}
CDG_DIR = ["CFPD", "GF", "Digital", "SCS"]
DFT_STAGE = ["pre-SOBC", "SOBC", "OBC", "FBC"]
BC_STAGE_DICT = {
    "Strategic Outline Case": "SOBC",
    "SOBC": "SOBC",
    "pre-Strategic Outline Case": "pre-SOBC",
    "pre-SOBC": "pre-SOBC",
    "Outline Business Case": "OBC",
    "OBC": "OBC",
    "Full Business Case": "FBC",
    "FBC": "FBC",
    # older returns that require cleaning
    "Pre - SOBC": "pre-SOBC",
    "Pre Strategic Outline Business Case": "pre_SOBC",
    None: None,
    "Other": "Other",
    "Other ": "Other",
    "To be confirmed": None,
    "To be confirmed ": None,
}
DCG_DATE = datetime.date(
    2021, 2, 22
)  # ipdc date. Python date format is Year, Month, day


def cdg_overall_dashboard(master: Master, wb: Workbook) -> Workbook:
    wb = load_workbook(wb)
    ws = wb.worksheets[0]

    for row_num in range(2, ws.max_row + 1):
        project_name = ws.cell(row=row_num, column=2).value
        if project_name in master.current_projects:
            """BC Stage"""
            bc_stage = master.master_data[0].data[project_name]["CDG approval point"]
            # ws.cell(row=row_num, column=4).value = convert_bc_stage_text(bc_stage)
            ws.cell(row=row_num, column=3).value = convert_bc_stage_text(bc_stage)
            try:
                bc_stage_lst_qrt = master.master_data[1].data[project_name][
                    "CDG approval point"
                ]
                if bc_stage != bc_stage_lst_qrt:
                    # ws.cell(row=row_num, column=4).font = Font(
                    #     name="Arial", size=10, color="00fc2525"
                    # )
                    ws.cell(row=row_num, column=3).font = Font(
                        name="Arial", size=10, color="00fc2525"
                    )
            except (KeyError, IndexError):
                pass

            """planning stage"""
            plan_stage = master.master_data[0].data[project_name]["Project stage"]
            # ws.cell(row=row_num, column=5).value = plan_stage
            ws.cell(row=row_num, column=4).value = plan_stage
            try:
                plan_stage_lst_qrt = master.master_data[1].data[project_name][
                    "Project stage"
                ]
                if plan_stage != plan_stage_lst_qrt:
                    # ws.cell(row=row_num, column=5).font = Font(
                    #     name="Arial", size=10, color="00fc2525"
                    # )
                    ws.cell(row=row_num, column=4).font = Font(
                        name="Arial", size=10, color="00fc2525"
                    )
            except (KeyError, IndexError):
                pass

            """Total WLC"""
            wlc_now = master.master_data[0].data[project_name]["Total Forecast"]
            # ws.cell(row=row_num, column=6).value = wlc_now
            ws.cell(row=row_num, column=5).value = wlc_now
            """WLC variance against lst quarter"""
            try:
                wlc_lst_quarter = master.master_data[1].data[project_name][
                    "Total Forecast"
                ]
                diff_lst_qrt = wlc_now - wlc_lst_quarter
                if float(diff_lst_qrt) > 0.49 or float(diff_lst_qrt) < -0.49:
                    # ws.cell(row=row_num, column=7).value = diff_lst_qrt
                    ws.cell(row=row_num, column=6).value = diff_lst_qrt
                else:
                    # ws.cell(row=row_num, column=7).value = "-"
                    ws.cell(row=row_num, column=6).value = "-"

                try:
                    percentage_change = ((wlc_now - wlc_lst_quarter) / wlc_now) * 100
                    if percentage_change > 5 or percentage_change < -5:
                        # ws.cell(row=row_num, column=7).font = Font(
                        #     name="Arial", size=10, color="00fc2525"
                        # )
                        ws.cell(row=row_num, column=6).font = Font(
                            name="Arial", size=10, color="00fc2525"
                        )
                except ZeroDivisionError:
                    pass

            except (KeyError, IndexError):
                ws.cell(row=row_num, column=6).value = "-"

            """WLC variance against baseline quarter"""
            bl = master.bl_index["quarter"][project_name][2]
            wlc_baseline = master.master_data[bl].data[project_name]["Total Forecast"]
            try:
                diff_bl = wlc_now - wlc_baseline
                if float(diff_bl) > 0.49 or float(diff_bl) < -0.49:
                    # ws.cell(row=row_num, column=8).value = diff_bl
                    ws.cell(row=row_num, column=7).value = diff_bl
                else:
                    # ws.cell(row=row_num, column=8).value = "-"
                    ws.cell(row=row_num, column=7).value = "-"
            except TypeError:  # exception is here as some projects e.g. Hs2 phase 2b have (real) written into historical totals
                pass

            try:
                percentage_change = ((wlc_now - wlc_baseline) / wlc_now) * 100
                if percentage_change > 5 or percentage_change < -5:
                    # ws.cell(row=row_num, column=8).font = Font(
                    #     name="Arial", size=10, color="00fc2525"
                    # )
                    ws.cell(row=row_num, column=7).font = Font(
                        name="Arial", size=10, color="00fc2525"
                    )

            except (
                ZeroDivisionError,
                TypeError,
            ):  # zerodivision error obvious, type error handling as above
                pass

            """vfm category now"""
            vfm_cat = master.master_data[0].data[project_name][
                "VfM Category single entry"
            ]
            # if (
            #     master.master_data[0].data[project_name]["VfM Category single entry"]
            #     is None
            # ):
            #     vfm_cat = (
            #         str(
            #             master.master_data[0].data[project_name][
            #                 "VfM Category lower range"
            #             ]
            #         )
            #         + " - "
            #         + str(
            #             master.master_data[0].data[project_name][
            #                 "VfM Category upper range"
            #             ]
            #         )
            #     )
            #     # ws.cell(row=row_num, column=10).value = vfm_cat
            #     ws.cell(row=row_num, column=8).value = vfm_cat
            #
            # else:
            #     vfm_cat = master.master_data[0].data[project_name][
            #         "VfM Category single entry"
            #     ]
            #     # ws.cell(row=row_num, column=10).value = vfm_cat
            ws.cell(row=row_num, column=8).value = vfm_cat

            """vfm category baseline"""
            bl_i = master.bl_index["quarter"][project_name][2]
            vfm_cat_baseline = master.master_data[bl_i].data[project_name][
                "VfM Category single entry"
            ]
            # try:
            #     if (
            #         master.master_data[bl_i].data[project_name][
            #             "VfM Category single entry"
            #         ]
            #         is None
            #     ):
            #         vfm_cat_baseline = (
            #             str(
            #                 master.master_data[bl_i].data[project_name][
            #                     "VfM Category lower range"
            #                 ]
            #             )
            #             + " - "
            #             + str(
            #                 master.master_data[bl_i].data[project_name][
            #                     "VfM Category upper range"
            #                 ]
            #             )
            #         )
            #         # ws.cell(row=row_num, column=11).value = vfm_cat_baseline
            #     else:
            #         vfm_cat_baseline = master.master_data[bl_i].data[project_name][
            #             "VfM Category single entry"
            #         ]
            #         # ws.cell(row=row_num, column=11).value = vfm_cat_baseline

            # except KeyError:
            #     try:
            #         vfm_cat_baseline = master.master_data[bl_i].data[project_name][
            #             "VfM Category single entry"
            #         ]
            #         # ws.cell(row=row_num, column=11).value = vfm_cat_baseline
            #     except KeyError:
            #         vfm_cat_baseline = master.master_data[bl_i].data[project_name][
            #             "VfM Category"
            #         ]
            #         # ws.cell(row=row_num, column=11).value = vfm_cat_baseline

            if vfm_cat != vfm_cat_baseline:
                if vfm_cat_baseline is None:
                    pass
                else:
                    ws.cell(row=row_num, column=8).font = Font(
                        name="Arial", size=8, color="00fc2525"
                    )

            current = master.master_data[0].data[project_name]["Project End Date"]
            try:
                last_quarter = master.master_data[1].data[project_name][
                    "Full Operations"
                ]
            except IndexError:
                pass
            bl = master.master_data[bl_i].data[project_name]["Project End Date"]
            #
            # abb = master.abbreviations[project_name]["abb"]
            # current = get_milestone_date(
            #     abb, milestones.milestone_dict, "current", " Full Operations"
            # )
            # last_quarter = get_milestone_date(
            #     abb, milestones.milestone_dict, "last", " Full Operations"
            # )
            # bl = get_milestone_date(
            #     abb, milestones.milestone_dict, "bl_one", " Full Operations"
            # )
            ws.cell(row=row_num, column=9).value = current
            if current is not None and current < DCG_DATE:
                ws.cell(row=row_num, column=9).value = "Completed"
            try:
                last_change = (current - last_quarter).days
                if last_change == 0:
                    ws.cell(row=row_num, column=10).value = "-"
                else:
                    ws.cell(row=row_num, column=10).value = plus_minus_days(last_change)
                if last_change is not None and last_change > 46:
                    ws.cell(row=row_num, column=10).font = Font(
                        name="Arial", size=10, color="00fc2525"
                    )
            except (TypeError, UnboundLocalError):
                pass
            try:
                bl_change = (current - bl).days
                if bl_change == 0:
                    ws.cell(row=row_num, column=11).value = "-"
                else:
                    ws.cell(row=row_num, column=11).value = plus_minus_days(bl_change)
                if bl_change is not None and bl_change > 85:
                    ws.cell(row=row_num, column=11).font = Font(
                        name="Arial", size=10, color="00fc2525"
                    )
            except TypeError:
                pass

            # last at/next at cdg information  removed
            try:
                ws.cell(row=row_num, column=12).value = concatenate_dates(
                    master.master_data[0].data[project_name]["Last date at CDG"],
                    DCG_DATE,
                )
                ws.cell(row=row_num, column=13).value = concatenate_dates(
                    master.master_data[0].data[project_name]["Next date at CDG"],
                    DCG_DATE,
                )
            except (KeyError, TypeError):
                print(
                    project_name
                    + " last at / next at ipdc data could not be calculated. Check data."
                )

            # """IPA DCA rating"""
            # ipa_dca = convert_rag_text(
            #     master.master_data[0].data[project_name]["GMPP - IPA DCA"]
            # )
            # ws.cell(row=row_num, column=15).value = ipa_dca
            # if ipa_dca == "None":
            #     ws.cell(row=row_num, column=15).value = ""

            """DCA rating - this quarter"""
            ws.cell(row=row_num, column=17).value = convert_rag_text(
                master.master_data[0].data[project_name]["Departmental DCA"]
            )
            """DCA rating - last qrt"""
            try:
                ws.cell(row=row_num, column=19).value = convert_rag_text(
                    master.master_data[1].data[project_name]["Departmental DCA"]
                )
            except (KeyError, IndexError):
                ws.cell(row=row_num, column=19).value = ""
            """DCA rating - 2 qrts ago"""
            try:
                ws.cell(row=row_num, column=20).value = convert_rag_text(
                    master.master_data[2].data[project_name]["Departmental DCA"]
                )
            except (KeyError, IndexError):
                ws.cell(row=row_num, column=20).value = ""
            """DCA rating - 3 qrts ago"""
            try:
                ws.cell(row=row_num, column=21).value = convert_rag_text(
                    master.master_data[3].data[project_name]["Departmental DCA"]
                )
            except (KeyError, IndexError):
                ws.cell(row=row_num, column=21).value = ""
            """DCA rating - baseline"""
            bl_i = master.bl_index["quarter"][project_name][2]
            ws.cell(row=row_num, column=23).value = convert_rag_text(
                master.master_data[bl_i].data[project_name]["Departmental DCA"]
            )

        """list of columns with conditional formatting"""
        list_columns = ["o", "q", "s", "t", "u", "w"]

        """same loop but the text is black. In addition these two loops go through the list_columns list above"""
        for column in list_columns:
            for i, dca in enumerate(rag_txt_list):
                text = black_text
                fill = fill_colour_list[i]
                dxf = DifferentialStyle(font=text, fill=fill)
                rule = Rule(
                    type="containsText", operator="containsText", text=dca, dxf=dxf
                )
                for_rule_formula = 'NOT(ISERROR(SEARCH("' + dca + '",' + column + "5)))"
                rule.formula = [for_rule_formula]
                ws.conditional_formatting.add(column + "5:" + column + "60", rule)

        for row_num in range(2, ws.max_row + 1):
            for col_num in range(5, ws.max_column + 1):
                if ws.cell(row=row_num, column=col_num).value == 0:
                    ws.cell(row=row_num, column=col_num).value = "-"

    return wb


# def convert_pdf_to_png():
#     pages = convert_from_path(root_path / "output/dandelion.pdf", 500)
#     for page in pages:
#         page.save(root_path / "output/dandelion.jpeg", "JPEG")


def cdg_compile_p_report(
    doc: Document,
    project_info: Dict[str, Union[str, int, date, float]],
    master: Master,
    project_name: str,
) -> Document:
    wd_heading(doc, project_info, project_name)
    key_contacts(doc, master, project_name)
    dca_table(doc, master, project_name)
    cdg_project_report_meta_data(doc, master, project_name)
    dca_narratives(doc, master, project_name)
    # costs = CostData(master, group=[project_name], baseline=["standard"])
    # benefits = BenefitsData(master, project_name)
    # milestones = MilestoneData(master, group=[project_name], baseline=["standard"])
    # project_report_meta_data(doc, costs, milestones, benefits, project_name)
    # change_word_doc_landscape(doc)
    # cost_profile = cost_profile_graph(costs, show="No")
    # put_matplotlib_fig_into_word(doc, cost_profile, transparent=False, size=8)
    # total_profile = total_costs_benefits_bar_chart(costs, benefits, show="No")
    # put_matplotlib_fig_into_word(doc, total_profile, transparent=False, size=8)
    # #  handling of no milestones within filtered period.
    # ab = master.abbreviations[project_name]["abb"]
    # try:
    #     # milestones.get_milestones()
    #     # milestones.get_chart_info()
    #     milestones.filter_chart_info(dates=["1/9/2020", "30/12/2022"])
    #     milestones_chart = milestone_chart(
    #         milestones,
    #         blue_line="ipdc_date",
    #         title=ab + " schedule (2021 - 22)",
    #         show="No",
    #     )
    #     put_matplotlib_fig_into_word(doc, milestones_chart, transparent=False, size=8)
    #     # print_out_project_milestones(doc, milestones, project_name)
    # except ValueError:  # extends the time period.
    #     milestones = MilestoneData(master, project_name)
    #     # milestones.get_milestones()
    #     # milestones.get_chart_info()
    #     milestones.filter_chart_info(dates=["1/9/2020", "30/12/2024"])
    #     milestones_chart = milestone_chart(
    #         milestones,
    #         blue_line="ipdc_date",
    #         title=ab + " schedule (2021 - 24)",
    #         show="No",
    #     )
    #     put_matplotlib_fig_into_word(doc, milestones_chart)
    # print_out_project_milestones(doc, milestones, project_name)
    # change_word_doc_portrait(doc)
    # project_scope_text(doc, master, project_name)
    return doc


def cdg_run_p_reports(master: Master, **kwargs) -> None:
    group = master.current_projects
    # group = get_group(master, str(master.current_quarter), kwargs)

    for p in group:
        print("Compiling summary for " + p)
        report_doc = open_word_doc(cdg_root_path / "input/summary_temp.docx")
        qrt = make_file_friendly(str(master.master_data[0].quarter))
        output = cdg_compile_p_report(
            report_doc, cdg_get_project_information(), master, p
        )
        abb = master.abbreviations[p]["abb"]
        output.save(
            cdg_root_path / "output/{}_report_{}.docx".format(abb, qrt)
        )  # add quarter here


def cdg_project_report_meta_data(
    doc: Document,
    master: Master,
    project_name: str,
):
    """Meta data table"""
    # doc.add_section(WD_SECTION_START.NEW_PAGE)
    # paragraph = doc.add_paragraph()
    # paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # paragraph.add_run("key project MI").bold = True

    """Costs meta data"""
    # this chuck is pretty messy because the data is messy
    run = doc.add_paragraph().add_run("Key meta data")
    font = run.font
    font.bold = True
    # font.underline = True
    t = doc.add_table(rows=1, cols=4)
    hdr_cells = t.rows[0].cells
    hdr_cells[0].text = "WLC:"
    try:
        hdr_cells[1].text = (
            "£"
            + str(round(master.master_data[0].data[project_name]["Total Forecast"]))
            + "m"
        )
    except TypeError:
        hdr_cells[1].text = "TBC"
    hdr_cells[2].text = "Business Case"
    hdr_cells[3].text = str(
        master.master_data[0].data[project_name]["CDG approval point"]
    )

    row_cells = t.add_row().cells
    row_cells[0].text = "Income:"
    row_cells[1].text = ""
    row_cells[2].text = "VFM:"
    row_cells[3].text = str(
        master.master_data[0].data[project_name]["VfM Category single entry"]
    )

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(t, column_widths)
    # make column keys bold
    make_columns_bold([t.columns[0], t.columns[2]])
    change_text_size([t.columns[0], t.columns[1], t.columns[2], t.columns[3]], 10)

    return doc


