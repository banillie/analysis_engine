from docx import Document
from docx.enum.section import WD_SECTION_START

from analysis_engine.segmentation import get_iter_list, get_group, get_correct_p_data
from analysis_engine.cleaning import convert_none_types
from analysis_engine.dictionaries import (
    DCA_KEYS,
    DCA_RATING_SCORES,
    STANDARDISE_COST_KEYS,
)
from analysis_engine.error_msgs import get_error_list, InputError


class DcaData:
    def __init__(self, master, **kwargs):
        self.master = master
        self.kwargs = kwargs
        self.report = kwargs["report"]
        self.quarters = self.master["quarter_list"]
        self.dca_dictionary = {}
        self.dca_changes = {}
        self.dca_count = {}
        self.get_dictionary()
        self.get_count()

    def get_dictionary(self) -> None:
        quarter_dict = {}
        if len(self.quarters) > 2:
            raise InputError(
                "Too many quarters entered. This analysis can have a maximum of two quarters. Program "
                "stopping. Please re-enter."
            )
        for tp in self.quarters:
            group = get_group(self.master, tp, **self.kwargs)
            type_dict = {}
            for conf_type in list(DCA_KEYS[self.report].keys()):  # confidence type
                dca_dict = {}
                for project_name in group:
                    p_data = get_correct_p_data(self.master, project_name, tp)
                    if p_data is None:
                        continue
                    dca_type = DCA_KEYS[self.report][conf_type]
                    if dca_type is None:
                        continue
                    colour = p_data[dca_type]
                    score = DCA_RATING_SCORES[p_data[dca_type]]
                    costs = convert_none_types(
                        p_data[STANDARDISE_COST_KEYS[self.report]["total"]]
                    )
                    dca_colour = [("DCA", colour)]
                    dca_score = [("DCA score", score)]
                    t = [("Type", dca_type)]
                    cost_amount = [("Costs", costs)]
                    quarter = [("Quarter", tp)]
                    abb = self.master["project_information"][project_name][
                        "Abbreviations"
                    ]
                    dca_dict[abb] = dict(
                        dca_colour + t + cost_amount + quarter + dca_score
                    )
                type_dict[conf_type] = dca_dict
            quarter_dict[tp] = type_dict

        self.dca_dictionary = quarter_dict

    def get_changes(self) -> None:
        """compiles dictionary of changes in dca ratings when provided with two quarter arguments"""
        current_quarter = self.master["quarter_list"][0]
        last_quarter = self.master["quarter_list"][1]
        c_dict = {}
        for conf_type in list(DCA_KEYS[self.report].keys()):  # confidence type
            lower_dict = {}
            for project_name in list(
                self.dca_dictionary[current_quarter][conf_type].keys()
            ):
                t = [("Type", conf_type)]
                try:
                    dca_one_colour = self.dca_dictionary[current_quarter][conf_type][
                        project_name
                    ]["DCA"]
                    dca_two_colour = self.dca_dictionary[last_quarter][conf_type][
                        project_name
                    ]["DCA"]
                    dca_one_score = self.dca_dictionary[current_quarter][conf_type][
                        project_name
                    ]["DCA score"]
                    dca_two_score = self.dca_dictionary[last_quarter][conf_type][
                        project_name
                    ]["DCA score"]
                    if dca_one_score == dca_two_score:
                        status = [("Status", "Same")]
                        change = [("Change", "Unchanged")]
                    if dca_one_score > dca_two_score:
                        status = [
                            (
                                "Status",
                                "Improved from "
                                + dca_two_colour
                                + " to "
                                + dca_one_colour,
                            )
                        ]
                        change = [("Change", "Up")]
                    if dca_one_score < dca_two_score:
                        status = [
                            (
                                "Status",
                                "Worsened from "
                                + dca_two_colour
                                + " to "
                                + dca_one_colour,
                            )
                        ]
                        change = [("Change", "Down")]
                except TypeError:  # This picks up None types
                    if dca_one_colour:  # if project not reporting dca previous quarter
                        status = [("Status", "entered at " + str(dca_one_colour))]
                        change = [("Change", "New entry")]
                    else:
                        status = [("Status", "Missing")]
                        change = [("Change", "Unknown")]
                except KeyError:  # This picks up projects not being present in the quarters being analysed.
                    status = [("Status", "entered at " + str(dca_one_colour))]
                    change = [("Change", "New entry")]

                lower_dict[project_name] = dict(t + status + change)

            c_dict[conf_type] = lower_dict
        self.dca_changes = c_dict

    def get_count(self) -> None:
        output_dict = {}
        error_list = []
        for quarter in self.dca_dictionary.keys():
            dca_dict = {}
            for i, dca_type in enumerate(list(self.dca_dictionary[quarter].keys())):
                clr = {}
                for x, colour in enumerate(list(DCA_RATING_SCORES.keys())):
                    count = 0
                    cost = 0
                    total = 0
                    cost_total = 0
                    for y, project in enumerate(
                        list(self.dca_dictionary[quarter][dca_type].keys())
                    ):
                        score = self.dca_dictionary[quarter][dca_type][project][
                            "DCA score"
                        ]
                        if score:
                            total += 1
                        try:
                            cost_total += self.dca_dictionary[quarter][dca_type][
                                project
                            ]["Costs"]
                        except TypeError:
                            error_list.append(
                                project
                                + " total costs for "
                                + str(quarter)
                                + " are in an incorrect data type and need changing"
                            )
                            pass
                        if (
                            self.dca_dictionary[quarter][dca_type][project]["DCA"]
                            == colour
                        ):
                            count += 1
                            try:
                                cost += self.dca_dictionary[quarter][dca_type][project][
                                    "Costs"
                                ]
                            except TypeError:  # error message above doesn't need repeating
                                pass

                    clr[colour] = {
                        "count": count,
                        "cost": cost,
                        # "ct": cost / cost_total,
                    }
                    clr["Total"] = {
                        "count": total,
                        "cost": cost_total,
                        # "ct": cost_total / cost_total,
                    }
                dca_dict[dca_type] = clr
            output_dict[quarter] = dca_dict

        error_list = get_error_list(error_list)
        for x in error_list:
            print(x)

        self.dca_count = output_dict


def dca_changes_into_word(dca_data: DcaData, document_path) -> Document:
    doc = Document(document_path)
    header = (
        "Showing changes between "
        + str(dca_data.quarters[0])
        + " and "
        + str(dca_data.quarters[1])
        + "."
    )
    top = doc.add_paragraph()
    top.add_run(header).bold = True

    for i, dca_type in enumerate(list(dca_data.dca_changes.keys())):
        if i != 0:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        else:
            pass
        title = dca_type + " " + "Confidence changes"
        top = doc.add_paragraph()
        top.add_run(title).bold = True

        doc.add_paragraph()
        sub_head = "Improvements"
        sub = doc.add_paragraph()
        sub.add_run(sub_head).bold = True
        count = 0
        for project_name in list(dca_data.dca_changes[dca_type].keys()):
            if dca_data.dca_changes[dca_type][project_name]["Change"] == "Up":
                doc.add_paragraph(
                    project_name
                    + " "
                    + dca_data.dca_changes[dca_type][project_name]["Status"]
                )
                count += 1
        total_line = str(count) + " project(s) in total improved"
        doc.add_paragraph(total_line)

        doc.add_paragraph()
        sub_head = "Decreases"
        sub = doc.add_paragraph()
        sub.add_run(sub_head).bold = True
        count = 0
        for project_name in list(dca_data.dca_changes[dca_type].keys()):
            if dca_data.dca_changes[dca_type][project_name]["Change"] == "Down":
                doc.add_paragraph(
                    project_name
                    + " "
                    + dca_data.dca_changes[dca_type][project_name]["Status"]
                )
                count += 1
        total_line = str(count) + " project(s) in total have decreased"
        doc.add_paragraph(total_line)

        doc.add_paragraph()
        sub_head = "Missing ratings"
        sub = doc.add_paragraph()
        sub.add_run(sub_head).bold = True
        count = 0
        for project_name in list(dca_data.dca_changes[dca_type].keys()):
            if dca_data.dca_changes[dca_type][project_name]["Change"] == "Unknown":
                doc.add_paragraph(
                    project_name
                    + " "
                    + dca_data.dca_changes[dca_type][project_name]["Status"]
                )
                count += 1
        total_line = str(count) + " project(s) in total are missing a rating"
        doc.add_paragraph(total_line)

        doc.add_paragraph()
        sub_head = "New Projects"
        sub = doc.add_paragraph()
        sub.add_run(sub_head).bold = True
        count = 0
        for project_name in list(dca_data.dca_changes[dca_type].keys()):
            if dca_data.dca_changes[dca_type][project_name]["Change"] == "New entry":
                doc.add_paragraph(
                    project_name
                    + " "
                    + dca_data.dca_changes[dca_type][project_name]["Status"]
                )
                count += 1
        total_line = str(count) + " new project(s)"
        doc.add_paragraph(total_line)

    return doc
