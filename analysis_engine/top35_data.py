import configparser
import datetime
import math
from collections import Counter, OrderedDict
from typing import List, Dict, Union
from datetime import date

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from openpyxl import load_workbook

from analysis_engine.data import (
    get_group,
    open_word_doc,
    make_file_friendly,
    compare_text_new_and_old,
    set_col_widths,
    make_columns_bold,
    get_iter_list,
    get_milestone_date,
    get_correct_p_data,
    COLOUR_DICT,
    dandelion_number_text,
    cal_group_angle,
    convert_rag_text,
    Master,
    logger,
    MilestoneData,
    milestone_info_handling,
    wd_heading, convert_date, get_project_info_data,
)
import platform
from pathlib import Path
from docx import Document, table
from dateutil import parser
from dateutil.parser import ParserError

# from docx.enum.section import WD_SECTION_START
from docx.shared import Cm, RGBColor, Pt


def _top_35_platform_docs_dir() -> Path:
    #  Cross plaform file path handling
    if platform.system() == "Linux":
        return Path.home() / "Documents" / "top_250"
    if platform.system() == "Darwin":
        return Path.home() / "Documents" / "top_250"
    else:
        return Path.home() / "Documents" / "top_250"


top35_root_path = _top_35_platform_docs_dir()


def top35_run_p_reports(master: Master, **kwargs) -> None:
    group = get_group(master, str(master.current_quarter), kwargs)
    for p in group:
        p_name = master.project_information[p]["Abbreviations"]
        print("Compiling summary for " + p_name)
        report_doc = open_word_doc(top35_root_path / "input/summary_temp.docx")
        output = compile_p_report(report_doc, master, p, **kwargs)
        output.save(
            top35_root_path / "output/{}_report.docx".format(p_name)
        )  # add quarter here


# def run_p_reports_single(master: Master, **kwargs) -> None:
#     group = get_group(master, str(master.current_quarter), kwargs)
#
#     report_doc = open_word_doc(top35_root_path / "input/summary_temp.docx")
#     for p in group:
#         print("Compiling summary for " + p)
#         qrt = make_file_friendly(str(master.master_data[0].quarter))
#         output = compile_p_report(report_doc, get_project_information(), master, p)
#         report_doc.add_section(WD_SECTION_START.NEW_PAGE)
#
#     output.save(
#         top35_root_path / "output/250_report.docx".format(p, qrt)
#     )  # add quarter here


# def run_pm_one_lines_single(master: Master, prog_info: Dict, **kwargs) -> None:
#     group = get_group(master, str(master.current_quarter), kwargs)
#
#     name_list = []
#     for p in group:
#         name_list.append((p, prog_info.data[p]["ID Number"]))
#
#     name_list.sort(key=lambda x: x[1])
#
#     report_doc = open_word_doc(top35_root_path / "input/summary_temp.docx")
#     for p in name_list:
#         p_master = master.master_data[0].data[p[0]]
#         p_name = prog_info.data[p[0]]["ID Number"]
#         output = deliverables(report_doc, p_master, p_name=p_name)
#         # output = pm_one_line(report_doc, p_master, prog_info.data[p[0]]["ID Number"])
#
#     output.save(top35_root_path / "output/deliverables.docx")  # add quarter here


def compile_p_report(
    doc: Document,
    master: Master,
    project_name: str,
    **kwargs,
) -> Document:
    # p_master = master.master_data[0].data[project_name]
    kwargs["group"] = [project_name]
    r_args = [doc, master, project_name]
    wd_heading(doc, **kwargs)
    key_contacts(*r_args)
    project_scope_text(*r_args)
    deliverables(*r_args)
    project_report_meta_data(*r_args)
    # doc.add_section(WD_SECTION_START.NEW_PAGE)
    dca_narratives(*r_args)
    # kwargs["group"] = [project_name]
    ms = MilestoneData(master, "ipdc_milestones", **kwargs)  # milestones
    print_out_project_milestones(doc, ms)
    cs = CentralSupportData(master, **kwargs)  # central support
    print_out_central_support(doc, cs)
    return doc


# def wd_heading(
#     doc: Document, project_info: Dict[str, Union[str, int]], project_name: str
# ) -> None:
#     """Function adds header to word doc"""
#     font = doc.styles["Normal"].font
#     font.name = "Arial"
#     font.size = Pt(12)
#
#     heading = str(project_info[project_name]["ID Number"])  # integrate into master
#     intro = doc.add_heading(str(heading), 0)
#     intro.alignment = 1
#     intro.bold = True


def key_contacts(doc: Document, master: Dict, project_name: str) -> None:
    """Function adds keys contact details"""
    p_master = master.master_data[0]["data"][project_name]
    sro_name = p_master["SRO NAME"]
    if sro_name is None:
        sro_name = "tbc"

    sro_email = p_master["SRO EMAIL"]
    if sro_email is None:
        sro_email = "email: tbc"

    # sro_phone = master.master_data[0].data[project_name]["SRO Phone No."]
    # if sro_phone == None:
    #     sro_phone = "phone number: tbc"

    # doc.add_paragraph(
    #     "SRO: " + str(sro_name) + ", " + str(sro_email) + ", " + str(sro_phone)
    # )
    doc.add_paragraph("SRO: " + str(sro_name) + ", " + str(sro_email))


def project_scope_text(doc: Document, master: Master, project_name: str) -> Document:
    doc.add_paragraph().add_run("Short project description").bold = True
    text_one = str(
        master.master_data[0]["data"][project_name]["SHORT PROJECT DESCRIPTION"]
    )
    try:
        text_two = str(
            master.master_data[1]["data"][project_name]["SHORT PROJECT DESCRIPTION"]
        )
    except KeyError:
        text_two = text_one
    compare_text_new_and_old(text_one, text_two, doc)
    return doc


def deliverables(doc: Document, master: Dict, project_name: str) -> Document:
    dels = [
        "TOP 3 PROJECT DELIVERABLES 1",  # deliverables
        "TOP 3 PROJECT DELIVERABLES 2",
        "TOP 3 PROJECT DELIVERABLES 3",
    ]

    doc.add_paragraph().add_run("Top 3 Deliverables").bold = True

    for i, d in enumerate(dels):
        text_one = master.master_data[0]["data"][project_name][d]
        try:
            text_two = master.master_data[1]["data"][project_name][d]
        except KeyError:
            text_two = text_one

        try:
            compare_text_new_and_old(text_one, text_two, doc)
        except AttributeError:  # this is necessary as string data contains NBSP and this removes them
            pass

    return doc


def make_text_red(cell: int, current, old) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if current != old:
                run.font.color.rgb = RGBColor(255, 0, 0)


def cell_colouring(word_table_cell: table.Table.cell, one, two) -> None:
    """Function that handles cell colouring for word documents"""

    try:
        if one != two:
            colour = parse_xml(r'<w:shd {} w:fill="cb1f00"/>'.format(nsdecls("w")))
            word_table_cell._tc.get_or_add_tcPr().append(colour)

    except TypeError:
        pass


def project_report_meta_data(
    doc: Document,
    master: Dict,
    project_name: str,
):
    p_master = master.master_data[0]["data"][project_name]
    try:
        p_master_last = master.master_data[1]["data"][project_name]
    except KeyError:
        p_master_last = p_master

    """Meta data table"""
    # doc.add_section(WD_SECTION_START.NEW_PAGE)
    # paragraph = doc.add_paragraph()
    # paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # paragraph.add_run("Key Info").bold = True

    doc.add_paragraph().add_run("Key Info").bold = True

    if p_master["WLC COMMENTS"] is not None:
        # doc.add_paragraph()
        doc.add_paragraph().add_run(
            "* Total costs comment. " + str(p_master["WLC COMMENTS"])
        )

    """Costs meta data"""
    t = doc.add_table(rows=1, cols=4)
    hdr_cells = t.rows[0].cells
    hdr_cells[0].text = "ON SCHEDULE:"
    on_time = str(p_master["PROJECT DEL TO CURRENT TIMINGS ?"])
    on_time_old = str(p_master_last["PROJECT DEL TO CURRENT TIMINGS ?"])
    hdr_cells[1].text = on_time
    # cell_colouring(hdr_cells[1], on_time, on_time_old)
    make_text_red(hdr_cells[1], on_time, on_time_old)
    hdr_cells[2].text = "ON GMPP:"
    on_gmpp = str(p_master["GMPP ID: IS THIS PROJECT ON GMPP"])
    on_gmpp_last = str(p_master_last["GMPP ID: IS THIS PROJECT ON GMPP"])
    hdr_cells[3].text = on_gmpp
    make_text_red(hdr_cells[3], on_gmpp, on_gmpp_last)
    # cell_colouring(hdr_cells[3], on_gmpp, on_gmpp_last)
    row_cells = t.add_row().cells
    row_cells[0].text = "ON BUDGET:"
    on_budget = str(p_master["PROJECT ON BUDGET?"])
    on_budget_last = str(p_master_last["PROJECT ON BUDGET?"])
    row_cells[1].text = on_budget
    make_text_red(row_cells[1], on_budget, on_budget_last)
    # cell_colouring(row_cells[1], on_budget, on_budget_last)
    row_cells[2].text = "TOTAL COST:"
    if p_master["WLC NON GOV"] is None or p_master["WLC NON GOV"] == 0:
        total = dandelion_number_text(p_master["WLC TOTAL"])
        total_last = dandelion_number_text(p_master_last["WLC TOTAL"])
    else:
        total = dandelion_number_text(p_master["WLC TOTAL"] + p_master["WLC NON GOV"])
        total_last = dandelion_number_text(p_master_last["WLC TOTAL"] + p_master_last["WLC NON GOV"])
    row_cells[3].text = total
    make_text_red(row_cells[3], total, total_last)
    # row_cells = t.add_row().cells
    # row_cells[0].text = "SRO CLEARANCE DATE:"
    # row_cells[1].text = str(p_master["DATE CLEARED BY SRO"])
    # row_cells[2].text = "MOST RECENT PERM SEC REVIEW:"
    # row_cells[3].text = str(p_master["DATE OF MOST RECENT PERM SEC REVIEW"])

    # set column width
    column_widths = (Cm(4), Cm(3), Cm(4), Cm(3))
    set_col_widths(t, column_widths)
    # make column keys bold
    make_columns_bold([t.columns[0], t.columns[2]])
    # change_text_size([t.columns[0], t.columns[1], t.columns[2], t.columns[3]], 10)
    # make_text_red([t.columns[1], t.columns[3]])

    return doc


def dca_narratives(doc: Document, master: Dict, project_name: str) -> None:
    doc.add_paragraph()
    # p = doc.add_paragraph()
    # text = "*Red text highlights changes in narratives from last quarter"
    # p.add_run(text).font.color.rgb = RGBColor(255, 0, 0)

    narrative_keys_list = [
        "SRO YOUR ON OFF TRACK ASSESSMENT:",
        "PRIMARY CONCERN & MITIGATING ACTIONS",
        "SHORT UPDATE FOR PM NOTE",
    ]

    headings_list = [
        "SRO assessment narrative",
        "SRO single biggest concern",
        "Short update for PM",
    ]

    for x in range(len(headings_list)):
        text_one = str(master.master_data[0]["data"][project_name][narrative_keys_list[x]])
        try:
            text_two = str(master.master_data[1]["data"][project_name][narrative_keys_list[x]])
        except KeyError:
            text_two = text_one
        doc.add_paragraph().add_run(str(headings_list[x])).bold = True
        compare_text_new_and_old(text_one, text_two, doc)


# def pm_one_line(doc: Document, p_master: Dict, p_name: str) -> None:
#     doc.add_paragraph()
#     # p = doc.add_paragraph()
#     # text = "*Red text highlights changes in narratives from last quarter"
#     # p.add_run(text).font.color.rgb = RGBColor(255, 0, 0)
#
#     narrative_keys_list = [
#         "SHORT UPDATE FOR PM NOTE",
#     ]
#
#     headings_list = [
#         "Short update for PM",
#     ]
#
#     for x in range(len(headings_list)):
#         try:  # overall try statement relates to data_bridge
#             text_one = str(
#                 p_master[narrative_keys_list[x]]
#             )
#             try:
#                 text_two = str(
#                     p_master[narrative_keys_list[x]]
#                 )
#             except (KeyError, IndexError):  # index error relates to data_bridge
#                 text_two = text_one
#         except KeyError:
#             break
#
#         doc.add_paragraph().add_run(str(p_name)).bold = True
#         compare_text_new_and_old(text_one, text_two, doc)
#
#     return doc
#


# def milestone_info_handling(output_list: list, t_list: list, kwargs) -> list:
#     """helper function for handling and cleaning up milestone date generated
#     via MilestoneDate class. Removes none type milestone names and non date
#     string values"""
#     if t_list[1][1] is None or t_list[1][1] == "Project - Business Case End Date":
#         pass
#     else:
#         try:
#             d = parser.parse(t_list[2][1], dayfirst=True)
#             t_list[3] = ("Date", d.date())
#             return output_list.append(t_list)
#         except TypeError:  # None Types
#             if "type" in kwargs:
#                 if kwargs["type"] == "central support":
#                     if isinstance(t_list[2][1], None):
#                         return output_list.append(t_list)
#         except ParserError:  # Non-date strings
#             pass
#

def cs_info_handling(output_list: list, t_list: list) -> list:
    """helper function for handling and cleaning up milestone date generated
    via MilestoneDate class. Removes none type milestone names and non date
    string values"""
    if t_list[1][1] is None:
        pass
    else:
        return output_list.append(t_list)


class CentralSupportData:
    def __init__(
        self,
        master: Master,
        # baseline_type: str = "ipdc_milestones",
        **kwargs,
    ):
        self.master = master
        self.group = []
        self.iter_list = []  # iteration list
        self.kwargs = kwargs
        self.baseline_type = None
        self.cs_dict = {}
        self.sorted_milestone_dict = {}
        self.max_date = None
        self.min_date = None
        self.schedule_change = {}
        self.schedule_key_last = None
        self.schedule_key_baseline = None
        self.get_cs_milestones()
        self.get_cs_chart_info()
        # self.calculate_schedule_changes()

    def get_cs_milestones(self) -> None:
        """
        Creates project milestone dictionaries for current, last_quarter, and
        baselines when provided with group and baseline type.
        """
        sp_dict = {}
        self.iter_list = get_iter_list(self.kwargs, self.master)
        for tp in self.iter_list:  # tp time period
            self.kwargs["tp"] = tp
            lower_dict = {}
            raw_list = []
            self.group = get_group(self.master, tp, self.kwargs)
            for project_name in self.group:
                project_list = []
                p_data = get_correct_p_data(
                    self.kwargs, self.master, self.baseline_type, project_name, tp
                )
                if p_data is None:
                    continue
                # i loops below removes None Milestone names and rejects non-datetime date values.
                p = self.master.abbreviations[project_name]["abb"]
                report = "Top 250"
                category = "Central Resource"
                for i in range(1, 17):
                    # these keys not present in all monthly masters.
                    try:
                        cs_response = p_data["R" + str(i) + " Central Response"]
                    except KeyError:
                        cs_response = None
                    try:
                        poc = p_data["R" + str(i) + " Point of Contact"]
                    except KeyError:
                        poc = None
                    # this not present in new entries
                    try:
                        secured = p_data["R" + str(i) + " secured"]
                    except KeyError:
                        secured = None

                    t = [
                        ("Project", p),
                        ("Requirement", p_data["R" + str(i) + " name"]),
                        # ("Type", "Approval"),
                        ("Escalated", p_data["R" + str(i) + " escalated to"]),
                        ("Date", convert_date(p_data["R" + str(i) + " needed by"])),
                        ("Type", p_data["R" + str(i) + " type"]),
                        ("Central Response", cs_response),
                        ("PoC", poc),
                        ("Secured", secured),
                        ("Report", report),
                        ("Cat", category),
                    ]
                    milestone_info_handling(project_list, t, **self.kwargs)

                # loop to stop keys names being the same. Done at project level.
                # not particularly concise code.
                upper_counter_list = []
                for entry in project_list:
                    upper_counter_list.append(entry[1][1])
                upper_count = Counter(upper_counter_list)
                lower_counter_list = []
                for entry in project_list:
                    if upper_count[entry[1][1]] > 1:
                        lower_counter_list.append(entry[1][1])
                        lower_count = Counter(lower_counter_list)
                        new_require_key = (
                            entry[1][1] + " (" + str(lower_count[entry[1][1]]) + ")"
                        )
                        entry[1] = ("Requirement", new_require_key)
                        raw_list.append(entry)
                    else:
                        raw_list.append(entry)

            # puts the list in chronological order
            sorted_list = sorted(raw_list, key=lambda k: (k[3][1] is None, k[3][1]))

            for r in range(len(sorted_list)):
                lower_dict["Requirement " + str(r)] = dict(sorted_list[r])

            sp_dict[tp] = lower_dict
        self.cs_dict = sp_dict

    def get_cs_chart_info(self) -> None:
        """returns data lists for matplotlib chart"""
        # Note this code could refactored so that it collects all milestones
        # reported across current, last and baseline. At the moment it only
        # uses milestones that are present in the current quarter.

        output_dict = {}
        for i in self.cs_dict:
            report = []
            category = []
            p_name = []
            key_names = []
            g_dates = []  # graph dates
            r_dates = []  # raw dates
            escalated = []
            type = []
            secured = []
            cr = []
            poc = []
            for v in self.cs_dict[self.iter_list[0]].values():
                p = None  # project
                mn = None  # milestone name
                d = None  # date
                for x in self.cs_dict[i].values():
                    if (
                        x["Project"] == v["Project"]
                        and x["Requirement"] == v["Requirement"]
                    ):
                        p = x["Project"]
                        mn = x["Requirement"]
                        # if len(self.group) == 1:
                        #     join = mn
                        # else:
                        #     join = p + ", " + mn
                        # if join not in key_names:  # stop duplicates
                        p_name.append(p)
                        key_names.append(mn)
                        d = x["Date"]
                        g_dates.append(d)
                        r_dates.append(d)
                        escalated.append(x["Escalated"])
                        type.append(x["Type"])
                        report.append(x["Report"])
                        category.append(x["Cat"])
                        try:
                            secured.append(x["Secured"])
                        except KeyError:
                            secured.append("NEW")
                        try:
                            poc.append(x["PoC"])
                        except KeyError:
                            poc.append("NEW")
                        try:
                            cr.append(x["Central Response"])
                        except KeyError:
                            cr.append("NEW")
                        break
                if p is None and mn is None and d is None:
                    p = v["Project"]
                    mn = v["Requirement"]
                    # if len(self.group) == 1:
                    #     join = mn
                    # else:
                    #     join = p + ", " + mn
                    # if join not in key_names:
                    p_name.append(p)
                    key_names.append(mn)
                    g_dates.append(v["Date"])
                    r_dates.append(None)
                    escalated.append(None)
                    type.append(None)
                    secured.append(None)
                    report.append(x["Report"])
                    category.append(x["Cat"])

            output_dict[i] = {
                "project": p_name,
                "names": key_names,
                "g_dates": g_dates,
                "r_dates": r_dates,
                "escalated": escalated,
                "type": type,
                "notes": cr,  # putting cr into notes for now
                "secured": secured,
                "poc": poc,
                "report": report,
                "cat": category,
            }

        self.sorted_milestone_dict = output_dict

    def filter_chart_info(self, **filter_kwargs):
        # bug handling required in the event that there are no milestones with the filter.
        # i.e. the filter returns no milestones.
        filtered_dict = {}
        if (
            "type" in filter_kwargs
            and "key" in filter_kwargs
            and "dates" in filter_kwargs
        ):
            start_date, end_date = zip(*filter_kwargs["dates"])
            start = parser.parse(start_date, dayfirst=True)
            end = parser.parse(end_date, dayfirst=True)
            for i, v in enumerate(self.cs_dict[self.iter_list[0]].values()):
                if v["Type"] in filter_kwargs["type"]:
                    if v["Milestone"] in filter_kwargs["keys"]:
                        if start.date() <= filter_kwargs["dates"] <= end.date():
                            filtered_dict["Milestone " + str(i)] = v
                            continue

        elif "type" in filter_kwargs and "key" in filter_kwargs:
            for i, v in enumerate(self.cs_dict[self.iter_list[0]].values()):
                if v["Type"] in filter_kwargs["type"]:
                    if v["Milestone"] in filter_kwargs["keys"]:
                        filtered_dict["Milestone " + str(i)] = v
                        continue

        elif "type" in filter_kwargs and "dates" in filter_kwargs:
            start_date, end_date = zip(filter_kwargs["dates"])
            start = parser.parse(start_date[0], dayfirst=True)
            end = parser.parse(end_date[0], dayfirst=True)
            for i, v in enumerate(self.cs_dict[self.iter_list[0]].values()):
                if v["Type"] in filter_kwargs["type"]:
                    if start.date() <= v["Date"] <= end.date():
                        filtered_dict["Milestone " + str(i)] = v
                        continue

        elif "key" in filter_kwargs and "dates" in filter_kwargs:
            start_date, end_date = zip(filter_kwargs["dates"])
            start = parser.parse(start_date[0], dayfirst=True)
            end = parser.parse(end_date[0], dayfirst=True)
            for i, v in enumerate(self.cs_dict[self.iter_list[0]].values()):
                if v["Milestone"] in filter_kwargs["key"]:
                    if start.date() <= v["Date"] <= end.date():
                        filtered_dict["Milestone " + str(i)] = v
                        continue

        elif "type" in filter_kwargs:
            for i, v in enumerate(self.cs_dict[self.iter_list[0]].values()):
                if v["Type"] in filter_kwargs["type"]:
                    filtered_dict["Milestone " + str(i)] = v
                    continue

        elif "key" in filter_kwargs:
            for i, v in enumerate(self.cs_dict[self.iter_list[0]].values()):
                if v["Milestone"] in filter_kwargs["key"]:
                    filtered_dict["Milestone " + str(i)] = v
                    continue

        elif "dates" in filter_kwargs:
            start_date, end_date = zip(filter_kwargs["dates"])
            start = parser.parse(start_date[0], dayfirst=True)
            end = parser.parse(end_date[0], dayfirst=True)
            for i, v in enumerate(self.cs_dict[self.iter_list[0]].values()):
                if start.date() <= v["Date"] <= end.date():
                    filtered_dict["Milestone " + str(i)] = v
                    continue

        output_dict = {}
        for dict in self.cs_dict.keys():
            if dict == self.iter_list[0]:
                output_dict[dict] = filtered_dict
            else:
                output_dict[dict] = self.cs_dict[dict]

        self.cs_dict = output_dict
        self.get_cs_chart_info()

    def calculate_schedule_changes(self) -> None:
        """calculates the changes in project schedules. If standard key for calculation
        not available it using the best next one available"""

        self.filter_chart_info(milestone_type=["Delivery", "Approval"])
        m_dict_keys = list(self.cs_dict.keys())

        def schedule_info(
            project_name: str,
            other_key_list: List[str],
            c_key_list: List[str],
            miles_dict: dict,
            dict_l_current: str,
            dict_l_other: str,
        ):
            output_dict = {}
            schedule_info = []
            for key in reversed(other_key_list):
                if key in c_key_list:
                    sop = get_milestone_date(
                        project_name, miles_dict, dict_l_other, " Start of Project"
                    )
                    if sop is None:
                        sop = get_milestone_date(
                            project_name, miles_dict, dict_l_current, other_key_list[0]
                        )
                        schedule_info.append(("start key", other_key_list[0]))
                    else:
                        schedule_info.append(("start key", " Start of Project"))
                    schedule_info.append(("start", sop))
                    schedule_info.append(("end key", key))
                    date = get_milestone_date(
                        project_name, miles_dict, dict_l_current, key
                    )
                    schedule_info.append(("end current date", date))
                    other_date = get_milestone_date(
                        project_name, miles_dict, dict_l_other, key
                    )
                    schedule_info.append(("end other date", other_date))
                    project_length = (other_date - sop).days
                    schedule_info.append(("project length", project_length))
                    change = (date - other_date).days
                    schedule_info.append(("change", change))
                    p_change = int((change / project_length) * 100)
                    schedule_info.append(("percent change", p_change))
                    output_dict[dict_l_other] = dict(schedule_info)
                    break

            return output_dict

        output_dict = {}
        for project_name in self.group:
            project_name = self.master.abbreviations[project_name]
            current_key_list = []
            last_key_list = []
            baseline_key_list = []
            for key in self.key_names:
                try:
                    p = key.split(",")[0]
                    milestone_key = key.split(",")[1]
                    if project_name == p:
                        if milestone_key != " Project - Business Case End Date":
                            current_key_list.append(milestone_key)
                except IndexError:
                    # patch of single project group. In this instance the project name
                    # is removed from the key_name via remove_project_name function as
                    # part of get chart info.
                    if len(self.group) == 1:
                        current_key_list.append(" " + key)
            for last_key in self.key_names_last:
                p = last_key.split(",")[0]
                milestone_key_last = last_key.split(",")[1]
                if project_name == p:
                    if milestone_key_last != " Project - Business Case End Date":
                        last_key_list.append(milestone_key_last)
            for baseline_key in self.key_names_baseline:
                p = baseline_key.split(",")[0]
                milestone_key_baseline = baseline_key.split(",")[1]
                if project_name == p:
                    if (
                        milestone_key_baseline
                        != " Project - Business Case End Date"
                        # and milestone_key_baseline != " Project End Date"
                    ):
                        baseline_key_list.append(milestone_key_baseline)

            b_dict = schedule_info(
                project_name,
                baseline_key_list,
                current_key_list,
                self.cs_dict,
                m_dict_keys[0],
                m_dict_keys[2],
            )
            l_dict = schedule_info(
                project_name,
                last_key_list,
                current_key_list,
                self.cs_dict,
                m_dict_keys[0],
                m_dict_keys[1],
            )
            lower_dict = {**b_dict, **l_dict}

            output_dict[project_name] = lower_dict

        self.schedule_change = output_dict


def print_out_project_milestones(
    doc: Document,
    milestones: MilestoneData,
) -> Document:
    # doc.add_section(WD_SECTION_START.NEW_PAGE)
    # table heading
    # ab = milestones.master.abbreviations[project_name]["abb"]
    doc.add_paragraph().add_run("Milestones").bold = True

    if not milestones.sorted_milestone_dict[milestones.iter_list[0]]["names"]:
        doc.add_paragraph().add_run("No milestones reported")
    else:
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Milestone"
        hdr_cells[1].text = "Date"
        hdr_cells[2].text = "Status"
        # hdr_cells[3].text = "Change from baseline"
        hdr_cells[3].text = "Notes"
        for i, m in enumerate(
            milestones.sorted_milestone_dict[milestones.iter_list[0]]["names"]
        ):
            row_cells = table.add_row().cells
            row_cells[0].text = m
            row_cells[1].text = milestones.sorted_milestone_dict[
                milestones.iter_list[0]
            ]["r_dates"][i].strftime("%d/%m/%Y")
            row_cells[2].text = milestones.sorted_milestone_dict[
                milestones.iter_list[0]
            ]["status"][i]
            try:
                row_cells[3].text = milestones.sorted_milestone_dict[
                    milestones.iter_list[0]
                ]["notes"][i]
                paragraph = row_cells[3].paragraphs[0]
                run = paragraph.runs
                font = run[0].font
                font.size = Pt(8)  # font size = 8
            except TypeError:
                pass

        table.style = "Table Grid"

        # column widths
        column_widths = (Cm(6), Cm(2.6), Cm(2.6), Cm(10))
        set_col_widths(table, column_widths)
        # make_columns_bold([table.columns[0], table.columns[3]])  # make keys bold
        # make_text_red([table.columns[1], table.columns[4]])  # make 'not reported red'

        # make_rows_bold(
        #     [table.rows[0]]
        # )  # makes top of table bold. Found function on stack overflow.
    return doc


def print_out_central_support(
    doc: Document,
    centrals: CentralSupportData,
) -> Document:

    doc.add_paragraph()
    doc.add_paragraph().add_run("Central govnt support requirements").bold = True

    if not centrals.sorted_milestone_dict[centrals.iter_list[0]]["names"]:
        doc.add_paragraph().add_run("No requirements reported")
    else:
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Requirement"
        hdr_cells[1].text = "Need by"
        hdr_cells[2].text = "Escalated to"
        hdr_cells[3].text = "Type"
        hdr_cells[4].text = "Type - other"
        hdr_cells[5].text = "Further details"

        for i, m in enumerate(centrals.sorted_milestone_dict[centrals.iter_list[0]]["names"]):
            row_cells = table.add_row().cells
            row_cells[0].text = m
            paragraph = row_cells[0].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size = Pt(8)  # font size = 8
            try:
                row_cells[1].text = centrals.sorted_milestone_dict[centrals.iter_list[0]][
                    "r_dates"
                ][i].strftime("%d/%m/%Y")
                paragraph = row_cells[1].paragraphs[0]
                run = paragraph.runs
                font = run[0].font
                font.size = Pt(9)
            except AttributeError:
                row_cells[1].text = "None"
            row_cells[2].text = str(
                centrals.sorted_milestone_dict[centrals.iter_list[0]]["escalated"][i]
            )
            paragraph = row_cells[2].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size = Pt(9)
            row_cells[3].text = str(
                centrals.sorted_milestone_dict[centrals.iter_list[0]]["type"][i]
            )
            paragraph = row_cells[3].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size = Pt(9)
            row_cells[4].text = str(
                centrals.sorted_milestone_dict[centrals.iter_list[0]]["notes"][i]
            )
            paragraph = row_cells[4].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size = Pt(8)  # font size = 8
            row_cells[5].text = str(
                centrals.sorted_milestone_dict[centrals.iter_list[0]]["secured"][i]
            )
            paragraph = row_cells[5].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size = Pt(9)  # font size = 8
            # try:
            #     row_cells[2].text = plus_minus_days(
            #         (
            #                 milestones.sorted_milestone_dict[milestones.iter_list[0]][
            #                     "r_dates"
            #                 ][i]
            #                 - milestones.sorted_milestone_dict[milestones.iter_list[1]][
            #                     "r_dates"
            #                 ][i]
            #         ).days
            #     )
            # except TypeError:
            #     row_cells[2].text = "Not reported"
            # try:
            #     row_cells[3].text = plus_minus_days(
            #         (
            #                 milestones.sorted_milestone_dict[milestones.iter_list[0]][
            #                     "r_dates"
            #                 ][i]
            #                 - milestones.sorted_milestone_dict[milestones.iter_list[2]][
            #                     "r_dates"
            #                 ][i]
            #         ).days
            #     )
            # except TypeError:
            #     row_cells[3].text = "Not reported"
            # try:
            #     row_cells[3].text = milestones.sorted_milestone_dict[
            #         milestones.iter_list[0]
            #     ]["notes"][i]
            #     paragraph = row_cells[3].paragraphs[0]
            #     run = paragraph.runs
            #     font = run[0].font
            #     font.size = Pt(8)  # font size = 8
            # except TypeError:
            #     pass

        table.style = "Table Grid"

        # column widths
        column_widths = (Cm(6), Cm(2), Cm(2), Cm(2), Cm(4), Cm(2))
        set_col_widths(table, column_widths)
        # column_widths = (Cm(6), Cm(2.6), Cm(2.6), Cm(10))
        # set_col_widths(table, column_widths)
    # make_columns_bold([table.columns[0], table.columns[3]])  # make keys bold
    # make_text_red([table.columns[1], table.columns[4]])  # make 'not reported red'

    # make_rows_bold(
    #     [table.rows[0]]
    # )  # makes top of table bold. Found function on stack overflow.
    return doc


# class DandelionData:
#     def __init__(self, master: Master, **kwargs):
#         self.master = master
#         self.kwargs = kwargs
#         self.baseline_type = "ipdc_costs"
#         self.group = []
#         self.iter_list = []
#         self.d_data = {}
#         self.get_data()
#
#     def get_data(self):
#         self.iter_list = get_iter_list(self.kwargs, self.master)
#         for (
#             tp
#         ) in self.iter_list:  # although tp is iterated only one can be handled for now.
#             #  for dandelion need groups of groups.
#             if "group" in self.kwargs:
#                 self.group = self.kwargs["group"]
#             elif "stage" in self.kwargs:
#                 self.group = self.kwargs["stage"]
#
#             if len(self.group) == 5:
#                 g_ang_l = [260, 310, 360, 50, 100]  # group angle list
#             if len(self.group) == 4:
#                 g_ang_l = [260, 326, 32, 100]
#             if len(self.group) == 3:
#                 g_ang_l = [280, 360, 80]
#             if len(self.group) == 2:
#                 g_ang_l = [290, 70]
#             if len(self.group) == 1:
#                 pass
#             g_d = {}  # group dictionary. first outer circle.
#             l_g_d = {}  # lower group dictionary
#
#             pf_wlc = get_dandelion_type_total(
#                 self.master, tp, self.group, self.kwargs
#             )  # portfolio wlc
#             if "pc" in self.kwargs:  # pc portfolio colour
#                 pf_colour = COLOUR_DICT[self.kwargs["pc"]]
#                 pf_colour_edge = COLOUR_DICT[self.kwargs["pc"]]
#             else:
#                 pf_colour = "#FFFFFF"
#                 pf_colour_edge = "grey"
#             pf_text = "Portfolio\n" + dandelion_number_text(
#                 pf_wlc
#             )  # option to specify pf name
#
#             ## center circle
#             g_d["portfolio"] = {
#                 "axis": (0, 0),
#                 "r": math.sqrt(pf_wlc),
#                 "colour": pf_colour,
#                 "text": pf_text,
#                 "fill": "solid",
#                 "ec": pf_colour_edge,
#                 "alignment": ("center", "center"),
#             }
#
#             ## first outer circle
#             for i, g in enumerate(self.group):
#                 self.kwargs["group"] = [g]
#                 g_wlc = get_dandelion_type_total(self.master, tp, g, self.kwargs)
#                 if len(self.group) > 1:
#                     y_axis = 0 + (
#                         (math.sqrt(pf_wlc) * 3.25) * math.sin(math.radians(g_ang_l[i]))
#                     )
#                     x_axis = 0 + (math.sqrt(pf_wlc) * 2.75) * math.cos(
#                         math.radians(g_ang_l[i])
#                     )
#                     g_text = g + "\n" + dandelion_number_text(g_wlc)  # group text
#                     if g_wlc == 0:
#                         g_wlc = pf_wlc / 20
#                     g_d[g] = {
#                         "axis": (y_axis, x_axis),
#                         "r": math.sqrt(g_wlc),
#                         "wlc": g_wlc,
#                         "colour": "#FFFFFF",
#                         "text": g_text,
#                         "fill": "dashed",
#                         "ec": "grey",
#                         "alignment": ("center", "center"),
#                         "angle": g_ang_l[i],
#                     }
#
#                 else:
#                     g_d = {}
#                     pf_wlc = g_wlc * 3
#                     g_text = g + "\n" + dandelion_number_text(g_wlc)  # group text
#                     if g_wlc == 0:
#                         g_wlc = 5
#                     g_d[g] = {
#                         "axis": (0, 0),
#                         "r": math.sqrt(g_wlc),
#                         "wlc": g_wlc,
#                         "colour": "#FFFFFF",
#                         "text": g_text,
#                         "fill": "dashed",
#                         "ec": "grey",
#                         "alignment": ("center", "center"),
#                     }
#
#             ## second outer circle
#             for i, g in enumerate(self.group):
#                 self.kwargs["group"] = [g]
#                 group = get_group(self.master, tp, self.kwargs)  # lower group
#                 p_list = []
#                 for p in group:
#                     self.kwargs["group"] = [p]
#                     p_value = get_dandelion_type_total(
#                         self.master, tp, p, self.kwargs
#                     )  # project wlc
#                     p_list.append((p_value, p))
#                 l_g_d[g] = list(reversed(sorted(p_list)))
#
#             for g in self.group:
#                 g_wlc = g_d[g]["wlc"]
#                 g_radius = g_d[g]["r"]
#                 g_y_axis = g_d[g]["axis"][0]  # group y axis
#                 g_x_axis = g_d[g]["axis"][1]  # group x axis
#                 try:
#                     p_values_list, p_list = zip(*l_g_d[g])
#                 except ValueError:  # handles no projects in l_g_d list
#                     continue
#                 if len(p_list) > 3 or len(self.group) == 1:
#                     ang_l = cal_group_angle(360, p_list, all=True)
#                 else:
#                     if len(p_list) == 1:
#                         ang_l = [g_d[g]["angle"]]
#                     if len(p_list) == 2:
#                         ang_l = [g_d[g]["angle"], g_d[g]["angle"] + 60]
#                     if len(p_list) == 3:
#                         ang_l = [
#                             g_d[g]["angle"],
#                             g_d[g]["angle"] + 60,
#                             g_d[g]["angle"] + 120,
#                         ]
#
#                 for i, p in enumerate(p_list):
#                     p_value = p_values_list[i]
#                     p_data = get_correct_p_data(
#                         self.kwargs, self.master, self.baseline_type, p, tp
#                     )
#                     # change confidence type here
#                     # SRO Schedule Confidence
#                     # Departmental DCA
#                     # SRO Benefits RAG
#                     # rag = p_data["Departmental DCA"]
#                     colour = COLOUR_DICT[convert_rag_text(None)]  # no rags for 250
#                     project_text = (
#                         self.master.abbreviations[p]["abb"]
#                         + "\n"
#                         + dandelion_number_text(p_value)
#                     )
#                     if p_value == 0:
#                         p_value = 200
#                     if p in self.master.dft_groups[tp]["GMPP"]:
#                         edge_colour = "#000000"  # edge of bubble
#                     else:
#                         edge_colour = colour
#
#                     # multi = math.sqrt(pf_wlc/g_wlc)  # multiplier
#                     # multi = (1 - (g_wlc / pf_wlc)) * 3
#                     try:
#                         if len(p_list) >= 14:
#                             multi = (pf_wlc / g_wlc) ** (1.0 / 2.0)  # square root
#                         else:
#                             multi = (pf_wlc / g_wlc) ** (1.0 / 3.0)  # cube root
#                         p_y_axis = g_y_axis + (g_radius * multi) * math.sin(
#                             math.radians(ang_l[i])
#                         )
#                         p_x_axis = g_x_axis + (g_radius * multi) * math.cos(
#                             math.radians(ang_l[i])
#                         )
#                     except ZeroDivisionError:
#                         p_y_axis = g_y_axis + 100 * math.sin(math.radians(ang_l[i]))
#                         p_x_axis = g_x_axis + 100 * math.cos(math.radians(ang_l[i]))
#
#                     if 185 >= ang_l[i] >= 175:
#                         text_angle = ("center", "top")
#                     if 5 >= ang_l[i] or 355 <= ang_l[i]:
#                         text_angle = ("center", "bottom")
#                     if 174 >= ang_l[i] >= 6:
#                         text_angle = ("left", "center")
#                     if 354 >= ang_l[i] >= 186:
#                         text_angle = ("right", "center")
#
#                     try:
#                         t_multi = (g_wlc / p_value) ** (1.0 / 4.0)
#                         # t_multi = (1 - (p_value/g_wlc)) * 2  # text multiplier
#                     except ZeroDivisionError:
#                         t_multi = 1
#                     yx_text_position = (
#                         p_y_axis
#                         + (math.sqrt(p_value) * t_multi)
#                         * math.sin(math.radians(ang_l[i])),
#                         p_x_axis
#                         + (math.sqrt(p_value) * t_multi)
#                         * math.cos(math.radians(ang_l[i])),
#                     )
#
#                     g_d[p] = {
#                         "axis": (p_y_axis, p_x_axis),
#                         "r": math.sqrt(p_value),
#                         "wlc": p_value,
#                         "colour": colour,
#                         "text": project_text,
#                         "fill": "solid",
#                         "ec": "grey",
#                         "alignment": text_angle,
#                         "tp": yx_text_position,
#                     }
#
#         self.d_data = g_d


# def get_dandelion_type_total(
#     master: Master, tp: str, g: str or List[str], kwargs
# ) -> int or str:  # Note no **kwargs as existing kwargs dict passed in
#     if "type" in kwargs:
#         if kwargs["type"] == "remaining":
#             cost = CostData(master, quarter=[tp], group=[g])  # group costs data
#             return cost.c_totals[tp]["prof"] + cost.c_totals[tp]["unprof"]
#         if kwargs["type"] == "spent":
#             cost = CostData(master, quarter=[tp], group=[g])  # group costs data
#             return cost.c_totals[tp]["spent"]
#         # if kwargs["type"] == "benefits":
#         #     benefits = BenefitsData(master, quarter=[tp], group=[g])
#         #     return benefits.b_totals[tp]["total"]
#
#     else:
#         cost = CostData(master, **kwargs)  # group costs data
#         return cost.c_totals[tp]["total"]
