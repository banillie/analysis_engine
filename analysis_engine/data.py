import configparser
import csv
import datetime
import difflib
import itertools
import json
import os
import pickle
import re
import typing
from collections import Counter
from typing import List, Dict, Union, Optional, Tuple, TextIO

# import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from datetime import date
from matplotlib.patches import RegularPolygon
from matplotlib.projections.polar import PolarAxes
from matplotlib.projections import register_projection
from matplotlib.spines import Spine
from matplotlib.transforms import Affine2D
from dateutil import parser
from dateutil.relativedelta import relativedelta
import numpy as np
from datamaps.api import project_data_from_master

# from dateutil.parser import ParserError
from docx import Document, table
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, Cm, RGBColor, Inches
from matplotlib import pyplot as plt
from matplotlib.patches import Wedge, Circle
from openpyxl import load_workbook, Workbook
from openpyxl.chart import BubbleChart, Reference, Series
from openpyxl.formatting import Rule
from openpyxl.styles import Font, PatternFill
from openpyxl.styles.differential import DifferentialStyle
from openpyxl.workbook import workbook
from textwrap import wrap
import logging

from analysis_engine.segmentation import get_iter_list, get_group, get_correct_p_data


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s: %(levelname)s - %(message)s",
    datefmt="%d-%b-%y %H:%M:%S",
)
logger = logging.getLogger(__name__)
# debug, info, warning, critical

CURRENT_LOG = []


class DuplicateFilter(logging.Filter):
    def filter(self, record):
        if record.msg not in CURRENT_LOG:
            CURRENT_LOG.append(record.msg)
            return True
        return False


logger.addFilter(DuplicateFilter())  # application of filter


class DataTypeError(Exception):
    pass


# def _cdg_platform_docs_dir() -> Path:
#     #  Cross plaform file path handling
#     if platform.system() == "Linux":
#         return Path.home() / "Documents" / "cdg"
#     if platform.system() == "Darwin":
#         return Path.home() / "Documents" / "cdg"
#     else:
#         return Path.home() / "Documents" / "cdg"


# cdg_root_path = _cdg_platform_docs_dir()


def calculate_arg_combinations(args_list: List):
    # code to help calculate the different combinations required for each
    # cli subcommands. not used in analysis engine.
    for x in range(0, len(args_list) + 1):
        for subset in itertools.combinations(args_list, x):
            print(subset)


# def get_master_data() -> List[
#     Dict[str, Union[str, int, datetime.date, float]]
# ]:  # how specify a list of dictionaries?
#     """Returns a list of dictionaries each containing quarter data"""
#     master_data_list = [
#         project_data_from_master(root_path / "core_data/master_4_2020.xlsx", 4, 2020),
#         project_data_from_master(root_path / "core_data/master_3_2020.xlsx", 3, 2020),
#         project_data_from_master(root_path / "core_data/master_2_2020.xlsx", 2, 2020),
#         project_data_from_master(root_path / "core_data/master_1_2020.xlsx", 1, 2020),
#         project_data_from_master(root_path / "core_data/master_4_2019.xlsx", 4, 2019),
#         project_data_from_master(root_path / "core_data/master_3_2019.xlsx", 3, 2019),
#         project_data_from_master(root_path / "core_data/master_2_2019.xlsx", 2, 2019),
#         project_data_from_master(root_path / "core_data/master_1_2019.xlsx", 1, 2019),
#         project_data_from_master(root_path / "core_data/master_4_2018.xlsx", 4, 2018),
#         project_data_from_master(root_path / "core_data/master_3_2018.xlsx", 3, 2018),
#         project_data_from_master(root_path / "core_data/master_2_2018.xlsx", 2, 2018),
#         project_data_from_master(root_path / "core_data/master_1_2018.xlsx", 1, 2018),
#         project_data_from_master(root_path / "core_data/master_4_2017.xlsx", 4, 2017),
#         project_data_from_master(root_path / "core_data/master_3_2017.xlsx", 3, 2017),
#         project_data_from_master(root_path / "core_data/master_2_2017.xlsx", 2, 2017),
#         project_data_from_master(root_path / "core_data/master_1_2017.xlsx", 1, 2017),
#         project_data_from_master(root_path / "core_data/master_4_2016.xlsx", 4, 2016),
#         project_data_from_master(root_path / "core_data/master_3_2016.xlsx", 3, 2016),
#     ]
#
#     return master_data_list

# def get_master_data(
#         confi_path: Path,
#         pi_path: Path,
#         get_data_function,
# ) -> List[Dict[str, Union[str, int, datetime.date, float]]]:
#     """Returns a list of dictionaries each containing quarter data"""
#     config = configparser.ConfigParser()
#     config.read(confi_path)
#     master_data_list = []
#     for key in config["MASTERS"]:
#         text = config["MASTERS"][key].split(", ")
#         year = text[2]
#         quarter = text[1]
#         m_path = str(pi_path) + text[0]
#         m = get_data_function(m_path, int(quarter), int(year))
#         master_data_list.append(m)
#
#     return list(reversed(master_data_list))
#


def get_master_data_json_test() -> Dict:
    # temp function for testing use of json data structure.
    return {
        "Q4 20/21": project_data_from_master(
            root_path / "core_data/master_4_2020.xlsx", 4, 2020
        )
    }


def get_master_data_file_paths() -> List[typing.TextIO]:
    file_list = [
        root_path / "core_data/master_2_2020.xlsx",
        root_path / "core_data/master_1_2020.xlsx",
        root_path / "core_data/master_4_2019.xlsx",
        root_path / "core_data/master_3_2019.xlsx",
        root_path / "core_data/master_2_2019.xlsx",
        root_path / "core_data/master_1_2019.xlsx",
        root_path / "core_data/master_4_2018.xlsx",
        root_path / "core_data/master_3_2018.xlsx",
        root_path / "core_data/master_2_2018.xlsx",
        root_path / "core_data/master_1_2018.xlsx",
        root_path / "core_data/master_4_2017.xlsx",
        root_path / "core_data/master_3_2017.xlsx",
        root_path / "core_data/master_2_2017.xlsx",
        root_path / "core_data/master_1_2017.xlsx",
        root_path / "core_data/master_4_2016.xlsx",
        root_path / "core_data/master_3_2016.xlsx",
    ]

    return file_list


def get_master_data_file_paths_fy_16_17() -> List[typing.TextIO]:
    file_list = [
        root_path / "core_data/master_4_2016.xlsx",
        root_path / "core_data/master_3_2016.xlsx",
    ]
    return file_list


def get_master_data_file_paths_fy_17_18() -> List[typing.TextIO]:
    file_list = [
        root_path / "core_data/master_4_2017.xlsx",
        root_path / "core_data/master_3_2017.xlsx",
        root_path / "core_data/master_2_2017.xlsx",
        root_path / "core_data/master_1_2017.xlsx",
    ]
    return file_list


def get_master_data_file_paths_fy_18_19() -> List[typing.TextIO]:
    file_list = [
        root_path / "core_data/master_4_2018.xlsx",
        root_path / "core_data/master_3_2018.xlsx",
        root_path / "core_data/master_2_2018.xlsx",
        root_path / "core_data/master_1_2018.xlsx",
    ]
    return file_list


def get_master_data_file_paths_fy_19_20() -> List[typing.TextIO]:
    file_list = [
        root_path / "core_data/master_4_2019.xlsx",
        root_path / "core_data/master_3_2019.xlsx",
        root_path / "core_data/master_2_2019.xlsx",
        root_path / "core_data/master_1_2019.xlsx",
    ]

    return file_list


def get_master_data_file_paths_fy_20_21() -> List[typing.TextIO]:
    file_list = [
        root_path / "core_data/master_2_2020.xlsx",
        root_path / "core_data/master_1_2020.xlsx",
    ]

    return file_list


def get_datamap_file_paths():
    pass


# def get_project_info_data(master_file: str) -> Dict:
#     # taken from datamaps project_data_from_master
#     wb = load_workbook(master_file)
#     ws = wb.active
#     for cell in ws["A"]:
#         # we don't want to clean None...
#         if cell.value is None:
#             continue
#         c = Cleanser(cell.value)
#         cell.value = c.clean()
#     p_dict = {}
#     for col in ws.iter_cols(min_col=2):
#         project_name = ""
#         o = OrderedDict()
#         for cell in col:
#             if cell.row == 1:
#                 project_name = cell.value
#                 p_dict[project_name] = o
#             else:
#                 val = ws.cell(row=cell.row, column=1).value
#                 if type(cell.value) == datetime:
#                     d_value = date(cell.value.year, cell.value.month, cell.value.day)
#                     p_dict[project_name][val] = d_value
#                 else:
#                     p_dict[project_name][val] = cell.value
#     # remove any "None" projects that were pulled from the master
#     try:
#         del p_dict[None]
#     except KeyError:
#         pass
#     return p_dict


# def get_project_information(
#         confi_path: Path,
#         pi_path: Path,
# ) -> Dict[str, Union[str, int]]:
#     """Returns dictionary containing all project meta data.
#     confi_path is ini file path.
#     pi_path is project_info path."""
#     config = configparser.ConfigParser()
#     config.read(confi_path)
#     path = str(pi_path) + config["PROJECT INFO"]["projects"]
#     return get_project_info_data(path)


def get_project_information_file_path() -> typing.TextIO:
    return root_path / "core_data/other/project_info.xlsx"


def get_gmpp_projects(project_info_dict) -> List[str]:
    # returns list of projects that are in gmpp
    output_list = []
    for p in project_info_dict.keys():
        if project_info_dict[p]["GMPP"]:
            output_list.append(p)

    return output_list


def get_error_list(seq: List[str]):
    seen = set()
    seen_add = seen.add
    return [x for x in seq if not (x in seen or seen_add(x))]


SALMON_FILL = PatternFill(
    start_color="FFFF8080", end_color="FFFF8080", fill_type="solid"
)
AMBER_FILL = PatternFill(start_color="FFBF00", end_color="FFBF00", fill_type="solid")
"""Store of different colours"""
ag_text = Font(color="00a5b700")  # text same colour as background
ag_fill = PatternFill(bgColor="00a5b700")
ar_text = Font(color="00f97b31")  # text same colour as background
ar_fill = PatternFill(bgColor="00f97b31")
red_text = Font(color="00fc2525")  # text same colour as background
red_fill = PatternFill(bgColor="00fc2525")
green_text = Font(color="0017960c")  # text same colour as background
green_fill = PatternFill(bgColor="0017960c")
amber_text = Font(color="00fce553")  # text same colour as background
amber_fill = PatternFill(bgColor="00fce553")

black_text = Font(color="00000000")
red_text = Font(color="FF0000")

darkish_grey_text = Font(color="002e4053")
darkish_grey_fill = PatternFill(bgColor="002e4053")
light_grey_text = Font(color="0085929e")
light_grey_fill = PatternFill(bgColor="0085929e")
greyblue_text = Font(color="85c1e9")
greyblue_fill = PatternFill(bgColor="85c1e9")

"""Conditional formatting, cell colouring and text colouring"""
# reference for column names when applying conditional fomatting
list_column_ltrs = [
    "a",
    "b",
    "c",
    "d",
    "e",
    "f",
    "g",
    "h",
    "i",
    "j",
    "k",
    "l",
    "m",
    "n",
    "o",
    "p",
    "q",
    "r",
    "s",
    "t",
    "u",
    "v",
    "w",
    "series_one",
    "series_two",
    "z",
]
# list of keys that have rag values for conditional formatting.
list_of_rag_keys = [
    "SRO Schedule Confidence",
    "Departmental DCA",
    "SRO Finance confidence",
    "SRO Benefits RAG",
    "GMPP - IPA DCA",
]

# lists of text and backfround colours and list of values for conditional formating rules.
rag_txt_colours = [ag_text, ar_text, red_text, green_text, amber_text]
rag_fill_colours = [ag_fill, ar_fill, red_fill, green_fill, amber_fill]
rag_txt_list_acroynms = ["A/G", "A/R", "R", "G", "A"]
rag_txt_list_full = ["Amber/Green", "Amber/Red", "Red", "Green", "Amber"]
gen_txt_colours = [darkish_grey_text, light_grey_text, greyblue_text]
gen_fill_colours = [darkish_grey_fill, light_grey_fill, greyblue_fill]
gen_txt_list = ["md", "pnr", "knc"]


# for project summary pages
SRO_CONF_TABLE_LIST = [
    "SRO DCA",
    "Finance DCA",
    "Benefits DCA",
    "Resourcing DCA",
    "Schedule DCA",
]
SRO_CONF_KEY_LIST = {
    "Departmental DCA": "SRO Overall DCA",
    "SRO Finance confidence": "SRO Cost DCA",
    "SRO Benefits RAG": "SRO Benefits DCA",
    "Overall Resource DCA - Now": "SRO Resource DCA",
    "SRO Schedule Confidence": "SRO Schedule DCA",
}


def get_ipdc_date(confi_path, board_str):
    config = configparser.ConfigParser()
    config.read(confi_path)
    return config["GLOBALS"][board_str]


LIST_OF_TITLES = [
    "ALL",
    "HE",
    "RAIL INFRASTRUCTURE",
    "RAIL FRANCHISING",
    "HS2",
    "HSMRPG",
    "AMIS (SARH2)",
    "ALL, NOT HS2,",
    "FBC Projects",
    "OBC Projects",
    "SOBC Projects",
]

#  list of different baseline types. hold at global level?
BASELINE_TYPES = {
    "Re-baseline this quarter": "quarter",
    "Re-baseline ALB/Programme milestones": "programme_milestones",
    "Re-baseline ALB/Programme cost": "programme_costs",
    "Re-baseline ALB/Programme benefits": "programme_benefits",
    "Re-baseline IPDC milestones": "ipdc_milestones",
    "Re-baseline IPDC cost": "ipdc_costs",
    "Re-baseline IPDC benefits": "ipdc_benefits",
    "Re-baseline HMT milestones": "hmt_milestones",
    "Re-baseline HMT cost": "hmt_costs",
    "Re-baseline HMT benefits": "hmt_benefits",
}
CDG_BASELINE_TYPES = {
    "BC approved by CDG Board in quarter": "quarter",
    "Costs re-baselined by CDG in quarter": "costs",
    "Schedule re-baselined by CDG in quarter": "schedule",
    "Benefits re-baselined by CDG in quarter": "benefits",
}
IPDC_BASELINE_TYPES = {
    "Re-baseline IPDC milestones": "ipdc_milestones",
    "Re-baseline IPDC cost": "ipdc_costs",
    "Re-baseline IPDC benefits": "ipdc_benefits",
}


# need both dft_group and dft_group_dictionary in the config

# DFT_GROUP_DICT = {
#     "AMIS": "Aviation",
#     "HSRG": "High Speed Rail Group",
#     "RIG": "Rail Infrastructure Group",
#     "RSS": "Rail Services",
#     "CDG": "Corporate Delivery Group",
#     "RPE": "Roads People and Environment",
# }
# CDG_DIR_DICT = {
#     "CFPD": "cfpd",
#     "GF": "gf",
#     "Digital": "D",
#     "SCS": "scs",
# }
YEAR_LIST = [
    # "16-17",
    # "17-18",
    # "18-19",
    # "19-20",
    # "20-21",
    "21-22",
    "22-23",
    "23-24",
    "24-25",
    "25-26",
    "26-27",
    "27-28",
    "28-29",
    "29-30",
    "30-31",
    "31-32",
    "32-33",
    "33-34",
    "34-35",
    "35-36",
    "36-37",
    "37-38",
    "38-39",
    "39-40",
]
COST_KEY_LIST = [
    " RDEL Forecast Total",
    " CDEL Forecast one off new costs",
    " Forecast Non-Gov",
]
COST_TYPE_KEY_LIST = [
    (
        "Pre-profile RDEL",
        "Pre-profile CDEL Forecast one off new costs",
        "Pre-profile Forecast Non-Gov",
    ),
    (
        "Total RDEL Forecast Total",
        "Total CDEL Forecast one off new costs",
        "Non-Gov Total Forecast",
    ),
    (
        "Unprofiled RDEL Forecast Total",
        "Unprofiled CDEL Forecast one off new costs",
        "Unprofiled Forecast Non-Gov",
    ),
]

SPENT_KEYS = [
    "Pre-profile RDEL",
    "Pre-profile CDEL Forecast one off new costs",
    "Pre-profile Forecast Non-Gov",
]
PROFILE_KEYS = [
    "Total RDEL Forecast Total",
    "Total CDEL Forecast one off new costs",
    "Non-Gov Total Forecast",
]
UNPROFILE_KEYS = [
    "Unprofiled RDEL Forecast Total",
    "Unprofiled CDEL Forecast one off new costs",
    "Unprofiled Forecast Non-Gov",
]

BEN_KEY_LIST = [
    "Pre-profile BEN Total",
    "Total BEN Forecast - Total Monetised Benefits",
    "Unprofiled Remainder BEN Forecast - Total Monetised Benefits",
]
BEN_TYPE_KEY_LIST = [
    (
        "Pre-profile BEN Forecast Gov Cashable",
        "Pre-profile BEN Forecast Gov Non-Cashable",
        "Pre-profile BEN Forecast - Economic (inc Private Partner)",
        "Pre-profile BEN Forecast - Disbenefit UK Economic",
    ),
    (
        "Total BEN Forecast - Gov. Cashable",
        "Total BEN Forecast - Gov. Non-Cashable",
        "Total BEN Forecast - Economic (inc Private Partner)",
        "Total BEN Forecast - Disbenefit UK Economic",
    ),
    (
        "Unprofiled Remainder BEN Forecast - Gov. Cashable",
        "Unprofiled Remainder BEN Forecast - Gov. Non-Cashable",
        "Unprofiled Remainder BEN Forecast - Economic (inc Private Partner)",
        "Unprofiled Remainder BEN Forecast - Disbenefit UK Economic",
    ),
]
# Matplotlib file formats
FILE_FORMATS = [
    "eps",
    "jpeg",
    "jpg",
    "pdf",
    "png",
    "ps",
    "raw",
    "rgba",
    "svg",
    "svgz",
    "tif",
    "tiff",
]
FIGURE_STYLE = {1: "half_horizontal", 2: "full_horizontal"}


def calculate_profiled(
    p: int or List[int], s: int or List[int], unpro: int or List[int]
) -> list:
    """small helper function to calculate the proper profiled amount. This is necessary as
    other wise 'profiled' would actually be the total figure.
    p = profiled list
    s = spent list
    unpro = unprofiled list"""
    if isinstance(p, list):
        f_profiled = []
        for y, amount in enumerate(p):
            t = amount - (s[y] + unpro[y])
            f_profiled.append(t)
        return f_profiled
    else:
        return p - (s + unpro)


# class JsonMaster:
#     def __init__(
#             self,
#             master_data: List[Dict[str, Union[str, int, datetime.date, float]]],
#             project_information: Dict[str, Union[str, int]],
#             all_groups,
#             **kwargs,
#     ) -> None:
#         self.master_data = master_data
#         self.project_information = project_information
#         self.all_groups = all_groups
#         self.all_projects = list(project_information.keys())
#         self.kwargs = kwargs
#         self.current_quarter = str(master_data[0].quarter)
#         self.current_projects = master_data[0].projects
#         self.abbreviations = {}
#         self.full_names = {}
#         self.bl_info = {}
#         self.bl_index = {}
#         self.dft_groups = {}
#         self.project_group = {}
#         self.project_stage = {}
#         self.pipeline_dict = {}
#         self.pipeline_list = []
#         self.quarter_list = []
#         self.get_quarter_list()
#         # self.get_baseline_data()
#         self.check_project_information()
#         self.get_project_abbreviations()
#         # self.check_baselines()
#         self.get_project_groups()
#         self.pipeline_projects_information()
#         self.get_current_tp()
#
#     def get_project_abbreviations(self) -> None:
#         """gets the abbreviations for all current projects.
#         held in the project info document"""
#         abb_dict = {}
#         fn_dict = {}
#         error_case = []
#         for p in self.all_projects:
#             abb = self.project_information[p]["Abbreviations"]
#             abb_dict[p] = {"abb": abb, "full name": p}
#             fn_dict[abb] = p
#             if abb is None:
#                 error_case.append(p)
#
#         if error_case:
#             for p in error_case:
#                 logger.critical("No abbreviation provided for " + p + ".")
#             raise ProjectNameError(
#                 "Abbreviations must be provided for all projects in project_info. Program stopping. Please amend"
#             )
#
#         self.abbreviations = abb_dict
#         self.full_names = fn_dict
#
#     def get_baseline_data(self) -> None:
#         """
#         Returns two dictionaries used to calculate baselines.
#         The current method is that each projects baseline list
#         is prefixed with current and last quarter index (including None
#         if necessary), as these need to be present at later stage
#         calculations. structure of output dict is
#         {bl_index: {bl_type: {proj_name: [baseline index list]}}}
#         """
#
#         # handles baselines across different datasets.
#         if "data_type" in self.kwargs:
#             if self.kwargs["data_type"] == "cdg":
#                 baseline_dict = CDG_BASELINE_TYPES
#         else:
#             baseline_dict = BASELINE_TYPES
#
#         baseline_info = {}
#         baseline_index = {}
#         for b_type in list(baseline_dict.keys()):
#             project_baseline_info = {}
#             project_baseline_index = {}
#             for name in self.current_projects:
#                 lower_list = []
#                 for i, master in list(enumerate(self.master_data)):
#                     quarter = str(master.quarter)
#                     if name in master.projects:
#                         approved_bc = master.data[name][b_type]
#                         if approved_bc == "Yes":
#                             lower_list.append((approved_bc, quarter, i))
#                     else:
#                         pass
#
#                 if name in self.master_data[1].projects:  # prefix for other bl data
#                     index_list = [0, 1]
#                 else:  # project not present last quarter so none
#                     index_list = [0, None]
#                 for x in lower_list:
#                     index_list.append(x[2])
#
#                 project_baseline_info[name] = list(lower_list)
#                 project_baseline_index[name] = list(index_list)
#
#             baseline_info[baseline_dict[b_type]] = project_baseline_info
#             baseline_index[baseline_dict[b_type]] = project_baseline_index
#
#         self.bl_info = baseline_info
#         self.bl_index = baseline_index
#
#     def check_project_information(self) -> None:
#         """Checks that project names in master are present/the same as in project info.
#         Stops the programme if not"""
#         error_cases = []
#         for p in self.current_projects:
#             if p not in self.all_projects:
#                 error_cases.append(p)
#
#         if error_cases:
#             for p in error_cases:
#                 logger.critical(p + " has not been found in the project_info document.")
#             try:
#                 m = str(self.master_data[0].month)
#             except KeyError:
#                 m = str(self.master_data[0].quarter)
#             raise ProjectNameError(
#                 "Project names in the "
#                 + m
#                 + " master and project_info must match. Program stopping. Please amend."
#             )
#         else:
#             logger.info("The latest master and project information match")
#
#     def check_baselines(self) -> None:  # check with team is required for IPDC.
#         """checks that projects have the correct baseline information. stops the
#         programme if baselines are missing"""
#
#         if "data_type" in self.kwargs:
#             if self.kwargs["data_type"] == "cdg":
#                 baseline_dict = CDG_BASELINE_TYPES
#         else:
#             baseline_dict = IPDC_BASELINE_TYPES
#
#         b_e_cases = []  # baseline error cases
#         b_v_e_cases = []  # baseline value error cases
#         for v in baseline_dict.values():
#             for p in self.current_projects:
#                 baselines = self.bl_index[v][p]
#                 if len(baselines) <= 2:
#                     b_e_cases.append(p)
#                     if v not in b_v_e_cases:
#                         b_v_e_cases.append(v)
#
#         if b_e_cases:
#             for i, b in enumerate(b_e_cases):
#                 logger.critical(
#                     b_e_cases[i]
#                     + " does not have a baseline point for "
#                     + b_v_e_cases[i]
#                     + " this could cause the programme to "
#                       "crash. Therefore the programme is stopping. "
#                       "Please amend the data for " + b_e_cases[i] + " so "
#                         " it has at least one baseline point for " + b_v_e_cases[i]
#                 )
#             raise ProjectNameError(  # should be Baselining Error or Initiation Error
#                 'Above issue(s) could cause a crash and require resolution. Program stopping'
#             )
#
#     def get_project_groups(self) -> None:
#         """gets the groups that projects are part of e.g. business case
#         stage or dft group"""
#
#         if "data_type" in self.kwargs:
#             if self.kwargs["data_type"] == "cdg":
#                 group_key = "Directorate"
#                 # group_dict = CDG_DIR_DICT
#                 approval = "Last Business Case (BC) achieved"
#             if self.kwargs["data_type"] == "top35":
#                 group_key = "Group"
#                 # group_dict = DFT_GROUP_DICT
#         else:
#             group_key = "Group"
#             # group_dict = DFT_GROUP_DICT
#             approval = "IPDC approval point"
#
#         raw_dict = {}
#         raw_list = []
#         group_list = []
#         stage_list = []
#         pn_e_cases = []  # project name error_cases
#         p_m_e_cases = []  # project master error cases
#         g_e_cases = []  # group error cases
#         for i, master in enumerate(self.master_data):
#             lower_dict = {}
#             for p in master.projects:
#                 try:
#                     dft_group = self.project_information[p][
#                         group_key
#                     ]
#                 except KeyError:
#                     pn_e_cases.append(p)
#                     p_m_e_cases.append(str(master.quarter))
#
#                 if dft_group is None or dft_group not in self.all_groups:
#                     g_e_cases.append(p)
#
#                 try:
#                     stage = BC_STAGE_DICT[master[p][approval]]
#                 except UnboundLocalError:  # top35 does not collect stage
#                     stage = "None"
#                 raw_list.append(("group", dft_group))
#                 raw_list.append(("stage", stage))
#                 lower_dict[p] = dict(raw_list)
#                 group_list.append(dft_group)
#                 stage_list.append(stage)
#
#             if pn_e_cases:
#                 for i, e in enumerate(pn_e_cases):
#                     logger.critical(
#                         f'Project name {pn_e_cases[i]} in master {p_m_e_cases[i]} not in project information '
#                         f'document. Make sure project names are consistent.'
#                     )
#                 raise ProjectNameError(
#                     'Above issue(s) could cause a crash and require resolution. Program stopping'
#                 )
#
#             if g_e_cases:
#                 for i in g_e_cases:
#                     logger.critical(
#                         str(i)
#                         + " does not have a recognised Group value in the project information document."
#                     )
#                 raise ProjectGroupError(
#                     'Above issue(s) could cause a crash and require resolution. Program stopping'
#                 )
#
#
#             try:
#                 raw_dict[str(master.month) + ", " + str(master.year)] = lower_dict
#             except KeyError:
#                 raw_dict[str(master.quarter)] = lower_dict
#
#         group_list = list(set(group_list))
#         stage_list = list(set(stage_list))
#
#         group_dict = {}
#         # for i, quarter in enumerate(list(raw_dict.keys())[:2]):  # just latest two quarters
#         for i, quarter in enumerate(list(raw_dict.keys())):
#             lower_g_dict = {}
#             for group_type in group_list:
#                 g_list = []
#                 for p in raw_dict[quarter].keys():
#                     p_group = raw_dict[quarter][p]["group"]
#                     if p_group == group_type:
#                         g_list.append(p)
#                 lower_g_dict[group_type] = g_list
#
#             gmpp_list = []
#             for p in self.master_data[i].projects:
#                 try:
#                     gmpp = self.project_information[p]["GMPP"]
#                 except KeyError:  # project name check happening in other places.
#                     gmpp = None
#                 if gmpp is not None:
#                     gmpp_list.append(p)
#                 lower_g_dict["GMPP"] = gmpp_list
#
#             group_dict[quarter] = lower_g_dict
#
#         stage_dict = {}
#         for quarter in list(raw_dict.keys())[:2]:  # just latest two quarters
#             lower_s_dict = {}
#             for stage_type in stage_list:
#                 s_list = []
#                 for p in raw_dict[quarter].keys():
#                     p_stage = raw_dict[quarter][p]["stage"]
#                     if p_stage == stage_type:
#                         s_list.append(p)
#                 if stage_type is None:
#                     if s_list:
#                         if "data_type" in self.kwargs:
#                             if self.kwargs["data_type"] == "cdg":
#                                 continue  # not actively using stages for cdg data yet so can pass
#                         if quarter == self.current_quarter:
#                             for x in s_list:
#                                 logger.critical(str(x) + " has no IPDC stage date")
#                                 raise ProjectStageError(
#                                     "Programme stopping as this could cause incomplete analysis"
#                                 )
#                         else:
#                             for x in s_list:
#                                 logger.warning(
#                                     "In "
#                                     + str(quarter)
#                                     + " master "
#                                     + str(x)
#                                     + " IPDC stage data is currently None. Please amend."
#                                 )
#                 lower_s_dict[stage_type] = s_list
#             stage_dict[quarter] = lower_s_dict
#
#         self.dft_groups = group_dict
#         self.project_stage = stage_dict
#
#     def get_quarter_list(self) -> None:
#         output_list = []
#         for master in self.master_data:
#             try:
#                 output_list.append(str(master.month) + ", " + str(master.year))
#             except KeyError:
#                 output_list.append(str(master.quarter))
#         self.quarter_list = output_list
#
#     def pipeline_projects_information(self) -> None:
#         pipeline_dict = {}
#         pipeline_list = []
#         total_wlc = 0
#         for p in self.all_projects:
#             if self.project_information[p]["Pipeline"] is not None:
#                 wlc = convert_none_types(self.project_information[p]["WLC"])
#                 pipeline_dict[p] = {
#                     "wlc": convert_none_types(self.project_information[p]["WLC"])
#                 }
#                 pipeline_list.append(p)
#                 total_wlc += wlc
#         pipeline_dict["pipeline"] = {"wlc": total_wlc}
#
#         self.pipeline_dict = pipeline_dict
#         self.pipeline_list = pipeline_list
#
#     def get_current_tp(self):
#         try:
#             self.current_quarter = (
#                     str(self.master_data[0].month) + ", " + str(self.master_data[0].year)
#             )
#         except KeyError:
#             self.current_quarter = self.master_data[0].quarter
#

# class Master:
#     def __init__(self, json_master: JsonMaster) -> None:
#         self.master_data = json_master["master_data"]
#         self.project_information = json_master["project_information"]
#         self.kwargs = json_master["kwargs"]
#         self.current_quarter = json_master["current_quarter"]
#         self.current_projects = json_master["current_projects"]
#         self.abbreviations = json_master["abbreviations"]
#         self.full_names = json_master["full_names"]
#         self.bl_info = json_master["bl_info"]
#         self.bl_index = json_master["bl_index"]
#         self.dft_groups = json_master["dft_groups"]
#         self.project_group = json_master["project_group"]
#         self.project_stage = json_master["project_stage"]
#         self.pipeline_dict = json_master["pipeline_dict"]
#         self.pipeline_list = json_master["pipeline_list"]
#         self.quarter_list = json_master["quarter_list"]


def convert_none_types(x):
    if x is None:
        return 0
    # data type handling. to complete.
    # if isinstance(x, str):
    #     raise DataTypeError("Correct data type in data")
    else:
        return x


class CostData:
    def __init__(self, master, **kwargs):
        self.master = master
        self.baseline_type = "ipdc_costs"
        self.kwargs = kwargs
        self.start_group = []
        self.group = []
        self.iter_list = []
        self.c_totals = {}
        self.c_bl_totals = {}
        self.c_profiles = {}
        self.profiles = {}
        # self.forecast_profiles = {}
        self.wlc_dict = {}
        self.wlc_change = {}
        # self.stack_p = {}
        self.get_cost_totals_new()

    def get_cost_totals_new(self) -> None:
        self.iter_list = get_iter_list(self.kwargs, self.master)
        lower_dict = {}
        # if "data_type" in self.kwargs:
        if self.kwargs["report"] == "cdg":
            for tp in self.iter_list:
                c_spent = 0
                c_remaining = 0
                c_total = 0
                in_achieved = 0
                in_remaining = 0
                in_total = 0
                self.group = get_group(self.master, tp, self.kwargs)
                for project_name in self.group:
                    p_data = get_correct_p_data(
                        self.master,
                        project_name,
                        tp,
                    )
                    c_spent += convert_none_types(p_data["Total Costs Spent"])
                    c_remaining += convert_none_types(p_data["Total Costs Remaining"])
                    c_total += convert_none_types(p_data["Total Costs"])
                    in_achieved += convert_none_types(p_data["Total Income Achieved"])
                    in_remaining += convert_none_types(p_data["Total Income Remaining"])
                    in_total += convert_none_types(p_data["Total Income"])

                lower_dict[tp] = {
                    "costs_spent": c_spent,
                    "costs_remaining": c_remaining,
                    "total": c_total,
                    "income_achieved": in_achieved,
                    "income_remaining": in_remaining,
                    "income_total": in_total,
                }

        if self.kwargs["report"] == "top_250":
            for tp in self.iter_list:
                rdel = 0
                cdel = 0
                c_total = 0
                non_gov = 0
                self.group = get_group(self.master, tp, self.kwargs)
                for project_name in self.group:
                    p_data = get_correct_p_data(
                        self.kwargs,
                        self.master,
                        self.baseline_type,
                        project_name,
                        tp,
                    )
                    rdel += convert_none_types(p_data["WLC GOV RDEL"])
                    cdel += convert_none_types(p_data["WLC GOV CDEL"])
                    c_total += convert_none_types(p_data["WLC TOTAL"])
                    non_gov += convert_none_types(p_data["WLC NON GOV"])

                lower_dict[tp] = {
                    "costs_rdel": rdel,
                    "costs_cdel": cdel,
                    "total": c_total + non_gov,
                    "non_gov": non_gov,
                }

        if self.kwargs["report"] == "ipdc":
            for tp in self.iter_list:
                self.group = get_group(self.master, tp, self.kwargs)
                rdel_total = []
                cdel_total = []
                ngov_total = []
                total = []
                for project_name in self.group:
                    p_data = get_correct_p_data(self.master, project_name, tp)
                    if p_data is None:
                        break
                    else:
                        rt = convert_none_types(p_data["Total RDEL Forecast Total"])
                        rdel_total.append(rt)
                        ct = convert_none_types(p_data["Total CDEL Forecast Total WLC"])
                        cdel_total.append(ct)
                        ng = convert_none_types(p_data["Non-Gov Total Forecast"])
                        ngov_total.append(ng)
                        t = convert_none_types(p_data["Total Forecast"])
                        # hard coded due to current use need.

                        if project_name in self.kwargs["remove income from totals"]:
                            try:
                                t = (
                                    t
                                    - p_data[
                                        "Total Forecast - Income both Revenue and Capital"
                                    ]
                                )
                            except KeyError:  # some older masters do have key.
                                pass
                        total.append(t)

                    # rdel_profiled.append(rt - (rs + ru))
                    # cdel_profiled.append(ct - (cs + cu))
                    # profiled.append(t - (s + u))

                lower_dict[tp] = {
                    # "cat_spent": [sum(rdel_spent), sum(cdel_spent)],
                    # "cat_prof": [sum(rdel_profiled), sum(cdel_profiled)],
                    # "cat_unprof": [sum(rdel_unprofiled), sum(cdel_unprofiled)],
                    # "spent": sum(spent),
                    # "prof": sum(profiled),
                    # "unprof": sum(unprofiled),
                    "total": sum(total),
                    "rdel": sum(rdel_total),
                    "cdel": sum(cdel_total) - sum(ngov_total),
                    "ngov": sum(ngov_total),
                }

        self.c_totals = lower_dict

    def get_cost_profile(self) -> None:
        """Returns several lists which contain the sum of different cost profiles for the group of project
        contained with the master"""
        self.iter_list = get_iter_list(self.kwargs, self.master)
        lower_dict = {}
        for tp in self.iter_list:
            yearly_profile = []
            rdel_yearly_profile = []
            cdel_yearly_profile = []
            ngov_yearly_profile = []
            self.group = get_group(self.master, tp, self.kwargs)
            for year in YEAR_LIST:
                cost_total = 0
                rdel_total = 0
                cdel_total = 0
                ngov_total = 0
                for cost_type in COST_KEY_LIST:
                    for p in self.group:
                        p_data = get_correct_p_data(
                            self.kwargs, self.master, self.baseline_type, p, tp
                        )
                        if p_data is None:
                            continue
                        try:
                            cost = p_data[year + cost_type]
                            if cost is None:
                                cost = 0
                            cost_total += cost
                        except KeyError:  # handles data across different financial years via proj_info
                            try:
                                cost = self.master.project_information[p][
                                    year + cost_type
                                ]
                            except KeyError:
                                cost = 0
                            if cost is None:
                                cost = 0
                            cost_total += cost

                        if cost_type == COST_KEY_LIST[0]:  # rdel
                            rdel_total += cost
                        if cost_type == COST_KEY_LIST[1]:  # cdel
                            cdel_total += cost
                        if cost_type == COST_KEY_LIST[2]:  # ngov
                            ngov_total += cost

                yearly_profile.append(cost_total)
                rdel_yearly_profile.append(rdel_total)
                cdel_yearly_profile.append(cdel_total)
                ngov_yearly_profile.append(ngov_total)
            lower_dict[tp] = {
                "prof": yearly_profile,
                "prof_ra": moving_average(yearly_profile, 2),
                "rdel": rdel_yearly_profile,
                "cdel": cdel_yearly_profile,
                "ngov": ngov_yearly_profile,
            }
        self.c_profiles = lower_dict

    def get_forecast_cost_profile(self) -> None:
        COST_CAT = [" RDEL ", " CDEL "]
        self.iter_list = get_iter_list(self.kwargs, self.master)
        tp_dict = {}
        for tp in self.iter_list:
            self.group = get_group(self.master, tp, self.kwargs)
            project_dict = {}
            list_total_total = []
            list_rdel_total = []
            list_cdel_total = []
            list_ngov_total = []
            list_std = []
            for p in self.group:
                RDEL_FORECAST_COST_KEYS = {
                    "Forecast one off new costs": [],
                    "Forecast recurring new costs": [],
                    "Forecast recurring old costs": [],
                    "Forecast Non Gov costs": [],
                    "Forecast Total": [],
                    "Forecast Income": [],
                }
                CDEL_FORECAST_COST_KEYS = {
                    "Forecast one off new costs": [],
                    "Forecast recurring new costs": [],
                    "Forecast recurring old costs": [],
                    " Forecast Non-Gov": [],
                    "Forecast Total WLC": [],
                    " Forecast - Income both Revenue and Capital": [],
                }
                p_data = get_correct_p_data(
                    self.kwargs, self.master, self.baseline_type, p, tp
                )
                if p_data is None:
                    continue
                rdel_std = convert_none_types(p_data["20-21 RDEL STD Total"])
                cdel_std = convert_none_types(p_data["20-21 CDEL STD Total"])
                list_std.append(rdel_std + cdel_std)
                for y in YEAR_LIST:
                    for cat in COST_CAT:
                        if cat == " RDEL ":
                            for k in RDEL_FORECAST_COST_KEYS.keys():
                                if y in ["16-17", "17-18", "18-19"]:
                                    try:
                                        rdel = convert_none_types(
                                            self.master.project_information[p][
                                                y + cat + k
                                            ]
                                        )
                                    except KeyError:
                                        rdel = 0
                                        print(y + cat + k + " not found.")
                                else:
                                    rdel = convert_none_types(p_data[y + cat + k])
                                RDEL_FORECAST_COST_KEYS[k].append(rdel)
                        if cat == " CDEL ":
                            for k in CDEL_FORECAST_COST_KEYS.keys():
                                if y in ["16-17", "17-18", "18-19"]:
                                    try:
                                        cdel = convert_none_types(
                                            self.master.project_information[p][
                                                y + cat + k
                                            ]
                                        )
                                    except KeyError:
                                        try:
                                            cdel = convert_none_types(
                                                self.master.project_information[p][
                                                    y + k
                                                ]
                                            )
                                        except KeyError:
                                            cdel = 0
                                            print(y + k + " not found.")
                                else:
                                    try:
                                        cdel = convert_none_types(p_data[y + cat + k])
                                    except KeyError:
                                        try:
                                            cdel = convert_none_types(p_data[y + k])
                                        except KeyError:
                                            # user messaging if necessary
                                            # print(tp + " " + y + k + ' could not be found. Check')
                                            cdel = 0
                                CDEL_FORECAST_COST_KEYS[k].append(cdel)

                    total_adding = [
                        RDEL_FORECAST_COST_KEYS["Forecast Total"],
                        CDEL_FORECAST_COST_KEYS["Forecast Total WLC"],
                    ]
                    year_total = [sum(x) for x in zip(*total_adding)]
                    ngov_adding = [
                        RDEL_FORECAST_COST_KEYS["Forecast Non Gov costs"],
                        CDEL_FORECAST_COST_KEYS[" Forecast Non-Gov"],
                    ]
                    ngov_total = [sum(x) for x in zip(*ngov_adding)]
                # adding individual project data to dict is not necessary
                # project_dict[p] = {
                #     "rdel": RDEL_FORECAST_COST_KEYS["Forecast Total"],
                #     "cdel": CDEL_FORECAST_COST_KEYS["Forecast Total WLC"],
                #     "ngov": ngov_total,
                #     "total": year_total,
                # }
                list_total_total.append(year_total)
                list_ngov_total.append(ngov_total)
                list_rdel_total.append(RDEL_FORECAST_COST_KEYS["Forecast Total"])
                list_cdel_total.append(CDEL_FORECAST_COST_KEYS["Forecast Total WLC"])

            project_dict["total"] = [sum(x) for x in zip(*list_total_total)]
            project_dict["rdel_total"] = [sum(x) for x in zip(*list_rdel_total)]
            project_dict["cdel_total"] = [sum(x) for x in zip(*list_cdel_total)]
            project_dict["ngov_total"] = [sum(x) for x in zip(*list_ngov_total)]
            project_dict["std"] = list_std

            self.profiles[tp] = project_dict

    def get_baseline_cost_profile(self) -> None:
        COST_CAT = [" RDEL ", " CDEL "]
        # self.iter_list = get_iter_list(self.kwargs, self.master)
        self.iter_list = [self.master.current_quarter]
        tp_dict = {}
        for tp in self.iter_list:
            self.group = get_group(self.master, tp, self.kwargs)
            project_dict = {}
            list_total_total = []
            list_rdel_total = []
            list_cdel_total = []
            list_ngov_total = []
            for p in self.group:
                # at moment bl only coming from current quarters data
                p_data = self.master.master_data[0]["data"][p]
                # p_data = get_correct_p_data(
                #     self.kwargs, self.master, self.baseline_type, p, tp
                # )
                if p_data is None:
                    continue
                for y in YEAR_LIST:
                    for cat in COST_CAT:
                        if cat == " RDEL ":
                            for k in RDEL_BL_COST_KEYS.keys():
                                if y in ["16-17", "17-18", "18-19"]:
                                    rdel = 0
                                else:
                                    rdel = convert_none_types(p_data[y + cat + k])
                                RDEL_BL_COST_KEYS[k].append(rdel)
                        if cat == " CDEL ":
                            for k in CDEL_BL_COST_KEYS.keys():
                                if y in ["16-17", "17-18", "18-19"]:
                                    cdel = 0
                                else:
                                    try:
                                        cdel = convert_none_types(p_data[y + cat + k])
                                    except KeyError:
                                        try:
                                            cdel = convert_none_types(p_data[y + k])
                                        except KeyError:
                                            # user messaging if necessary
                                            # print(tp + " " + y + k + ' could not be found. Check')
                                            cdel = 0
                                CDEL_BL_COST_KEYS[k].append(cdel)
                    total_adding = [
                        RDEL_BL_COST_KEYS["BL Total"],
                        CDEL_BL_COST_KEYS["BL WLC"],
                    ]
                    year_total = [sum(x) for x in zip(*total_adding)]
                    ngov_adding = [
                        RDEL_BL_COST_KEYS["BL Non Gov costs"],
                        CDEL_BL_COST_KEYS[" BL Non-Gov"],
                    ]
                    ngov_total = [sum(x) for x in zip(*ngov_adding)]
                # project_dict[p] = {
                #     "rdel": RDEL_BL_COST_KEYS["BL Total"],
                #     "cdel": CDEL_BL_COST_KEYS["BL WLC"],
                #     "total": year_total,
                # }
                list_total_total.append(year_total)
                list_ngov_total.append(ngov_total)
                list_rdel_total.append(RDEL_BL_COST_KEYS["BL Total"])
                list_cdel_total.append(CDEL_BL_COST_KEYS["BL WLC"])
            project_dict["total"] = [sum(x) for x in zip(*list_total_total)]
            project_dict["rdel_total"] = [sum(x) for x in zip(*list_rdel_total)]
            project_dict["cdel_total"] = [sum(x) for x in zip(*list_cdel_total)]
            project_dict["ngov_total"] = [sum(x) for x in zip(*list_ngov_total)]
            # tp_dict["baseline"] = project_dict
            self.profiles["baseline"] = project_dict

    def get_wlc_data(self) -> None:
        """
        calculates the quarters total wlc change
        """
        self.iter_list = get_iter_list(self.kwargs, self.master)
        wlc_dict = {}
        for tp in self.iter_list:
            self.group = get_group(self.master, tp, self.kwargs)
            wlc_dict = {}
            p_total = 0  # portfolio total

            for i, g in enumerate(self.group):
                l_group = get_group(self.master, tp, self.kwargs, i)  # lower group
                g_total = 0
                l_g_l = []  # lower group list
                for p in l_group:
                    p_data = get_correct_p_data(
                        self.kwargs, self.master, self.baseline_type, p, tp
                    )
                    wlc = p_data["Total Forecast"]
                    if isinstance(wlc, (float, int)) and wlc is not None and wlc != 0:
                        if wlc > 50000:
                            logger.info(
                                tp
                                + ", "
                                + str(p)
                                + " is £"
                                + str(round(wlc))
                                + " please check this is correct. For now analysis_engine has recorded it as £0"
                            )
                        # wlc_dict[p] = wlc
                    if wlc == 0:
                        logger.info(
                            tp
                            + ", "
                            + str(p)
                            + " wlc is currently £"
                            + str(wlc)
                            + " note this is key information that should be provided by the project"
                        )
                        # wlc_dict[p] = wlc
                    if wlc is None:
                        logger.info(
                            tp
                            + ", "
                            + str(p)
                            + " wlc is currently None note this is key information that should be provided by the project"
                        )
                        wlc = 0

                    l_g_l.append((wlc, p))
                    g_total += wlc

                wlc_dict[g] = list(reversed(sorted(l_g_l)))
                p_total += g_total

            wlc_dict["total"] = p_total
            wlc_dict[tp] = wlc_dict

        self.wlc_dict = wlc_dict

    def calculate_wlc_change(self) -> None:
        wlc_change_dict = {}
        for i, tp in enumerate(self.wlc_dict.keys()):
            p_wlc_change_dict = {}
            for p in self.wlc_dict[tp].keys():
                wlc_one = self.wlc_dict[tp][p]
                try:
                    wlc_two = self.wlc_dict[self.iter_list[i + 1]][p]
                    try:
                        percentage_change = int(((wlc_one - wlc_two) / wlc_one) * 100)
                        p_wlc_change_dict[p] = percentage_change
                    except ZeroDivisionError:
                        logger.info(
                            "As "
                            + str(p)
                            + " has no wlc total figure for "
                            + tp
                            + " change has been calculated as zero"
                        )
                except IndexError:  # handles NoneTypes.
                    pass

            wlc_change_dict[tp] = p_wlc_change_dict

        self.wlc_change = wlc_change_dict


def get_cost_baseline_keys():
    wb = load_workbook("/home/will/Downloads/cost_baselines.xlsx")
    ws = wb.active

    ws_keys = []
    for x in range(1, ws.max_row + 1):
        v = ws.cell(row=x, column=1).value
        if v is not None:
            ws_keys.append(v)

    RDEL_BL_COST_KEYS = [
        "BL one off new costs",
        "BL recurring new costs",
        "BL recurring old costs",
        "BL Non Gov costs",
        "BL Total",
        "BL Income",
    ]

    CDEL_BL_COST_KEYS = [
        "BL one off new costs",
        "BL recurring new costs",
        "BL recurring old costs",
        " BL Non-Gov",
        "BL WLC",
        " BL Income both Revenue and Capital",
    ]

    YEAR_LIST

    COST_CAT = [" RDEL ", " CDEL "]

    wb = Workbook()
    ws = wb.active

    for i, key in enumerate(ws_keys):
        ws.cell(row=i + 1, column=1).value = key

    created_list = []
    for cat in COST_CAT:
        if cat == " RDEL ":
            for y in YEAR_LIST:
                for k in RDEL_BL_COST_KEYS:
                    created_list.append(y + cat + k)
        if cat == " CDEL ":
            for y in YEAR_LIST:
                for k in CDEL_BL_COST_KEYS:
                    if k == " BL Non-Gov" or k == " BL Income both Revenue and Capital":
                        created_list.append(y + k)
                    else:
                        created_list.append(y + cat + k)

    for i in range(1, ws.max_row + 1):
        v = ws.cell(row=i + 1, column=1).value
        if v in created_list:
            ws.cell(row=i + 1, column=2).value = v

    wb.save("/home/will/Downloads/messy_bs_keys.xlsx")

    return created_list


def get_cost_forecast_keys():
    wb = load_workbook("/home/will/Downloads/cost_baselines.xlsx")
    ws = wb.active

    ws_keys = []
    for x in range(1, ws.max_row + 1):
        v = ws.cell(row=x, column=1).value
        if v is not None:
            ws_keys.append(v)

    RDEL_FORECAST_COST_KEYS = [
        "Forecast one off new costs",
        "Forecast recurring new costs",
        "Forecast recurring old costs",
        "Forecast Non Gov costs",
        "Forecast Total",
        "Forecast Income",
    ]

    CDEL_FORECAST_COST_KEYS = [
        "Forecast one off new costs",
        "Forecast recurring new costs",
        "Forecast recurring old costs",
        " Forecast Non-Gov",
        "Forecast Total WLC",
        " Forecast - Income both Revenue and Capital",
    ]

    COST_CAT = [" RDEL ", " CDEL "]

    wb = Workbook()
    ws = wb.active

    for i, key in enumerate(ws_keys):
        ws.cell(row=i + 1, column=1).value = key

    created_list = []
    for cat in COST_CAT:
        if cat == " RDEL ":
            for y in YEAR_LIST:
                for k in RDEL_FORECAST_COST_KEYS:
                    created_list.append(y + cat + k)
        if cat == " CDEL ":
            for y in YEAR_LIST:
                for k in CDEL_FORECAST_COST_KEYS:
                    if (
                        k == " Forecast Non-Gov"
                        or k == " Forecast - Income both Revenue and Capital"
                    ):
                        created_list.append(y + k)
                    else:
                        created_list.append(y + cat + k)

    for i in range(1, ws.max_row + 1):
        v = ws.cell(row=i + 1, column=1).value
        if v in created_list:
            ws.cell(row=i + 1, column=2).value = v

    wb.save("/home/will/Downloads/messy_af_keys.xlsx")

    return created_list


def get_sp_data(master, **kwargs) -> Dict[str, float]:
    if "group" in kwargs:
        group = kwargs["group"]
    elif "stage" in kwargs:
        group = kwargs["stage"]
        del kwargs["stage"]

    iter_list = get_iter_list(kwargs, master)
    tp_dict = {}

    for tp in iter_list:
        sp_dict = {}
        kwargs["quarter"] = [tp]
        if "type" in kwargs:
            cat_list = ["cdel", "rdel", "ngov"]
            kwargs["group"] = [group]
            c = CostData(master, **kwargs)
            c.get_cost_profile()
            for cat in cat_list:
                sp_dict[cat] = c.c_profiles[tp][cat]
        else:
            if len(group) > 1:
                for i, g in enumerate(group):
                    kwargs["group"] = [g]
                    c = CostData(master, **kwargs)
                    c.get_cost_profile()
                    sp_dict[g] = c.c_profiles[tp]["prof"]
            else:
                for i, g in enumerate(group):
                    kwargs["group"] = [g]
                    low_group = get_group(master, tp, kwargs)  # lower group
                    for p in low_group:
                        kwargs["group"] = [p]
                        c = CostData(master, **kwargs)
                        c.get_cost_profile()
                        sp_dict[master.abbreviations[p]["abb"]] = c.c_profiles[tp][
                            "prof"
                        ]

        tp_dict[tp] = sp_dict

    return tp_dict


def sort_group_by_key(
    key: str, group_idx: int, master, baseline_type: str, tp: str, kwargs
) -> List:  # no ** in front as passing in existing kwargs dict
    """
    Helper function. orders projects by key value e.g. total forecast
    """
    output_list = []
    group = get_group(master, tp, kwargs, group_idx)  # lower group
    for p in group:
        p_data = get_correct_p_data(kwargs, master, baseline_type, p, tp)
        value = p_data[key]
        if value is None:
            value = 0

        output_list.append((value, p))

    return list(reversed(sorted(output_list)))


def put_cost_totals_into_wb(costs: CostData) -> workbook:
    wb = Workbook()
    ws = wb.active

    for i, p in enumerate(costs.c_totals.keys()):
        ws.cell(row=i + 2, column=1).value = p
        ws.cell(row=i + 2, column=2).value = costs.c_totals[p]["spent"]
        ws.cell(row=i + 2, column=3).value = costs.c_totals[p]["profiled"]
        ws.cell(row=i + 2, column=4).value = costs.c_totals[p]["unprofiled"]
        s = int(
            costs.c_totals[p]["spent"]
            + costs.c_totals[p]["profiled"]
            + costs.c_totals[p]["unprofiled"]
        )
        ws.cell(row=i + 2, column=5).value = s
        total = int(costs.c_totals[p]["total"])
        ws.cell(row=i + 2, column=6).value = total
        if s != total:
            ws.cell(row=i + 2, column=5).fill = SALMON_FILL

    ws.cell(row=1, column=1).value = "project"
    ws.cell(row=1, column=2).value = "spent"
    ws.cell(row=1, column=3).value = "profiled"
    ws.cell(row=1, column=4).value = "unprofiled"
    ws.cell(row=1, column=5).value = "sum"
    ws.cell(row=1, column=6).value = "total wlc"

    return wb


def moving_average(x, w):
    return np.convolve(x, np.ones(w), "valid") / w


def milestone_info_handling(output_list: list, t_list: list, **kwargs) -> list:
    """helper function for handling and cleaning up milestone date generated
    via MilestoneDate class. Removes none type milestone names and non date
    string values"""

    if t_list[1][1] is None or t_list[1][1] == "Project - Business Case End Date":
        pass
    else:
        try:
            t_list[3] = ("Date", t_list[3][1].date())
            return output_list.append(t_list)
        except AttributeError:  # Non-datetime values
            if "type" in kwargs:
                if kwargs["type"] == "central support":
                    if t_list[3][1] is None:
                        return output_list.append(t_list)
                    else:
                        logger.info(
                            t_list[0][1]
                            + ": incorrect date format for entry '"
                            + t_list[1][1]
                            + ""
                            "', requires amending or will not be included. "
                            + str(kwargs["tp"])
                            + " data."
                        )
            if t_list[3][1] is None:
                pass
            else:
                logger.info(
                    t_list[0][1]
                    + ": incorrect date format for entry '"
                    + t_list[1][1]
                    + ""
                    "', requires amending or will not be included. "
                    + str(kwargs["tp"])
                    + " data."
                )


def remove_none_types(input_list):
    return [x for x in input_list if x is not None]


def get_milestone_date(
    project_name: str,
    milestone_dictionary: Dict[str, Union[datetime.date, str]],
    quarter_bl: str,
    milestone_name: str,
) -> datetime:
    m_dict = milestone_dictionary[quarter_bl]
    for k in m_dict.keys():
        if m_dict[k]["Project"] == project_name:
            if m_dict[k]["Milestone"] == milestone_name:
                # removed milestone_name[1:]
                return m_dict[k]["Date"]


def get_milestone_notes(
    project_name: str,
    milestone_dictionary: Dict[str, Union[datetime.date, str]],
    tp: str,  # time period
    milestone_name: str,
) -> datetime:
    m_dict = milestone_dictionary[tp]
    for k in m_dict.keys():
        if m_dict[k]["Project"] == project_name:
            if m_dict[k]["Milestone"] == milestone_name:
                return m_dict[k]["Notes"]


class MilestoneData:
    def __init__(
        self,
        master,
        baseline_type: str = "ipdc_milestones",
        **kwargs,
    ):
        self.master = master
        self.group = []
        self.iter_list = []  # iteration list
        self.kwargs = kwargs
        self.baseline_type = baseline_type
        self.milestone_dict = {}
        self.sorted_milestone_dict = {}
        self.max_date = None
        self.min_date = None
        self.schedule_change = {}
        self.schedule_key_last = None
        self.schedule_key_baseline = None
        self.get_milestones()
        self.get_chart_info()
        # self.calculate_schedule_changes()

    def get_milestones(self) -> None:
        """
        Creates project milestone dictionaries for current, last_quarter, and
        baselines when provided with group and baseline type.
        """
        m_dict = {}
        self.iter_list = get_iter_list(self.kwargs, self.master)
        for tp in self.iter_list:  # tp time period
            self.kwargs["tp"] = tp
            lower_dict = {}
            raw_list = []
            self.group = get_group(self.master, tp, self.kwargs)
            for project_name in self.group:
                project_milestones = []
                p_data = get_correct_p_data(
                    self.kwargs, self.master, self.baseline_type, project_name, tp
                )
                if p_data is None:
                    continue
                # i loops below removes None Milestone names and rejects non-datetime date values.
                p = self.master.abbreviations[project_name]["abb"]
                category = "Milestone"
                if "data_type" in self.kwargs:
                    if self.kwargs["data_type"] == "top35":
                        report = "Top 250"
                        for i in range(1, 30):
                            entry = [
                                ("Project", p),
                                ("Milestone", p_data["MM" + str(i) + " name"]),
                                ("Notes", p_data["MM" + str(i) + " Comment"]),
                                ("Date", convert_date(p_data["MM" + str(i) + " date"])),
                                ("Status", p_data["MM" + str(i) + " status"]),
                                ("Report", report),
                                ("Cat", category),
                            ]
                            milestone_info_handling(
                                project_milestones, entry, **self.kwargs
                            )
                    if self.kwargs["data_type"] == "cdg":
                        report = "CDG"
                        for i in range(1, 15):
                            entry = [
                                ("Project", p),
                                ("Milestone", p_data["MM" + str(i)]),
                                ("Notes", p_data["MM" + str(i) + " NOTES"]),
                                ("Date", convert_date(p_data["MM" + str(i) + " DATE"])),
                                ("Status", p_data["MM" + str(i) + " STATUS"]),
                                ("Report", report),
                                ("Cat", category),
                            ]
                            milestone_info_handling(
                                project_milestones, entry, **self.kwargs
                            )

                else:
                    report = "IPDC/GMPP"
                    for i in range(1, 50):
                        try:
                            entry = [
                                ("Project", p),
                                ("Milestone", p_data["Approval MM" + str(i)]),
                                ("Type", "Approval"),
                                (
                                    "Date",
                                    convert_date(
                                        p_data[
                                            "Approval MM"
                                            + str(i)
                                            + " Forecast / Actual"
                                        ]
                                    ),
                                ),
                                ("Notes", p_data["Approval MM" + str(i) + " Notes"]),
                                ("Report", report),
                                ("Cat", category),
                            ]
                            milestone_info_handling(
                                project_milestones, entry, **self.kwargs
                            )
                            entry = [
                                ("Project", p),
                                ("Milestone", p_data["Assurance MM" + str(i)]),
                                ("Type", "Assurance"),
                                (
                                    "Date",
                                    convert_date(
                                        p_data[
                                            "Assurance MM"
                                            + str(i)
                                            + " Forecast - Actual"
                                        ]
                                    ),
                                ),
                                ("Notes", p_data["Assurance MM" + str(i) + " Notes"]),
                                ("Report", report),
                                ("Cat", category),
                            ]
                            milestone_info_handling(
                                project_milestones, entry, **self.kwargs
                            )
                        except KeyError:  # handles inconsistent keys naming for approval milestones.
                            try:
                                entry = [
                                    ("Project", p),
                                    ("Milestone", p_data["Approval MM" + str(i)]),
                                    ("Type", "Approval"),
                                    (
                                        "Date",
                                        convert_date(
                                            p_data[
                                                "Approval MM"
                                                + str(i)
                                                + " Forecast - Actual"
                                            ]
                                        ),
                                    ),
                                    (
                                        "Notes",
                                        p_data["Approval MM" + str(i) + " Notes"],
                                    ),
                                    ("Report", report),
                                    ("Cat", category),
                                ]
                                milestone_info_handling(
                                    project_milestones, entry, **self.kwargs
                                )
                            except KeyError:
                                pass

                    # handles inconsistent number of Milestone. could be incorporated above.
                    for i in range(18, 67):
                        try:
                            entry = [
                                ("Project", p),
                                ("Milestone", p_data["Project MM" + str(i)]),
                                ("Type", "Delivery"),
                                (
                                    "Date",
                                    convert_date(
                                        p_data[
                                            "Project MM" + str(i) + " Forecast - Actual"
                                        ]
                                    ),
                                ),
                                ("Notes", p_data["Project MM" + str(i) + " Notes"]),
                                ("Report", report),
                                ("Cat", category),
                            ]
                            milestone_info_handling(
                                project_milestones, entry, **self.kwargs
                            )
                        except KeyError:
                            pass

                    # change in Q3. Some milestones collected via HMT approval section.
                    # this loop picks them up
                    for i in range(1, 4):
                        try:
                            entry = [
                                ("Project", p),
                                ("Milestone", p_data["HMT Approval " + str(i)]),
                                ("Type", "Approval"),
                                (
                                    "Date",
                                    convert_date(
                                        p_data[
                                            "HMT Approval "
                                            + str(i)
                                            + " Forecast / Actual"
                                        ]
                                    ),
                                ),
                                ("Notes", p_data["HMT Approval " + str(i) + " Notes"]),
                                ("Report", report),
                                ("Cat", category),
                            ]
                            milestone_info_handling(
                                project_milestones, entry, **self.kwargs
                            )
                        except KeyError:
                            pass

                # loop to stop keys names being the same. Done at project level.
                # not particularly concise code.
                upper_counter_list = []
                for entry in project_milestones:
                    upper_counter_list.append(entry[1][1])
                upper_count = Counter(upper_counter_list)
                lower_counter_list = []
                for entry in project_milestones:
                    if upper_count[entry[1][1]] > 1:
                        lower_counter_list.append(entry[1][1])
                        lower_count = Counter(lower_counter_list)
                        new_milestone_key = (
                            entry[1][1] + " (" + str(lower_count[entry[1][1]]) + ")"
                        )
                        entry[1] = ("Milestone", new_milestone_key)
                        raw_list.append(entry)
                    else:
                        raw_list.append(entry)

            # puts the list in chronological order
            sorted_list = sorted(raw_list, key=lambda k: (k[3][1] is None, k[3][1]))

            for r in range(len(sorted_list)):
                lower_dict["Milestone " + str(r)] = dict(sorted_list[r])

            m_dict[tp] = lower_dict
            # HERE POTENTIAL TO MERGE DICTS
        self.milestone_dict = m_dict

    def get_chart_info(self) -> None:
        """returns data lists for matplotlib chart"""
        # Note this code could refactored so that it collects all milestones
        # reported across current, last and baseline. At the moment it only
        # uses milestones that are present in the current quarter.

        output_dict = {}
        for i in self.milestone_dict:
            report = []
            category = []
            p_names = []
            key_names = []
            g_dates = []  # graph dates
            r_dates = []  # raw dates
            notes = []
            status = []
            for v in self.milestone_dict[self.iter_list[0]].values():
                p = None  # project
                mn = None  # milestone name
                d = None  # date
                for x in self.milestone_dict[i].values():
                    if (
                        x["Project"] == v["Project"]
                        and x["Milestone"] == v["Milestone"]
                    ):
                        p = x["Project"]
                        mn = x["Milestone"]
                        p_names.append(p)
                        key_names.append(mn)
                        d = x["Date"]
                        g_dates.append(d)
                        r_dates.append(d)
                        notes.append(x["Notes"])
                        report.append(x["Report"])
                        category.append(x["Cat"])
                        try:
                            status.append(x["Status"])
                        except KeyError:
                            pass
                        break
                if p is None and mn is None and d is None:
                    p = v["Project"]
                    mn = v["Milestone"]
                    p_names.append(p)
                    key_names.append(mn)
                    g_dates.append(v["Date"])
                    r_dates.append(None)
                    notes.append(None)
                    status.append(None)
                    report.append(x["Report"])
                    category.append(x["Cat"])

            output_dict[i] = {
                "project": p_names,
                "names": key_names,
                "g_dates": g_dates,
                "r_dates": r_dates,
                "notes": notes,
                "status": status,  # only present for top35
                "report": report,
                "cat": category,
            }

        self.sorted_milestone_dict = output_dict

    # def get_chart_info_old(self) -> None:
    #     """returns data lists for matplotlib chart"""
    #     # Note this code could refactored so that it collects all milestones
    #     # reported across current, last and baseline. At the moment it only
    #     # uses milestones that are present in the current quarter.
    #     key_names = []
    #     key_names_last = []
    #     keys_names_baseline = []
    #     md_current = []
    #     md_last = []
    #     md_last_po = []  # po is for printout
    #     md_baseline = []
    #     md_baseline_po = []
    #     md_baseline_two_po = []
    #     md_baseline_two = []
    #     type_list = []
    #
    #     for m in self.milestone_dict[self.iter_list[0]].values():
    #         m_project = m["Project"]
    #         m_name = m["Milestone"]
    #         m_date = m["Date"]
    #         m_type = m["Type"]
    #         key_names.append(m_project + ", " + m_name)
    #         md_current.append(m_date)
    #         type_list.append(m_type)
    #
    #         # In two loops below NoneType has to be replaced with a datetime object
    #         # due to matplotlib being unable to handle NoneTypes when milestone_chart
    #         # is created. Haven't been able to find a solution to this.
    #         try:
    #             m_last_date = None
    #             for m_last in self.milestone_dict[self.iter_list[1]].values():
    #                 if m_last["Project"] == m_project:
    #                     if m_last["Milestone"] == m_name:
    #                         key_names_last.append(m_project + ", " + m_name)
    #                         m_last_date = m_last["Date"]
    #                         md_last.append(m_last_date)
    #                         md_last_po.append(m_last_date)
    #                         continue
    #             if m_last_date is None:
    #                 md_last.append(m_date)
    #                 md_last_po.append(None)
    #
    #             m_bl_date = None
    #             for m_bl in self.milestone_dict[self.iter_list[2]].values():
    #                 if m_bl["Project"] == m_project:
    #                     if m_bl["Milestone"] == m_name:
    #                         keys_names_baseline.append(m_project + ", " + m_name)
    #                         m_bl_date = m_bl["Date"]
    #                         md_baseline.append(m_bl_date)
    #                         md_baseline_po.append(m_bl_date)
    #                         continue
    #             if m_bl_date is None:
    #                 md_baseline.append(m_date)
    #                 md_baseline_po.append(None)
    #
    #             m_bl_two_date = None
    #             for m_bl_two in self.milestone_dict[self.iter_list[3]].values():
    #                 if m_bl_two["Project"] == m_project:
    #                     if m_bl_two["Milestone"] == m_name:
    #                         m_bl_two_date = m_bl_two["Date"]
    #                         md_baseline_two.append(m_bl_two_date)
    #                         md_baseline_two_po.append(m_bl_two_date)
    #                         continue
    #             if m_bl_two_date is None:
    #                 md_baseline_two.append(m_date)
    #                 md_baseline_two_po.append(None)
    #
    #         except IndexError:
    #             pass
    #
    #     if len(self.group) == 1:
    #         key_names = remove_project_name_from_milestone_key(
    #             self.master.abbreviations[self.group[0]]["abb"], key_names
    #         )
    #     else:
    #         pass
    #
    #     self.key_names = key_names
    #     self.key_names_last = key_names_last
    #     self.key_names_baseline = keys_names_baseline
    #     self.md_current = md_current
    #     self.md_last = md_last
    #     self.md_last_po = md_last_po
    #     self.md_baseline = md_baseline
    #     self.md_baseline_po = md_baseline_po
    #     self.md_baseline_two = md_baseline_two
    #     self.md_baseline_two_po = md_baseline_two_po
    #     self.type_list = type_list
    #     self.max_date = max(
    #         remove_none_types(self.md_current)
    #         + remove_none_types(self.md_last)
    #         + remove_none_types(self.md_baseline)
    #     )
    #     self.min_date = min(
    #         remove_none_types(self.md_current)
    #         + remove_none_types(self.md_last)
    #         + remove_none_types(self.md_baseline)
    #     )

    def filter_chart_info(self, **filter_kwargs):
        # bug handling required in the event that there are no milestones with the filter.
        # i.e. the filter returns no milestones.
        filtered_dict = {}
        if (
            "type" in filter_kwargs
            and "key" in filter_kwargs
            and "dates" in filter_kwargs
        ):
            start_date, end_date = zip(*filter_kwargs["dates"])
            start = parser.parse(start_date, dayfirst=True)
            end = parser.parse(end_date, dayfirst=True)
            for i, v in enumerate(self.milestone_dict[self.iter_list[0]].values()):
                if v["Type"] in filter_kwargs["type"]:
                    if v["Milestone"] in filter_kwargs["keys"]:
                        if start.date() <= filter_kwargs["dates"] <= end.date():
                            filtered_dict["Milestone " + str(i)] = v
                            continue

        elif "type" in filter_kwargs and "key" in filter_kwargs:
            for i, v in enumerate(self.milestone_dict[self.iter_list[0]].values()):
                if v["Type"] in filter_kwargs["type"]:
                    if v["Milestone"] in filter_kwargs["keys"]:
                        filtered_dict["Milestone " + str(i)] = v
                        continue

        elif "type" in filter_kwargs and "dates" in filter_kwargs:
            start_date, end_date = zip(filter_kwargs["dates"])
            start = parser.parse(start_date[0], dayfirst=True)
            end = parser.parse(end_date[0], dayfirst=True)
            for i, v in enumerate(self.milestone_dict[self.iter_list[0]].values()):
                if v["Type"] in filter_kwargs["type"]:
                    if start.date() <= v["Date"] <= end.date():
                        filtered_dict["Milestone " + str(i)] = v
                        continue

        elif "key" in filter_kwargs and "dates" in filter_kwargs:
            start_date, end_date = zip(filter_kwargs["dates"])
            start = parser.parse(start_date[0], dayfirst=True)
            end = parser.parse(end_date[0], dayfirst=True)
            for i, v in enumerate(self.milestone_dict[self.iter_list[0]].values()):
                if v["Milestone"] in filter_kwargs["key"]:
                    if start.date() <= v["Date"] <= end.date():
                        filtered_dict["Milestone " + str(i)] = v
                        continue

        elif "type" in filter_kwargs:
            for i, v in enumerate(self.milestone_dict[self.iter_list[0]].values()):
                if v["Type"] in filter_kwargs["type"]:
                    filtered_dict["Milestone " + str(i)] = v
                    continue

        elif "key" in filter_kwargs:
            for i, v in enumerate(self.milestone_dict[self.iter_list[0]].values()):
                if v["Milestone"] in filter_kwargs["key"]:
                    filtered_dict["Milestone " + str(i)] = v
                    continue

        elif "dates" in filter_kwargs:
            start_date, end_date = zip(filter_kwargs["dates"])
            start = parser.parse(start_date[0], dayfirst=True)
            end = parser.parse(end_date[0], dayfirst=True)
            for i, v in enumerate(self.milestone_dict[self.iter_list[0]].values()):
                if start.date() <= v["Date"] <= end.date():
                    filtered_dict["Milestone " + str(i)] = v
                    continue

        output_dict = {}
        for dict in self.milestone_dict.keys():
            if dict == self.iter_list[0]:
                output_dict[dict] = filtered_dict
            else:
                output_dict[dict] = self.milestone_dict[dict]

        self.milestone_dict = output_dict
        self.get_chart_info()

    def calculate_schedule_changes(self) -> None:
        """calculates the changes in project schedules. If standard key for calculation
        not available it using the best next one available"""

        self.filter_chart_info(milestone_type=["Delivery", "Approval"])
        m_dict_keys = list(self.milestone_dict.keys())

        def schedule_info(
            project_name: str,
            other_key_list: List[str],
            c_key_list: List[str],
            miles_dict: dict,
            dict_l_current: str,
            dict_l_other: str,
        ):
            output_dict = {}
            schedule_info = []
            for key in reversed(other_key_list):
                if key in c_key_list:
                    sop = get_milestone_date(
                        project_name, miles_dict, dict_l_other, " Start of Project"
                    )
                    if sop is None:
                        sop = get_milestone_date(
                            project_name, miles_dict, dict_l_current, other_key_list[0]
                        )
                        schedule_info.append(("start key", other_key_list[0]))
                    else:
                        schedule_info.append(("start key", " Start of Project"))
                    schedule_info.append(("start", sop))
                    schedule_info.append(("end key", key))
                    date = get_milestone_date(
                        project_name, miles_dict, dict_l_current, key
                    )
                    schedule_info.append(("end current date", date))
                    other_date = get_milestone_date(
                        project_name, miles_dict, dict_l_other, key
                    )
                    schedule_info.append(("end other date", other_date))
                    project_length = (other_date - sop).days
                    schedule_info.append(("project length", project_length))
                    change = (date - other_date).days
                    schedule_info.append(("change", change))
                    p_change = int((change / project_length) * 100)
                    schedule_info.append(("percent change", p_change))
                    output_dict[dict_l_other] = dict(schedule_info)
                    break

            return output_dict

        output_dict = {}
        for project_name in self.group:
            project_name = self.master.abbreviations[project_name]
            current_key_list = []
            last_key_list = []
            baseline_key_list = []
            for key in self.key_names:
                try:
                    p = key.split(",")[0]
                    milestone_key = key.split(",")[1]
                    if project_name == p:
                        if milestone_key != " Project - Business Case End Date":
                            current_key_list.append(milestone_key)
                except IndexError:
                    # patch of single project group. In this instance the project name
                    # is removed from the key_name via remove_project_name function as
                    # part of get chart info.
                    if len(self.group) == 1:
                        current_key_list.append(" " + key)
            for last_key in self.key_names_last:
                p = last_key.split(",")[0]
                milestone_key_last = last_key.split(",")[1]
                if project_name == p:
                    if milestone_key_last != " Project - Business Case End Date":
                        last_key_list.append(milestone_key_last)
            for baseline_key in self.key_names_baseline:
                p = baseline_key.split(",")[0]
                milestone_key_baseline = baseline_key.split(",")[1]
                if project_name == p:
                    if (
                        milestone_key_baseline
                        != " Project - Business Case End Date"
                        # and milestone_key_baseline != " Project End Date"
                    ):
                        baseline_key_list.append(milestone_key_baseline)

            b_dict = schedule_info(
                project_name,
                baseline_key_list,
                current_key_list,
                self.milestone_dict,
                m_dict_keys[0],
                m_dict_keys[2],
            )
            l_dict = schedule_info(
                project_name,
                last_key_list,
                current_key_list,
                self.milestone_dict,
                m_dict_keys[0],
                m_dict_keys[1],
            )
            lower_dict = {**b_dict, **l_dict}

            output_dict[project_name] = lower_dict

        self.schedule_change = output_dict


# class CombinedData:
#     def __init__(self, wb, pfm_milestone_data):
#         self.wb = wb
#         self.pfm_milestone_data = pfm_milestone_data
#         # self.project_current = {}
#         # self.project_last = {}
#         # self.project_baseline = {}
#         # self.project_baseline_two = {}
#         self.group_current = {}
#         self.group_last = {}
#         self.group_baseline = {}
#         self.group_baseline_two = {}
#         self.combined_tuple_list_forecast = []
#         self.combined_tuple_list_baseline = []
#         self.combine_mi_pfm_data()
#
#     def combine_mi_pfm_data(self):
#         """
#         coverts data from MI system into usable format for graphical outputs
#         """
#         ws = self.wb.active
#
#         mi_milestone_name_list = []  # handles duplicates
#         mi_tuple_list_forecast = []
#         mi_tuple_list_baseline = []
#         for r in range(4, ws.max_row + 1):
#             mi_milestone_key_name_raw = ws.cell(row=r, column=3).value
#             mi_milestone_key_name = "MI, " + mi_milestone_key_name_raw
#             forecast_date = ws.cell(row=r, column=8).value
#             baseline_date = ws.cell(row=r, column=9).value
#             notes = ws.cell(row=r, column=10).value
#             if mi_milestone_key_name not in mi_milestone_name_list:
#                 mi_milestone_name_list.append(mi_milestone_key_name)
#                 mi_tuple_list_forecast.append(
#                     (mi_milestone_key_name, forecast_date.date(), notes)
#                 )
#                 mi_tuple_list_baseline.append(
#                     (mi_milestone_key_name, baseline_date.date(), notes)
#                 )
#             else:
#                 for i in range(
#                         2, 15
#                 ):  # alters duplicates by adding number to end of keys
#                     mi_altered_milestone_key_name = mi_milestone_key_name + " " + str(i)
#                     if mi_altered_milestone_key_name in mi_milestone_name_list:
#                         continue
#                     else:
#                         mi_tuple_list_forecast.append(
#                             (mi_altered_milestone_key_name, forecast_date.date(), notes)
#                         )
#                         mi_tuple_list_baseline.append(
#                             (mi_altered_milestone_key_name, baseline_date.date(), notes)
#                         )
#                         mi_milestone_name_list.append(mi_altered_milestone_key_name)
#                         break
#
#         mi_tuple_list_forecast = sorted(
#             mi_tuple_list_forecast, keys=lambda k: (k[1] is None, k[1])
#         )  # put the list in chronological order
#         mi_tuple_list_baseline = sorted(
#             mi_tuple_list_baseline, keys=lambda k: (k[1] is None, k[1])
#         )  # put the list in chronological order
#
#         pfm_tuple_list_forecast = []
#         pfm_tuple_list_baseline = []
#         for data in self.pfm_milestone_data.ordered_list_bl_two:
#             pfm_tuple_list_forecast.append(("PfM, " + data[0], data[1], data[2]))
#         for data in self.pfm_milestone_data.group_choronological_list_baseline:
#             pfm_tuple_list_baseline.append(("PfM, " + data[0], data[1], data[2]))
#
#         combined_tuple_list_forecast = mi_tuple_list_forecast + pfm_tuple_list_forecast
#         combined_tuple_list_baseline = mi_tuple_list_baseline + pfm_tuple_list_baseline
#
#         combined_tuple_list_forecast = sorted(
#             combined_tuple_list_forecast, keys=lambda k: (k[1] is None, k[1])
#         )  # put the list in chronological order
#         combined_tuple_list_baseline = sorted(
#             combined_tuple_list_baseline, keys=lambda k: (k[1] is None, k[1])
#         )  # put the list in chronological order
#
#         milestone_dict_forecast = {}
#         for series_one in combined_tuple_list_forecast:
#             if series_one[0] is not None:
#                 milestone_dict_forecast[series_one[0]] = {series_one[1]: series_one[2]}
#         milestone_dict_baseline = {}
#         for series_one in combined_tuple_list_baseline:
#             if series_one[0] is not None:
#                 milestone_dict_baseline[series_one[0]] = {series_one[1]: series_one[2]}
#
#         self.group_current = milestone_dict_forecast
#         self.group_last = {}
#         self.group_baseline = milestone_dict_baseline
#         self.group_baseline_two = {}
#         self.combined_tuple_list_forecast = combined_tuple_list_forecast
#         self.combined_tuple_list_baseline = combined_tuple_list_baseline
#
#
# class MilestoneCharts:
#     def __init__(
#             self,
#             latest_milestone_names,
#             latest_milestone_dates,
#             last_milestone_dates,
#             baseline_milestone_dates,
#             graph_title,
#             ipdc_date,
#     ):
#         self.latest_milestone_names = latest_milestone_names
#         self.latest_milestone_dates = latest_milestone_dates
#         self.last_milestone_dates = last_milestone_dates
#         self.baseline_milestone_dates = baseline_milestone_dates
#         self.graph_title = graph_title
#         self.ipdc_date = ipdc_date
#         # self.milestone_swimlane_charts()
#         self.build_charts()
#
#     def milestone_swimlane_charts(self):
#         # build scatter chart
#         fig, ax1 = plt.subplots()
#         fig.suptitle(self.graph_title, fontweight="bold")  # title
#         # set fig size
#         fig.set_figheight(4)
#         fig.set_figwidth(8)
#
#         ax1.scatter(
#             self.baseline_milestone_dates, self.latest_milestone_names, label="Baseline"
#         )
#         ax1.scatter(
#             self.last_milestone_dates, self.latest_milestone_names, label="Last Qrt"
#         )
#         ax1.scatter(
#             self.latest_milestone_dates, self.latest_milestone_names, label="Latest Qrt"
#         )
#
#         # format the series_one ticks
#         years = mdates.YearLocator()  # every year
#         months = mdates.MonthLocator()  # every month
#         years_fmt = mdates.DateFormatter("%Y")
#         months_fmt = mdates.DateFormatter("%b")
#
#         # calculate the length of the time period covered in chart. Not perfect as baseline dates can distort.
#         try:
#             td = (self.latest_milestone_dates[-1] - self.latest_milestone_dates[0]).days
#             if td <= 365 * 3:
#                 ax1.xaxis.set_major_locator(years)
#                 ax1.xaxis.set_minor_locator(months)
#                 ax1.xaxis.set_major_formatter(years_fmt)
#                 ax1.xaxis.set_minor_formatter(months_fmt)
#                 plt.setp(ax1.xaxis.get_minorticklabels(), rotation=45)
#                 plt.setp(
#                     ax1.xaxis.get_majorticklabels(), rotation=45, weight="bold"
#                 )  # milestone_swimlane_charts(key_name,
#                 #                           current_m_data,
#                 #                           last_m_data,
#                 #                           baseline_m_data,
#                 #                           'All Milestones')
#                 # scaling series_one axis
#                 # series_one axis value to no more than three months after last latest milestone date, or three months
#                 # before first latest milestone date. Hack, can be improved. Text highlights movements off chart.
#                 x_max = self.latest_milestone_dates[-1] + timedelta(days=90)
#                 x_min = self.latest_milestone_dates[0] - timedelta(days=90)
#                 for date in self.baseline_milestone_dates:
#                     if date > x_max:
#                         ax1.set_xlim(x_min, x_max)
#                         plt.figtext(
#                             0.98,
#                             0.03,
#                             "Check full schedule to see all milestone movements",
#                             horizontalalignment="right",
#                             fontsize=6,
#                             fontweight="bold",
#                         )
#                     if date < x_min:
#                         ax1.set_xlim(x_min, x_max)
#                         plt.figtext(
#                             0.98,
#                             0.03,
#                             "Check full schedule to see all milestone movements",
#                             horizontalalignment="right",
#                             fontsize=6,
#                             fontweight="bold",
#                         )
#             else:
#                 ax1.xaxis.set_major_locator(years)
#                 ax1.xaxis.set_minor_locator(months)
#                 ax1.xaxis.set_major_formatter(years_fmt)
#                 plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, weight="bold")
#         except IndexError:  # if milestone dates list is empty:
#             pass
#
#         ax1.legend()  # insert legend
#
#         # reverse series_two axis so order is earliest to oldest
#         ax1 = plt.gca()
#         ax1.set_ylim(ax1.get_ylim()[::-1])
#         ax1.tick_params(axis="series_two", which="major", labelsize=7)
#         ax1.yaxis.grid()  # horizontal lines
#         ax1.set_axisbelow(True)
#         # ax1.get_yaxis().set_visible(False)
#
#         # for i, txt in enumerate(latest_milestone_names):
#         #     ax1.annotate(txt, (i, latest_milestone_dates[i]))
#
#         # Add line of IPDC date, but only if in the time period
#         try:
#             if (
#                     self.latest_milestone_dates[0]
#                     <= self.ipdc_date
#                     <= self.latest_milestone_dates[-1]
#             ):
#                 plt.axvline(self.ipdc_date)
#                 plt.figtext(
#                     0.98,
#                     0.01,
#                     "Line represents when IPDC will discuss Q1 20_21 portfolio management report",
#                     horizontalalignment="right",
#                     fontsize=6,
#                     fontweight="bold",
#                 )
#         except IndexError:
#             pass
#
#         # size of chart and fit
#         fig.canvas.draw()
#         fig.tight_layout(rect=[0, 0.03, 1, 0.95])  # for title
#
#         fig.savefig(
#             root_path / "output/{}.png".format(self.graph_title), bbox_inches="tight"
#         )
#
#         # plt.close() #automatically closes figure so don't need to do manually.
#
#     def build_charts(self):
#
#         # add \n to series_two axis labels and cut down if two long
#         # labels = ['\n'.join(wrap(l, 40)) for l in latest_milestone_names]
#         labels = self.latest_milestone_names
#         final_labels = []
#         for l in labels:
#             if len(l) > 40:
#                 final_labels.append(l[:35])
#             else:
#                 final_labels.append(l)
#
#         # Chart
#         no_milestones = len(self.latest_milestone_names)
#
#         if no_milestones <= 30:
#             (
#                 np.array(final_labels),
#                 np.array(self.latest_milestone_dates),
#                 np.array(self.last_milestone_dates),
#                 np.array(self.baseline_milestone_dates),
#                 self.graph_title,
#                 self.ipdc_date,
#             )
#
#         if 31 <= no_milestones <= 60:
#             half = int(no_milestones / 2)
#             MilestoneCharts(
#                 np.array(final_labels[:half]),
#                 np.array(self.latest_milestone_dates[:half]),
#                 np.array(self.last_milestone_dates[:half]),
#                 np.array(self.baseline_milestone_dates[:half]),
#                 self.graph_title,
#                 self.ipdc_date,
#             )
#             title = self.graph_title + " cont."
#             MilestoneCharts(
#                 np.array(final_labels[half:no_milestones]),
#                 np.array(self.latest_milestone_dates[half:no_milestones]),
#                 np.array(self.last_milestone_dates[half:no_milestones]),
#                 np.array(self.baseline_milestone_dates[half:no_milestones]),
#                 title,
#                 self.ipdc_date,
#             )
#
#         if 61 <= no_milestones <= 90:
#             third = int(no_milestones / 3)
#             MilestoneCharts(
#                 np.array(final_labels[:third]),
#                 np.array(self.latest_milestone_dates[:third]),
#                 np.array(self.last_milestone_dates[:third]),
#                 np.array(self.baseline_milestone_dates[:third]),
#                 self.graph_title,
#                 self.ipdc_date,
#             )
#             title = self.graph_title + " cont. 1"
#             MilestoneCharts(
#                 np.array(final_labels[third: third * 2]),
#                 np.array(self.latest_milestone_dates[third: third * 2]),
#                 np.array(self.last_milestone_dates[third: third * 2]),
#                 np.array(self.baseline_milestone_dates[third: third * 2]),
#                 title,
#                 self.ipdc_date,
#             )
#             title = self.graph_title + " cont. 2"
#             MilestoneCharts(
#                 np.array(final_labels[third * 2: no_milestones]),
#                 np.array(self.latest_milestone_dates[third * 2: no_milestones]),
#                 np.array(self.last_milestone_dates[third * 2: no_milestones]),
#                 np.array(self.baseline_milestone_dates[third * 2: no_milestones]),
#                 title,
#                 self.ipdc_date,
#             )
#         pass
#


def put_milestones_into_wb(milestones: MilestoneData) -> Workbook:
    wb = Workbook()
    ws = wb.active

    row_num = 2
    ms_names = milestones.sorted_milestone_dict[milestones.iter_list[0]]["names"]

    for i, m in enumerate(ms_names):
        for x, tp in enumerate(milestones.iter_list):
            ms_date = milestones.sorted_milestone_dict[tp]["r_dates"][i]
            ws.cell(row=row_num + i, column=5 + x).value = ms_date
            ws.cell(row=row_num + i, column=5 + x).number_format = "dd/mm/yy"

    for i, m in enumerate(ms_names):  # want the latest notes.
        ws.cell(row=row_num + i, column=1).value = milestones.sorted_milestone_dict[
            milestones.iter_list[0]
        ]["project"][
            i
        ]  # project name
        ws.cell(row=row_num + i, column=2).value = milestones.sorted_milestone_dict[
            milestones.iter_list[0]
        ]["report"][
            i
        ]  # milestone
        ws.cell(row=row_num + i, column=3).value = milestones.sorted_milestone_dict[
            milestones.iter_list[0]
        ]["cat"][
            i
        ]  # milestone
        ws.cell(row=row_num + i, column=4).value = milestones.sorted_milestone_dict[
            milestones.iter_list[0]
        ]["names"][
            i
        ]  # milestone
        try:
            notes = milestones.sorted_milestone_dict[milestones.iter_list[0]]["notes"][
                i
            ]
            ws.cell(row=row_num + i, column=len(milestones.iter_list) + 9).value = notes
        except KeyError:  # "notes" not in central support dict
            pass
        try:  # only present in top250 milestones
            status = milestones.sorted_milestone_dict[milestones.iter_list[0]][
                "status"
            ][i]
            ws.cell(
                row=row_num + i, column=len(milestones.iter_list) + 5
            ).value = status
        except (IndexError, KeyError):
            # IndexError for ipdc rpting. "status" is in dict, but list is empty.
            # KeyError for top250 central support, which does not have "status" in dict.
            pass
        try:  # only present in top250 central support
            escalated = milestones.sorted_milestone_dict[milestones.iter_list[0]][
                "escalated"
            ][i]
            cs_type = milestones.sorted_milestone_dict[milestones.iter_list[0]]["type"][
                i
            ]
            cs_response = milestones.sorted_milestone_dict[milestones.iter_list[0]][
                "cs_response"
            ][i]
            secured = milestones.sorted_milestone_dict[milestones.iter_list[0]][
                "secured"
            ][i]
            ws.cell(
                row=row_num + i, column=len(milestones.iter_list) + 6
            ).value = escalated
            ws.cell(
                row=row_num + i, column=len(milestones.iter_list) + 7
            ).value = cs_type
            ws.cell(
                row=row_num + i, column=len(milestones.iter_list) + 8
            ).value = secured
        except KeyError:  # all above not in ipdc or top250 milestones
            pass

    ws.cell(row=1, column=1).value = "Project"
    ws.cell(row=1, column=2).value = "Report"
    ws.cell(row=1, column=3).value = "Type"
    ws.cell(row=1, column=4).value = "Milestone"
    for x, tp in enumerate(milestones.iter_list):
        ws.cell(row=1, column=5 + x).value = tp
    ws.cell(row=1, column=len(milestones.iter_list) + 5).value = "Status (top 250 ms)"
    ws.cell(
        row=1, column=len(milestones.iter_list) + 6
    ).value = "Escalated (top 250 cs)"
    ws.cell(row=1, column=len(milestones.iter_list) + 7).value = "Type (top 250 cs)"
    ws.cell(row=1, column=len(milestones.iter_list) + 8).value = "Secured (top 250 cs)"
    ws.cell(
        row=1, column=len(milestones.iter_list) + 9
    ).value = "Notes / Central Response (top 250 cs)"

    return wb


def vfm_matplotlib_graph(labels, current_qrt, last_qrt, title):
    #  Need to split this strings over two lines on series_one axis
    for n, i in enumerate(labels):
        if i == "Very High and Financially Positive":
            labels[n] = "Very High and \n Financially Positive"
        if i == "Economically Positive":
            labels[n] = "Economically \n Positive"

    x = np.arange(len(labels))  # the label locations
    width = 0.35  # the width of the bars

    fig, ax = plt.subplots()
    rects_one = ax.bar(x - width / 2, current_qrt, width, label="This quarter")
    rects_two = ax.bar(x + width / 2, last_qrt, width, label="Last quarter")

    # Add some text for labels, title and custom series_one-axis tick labels, etc.
    # ax.set_ylabel('Number')
    ax.set_title(title)
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    # Rotate the tick labels and set their alignment.
    # plt.setp(ax.get_xticklabels(), alignment=)
    ax.legend()

    def autolabel(rects):
        """Attach a text label above each bar in *rects*, displaying its height."""
        for rect in rects:
            height = rect.get_height()
            ax.annotate(
                "{}".format(height),
                xy=(rect.get_x() + rect.get_width() / 2, height),
                xytext=(0, 3),  # 3 points vertical offset
                textcoords="offset points",
                ha="center",
                va="bottom",
            )

    autolabel(rects_one)
    autolabel(rects_two)

    fig.tight_layout()

    fig.savefig(root_path / "output/{}.png".format(title), bbox_inches="tight")

    # plt.show()


def set_figure_size(graph_type: str) -> Tuple[int, int]:
    if graph_type == "half_horizontal":
        return 11.69, 5.10
    if graph_type == "full_horizontal":
        return 11.69, 8.20


def cost_profile_into_wb(costs: CostData) -> Workbook:
    wb = Workbook()
    ws = wb.active

    row_num = 2
    for x, tp in enumerate(costs.iter_list):
        ws.cell(row=1, column=1).value = "F/Y"
        ws.cell(row=1, column=2 + x).value = tp
        for i, cv in enumerate(costs.c_profiles[tp]["prof"]):  # cv cost value
            ws.cell(row=row_num + i, column=1).value = YEAR_LIST[i]
            ws.cell(row=row_num + i, column=2 + x).value = cv

    return wb


def set_fig_size(kwargs, fig: plt.figure) -> plt.figure:
    if "fig_size" in kwargs:
        fig.set_size_inches(set_figure_size(kwargs["fig_size"]))
    else:
        fig.set_size_inches(set_figure_size(FIGURE_STYLE[2]))

    return fig


def get_chart_title(
    # master: Master,
    # title_end: str,
    **c_kwargs,  # chart kwargs
) -> str:
    if "title" in c_kwargs:
        title = c_kwargs["title"]
    # elif "group" in c_kwargs:
    #     if set(c_kwargs["group"]) == set(DFT_GROUP):
    #         title = "Portfolio " + title_end
    #     elif set(c_kwargs["group"]) == set(master.current_projects):
    #         title = "Portfolio " + title_end
    #     # elif c.group["group"] == data_class.master.current_projects:
    #     #     title = "Portfolio " + title_end
    #     elif len(c_kwargs["group"]) == 1:
    #         try:
    #             title = (
    #                     master.abbreviations[c_kwargs["group"][0]]["abb"] + " " + title_end
    #             )
    #         except KeyError:
    #             title = c_kwargs["group"][0] + " " + title_end
    #     else:
    #         logger.info("Please provide a title for this chart using --title.")
    #         title = None
    #
    # elif (
    #         "stage" in c_kwargs
    # ):  # not clear when this loop would be used. leaving in for now.,
    #     if set(c_kwargs["stage"]) == set(DFT_STAGE):
    #         title = "Portfolio " + title_end
    #     elif set(c_kwargs["stage"]) == set(master.current_projects):
    #         title = "Portfolio " + title_end
    #     elif len(c_kwargs["stage"]) == 1:
    #         try:
    #             title = (
    #                     master.abbreviations[c_kwargs["stage"][0]]["abb"] + " " + title_end
    #             )
    #         except KeyError:
    #             title = c_kwargs["stage"][0] + " " + title_end
    #     else:
    #         logger.info("Please provide a title for this chart using --title.")
    #         title = None
    else:
        logger.info("Please provide a title for this chart using --title.")
        title = None

    return title


def cost_profile_graph(costs: CostData, master, **kwargs) -> plt.figure:
    """Compiles a matplotlib line chart for costs of GROUP of projects contained within cost_master class"""

    fig, (ax1) = plt.subplots(1)  # two subplots for this chart
    fig = set_fig_size(kwargs, fig)
    title = get_chart_title(master, "cost profile trend", **kwargs)  # title
    plt.suptitle(title, fontweight="bold", fontsize=20)

    # Overall cost profile chart
    for i, iter in enumerate(costs.iter_list):
        ax1.plot(
            YEAR_LIST[:-1],
            np.array(costs.c_profiles[iter]["prof_ra"]),
            label=iter,
            linewidth=5.0,
            marker="o",
            zorder=10 - i,
        )

    # Chart styling
    plt.xticks(rotation=45, size=16)
    plt.yticks(size=16)
    # ax1.tick_params(axis="series_one", which="major")  # matplotlib version issue
    ax1.set_ylabel("Cost (£m)")
    ax1.set_xlabel("Financial Year")
    xlab1 = ax1.xaxis.get_label()
    xlab1.set_style("italic")
    xlab1.set_size(20)
    ylab1 = ax1.yaxis.get_label()
    ylab1.set_style("italic")
    ylab1.set_size(20)
    ax1.grid(color="grey", linestyle="-", linewidth=0.2)
    ax1.legend(prop={"size": 20})

    fig.tight_layout(rect=[0, 0.03, 1, 0.95])  # size/fit of chart

    if "chart" in kwargs:
        if kwargs["chart"]:
            plt.show()

    return fig


def cost_profile_graph_new(costs: CostData, master, **kwargs) -> plt.figure:
    """Compiles a matplotlib line chart for costs of GROUP of projects contained within cost_master class"""

    fig, (ax1) = plt.subplots(1)  # two subplots for this chart
    fig = set_fig_size(kwargs, fig)
    title = get_chart_title(master, "cost profile trend", **kwargs)  # title
    plt.suptitle(title, fontweight="bold", fontsize=20)

    tidy_label = {
        str(master.current_quarter): "Forecast",
        "baseline": "Baseline",
        "current": "Forecast",
    }

    # Overall cost profile chart
    for i, iter in enumerate(list(costs.profiles.keys())):
        try:
            label = tidy_label[iter]
        except KeyError:
            label = iter + " (Forecast)"
        ax1.plot(
            YEAR_LIST,
            np.array(costs.profiles[iter]["total"]),
            label=label,
            linewidth=5.0,
            marker="o",
            zorder=10 - i,
        )

    # Chart styling
    plt.xticks(rotation=45, size=16)
    plt.yticks(size=16)
    # ax1.tick_params(axis="series_one", which="major")  # matplotlib version issue
    ax1.set_ylabel("Cost (£m)")
    ax1.set_xlabel("Financial Year")
    xlab1 = ax1.xaxis.get_label()
    xlab1.set_style("italic")
    xlab1.set_size(20)
    ylab1 = ax1.yaxis.get_label()
    ylab1.set_style("italic")
    ylab1.set_size(20)
    ax1.grid(color="grey", linestyle="-", linewidth=0.2)
    ax1.legend(prop={"size": 20})

    fig.tight_layout(rect=[0, 0.03, 1, 0.95])  # size/fit of chart

    if "chart" in kwargs:
        if kwargs["chart"]:
            plt.show()

    return fig


def cost_profile_baseline_graph(
    cost_master: CostData, *title: Tuple[Optional[str]]
) -> plt.figure:
    """Compiles a matplotlib line chart for costs of GROUP of projects contained within cost_master class.
    As as default last quarters profile is not included. It creates two plots. First plot shows overall
    profile in current, last quarters anb baseline form. Second plot shows rdel, cdel, and 'non-gov' cost profile"""

    fig, (ax1, ax2) = plt.subplots(2)  # two subplots for this chart

    """cost profile charts"""
    if len(cost_master.entity) == 1:
        fig.suptitle(cost_master.entity[0] + " Cost Profile", fontweight="bold")
    else:
        fig.suptitle(title[0] + " Cost Profile", fontweight="bold")  # title

    # Overall cost profile chart
    if (
        sum(cost_master.baseline_profile_three) != 0
    ):  # handling in the event that group of projects have no baseline profile.
        ax1.plot(
            YEAR_LIST,
            np.array(cost_master.baseline_profile_three),  # baseline profile
            label="Baseline 3",
            linewidth=3.0,
            marker="o",
        )
    else:
        pass
    if (
        sum(cost_master.baseline_profile_two) != 0
    ):  # handling in the event that group of projects have no baseline profile.
        ax1.plot(
            YEAR_LIST,
            np.array(cost_master.baseline_profile_two),  # baseline profile
            label="Baseline 2",
            linewidth=3.0,
            marker="o",
        )
    else:
        pass
    if (
        sum(cost_master.baseline_profile_one) != 0
    ):  # handling in the event that group of projects have no last quarter profile
        ax1.plot(
            YEAR_LIST,
            np.array(cost_master.baseline_profile_one),  # last quarter profile
            label="Baseline 1",
            linewidth=3.0,
            marker="o",
        )
    else:
        pass
    ax1.plot(
        YEAR_LIST,
        np.array(cost_master.current_profile),  # current profile
        label="Latest",
        linewidth=3.0,
        marker="o",
    )

    # Chart styling
    ax1.tick_params(axis="series_one", which="major", labelsize=6, rotation=45)
    ax1.set_ylabel("Cost (£m)")
    ylab1 = ax1.yaxis.get_label()
    ylab1.set_style("italic")
    ylab1.set_size(8)
    ax1.grid(color="grey", linestyle="-", linewidth=0.2)
    ax1.legend(prop={"size": 6})
    ax1.set_title(
        "Fig 1 - cost profile changes", loc="left", fontsize=8, fontweight="bold"
    )

    # plot rdel, cdel, non-gov chart data
    if (
        sum(cost_master.ngov_profile) != 0
    ):  # if statement as most projects don't have ngov cost.
        ax2.plot(
            YEAR_LIST,
            np.array(cost_master.ngov_profile),
            label="Non-Gov",
            linewidth=3.0,
            marker="o",
        )
    ax2.plot(
        YEAR_LIST,
        np.array(cost_master.cdel_profile),
        label="CDEL",
        linewidth=3.0,
        marker="o",
    )
    ax2.plot(
        YEAR_LIST,
        np.array(cost_master.rdel_profile),
        label="RDEL",
        linewidth=3.0,
        marker="o",
    )

    # rdel/cdel profile chart styling
    ax2.tick_params(axis="series_one", which="major", labelsize=6, rotation=45)
    ax2.set_xlabel("Financial Years")
    ax2.set_ylabel("Cost (£m)")
    xlab2 = ax2.xaxis.get_label()
    ylab2 = ax2.yaxis.get_label()
    xlab2.set_style("italic")
    xlab2.set_size(8)
    ylab2.set_style("italic")
    ylab2.set_size(8)
    ax2.grid(color="grey", linestyle="-", linewidth=0.2)
    ax2.legend(prop={"size": 6})
    ax2.set_title(
        "Fig 2 - current cost type profile", loc="left", fontsize=8, fontweight="bold"
    )

    # plt.show()

    return fig


def spent_calculation(
    master: Dict[str, Union[str, datetime.date, int, float]], project: str
) -> int:
    keys = [
        "Pre-profile RDEL",
        "20-21 RDEL STD Total",
        "Pre-profile CDEL Forecast one off new costs",
        "20-21 CDEL STD Total",
        "Pre-profile Forecast Non-Gov",
        "20-21 CDEL STD Non Gov costs",
    ]

    total = 0
    for k in keys:
        try:
            total += master[project][k]
        except TypeError:  # None types
            pass

    return total


def open_word_doc(wd_path: str) -> Document:
    """Function stores an empty word doc as a variable"""
    return Document(wd_path)


def get_word_doc() -> Document():
    """returns the summary temp doc"""
    wd_path = root_path / "input/summary_temp.docx"
    return open_word_doc(wd_path)


def key_contacts(doc: Document, master, **kwargs) -> None:
    data = master.master_data[0]["data"][kwargs["full_name"]]
    """Function adds keys contact details"""
    sro_name = data["Senior Responsible Owner (SRO)"]
    if sro_name is None:
        sro_name = "tbc"

    sro_email = data["Senior Responsible Owner (SRO) - Email"]
    if sro_email is None:
        sro_email = "email: tbc"

    sro_phone = data["SRO Phone No."]
    if sro_phone == None:
        sro_phone = "phone number: tbc"

    doc.add_paragraph(
        "SRO: " + str(sro_name) + ", " + str(sro_email) + ", " + str(sro_phone)
    )

    pd_name = data["Project Director (PD)"]
    if pd_name is None:
        pd_name = "TBC"

    pd_email = data["Project Director (PD) - Email"]
    if pd_email is None:
        pd_email = "email: tbc"

    pd_phone = data["PD Phone No."]
    if pd_phone is None:
        pd_phone = "phone: tbc"

    doc.add_paragraph(
        "PD: " + str(pd_name) + ", " + str(pd_email) + ", " + str(pd_phone)
    )

    # contact_name = data["Working Contact Name"]
    # if contact_name is None:
    #     contact_name = "tbc"

    # contact_email = data["Working Contact Email"]
    # if contact_email is None:
    #     contact_email = "email: tbc"
    #
    # contact_phone = data["Working Contact Telephone"]
    # if contact_phone is None:
    #     contact_phone = "phone: tbc"
    #
    # doc.add_paragraph(
    #     "PfM reporting lead: "
    #     + str(contact_name)
    #     + ", "
    #     + str(contact_email)
    #     + ", "
    #     + str(contact_phone)
    # )


def dca_table(doc: Document, master, **kwargs) -> None:
    """Creates SRO confidence table"""

    # doc.add_paragraph()
    p = doc.add_paragraph()
    text = "* Note in Q2 2021/22 DCA ratings moved to a three point scale."
    p.add_run(text)
    # .font.color.rgb = RGBColor(255, 0, 0)

    w_table = doc.add_table(rows=1, cols=5)
    hdr_cells = w_table.rows[0].cells
    hdr_cells[0].text = "Delivery confidence"
    hdr_cells[1].text = "This quarter"
    hdr_cells[2].text = str(master.master_data[1]["quarter"])
    hdr_cells[3].text = str(master.master_data[2]["quarter"])
    hdr_cells[4].text = str(master.master_data[3]["quarter"])

    for x, dca_key in enumerate(list(SRO_CONF_KEY_LIST.keys())):
        row_cells = w_table.add_row().cells
        row_cells[0].text = SRO_CONF_KEY_LIST[dca_key]
        for i, m in enumerate(master.master_data[:4]):  # last four masters taken
            try:
                rating = convert_rag_text(m["data"][kwargs["full_name"]][dca_key])
                row_cells[i + 1].text = rating
                cell_colouring(row_cells[i + 1], rating)
            except (KeyError, TypeError):
                row_cells[i + 1].text = "N/A"

    w_table.style = "Table Grid"
    make_rows_bold([w_table.rows[0]])  # makes top of table bold.
    # make_columns_bold([table.columns[0]]) #right cells in table bold
    column_widths = (Cm(4.4), Cm(2.8), Cm(2.8), Cm(2.8), Cm(2.8))
    set_col_widths(w_table, column_widths)


def dca_narratives(doc: Document, master, **kwargs) -> None:
    """Places all narratives into document and checks for differences between
    current and last quarter"""

    # doc.add_paragraph()
    # p = doc.add_paragraph()
    # text = "*Red text highlights changes in narratives from last quarter"
    # p.add_run(text).font.color.rgb = RGBColor(255, 0, 0)

    headings_list = [
        "SRO delivery confidence narrative",
        "Financial cost narrative",
        "Financial comparison with last quarter",
        "Financial comparison with baseline",
        "Benefits Narrative",
        "Benefits comparison with last quarter",
        "Benefits comparison with baseline",
        "Milestone narrative",
    ]

    narrative_keys_list = [
        "Departmental DCA Narrative",
        "Project Costs Narrative",
        "Cost comparison with last quarters cost narrative",
        "Cost comparison within this quarters cost narrative",
        "Benefits Narrative",
        "Ben comparison with last quarters cost - narrative",
        "Ben comparison within this quarters cost - narrative",
        "Milestone Commentary",
    ]

    if kwargs["type"] == "short":
        headings_list = headings_list[:1]

    for x in range(len(headings_list)):
        try:  # overall try statement relates to data_bridge
            text_one = str(
                master.master_data[0]["data"][kwargs["full_name"]][
                    narrative_keys_list[x]
                ]
            )
            try:
                text_two = str(
                    master.master_data[1]["data"][kwargs["full_name"]][
                        narrative_keys_list[x]
                    ]
                )
            except (KeyError, IndexError):  # index error relates to data_bridge
                text_two = text_one
        except KeyError:
            break

        doc.add_paragraph().add_run(str(headings_list[x])).bold = True

        # There are two options here for comparing text. Have left this for now.
        # compare_text_showall(dca_a, dca_b, doc)
        compare_text_new_and_old(text_one, text_two, doc)


def forward_look(doc: Document, master, **kwargs) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    text = "*Red text highlights changes in narratives from last quarter"
    p.add_run(text).font.color.rgb = RGBColor(255, 0, 0)

    fwd_look = str(
        master.master_data[0]["data"][kwargs["full_name"]][
            "SRO Forward Look Assessment"
        ]
    )

    doc.add_paragraph().add_run(
        "SRO Forward Look Assessment: " + str(fwd_look).upper()
    ).bold = True


def forward_look_narrative(doc: Document, master, **kwargs) -> None:
    headings_list = [
        "SRO Forward Look Narrative",
    ]

    narrative_keys_list = ["SRO Forward Look Narrative"]

    for x in range(len(headings_list)):
        try:  # overall try statement relates to data_bridge
            text_one = str(
                master.master_data[0]["data"][kwargs["full_name"]][
                    narrative_keys_list[x]
                ]
            )
            try:
                text_two = str(
                    master.master_data[1]["data"][kwargs["full_name"]][
                        narrative_keys_list[x]
                    ]
                )
            except (KeyError, IndexError):  # index error relates to data_bridge
                text_two = text_one
        except KeyError:
            break

        doc.add_paragraph().add_run(str(headings_list[x])).bold = True

        # There are two options here for comparing text. Have left this for now.
        # compare_text_showall(dca_a, dca_b, doc)
        compare_text_new_and_old(text_one, text_two, doc)


def convert_rag_text(dca_rating: str) -> str:
    """Converts RAG name into a acronym"""

    if dca_rating == "Green":
        return "G"
    elif dca_rating == "Amber/Green":
        return "A/G"
    elif dca_rating == "Amber":
        return "A"
    elif dca_rating == "Amber/Red":
        return "A/R"
    elif dca_rating == "Red":
        return "R"
    else:
        return ""


def set_col_widths(word_table: table, widths: list) -> None:
    """This function sets the width of table in a word document"""
    for row in word_table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def compare_text_new_and_old(text_1: str, text_2: str, doc: Document) -> None:
    """compares two sets of text and highlights differences in red text."""

    comp = difflib.Differ()
    diff = list(comp.compare(text_2.split(), text_1.split()))
    diff = [x for x in diff if x[0] != "-"]  # remove all deleted text
    diff = [x for x in diff if x[0] != "?"]  # remove ?. not sure what these represent.
    y = doc.add_paragraph()
    for i, text in enumerate(diff):
        # f = len(diff) - 1
        # if i < f:
        #     a = i - 1
        # else:
        #     a = i

        if text[0:3] == "  |" or text[0:3] == "+ |":
            # j = i + 1
            # if diff[i][:3] == "  |"
            # if diff[i][0:3] and diff[a][0:3] == "  |":
            y = doc.add_paragraph()
            # else:
            #     pass
        # elif text[0:3] == "+ |":
        #     if diff[i][0:3] and diff[a][0:3] == "+ |":
        #         y = doc.add_paragraph()
        #     else:
        #         pass
        # if text[0:3] == "- |":
        #     pass
        # if text[0:3] == "  -":
        #     y = doc.add_paragraph()
        #     g = diff[i][2]
        #     y.add_run(g)
        # elif text[0:3] == "  •":
        #     y = doc.add_paragraph()
        #     g = text[2]
        #     y.add_run(g)
        if text[0] == " ":
            # total_nc += 1
            if i == 0:
                y.add_run(text[2:])
            # if total_nc == 1:
            #     y.add_run(text[2:])
            else:
                y.add_run(text[1:])
        if text[0] == "+":
            # w = len(diff[i])
            # g = diff[i][1:w]
            # total_plus += 1  # new text might not be first
            if i == 0:
                y.add_run(text[2:]).font.color.rgb = RGBColor(255, 0, 0)
            else:
                y.add_run(text[1:]).font.color.rgb = RGBColor(255, 0, 0)
        # if diff[i][0] == "-":
        #     pass
        # if diff[i][0] == "?":
        #     pass
        # else:
        #     if diff[i] != "+ |":
        #         y.add_run(diff[i][1:])


def make_file_friendly(quarter_str: str) -> str:
    """Converts datamaps.api project_data_from_master quarter data into a string to use when
    saving output files. Courtesy of M Lemon."""
    regex = r"Q(\d) (\d+)\/(\d+)"
    return re.sub(regex, r"Q\1_\2_\3", quarter_str)


def total_costs_benefits_bar_chart(
    costs: CostData, ben: BenefitsData, master, **kwargs
) -> plt.figure:
    """compiles a matplotlib bar chart which shows total project costs"""
    fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2)  # four sub plots

    try:
        fig_size = kwargs["fig_size"]
        fig.set_size_inches(set_figure_size(fig_size))
    except KeyError:
        fig.set_size_inches(set_figure_size(FIGURE_STYLE[2]))
        # pass

    title = get_chart_title(master, " totals", **kwargs)
    plt.suptitle(title, fontweight="bold", fontsize=25)
    plt.xticks(size=12)
    plt.yticks(size=10)

    # Y AXIS SCALE MAX

    highest_int = max(
        [
            costs.c_totals[costs.iter_list[0]]["total"],
            ben.b_totals[ben.iter_list[0]]["total"],
        ]
    )
    y_max = highest_int + percentage(5, highest_int)
    ax1.set_ylim(0, y_max)

    # COST SPENT, PROFILED AND UNPROFILED
    labels = costs.iter_list
    spent = []
    prof = []
    unprof = []
    for x in labels:
        spent.append(costs.c_totals[x]["spent"])
        prof.append(costs.c_totals[x]["prof"])
        unprof.append(costs.c_totals[x]["unprof"])

    width = 0.5
    ax1.bar(labels, np.array(spent), width, label="Spent")
    ax1.bar(
        labels,
        np.array(prof),
        width,
        bottom=np.array(spent),
        label="Profiled",
    )
    ax1.bar(
        labels,
        np.array(unprof),
        width,
        bottom=np.array(spent) + np.array(prof),
        label="Unprofiled",
    )
    ax1.legend(prop={"size": 10})
    ax1.xaxis.set_tick_params(labelsize=12)
    ax1.yaxis.set_tick_params(labelsize=12)
    ax1.set_ylabel("Cost (£m)")
    ylab1 = ax1.yaxis.get_label()
    ylab1.set_style("italic")
    ylab1.set_size(12)
    # ax1.tick_params(axis="series_one", which="major")
    # ax1.tick_params(axis="series_two", which="major")
    ax1.set_title(
        "Fig 1. Change in total cost",
        loc="left",
        fontsize=12,
        fontweight="bold",
    )

    # RDEL, CDEL AND NON-GOV TOTALS BAR CHART
    labels = ["RDEL", "CDEL", "Non Gov"]
    width = 0.5
    ax2.bar(
        labels,
        np.array(costs.c_totals[costs.iter_list[0]]["cat_spent"]),
        width,
        label="Spent",
    )
    ax2.bar(
        labels,
        np.array(costs.c_totals[costs.iter_list[0]]["cat_prof"]),
        width,
        bottom=np.array(costs.c_totals[costs.iter_list[0]]["cat_spent"]),
        label="Profiled",
    )
    ax2.bar(
        labels,
        np.array(costs.c_totals[costs.iter_list[0]]["cat_unprof"]),
        width,
        bottom=np.array(costs.c_totals[costs.iter_list[0]]["cat_spent"])
        + np.array(costs.c_totals[costs.iter_list[0]]["cat_prof"]),
        label="Unprofiled",
    )
    ax2.legend(prop={"size": 10})
    ax2.xaxis.set_tick_params(labelsize=12)
    ax2.yaxis.set_tick_params(labelsize=12)
    ax2.set_ylabel("Costs (£m)")
    ylab3 = ax2.yaxis.get_label()
    ylab3.set_style("italic")
    ylab3.set_size(12)
    # ax2.tick_params(axis="series_one", which="major", labelsize=6)
    # ax2.tick_params(axis="series_two", which="major", labelsize=6)
    ax2.set_title(
        "Fig 2. Current cost breakdown",
        loc="left",
        fontsize=12,
        fontweight="bold",
    )

    ax2.set_ylim(0, y_max)  # scale series_two axis max

    # BENEFITS SPENT, PROFILED AND UNPROFILED
    labels = ben.iter_list
    delivered = []
    prof = []
    unprof = []
    for x in labels:
        delivered.append(ben.b_totals[x]["delivered"])
        prof.append(ben.b_totals[x]["prof"])
        unprof.append(ben.b_totals[x]["unprof"])

    width = 0.5
    ax3.bar(labels, np.array(delivered), width, label="delivered")
    ax3.bar(
        labels,
        np.array(prof),
        width,
        bottom=np.array(delivered),
        label="profiled",
    )
    ax3.bar(
        labels,
        np.array(unprof),
        width,
        bottom=np.array(delivered) + np.array(prof),
        label="unprofiled",
    )
    ax3.legend(prop={"size": 10})
    ax3.xaxis.set_tick_params(labelsize=12)
    ax3.yaxis.set_tick_params(labelsize=12)
    ax3.set_ylabel("Benefits (£m)")
    ylab3 = ax3.yaxis.get_label()
    ylab3.set_style("italic")
    ylab3.set_size(12)
    # ax3.tick_params(axis="series_one", which="major", labelsize=6)
    # ax3.tick_params(axis="series_two", which="major", labelsize=6)
    ax3.set_title(
        "Fig 3. Change in total benefits",
        loc="left",
        fontsize=12,
        fontweight="bold",
    )

    ax3.set_ylim(0, y_max)

    # BENEFITS CASHABLE, NON-CASHABLE, ECONOMIC, DISBENEFIT BAR CHART
    labels = ["Cashable", "Non-Cashable", "Economic", "Disbenefit"]
    width = 0.5
    ax4.bar(
        labels,
        np.array(ben.b_totals[ben.iter_list[0]]["cat_spent"]),
        width,
        label="Delivered",
    )
    ax4.bar(
        labels,
        np.array(ben.b_totals[ben.iter_list[0]]["cat_prof"]),
        width,
        bottom=np.array(ben.b_totals[ben.iter_list[0]]["cat_spent"]),
        label="Profiled",
    )
    ax4.bar(
        labels,
        np.array(ben.b_totals[ben.iter_list[0]]["cat_unprof"]),
        width,
        bottom=np.array(ben.b_totals[ben.iter_list[0]]["cat_spent"])
        + np.array(ben.b_totals[ben.iter_list[0]]["cat_prof"]),
        label="Unprofiled",
    )
    ax4.legend(prop={"size": 10})
    ax4.xaxis.set_tick_params(labelsize=12)
    ax4.yaxis.set_tick_params(labelsize=12)
    ax4.set_ylabel("Benefits (£m)")
    ylab4 = ax4.yaxis.get_label()
    ylab4.set_style("italic")
    ylab4.set_size(12)
    # ax4.tick_params(axis="series_one", which="major", labelsize=6)
    # ax4.tick_params(axis="series_two", which="major", labelsize=6)
    ax4.set_title(
        "Fig 4. Current benefit breakdown",
        loc="left",
        fontsize=12,
        fontweight="bold",
    )

    check_min = (
        ben.b_totals[ben.iter_list[0]]["cat_spent"][3]
        + ben.b_totals[ben.iter_list[0]]["cat_prof"][3]
        + ben.b_totals[ben.iter_list[0]]["cat_unprof"][3]
    )

    if check_min == 0:
        ax4.set_ylim(0, y_max)
    else:  # for negative benefits
        y_min = check_min + percentage(40, check_min)  # arbitrary 40 percent
        ax4.set_ylim(y_min, y_max + abs(y_min))

    fig.tight_layout(rect=[0, 0.03, 1, 0.95])  # size/fit of chart

    if "chart" in kwargs:
        if kwargs["chart"]:
            plt.show()

    return fig


def check_baselines(master) -> None:
    """checks that projects have the correct baseline information. stops the
    programme if baselines are missing"""

    for v in BASELINE_TYPES.values():
        for p in master.current_projects:
            baselines = master.bl_index[v][p]
            if len(baselines) <= 2:
                print(
                    p
                    + " does not have a baseline point for "
                    + v
                    + " this could cause the programme to"
                    "crash. Therefore the programme is stopping. "
                    "Please amend the data for " + p + " so that "
                    " it has at least one baseline point for " + v
                )
                break
        else:
            continue
        break


def percentage(percent: int, whole: float) -> int:
    return round((percent * whole) / 100.0)


def get_old_fy_cost_data(
    master_file: typing.TextIO, project_id_wb: typing.TextIO
) -> None:
    """
    Gets all old financial data from a specified master and places into project id document.
    """
    master = project_data_from_master(
        master_file, 1, 2010
    )  # random year specified as not in use
    wb = load_workbook(project_id_wb)
    ws = wb.active

    for i in range(1, ws.max_column + 1):
        project_name = ws.cell(row=1, column=1 + i).value
        for row_num in range(2, ws.max_row + 1):
            key = ws.cell(row=row_num, column=1).value
            try:
                if key in master.data[project_name].keys():
                    ws.cell(row=row_num, column=1 + i).value = master.data[
                        project_name
                    ][key]
            except KeyError:  # project might not be present in quarter
                pass

    wb.save(project_id_wb)


def run_get_old_fy_data(master_files_list: list, project_id_wb: typing.TextIO) -> None:
    for f in reversed(
        master_files_list
    ):  # reversed so it gets the latest data in masters
        get_old_fy_cost_data(f, project_id_wb)


def place_old_fy_data_into_master_wb(
    master_file: typing.TextIO, project_id_wb: typing.TextIO
) -> None:
    """
    places all old financial year data into master files.
    """
    id_master = project_data_from_master(
        project_id_wb, 2, 2020
    )  # random year specify as not used
    wb = load_workbook(master_file)
    ws = wb.active

    for i in range(1, ws.max_column + 1):
        project_name = ws.cell(row=1, column=1 + i).value
        for row_num in range(2, ws.max_row + 1):
            key = ws.cell(row=row_num, column=1).value
            try:
                if key in id_master.data[project_name].keys():
                    ws.cell(row=row_num, column=1 + i).value = id_master.data[
                        project_name
                    ][key]
            except KeyError:  # project might not be present in quarter
                pass

    wb.save(master_file)


def run_place_old_fy_data_into_masters(
    master_files_list: list, project_id_wb: typing.TextIO
) -> None:
    for f in master_files_list:
        place_old_fy_data_into_master_wb(f, project_id_wb)


def put_key_change_master_into_dict(key_change_file: typing.TextIO) -> Dict[str, str]:
    """
    places keys information i.e. keys old and new names from wb into a python dictionary
    """
    wb = load_workbook(key_change_file)
    ws = wb.active

    output_dict = {}
    for x in range(1, ws.max_row + 1):
        key = ws.cell(row=x, column=1).value
        codename = ws.cell(row=x, column=2).value
        output_dict[key] = codename

    return output_dict


def alter_wb_master_file_key_names(
    master_file: typing.TextIO, key_change_dict: Dict[str, str]
) -> workbook:
    """
    places altered keys names, from the keys change master dictionary, into master wb(s).
    """
    wb = load_workbook(master_file)
    ws = wb.active

    for row_num in range(2, ws.max_row + 1):
        for (
            key
        ) in key_change_dict.keys():  # changes stored in the altered keys change log wb
            if ws.cell(row=row_num, column=1).value == key:
                ws.cell(row=row_num, column=1).value = key_change_dict[key]
        for year in YEAR_LIST:  # changes to yearly profile keys
            if ws.cell(row=row_num, column=1).value == year + " CDEL Forecast Total":
                ws.cell(row=row_num, column=1).value = (
                    year + " CDEL Forecast one off new costs"
                )

    return wb.save(master_file)


def run_change_keys(master_files_list: list, key_dict: Dict[str, str]) -> None:
    """
    runs code which replaces old keys names with new names in master excel workbooks.
    """
    for f in master_files_list:
        alter_wb_master_file_key_names(f, key_dict)


# keys_dict = put_key_change_master_into_dict(root_path / "core_data/analysis_engine/keys_to_change.xlsx")
# run_change_keys(get_master_data_file_paths_fy_16_17(), keys_dict)


def string_conversion(name):
    if isinstance(name, str):
        return [name]
    else:
        return name


def compare_masters(files: List[typing.TextIO], projects: List[str] or str) -> workbook:
    """Takes two masters and compares if there have been any changes to project data values,
    as well as any new data keys.
    files = list of file paths to masters. Only last two masters are used
    projects = list of those projects that require data checking"""

    projects = string_conversion(projects)

    last_master = project_data_from_master(files[1], 4, 2090)

    wb = load_workbook(files[0])
    ws = wb.active

    project_count = []
    change_count = 0
    new_key_count = 0

    for row_num in range(2, ws.max_row + 1):
        key = ws.cell(row=row_num, column=1).value
        for column_num in range(2, ws.max_column + 1):
            project_name = ws.cell(row=1, column=column_num).value
            if project_name not in projects:
                pass
            else:
                wb_value = ws.cell(row=row_num, column=column_num).value
                try:
                    dict_value = last_master.data[project_name][key]
                    if wb_value == dict_value:
                        pass
                    else:
                        ws.cell(row=row_num, column=column_num).fill = PatternFill(
                            start_color="ffba00", end_color="ffba00", fill_type="solid"
                        )
                        project_count.append(project_name)
                        change_count += 1
                except KeyError:
                    if (
                        project_name in last_master.projects
                    ):  # keys error due to keys not being present.
                        ws.cell(row=row_num, column=1).fill = PatternFill(
                            start_color="ffba00", end_color="ffba00", fill_type="solid"
                        )
                    else:  # keys error due to project not being present.
                        ws.cell(row=1, column=column_num).fill = PatternFill(
                            start_color="ffba00", end_color="ffba00", fill_type="solid"
                        )
    # separate lop to calculate this
    for row_num in range(2, ws.max_row + 1):
        key = ws.cell(row=row_num, column=1).value
        if key not in list(last_master.data[projects[0]].keys()):
            new_key_count += 1

    count_ws = wb.create_sheet("Count", 1)

    count_ws.cell(row=1, column=1).value = "No of changes"
    count_ws.cell(row=1, column=2).value = change_count
    count_ws.cell(row=2, column=1).value = "No of new keys"
    count_ws.cell(row=2, column=2).value = new_key_count

    project_count = Counter(project_count)
    for i, p in enumerate(project_count.keys()):
        count_ws.cell(row=i + 3, column=1).value = p
        count_ws.cell(row=i + 3, column=2).value = project_count[p]

    return wb


def totals_chart(costs: CostData, benefits: BenefitsData, **kwargs) -> None:
    """Small function to hold together code to create and save a total_costs_benefits_bar_chart"""
    if kwargs == {}:
        f = total_costs_benefits_bar_chart(costs, benefits)
        f.savefig(root_path / "output/{}_profile.png".format(costs.project_group[0]))
    else:
        if list(kwargs.keys()) == ["title", "fig_size"]:
            f = total_costs_benefits_bar_chart(
                costs, benefits, fig_size=kwargs["fig_size"], title=kwargs["title"]
            )
            f.savefig(root_path / "output/{}_profile.png".format(str(kwargs["title"])))
        if list(kwargs.keys()) == ["fig_size", "title"]:
            f = total_costs_benefits_bar_chart(
                costs, benefits, fig_size=kwargs["fig_size"], title=kwargs["title"]
            )
            f.savefig(root_path / "output/{}_profile.png".format(str(kwargs["title"])))
        if list(kwargs.keys()) == ["title"]:
            f = total_costs_benefits_bar_chart(costs, benefits, title=kwargs["title"])
            f.savefig(root_path / "output/{}_profile.png".format(str(kwargs["title"])))
        if list(kwargs.keys()) == ["fig_size"]:
            f = total_costs_benefits_bar_chart(
                costs, benefits, fig_size=kwargs["fig_size"]
            )
            f.savefig(
                root_path / "output/{}_profile.png".format(costs.project_group[0])
            )


def standard_profile(costs: CostData, **kwargs):
    """Small function to hold together code to create and save a cost_profile_graph"""
    if kwargs == {}:
        f = cost_profile_graph(costs)
        f.savefig(root_path / "output/{}_profile.png".format(costs.project_group[0]))
    else:
        if list(kwargs.keys()) == ["title", "fig_size"]:
            f = cost_profile_graph(
                costs, fig_size=kwargs["fig_size"], title=kwargs["title"]
            )
            f.savefig(root_path / "output/{}_profile.png".format(str(kwargs["title"])))
        if list(kwargs.keys()) == ["fig_size", "title"]:
            f = cost_profile_graph(
                costs, fig_size=kwargs["fig_size"], title=kwargs["title"]
            )
            f.savefig(root_path / "output/{}_profile.png".format(str(kwargs["title"])))
        if list(kwargs.keys()) == ["title"]:
            f = cost_profile_graph(costs, title=kwargs["title"])
            f.savefig(root_path / "output/{}_profile.png".format(str(kwargs["title"])))
        if list(kwargs.keys()) == ["fig_size"]:
            f = cost_profile_graph(costs, fig_size=kwargs["fig_size"])
            f.savefig(
                root_path / "output/{}_profile.png".format(costs.project_group[0])
            )


def save_graph(fig: plt.figure, file_name: str, **kwargs) -> None:
    """Generic function for saving matplotlib figure into a word document"""
    if "orientation" in list(kwargs.keys()):
        if kwargs["orientation"] == "landscape":
            fig.savefig("temp_file.png")
            doc = open_word_doc(root_path / "input/summary_temp_landscape.docx")
            doc.add_picture("temp_file.png", width=Inches(8))
            doc.save(root_path / "output/{}.docx".format(file_name))
            os.remove("temp_file.png")
        if kwargs["orientation"] == "portrait":
            fig.savefig("temp_file.png")
            doc = open_word_doc(root_path / "input/summary_temp.docx")
            doc.add_picture("temp_file.png", width=Inches(8))
            doc.save(root_path / "output/{}.docx".format(file_name))
            os.remove("temp_file.png")
    else:
        fig.savefig("temp_file.png")
        doc = open_word_doc(root_path / "input/summary_temp_landscape.docx")
        doc.add_picture("temp_file.png", width=Inches(8))
        doc.save(root_path / "output/{}.docx".format(file_name))
        os.remove("temp_file.png")


# from stackoverflow.
def do_mask(x: List[datetime.date], y: List[datetime.date]):
    """
    helper function for putting series of datetime.date values with NoneType into
    matplotlib.
    """
    mask = None
    mask = ~(x == None)
    return np.array(x)[mask], np.array(y)[mask]


def handle_long_keys(key_names: List[str]) -> List[str]:
    # helper function for milestone chart
    labels = ["\n".join(wrap(l, 40)) for l in key_names]
    final_labels = []
    for l in labels:
        if len(l) > 70:
            final_labels.append(l[:70])
        else:
            final_labels.append(l)
    return final_labels


def milestone_chart(
    milestones: MilestoneData,
    master,
    **kwargs,
) -> plt.figure:
    fig, ax1 = plt.subplots()
    fig = set_fig_size(kwargs, fig)

    title = get_chart_title(master, "schedule", **kwargs)
    plt.suptitle(title, fontweight="bold", fontsize=20)

    project = milestones.sorted_milestone_dict[milestones.iter_list[0]]["project"]
    ms_names = milestones.sorted_milestone_dict[milestones.iter_list[0]]["names"]

    def merge_project_milestone_name(no):
        return project[no] + ", " + ms_names[no]

    combined = [merge_project_milestone_name(i) for i, p in enumerate(project)]

    # if len(milestones.group) == 1:
    #     pn = milestones.master.abbreviations[milestones.group[0]]["abb"]
    #     ms_names = remove_project_name_from_milestone_key(pn, ms_names)

    ms_names = handle_long_keys(combined)

    # for i in reversed(milestones.iter_list):
    #     ax1.scatter(
    #         milestones.sorted_milestone_dict[i]["g_dates"],
    #         ms_names,
    #         label=i,
    #         s=200,
    #     )

    # alpha_list = [1, 0.8, 0.6, 0.4, 0.2]
    # c_list = ["#436f70", "#4c7e80", "#5f9ea0"]

    for i, tp in enumerate(milestones.iter_list):
        m = [
            x for x in milestones.sorted_milestone_dict[tp]["g_dates"] if x is not None
        ]
        ax1.scatter(
            m,
            # milestones.sorted_milestone_dict[tp]["g_dates"],
            ms_names[0 : len(m)],
            label=tp,
            s=200,
            # color=c_list[i],
            # alpha=alpha_list[i],
            zorder=20 - i,
        )

    ax1.legend(prop={"size": 14})  # insert legend
    plt.yticks(size=10)
    # reverse series_two axis so order is earliest to oldest
    ax1 = plt.gca()
    ax1.set_ylim(ax1.get_ylim()[::-1])
    ax1.yaxis.grid()  # horizontal lines
    ax1.set_axisbelow(True)

    # ax1.scatter(*do_mask(milestone_data.md_current, milestone_data.key_names), label="Current", zorder=10, c='g')
    # ax1.scatter(*do_mask(milestone_data.md_last, milestone_data.key_names), label="Last quarter", zorder=5, c='orange')
    # ax1.scatter(*do_mask(milestone_data.md_baseline, milestone_data.key_names), label="Baseline", zorder=1, c='b')

    years = mdates.YearLocator()  # every year
    months = mdates.MonthLocator()  # every month
    weeks = mdates.WeekdayLocator()
    years_fmt = mdates.DateFormatter("%Y")
    months_fmt = mdates.DateFormatter("%b")
    weeks_fmt = mdates.DateFormatter("%d")

    max_date = calculate_max_min_date(milestones, value="max")
    min_date = calculate_max_min_date(milestones, value="min")
    td = (max_date - min_date).days
    if td >= 365 * 3:
        ax1.xaxis.set_major_locator(years)
        ax1.xaxis.set_minor_locator(months)
        ax1.xaxis.set_major_formatter(years_fmt)
        plt.setp(ax1.xaxis.get_minorticklabels(), rotation=45, size=12)
        plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, weight="bold", size=14)
    elif 365 * 3 >= td >= 90:
        ax1.xaxis.set_major_locator(years)
        ax1.xaxis.set_minor_locator(months)
        ax1.xaxis.set_major_formatter(years_fmt)
        ax1.xaxis.set_minor_formatter(months_fmt)
        plt.setp(ax1.xaxis.get_minorticklabels(), rotation=45, size=12)
        plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, weight="bold", size=14)
    else:
        ax1.xaxis.set_major_locator(months)
        ax1.xaxis.set_minor_locator(weeks)
        ax1.xaxis.set_major_formatter(months_fmt)
        ax1.xaxis.set_minor_formatter(weeks_fmt)
        plt.setp(ax1.xaxis.get_minorticklabels(), rotation=45, size=12)
        plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, weight="bold", size=14)

    if "show_keys" in kwargs:
        if kwargs["show_keys"] == "no":
            ax1.get_yaxis().set_visible(False)

    # Add line of analysis_engine date, but only if in the time period
    if "blue_line" in kwargs:
        blue_line = kwargs["blue_line"]
        if blue_line == "Today":
            if min_date <= datetime.date.today() <= max_date:
                plt.axvline(datetime.date.today())
                plt.figtext(
                    0.98,
                    0.01,
                    "Line represents date chart compiled",
                    horizontalalignment="right",
                    fontsize=10,
                    fontweight="bold",
                )
        elif blue_line == "IPDC":
            IPDC_DATE = parser.parse(
                get_ipdc_date(
                    str(root_path) + "/core_data/ipdc_config.ini", "ipdc_date"
                ),
                dayfirst=True,
            ).date()
            if min_date <= IPDC_DATE <= max_date:
                plt.axvline(IPDC_DATE)
                plt.figtext(
                    0.98,
                    0.01,
                    "Line represents IPDC date",
                    horizontalalignment="right",
                    fontsize=10,
                    fontweight="bold",
                )
        elif blue_line == "CDG":
            CDG_DATE = parser.parse(
                get_ipdc_date(
                    str(cdg_root_path) + "/core_data/cdg_config.ini", "cdg_date"
                ),
                dayfirst=True,
            ).date()
            if min_date <= CDG_DATE <= max_date:
                plt.axvline(CDG_DATE)
                plt.figtext(
                    0.98,
                    0.01,
                    "Line represents CDG Board date",
                    horizontalalignment="right",
                    fontsize=10,
                    fontweight="bold",
                )
        elif isinstance(blue_line, str):
            line_date = parser.parse(blue_line, dayfirst=True)
            if min_date <= line_date.date() <= max_date:
                plt.axvline(line_date.date())
                plt.figtext(
                    0.98,
                    0.01,
                    "Line represents " + blue_line,
                    horizontalalignment="right",
                    fontsize=10,
                    fontweight="bold",
                )

    # size of chart and fit
    fig.tight_layout(rect=[0, 0.03, 1, 0.95])  # for title

    if "chart" in kwargs:
        if kwargs["chart"]:
            plt.show()

    return fig


# def compile_all_profiles():
#     report_doc = open_word_doc(wd_path)
#     for i, p in enumerate(LIST_OF_GROUPS):
#         costs.get_cost_profile(p, 'ipdc_costs')
#         graph = cost_profile_graph(costs, LIST_OF_TITLES[i])
#         put_matplotlib_fig_into_word(report_doc, graph)
#         report_doc.save(root_path / "output/different_cost_profiles.docx")
#
#
# def compile_all_totals():
#     report_doc = open_word_doc(wd_path)
#     for i, p in enumerate(LIST_OF_GROUPS):
#         costs.get_cost_totals(p, 'ipdc_costs')
#         benefits.get_ben_totals(p, 'ipdc_benefits')
#         graph = total_costs_benefits_bar_chart(costs, benefits, LIST_OF_TITLES[i])
#         put_matplotlib_fig_into_word(report_doc, graph)
#         report_doc.save(root_path / "output/different_total_cost_profiles.docx")

# def build_charts(latest_milestone_names,
#                  latest_milestone_dates,
#                  last_milestone_dates,
#                  baseline_milestone_dates,
#                  baseline_milestone_dates_two,
#                  graph_title,
#                  ipdc_date,
#                  no_of_labels):
#     """
#     calculates how many graphical outputs to produced
#     based on number of milestones. Milestone keys names,
#     dates, graph title, date for blue line to represent,
#     and number of labels to have on each graph
#     are passed in.
#     """
#
#     # axis labels are reduced if two long
#     labels = latest_milestone_names
#     final_labels = []
#     for l in labels:
#         if len(l) >= 40:
#             final_labels.append(l[:40])
#         else:
#             final_labels.append(l)
#
#     #  Charts are built
#     no_milestones = len(latest_milestone_names)
#
#     if no_milestones <= no_of_labels:
#         milestone_swimlane_charts(np.array(final_labels), np.array(latest_milestone_dates),
#                                   np.array(last_milestone_dates),
#                                   np.array(baseline_milestone_dates),
#                                   np.array(baseline_milestone_dates_two),
#                                   graph_title, ipdc_date)
#
#     if no_of_labels + 1 <= no_milestones <= no_of_labels*2:
#         half = int(no_milestones / 2)
#         milestone_swimlane_charts(np.array(final_labels[:half]),
#                                   np.array(latest_milestone_dates[:half]),
#                                   np.array(last_milestone_dates[:half]),
#                                   np.array(baseline_milestone_dates[:half]),
#                                   np.array(baseline_milestone_dates_two[:half]),
#                                   graph_title, ipdc_date)
#         title = graph_title + ' cont.'
#         milestone_swimlane_charts(np.array(final_labels[half:no_milestones]),
#                                   np.array(latest_milestone_dates[half:no_milestones]),
#                                   np.array(last_milestone_dates[half:no_milestones]),
#                                   np.array(baseline_milestone_dates[half:no_milestones]),
#                                   np.array(baseline_milestone_dates_two[half:no_milestones]),
#                                   title, ipdc_date)
#
#     if (no_of_labels*2) + 1 <= no_milestones <= no_of_labels*3:
#         third = int(no_milestones / 3)
#         milestone_swimlane_charts(np.array(final_labels[:third]),
#                                   np.array(latest_milestone_dates[:third]),
#                                   np.array(last_milestone_dates[:third]),
#                                   np.array(baseline_milestone_dates[:third]),
#                                   np.array(baseline_milestone_dates_two[:third]),
#                                   graph_title, ipdc_date)
#         title = graph_title + ' cont. 1'
#         milestone_swimlane_charts(np.array(final_labels[third:third * 2]),
#                                   np.array(latest_milestone_dates[third:third * 2]),
#                                   np.array(last_milestone_dates[third:third * 2]),
#                                   np.array(baseline_milestone_dates[third:third * 2]),
#                                   np.array(baseline_milestone_dates_two[third:third * 2]),
#                                   title, ipdc_date)
#         title = graph_title + ' cont. 2'
#         milestone_swimlane_charts(np.array(final_labels[third * 2:no_milestones]),
#                                   np.array(latest_milestone_dates[third * 2:no_milestones]),
#                                   np.array(last_milestone_dates[third * 2:no_milestones]),
#                                   np.array(baseline_milestone_dates[third * 2:no_milestones]),
#                                   np.array(baseline_milestone_dates_two[third * 2:no_milestones]),
#                                   title, ipdc_date)
#     pass


DCA_RATING_SCORES = {
    "Green": 5,
    "Amber/Green": 4,
    "Amber": 3,
    "Amber/Red": 2,
    "Red": 1,
    None: None,
}


# unnecessary not in use. data straight to get group
# def get_group_to_iterate(class_kwargs, master: Master, tp: str) -> List[str]:
#     if "baseline" in class_kwargs:
#         return get_group(
#             master, tp, class_kwargs
#         )
#     elif "quarter" in class_kwargs:
#         return get_group(master, tp, class_kwargs)


def get_tp_index(master, tp: str, class_kwargs):
    if "baseline" in class_kwargs:
        return 0  # baseline uses latest project group only
    elif "quarter" in class_kwargs:
        return master.quarter_list.index(tp)


bl_iter_list = ["current", "last", "bl_one", "bl_two", "bl_three"]

CDG_DATA_KEY_DICT = {
    "IPDC approval point": "Last Business Case (BC) achieved",
    "Total Forecast": "Total Costs",
    "Departmental DCA": "Overall Delivery Confidence",
    "Senior Responsible Owner (SRO)": "SRO",
    "Senior Responsible Owner (SRO) - Email": "SRO email",
    "Total Forecast": "Total Costs",
}


class DcaData:
    def __init__(self, master, **kwargs):
        self.master = master
        self.kwargs = kwargs
        # self.conf_list = list(DCA_KEYS.keys())
        self.group = []
        self.baseline_type = "ipdc_costs"
        self.iter_list = []
        self.dca_dictionary = {}
        self.dca_changes = {}
        self.dca_count = {}
        self.get_dictionary()
        self.get_count()
        # self.get_changes()  # call when needed

    def get_dictionary(self) -> None:
        self.iter_list = get_iter_list(self.kwargs, self.master)
        quarter_dict = {}

        if "data_type" in self.kwargs:
            if self.kwargs["data_type"] == "cdg":
                self.dca_keys = CPG_DCA_KEYS
        else:
            self.dca_keys = DCA_KEYS

        # if "conf_type" in self.kwargs:  # option here to change confidence types
        # if self.kwargs["conf_type"] == "sro_three":
        #     # key value to be changed to 'GMPP - SRO DCA' when data available.
        #     self.dca_keys = {
        #         "sro_three": "GMPP - SRO DCA",
        #     }
        # else:
        # self.dca_keys = {
        #     self.kwargs["conf_type"]: DCA_KEYS[self.kwargs["conf_type"]],
        # }

        for tp in self.iter_list:
            self.group = get_group(self.master, tp, self.kwargs)
            type_dict = {}
            for conf_type in list(self.dca_keys.keys()):  # confidence type
                dca_dict = {}
                try:
                    for project_name in self.group:
                        p_data = get_correct_p_data(
                            self.kwargs,
                            self.master,
                            self.baseline_type,
                            project_name,
                            tp,
                        )
                        if p_data is None:
                            continue
                        dca_type = self.dca_keys[conf_type]
                        colour = p_data[dca_type]
                        score = DCA_RATING_SCORES[p_data[dca_type]]
                        if "data_type" in self.kwargs:
                            if self.kwargs["data_type"] == "cdg":
                                costs = p_data[CDG_DATA_KEY_DICT["Total Forecast"]]
                        else:
                            costs = p_data["Total Forecast"]
                        dca_colour = [("DCA", colour)]
                        dca_score = [("DCA score", score)]
                        t = [("Type", dca_type)]
                        cost_amount = [("Costs", costs)]
                        quarter = [("Quarter", tp)]
                        dca_dict[self.master.abbreviations[project_name]["abb"]] = dict(
                            dca_colour + t + cost_amount + quarter + dca_score
                        )
                    type_dict[conf_type] = dca_dict
                except KeyError:  # handles dca_type e.g. schedule confidence key not present
                    pass

            quarter_dict[tp] = type_dict

        self.dca_dictionary = quarter_dict

    def get_changes(self) -> None:
        """compiles dictionary of changes in dca ratings when provided with two quarter arguments"""
        c_dict = {}
        for conf_type in list(self.dca_keys.keys()):  # confidence type
            lower_dict = {}
            for project_name in list(
                self.dca_dictionary[self.iter_list[0]][conf_type].keys()
            ):
                t = [("Type", conf_type)]
                try:
                    dca_one_colour = self.dca_dictionary[self.iter_list[0]][conf_type][
                        project_name
                    ]["DCA"]
                    dca_two_colour = self.dca_dictionary[self.iter_list[1]][conf_type][
                        project_name
                    ]["DCA"]
                    dca_one_score = self.dca_dictionary[self.iter_list[0]][conf_type][
                        project_name
                    ]["DCA score"]
                    dca_two_score = self.dca_dictionary[self.iter_list[1]][conf_type][
                        project_name
                    ]["DCA score"]
                    if dca_one_score == dca_two_score:
                        status = [("Status", "Same")]
                        change = [("Change", "Unchanged")]
                    if dca_one_score > dca_two_score:
                        status = [
                            (
                                "Status",
                                "Improved from "
                                + dca_two_colour
                                + " to "
                                + dca_one_colour,
                            )
                        ]
                        change = [("Change", "Up")]
                    if dca_one_score < dca_two_score:
                        status = [
                            (
                                "Status",
                                "Worsened from "
                                + dca_two_colour
                                + " to "
                                + dca_one_colour,
                            )
                        ]
                        change = [("Change", "Down")]
                except TypeError:  # This picks up None types
                    if dca_one_colour:  # if project not reporting dca previous quarter
                        status = [("Status", "entered at " + str(dca_one_colour))]
                        change = [("Change", "New entry")]
                    else:
                        status = [("Status", "Missing")]
                        change = [("Change", "Unknown")]
                except KeyError:  # This picks up projects not being present in the quarters being analysed.
                    status = [("Status", "entered at " + str(dca_one_colour))]
                    change = [("Change", "New entry")]

                lower_dict[project_name] = dict(t + status + change)

            c_dict[conf_type] = lower_dict
        self.dca_changes = c_dict

    def get_count(self) -> None:
        """Returns dictionary containing a count of dcas"""
        output_dict = {}
        error_list = []
        for quarter in self.dca_dictionary.keys():  # HERE
            dca_dict = {}
            for i, dca_type in enumerate(list(self.dca_dictionary[quarter].keys())):
                clr = {}
                for x, colour in enumerate(list(DCA_RATING_SCORES.keys())):
                    count = 0
                    cost = 0
                    total = 0
                    cost_total = 0
                    for y, project in enumerate(
                        list(self.dca_dictionary[quarter][dca_type].keys())
                    ):
                        total += 1
                        try:
                            cost_total += self.dca_dictionary[quarter][dca_type][
                                project
                            ]["Costs"]
                        except TypeError:
                            error_list.append(
                                project
                                + " total costs for "
                                + str(quarter)
                                + " are in an incorrect data type and need changing"
                            )
                            pass
                        if (
                            self.dca_dictionary[quarter][dca_type][project]["DCA"]
                            == colour
                        ):
                            count += 1
                            try:
                                cost += self.dca_dictionary[quarter][dca_type][project][
                                    "Costs"
                                ]
                            except TypeError:  # error message above doesn't need repeating
                                pass

                    clr[colour] = {
                        "count": count,
                        "cost": cost,
                        "ct": cost / cost_total,
                    }
                    clr["Total"] = {
                        "count": total,
                        "cost": cost_total,
                        "ct": cost_total / cost_total,
                    }
                dca_dict[dca_type] = clr
            output_dict[quarter] = dca_dict

        error_list = get_error_list(error_list)
        for x in error_list:
            print(x)

        self.dca_count = output_dict


def dca_changes_into_word(dca_data: DcaData, doc: Document) -> Document:
    header = (
        "Showing changes between "
        + str(dca_data.iter_list[0])
        + " and "
        + str(dca_data.iter_list[1])
        + "."
    )
    top = doc.add_paragraph()
    top.add_run(header).bold = True

    for i, dca_type in enumerate(list(dca_data.dca_changes.keys())):
        if i != 0:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        else:
            pass
        title = dca_type + " " + "Confidence changes"
        top = doc.add_paragraph()
        top.add_run(title).bold = True

        doc.add_paragraph()
        sub_head = "Improvements"
        sub = doc.add_paragraph()
        sub.add_run(sub_head).bold = True
        count = 0
        for project_name in list(dca_data.dca_changes[dca_type].keys()):
            if dca_data.dca_changes[dca_type][project_name]["Change"] == "Up":
                doc.add_paragraph(
                    project_name
                    + " "
                    + dca_data.dca_changes[dca_type][project_name]["Status"]
                )
                count += 1
        total_line = str(count) + " project(s) in total improved"
        doc.add_paragraph(total_line)

        doc.add_paragraph()
        sub_head = "Decreases"
        sub = doc.add_paragraph()
        sub.add_run(sub_head).bold = True
        count = 0
        for project_name in list(dca_data.dca_changes[dca_type].keys()):
            if dca_data.dca_changes[dca_type][project_name]["Change"] == "Down":
                doc.add_paragraph(
                    project_name
                    + " "
                    + dca_data.dca_changes[dca_type][project_name]["Status"]
                )
                count += 1
        total_line = str(count) + " project(s) in total have decreased"
        doc.add_paragraph(total_line)

        doc.add_paragraph()
        sub_head = "Missing ratings"
        sub = doc.add_paragraph()
        sub.add_run(sub_head).bold = True
        count = 0
        for project_name in list(dca_data.dca_changes[dca_type].keys()):
            if dca_data.dca_changes[dca_type][project_name]["Change"] == "Unknown":
                doc.add_paragraph(
                    project_name
                    + " "
                    + dca_data.dca_changes[dca_type][project_name]["Status"]
                )
                count += 1
        total_line = str(count) + " project(s) in total are missing a rating"
        doc.add_paragraph(total_line)

        doc.add_paragraph()
        sub_head = "New Projects"
        sub = doc.add_paragraph()
        sub.add_run(sub_head).bold = True
        count = 0
        for project_name in list(dca_data.dca_changes[dca_type].keys()):
            if dca_data.dca_changes[dca_type][project_name]["Change"] == "New entry":
                doc.add_paragraph(
                    project_name
                    + " "
                    + dca_data.dca_changes[dca_type][project_name]["Status"]
                )
                count += 1
        total_line = str(count) + " new project(s)"
        doc.add_paragraph(total_line)

    return doc


VFM_LIST = [
    "NPV for all projects and NPV for programmes if available",
    "Adjusted Benefits Cost Ratio (BCR)",
    "Initial Benefits Cost Ratio (BCR)",
    "VfM Category single entry",
    "VfM Category lower range",
    "VfM Category upper range",
    "Present Value Cost (PVC)",
    "Present Value Benefit (PVB)",
    "Benefits Narrative",
]

VFM_CAT = [
    "Poor",
    "Low",
    "Medium",
    "High",
    "Very High",
    "Very High and Financially Positive",
    "Economically Positive",
    "Total",
]


class VfMData:
    def __init__(
        self,
        master,
        **kwargs,
    ):
        self.master = master
        self.iter_list = []
        self.baseline_type = "ipdc_benefits"
        self.kwargs = kwargs
        self.group = []
        self.vfm_dictionary = {}
        self.vfm_cat_count = {}
        self.vfm_cat_pvc = {}
        self.get_dictionary()
        self.get_count()

    def get_dictionary(self) -> None:
        quarter_dict = {}
        self.iter_list = get_iter_list(self.kwargs, self.master)
        for tp in self.iter_list:
            project_dict = {}
            self.group = get_group(self.master, tp, self.kwargs)
            for p in self.group:
                p_data = get_correct_p_data(
                    self.kwargs, self.master, self.baseline_type, p, tp
                )
                if p_data is None:
                    continue
                vfm_list = []
                vfm = ("Group", self.master.project_information[p]["Group"])
                vfm_list.append(vfm)
                for vfm_type in VFM_LIST:
                    try:
                        vfm = (
                            vfm_type,
                            p_data[vfm_type],
                        )
                        vfm_list.append(vfm)
                    except KeyError:  # vfm range keys not in all masters
                        pass

                project_dict[p] = dict(vfm_list)
            quarter_dict[tp] = project_dict

        self.vfm_dictionary = quarter_dict

    def get_count(self) -> None:
        """Returns dictionary containing a count of vfm categories and pvc totals"""
        count_output_dict = {}
        pvc_output_dict = {}
        error_list = []
        for i, quarter in enumerate(self.vfm_dictionary.keys()):
            pvc_list = []
            cat_list = []
            for cat in VFM_CAT:
                cat_pvc_count = 0
                total_pvc_count = 0
                cat_count = 0
                total_count = 0
                for y, project in enumerate(list(self.vfm_dictionary[quarter].keys())):
                    proj_cat = self.vfm_dictionary[quarter][project][
                        "VfM Category single entry"
                    ]
                    try:
                        project_pvc = self.vfm_dictionary[quarter][project][
                            "Present Value Cost (PVC)"
                        ]
                        total_pvc_count += project_pvc
                        if proj_cat == cat:
                            cat_pvc_count += project_pvc
                    except TypeError:
                        if project_pvc is not None:
                            error_list.append(
                                quarter + " " + project + " PVC data needs checking"
                            )
                            pass
                    # proj_cat = self.vfm_dictionary[quarter][project][
                    #     "VfM Category single entry"
                    # ]
                    if proj_cat is not None:
                        total_count += 1
                        if proj_cat == cat:
                            cat_count += 1
                    if proj_cat is None:
                        # if i == 0:
                        error_list.append(
                            quarter + " " + project + " VfM Category is None"
                        )

                pvc_list.append((cat, cat_pvc_count))
                cat_list.append((cat, cat_count))
            pvc_list.append(("Total", total_pvc_count))
            cat_list.append(("Total", total_count))
            pvc_output_dict[quarter] = dict(pvc_list)
            count_output_dict[quarter] = dict(cat_list)

        #  Data handling. should only print out errors in quarters being analysed.
        error_list = get_error_list(error_list)
        for x in error_list:
            print(x)

        self.vfm_cat_pvc = pvc_output_dict
        self.vfm_cat_count = count_output_dict


def vfm_into_excel(vfm_data: VfMData) -> workbook:
    wb = Workbook()

    for quarter in vfm_data.vfm_dictionary.keys():
        start_row = 3
        ws = wb.create_sheet(
            make_file_friendly(quarter)
        )  # creating worksheets. names restricted to 30 characters.
        ws.title = make_file_friendly(quarter)  # title of worksheet
        for i, project_name in enumerate(list(vfm_data.vfm_dictionary[quarter].keys())):
            abb = vfm_data.master.abbreviations[project_name]["abb"]
            ws.cell(row=start_row + i, column=2).value = abb
            for x, key in enumerate(
                list(vfm_data.vfm_dictionary[quarter][project_name].keys())
            ):
                ws.cell(row=2, column=3 + x).value = key
                ws.cell(
                    row=start_row + i, column=3 + x
                ).value = vfm_data.vfm_dictionary[quarter][project_name][key]

        ws.cell(row=2, column=2).value = "Project/Programme"

    start_row = 4
    ws = wb.create_sheet("Count")
    ws.title = "Count"
    for x, quarter in enumerate(vfm_data.vfm_dictionary.keys()):
        ws.cell(row=3, column=3 + x).value = quarter
        ws.cell(row=3 + 12, column=3 + x).value = quarter
        for i, cat in enumerate(VFM_CAT):
            ws.cell(row=start_row + i, column=2).value = cat
            ws.cell(row=start_row + i + 12, column=2).value = cat
            try:
                ws.cell(row=start_row + i, column=3 + x).value = vfm_data.vfm_cat_pvc[
                    quarter
                ][cat]
                ws.cell(
                    row=start_row + i + 12, column=3 + x
                ).value = vfm_data.vfm_cat_count[quarter][cat]
            except KeyError:
                pass

    ws.cell(row=2, column=2).value = "PVC total per category"
    ws.cell(row=3, column=2).value = "Category"
    ws.cell(row=2 + 12, column=2).value = "Category count"
    ws.cell(row=3 + 12, column=2).value = "Category"

    wb.remove(wb["Sheet"])
    return wb


# for speed_dial analysis_engine
# degree_range, rot_text, and gauge all in early development. Code taken from
# http://nicolasfauchereau.github.io/climatecode/posts/drawing-a-gauge-with-matplotlib/
def degree_range(n):
    start = np.linspace(-30, 210, n + 1, endpoint=True)[0:-1]
    end = np.linspace(-30, 210, n + 1, endpoint=True)[1::]
    mid_points = start + ((end - start) / 2.0)
    return np.c_[start, end], mid_points


def rot_text(ang):
    rotation = np.degrees(np.radians(ang) * np.pi / np.pi - np.radians(90))
    return rotation


def gauge(
    labels: List[str],
    total: str,
    arrow_one: float,
    arrow_two: float,
    up: str,
    down: str,
    title: str,
):
    no = len(labels)
    fig, ax = plt.subplots(facecolor=FACE_COLOUR)
    fig.set_size_inches(18.5, 10.5)
    ax.set_facecolor(FACE_COLOUR)  # TBC if face colour is required
    ang_range, mid_points = degree_range(no)
    if no == 3:
        colours = [
            COLOUR_DICT["R"],
            COLOUR_DICT["A"],
            COLOUR_DICT["G"],
        ]
    else:
        colours = [
            COLOUR_DICT["R"],
            COLOUR_DICT["A/R"],
            COLOUR_DICT["A"],
            COLOUR_DICT["A/G"],
            COLOUR_DICT["G"],
        ]

    patches = []
    for ang, c in zip(ang_range, colours):
        patches.append(
            Wedge(
                (0.0, 0.0),
                0.4,
                *ang,
                width=0.15,
                facecolor=c,
                lw=2,
                ec="black",
                zorder=2,
            )
        )

    [ax.add_patch(p) for p in patches]

    for mid, lab in zip(mid_points, reversed(labels)):
        ax.text(
            0.325 * np.cos(np.radians(mid)),
            0.325 * np.sin(np.radians(mid)),
            lab,
            horizontalalignment="center",
            verticalalignment="center",
            fontsize=35,
            fontweight="bold",
            fontname=FONT_TYPE,
            rotation=rot_text(mid),
            zorder=4,
        )

    ## title
    plt.suptitle(
        title.upper() + " Confidence",
        fontweight="bold",
        fontsize=50,
        fontname=FONT_TYPE,
    )

    ax.text(
        0,
        -0.1,
        total,
        horizontalalignment="center",
        verticalalignment="center",
        fontsize=60,
        fontweight="bold",
        fontname=FONT_TYPE,
        zorder=1,
    )

    ax.text(
        0,
        -0.17,
        "SRO ratings",
        horizontalalignment="center",
        verticalalignment="center",
        fontsize=45,
        fontweight="bold",
        fontname=FONT_TYPE,
        zorder=1,
    )

    def get_arrow_point(score: float):
        return (240 * score) - 120

    ax.annotate(
        "",
        xy=(
            (0.275 * np.sin(np.radians(get_arrow_point(arrow_two)))),
            (0.275 * np.cos(np.radians(get_arrow_point(arrow_two)))),
        ),
        xytext=(0, 0),
        arrowprops=dict(
            arrowstyle="wedge", color="grey", fill=True, alpha=0.5, linewidth=20  #
        ),
    )

    ax.annotate(
        "",
        xy=(
            (0.275 * np.sin(np.radians(get_arrow_point(arrow_one)))),
            (0.275 * np.cos(np.radians(get_arrow_point(arrow_one)))),
        ),
        xytext=(0, 0),
        arrowprops=dict(arrowstyle="wedge", linewidth=20),
    )

    ax.add_patch(Circle((0, 0), radius=0.02, facecolor="k"))
    ax.add_patch(Circle((0, 0), radius=0.01, facecolor="w", zorder=11))

    ## arrows around the dial
    if arrow_one != arrow_two:  # only done if two quarters data available.
        plt.annotate(
            "",
            xy=(-0.2, 0.4),
            xycoords="data",
            xytext=(-0.4, 0.2),
            textcoords="data",
            arrowprops=dict(
                arrowstyle="->", connectionstyle="arc3, rad=-0.3", linewidth=4
            ),
        )
        ax.text(-0.35, 0.35, down, fontsize=30, fontname=FONT_TYPE)

        ax.text(0.35, 0.35, up, fontsize=30, fontname=FONT_TYPE)

        plt.annotate(
            "",
            xy=(0.4, 0.2),
            xycoords="data",
            xytext=(0.2, 0.4),
            textcoords="data",
            arrowprops=dict(
                arrowstyle="<-", connectionstyle="arc3, rad=-0.3", linewidth=4
            ),
        )

    plt.axis("scaled")
    plt.axis("off")
    # plt.show()

    return fig


def build_speedials(dca_data: DcaData, doc, **kwargs) -> None:
    """function for build speed dial graphical output"""
    for conf_type in dca_data.dca_count[dca_data.iter_list[0]]:
        c_count = []
        l_count = []
        if dca_data.kwargs["rag_number"] == "3":
            rag_no = {"Green": 5, "Amber": 3, "Red": 1}
        if dca_data.kwargs["rag_number"] == "5":
            rag_no = {
                "Green": 5,
                "Amber/Green": 4,
                "Amber": 3,
                "Amber/Red": 2,
                "Red": 1,
            }
        for colour in rag_no.keys():
            c_no = dca_data.dca_count[dca_data.iter_list[0]][conf_type][colour]["count"]
            c_count.append(c_no)
            try:  # to capture reporting process with only one quarters data
                l_no = dca_data.dca_count[dca_data.iter_list[1]][conf_type][colour][
                    "count"
                ]
                l_count.append(l_no)
            except KeyError:
                l_count.append(c_no)
            except IndexError:
                pass
        c_total = dca_data.dca_count[dca_data.iter_list[0]][conf_type]["Total"]["count"]

        up = 0
        down = 0

        if bool(dca_data.dca_changes):  # empty dictionaries evaluation to false
            for p in dca_data.dca_changes[conf_type]:
                change = dca_data.dca_changes[conf_type][p]["Change"]
                if change == "Up":
                    up += 1
                if change == "Down":
                    down += 1

        if dca_data.kwargs["rag_number"] == "3":
            rate = [0, 0.5, 1]
        if dca_data.kwargs["rag_number"] == "5":
            rate = [0, 0.25, 0.5, 0.75, 1]
        dial_one = np.average(rate, weights=c_count)
        try:
            dial_two = np.average(rate, weights=l_count)
        except TypeError:  # no previous quarter data
            dial_two = dial_one

        graph = gauge(
            c_count,
            str(c_total),
            dial_one,
            dial_two,
            str(up),
            str(down),
            title=conf_type,
        )

        put_matplotlib_fig_into_word(doc, graph, size=7.5)


def sort_projects_by_dca(
    master_data: List[Dict[str, Union[str, int, datetime.date, float]]],
    projects: List[str] or str,
) -> List[str]:
    # returns a list of projects sorted by dca rag rating
    rag_list = []
    for project_name in projects:
        rag = master_data["data"][project_name]["Departmental DCA"]
        rag_list.append((project_name, rag))

    rag_list_sorted = sorted(rag_list, key=lambda x: x[1])

    return rag_list_sorted


def cost_schedule_scatter_chart_matplotlib(milestones: MilestoneData, costs: CostData):
    sc_list = []
    cc_list = []
    volume_list = []
    colour_list = []
    for project_name in milestones.group:
        ab = milestones.master.abbreviations[project_name]["abb"]
        sc = milestones.schedule_change[ab]["bl_one"][
            "percent change"
        ]  # sc schedule change
        cc = costs.wlc_change[project_name]["baseline one"]  # cc cost change
        volume = costs.master.master_data[0].data[project_name]["Total Forecast"]
        colour = COLOUR_DICT[
            convert_rag_text(
                costs.master.master_data[0].data[project_name]["Departmental DCA"]
            )
        ]
        if sc > 5 or cc > 5:
            sc_list.append(sc)
            cc_list.append(cc)
            volume_list.append(np.sqrt(volume) * 5)
            colour_list.append(colour)
        else:
            pass

    fig, ax = plt.subplots()
    ax.scatter(sc_list, cc_list, c=colour_list, s=volume_list)
    # , alpha=0.5)

    ax.set_xlabel("Schedule", fontsize=15)
    ax.set_ylabel("Costs", fontsize=15)
    ax.set_title("Volume and percent change")
    plt.ylim(-75, 75)
    plt.xlim(-60, 60)

    ax.grid(True)
    fig.tight_layout()

    # plt.show()


def cost_schedule_scatter_chart_excel(ws, rag_count):
    chart = BubbleChart()
    chart.style = 18  # use a preset style

    # add the first series of data
    amber_stop = 2 + rag_count["Amber"]
    xvalues = Reference(ws, min_col=3, min_row=3, max_row=amber_stop)
    yvalues = Reference(ws, min_col=4, min_row=3, max_row=amber_stop)
    size = Reference(ws, min_col=5, min_row=3, max_row=amber_stop)
    series = Series(values=yvalues, xvalues=xvalues, zvalues=size, title="Amber")
    chart.series.append(series)
    series.graphicalProperties.solidFill = "fce553"

    # add the second
    amber_g_stop = amber_stop + rag_count["Amber/Green"]
    xvalues = Reference(ws, min_col=3, min_row=amber_stop + 1, max_row=amber_g_stop)
    yvalues = Reference(ws, min_col=4, min_row=amber_stop + 1, max_row=amber_g_stop)
    size = Reference(ws, min_col=5, min_row=amber_stop + 1, max_row=amber_g_stop)
    series = Series(values=yvalues, xvalues=xvalues, zvalues=size, title="Amber/Green")
    chart.series.append(series)
    series.graphicalProperties.solidFill = "a5b700"

    amber_r_stop = amber_g_stop + rag_count["Amber/Red"]
    xvalues = Reference(ws, min_col=3, min_row=amber_g_stop + 1, max_row=amber_r_stop)
    yvalues = Reference(ws, min_col=4, min_row=amber_g_stop + 1, max_row=amber_r_stop)
    size = Reference(ws, min_col=5, min_row=amber_g_stop + 1, max_row=amber_r_stop)
    series = Series(values=yvalues, xvalues=xvalues, zvalues=size, title="Amber/Red")
    chart.series.append(series)
    series.graphicalProperties.solidFill = "f97b31"

    green_stop = amber_r_stop + rag_count["Green"]
    xvalues = Reference(ws, min_col=3, min_row=amber_r_stop + 1, max_row=green_stop)
    yvalues = Reference(ws, min_col=4, min_row=amber_r_stop + 1, max_row=green_stop)
    size = Reference(ws, min_col=5, min_row=amber_r_stop + 1, max_row=green_stop)
    series = Series(values=yvalues, xvalues=xvalues, zvalues=size, title="Green")
    chart.series.append(series)
    series.graphicalProperties.solidFill = "17960c"

    red_stop = green_stop + rag_count["Red"]
    xvalues = Reference(ws, min_col=3, min_row=green_stop + 1, max_row=red_stop)
    yvalues = Reference(ws, min_col=4, min_row=green_stop + 1, max_row=red_stop)
    size = Reference(ws, min_col=5, min_row=green_stop + 1, max_row=red_stop)
    series = Series(values=yvalues, xvalues=xvalues, zvalues=size, title="Red")
    chart.series.append(series)
    series.graphicalProperties.solidFill = "cb1f00"

    ws.add_chart(chart, "L2")

    return ws


def cost_v_schedule_chart_into_wb(milestones: MilestoneData, costs: CostData):
    wb = Workbook()
    ws = wb.active

    rags = []
    for project_name in milestones.group:
        rag = milestones.master.master_data[0].data[project_name]["Departmental DCA"]
        if rag is not None:
            rags.append((project_name, rag))
        else:
            print(
                project_name
                + " has not reported an SRO RAG Confidence so it will not be included. Check data"
            )

    rags = sorted(rags, key=lambda x: x[1])

    rag_c = Counter(x[1] for x in rags)  # rag_c is rag_count

    ws.cell(row=2, column=2).value = "Project Name"
    ws.cell(row=2, column=3).value = "Schedule change"
    ws.cell(row=2, column=4).value = "WLC Change"
    ws.cell(row=2, column=5).value = "WLC"
    ws.cell(row=2, column=6).value = "DCA"
    ws.cell(row=2, column=7).value = "Start key"
    ws.cell(row=2, column=8).value = "End key"

    for x, project_name in enumerate(rags):
        ab = milestones.master.abbreviations[project_name[0]]["abb"]
        ws.cell(row=x + 3, column=2).value = ab
        ws.cell(row=x + 3, column=3).value = milestones.schedule_change[ab]["bl_one"][
            "percent change"
        ]
        ws.cell(row=x + 3, column=4).value = costs.wlc_change[project_name[0]][
            "baseline one"
        ]
        ws.cell(row=x + 3, column=5).value = costs.master.master_data[0].data[
            project_name[0]
        ]["Total Forecast"]
        ws.cell(row=x + 3, column=6).value = costs.master.master_data[0].data[
            project_name[0]
        ]["Departmental DCA"]
        ws.cell(row=x + 3, column=7).value = milestones.schedule_change[ab]["bl_one"][
            "start key"
        ]
        ws.cell(row=x + 3, column=8).value = milestones.schedule_change[ab]["bl_one"][
            "end key"
        ]

    cost_schedule_scatter_chart_excel(ws, rag_c)
    cost_schedule_scatter_chart_matplotlib(milestones, costs)

    return wb


def make_columns_bold(columns: list) -> None:
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def change_text_size(columns: list, size: int) -> None:
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(size)


def make_text_red(columns: list) -> None:
    for column in columns:
        for cell in column.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text == "Not reported":
                        run.font.color.rgb = RGBColor(255, 0, 0)


def print_out_project_milestones(
    doc: Document, milestones: MilestoneData, project_name: str
) -> Document:
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    # table heading
    ab = milestones.master.abbreviations[project_name]["abb"]
    doc.add_paragraph().add_run(str(ab + " milestone table (2021 - 22)")).bold = True

    ab = milestones.master.abbreviations[project_name]["abb"]

    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Milestone"
    hdr_cells[1].text = "Date"
    hdr_cells[2].text = "Change from last quarter"
    hdr_cells[3].text = "Change from baseline"
    hdr_cells[4].text = "Notes"

    for i, m in enumerate(
        milestones.sorted_milestone_dict[milestones.iter_list[0]]["names"]
    ):
        row_cells = table.add_row().cells
        row_cells[0].text = m
        row_cells[1].text = milestones.sorted_milestone_dict[milestones.iter_list[0]][
            "r_dates"
        ][i].strftime("%d/%m/%Y")
        try:
            row_cells[2].text = plus_minus_days(
                (
                    milestones.sorted_milestone_dict[milestones.iter_list[0]][
                        "r_dates"
                    ][i]
                    - milestones.sorted_milestone_dict[milestones.iter_list[1]][
                        "r_dates"
                    ][i]
                ).days
            )
        except TypeError:
            row_cells[2].text = "Not reported"
        try:
            row_cells[3].text = plus_minus_days(
                (
                    milestones.sorted_milestone_dict[milestones.iter_list[0]][
                        "r_dates"
                    ][i]
                    - milestones.sorted_milestone_dict[milestones.iter_list[2]][
                        "r_dates"
                    ][i]
                ).days
            )
        except TypeError:
            row_cells[3].text = "Not reported"
        try:
            row_cells[4].text = milestones.sorted_milestone_dict[
                milestones.iter_list[0]
            ]["notes"][i]
            paragraph = row_cells[4].paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size = Pt(8)  # font size = 8
        except TypeError:
            pass

        # try:
        #     if milestone_filter_start_date <= milestone_date:  # filter based on date
        #         row_cells = table.add_row().cells
        #         row_cells[0].text = milestone
        #         if milestone_date is None:
        #             row_cells[1].text = 'No date'
        #         else:
        #             row_cells[1].text = milestone_date.strftime("%d/%m/%Y")
        #         b_one_value = first_diff_data[project_name][milestone]
        #         row_cells[2].text = plus_minus_days(b_one_value)
        #         b_two_value = second_diff_data[project_name][milestone]
        #         row_cells[3].text = plus_minus_days(b_two_value)
        #
        #         notes = p_current_milestones[project_name][milestone][milestone_date]
        #         row_cells[4].text = str(notes)
        #         # trying to high changes to narratuve in red text
        #         # if milestone in p_last_milestones[project_name].keys():
        #         #     last_milestone_date = p_last_milestones[project_name][milestone]
        #         #     last_note = p_last_milestones[project_name][milestone][last_milestone_date]
        #         #     row_cells[4] = compare_text_newandold(notes, last_note, doc)
        #         # elif milestone not in p_last_milestones[project_name].keys():
        #         #     row_cells[4].text = str(notes)
        #
        #         paragraph = row_cells[4].paragraphs[0]
        #         run = paragraph.runs
        #         font = run[0].font
        #         font.size = Pt(8)  # font size = 8
        #
        #
        # except TypeError:  # this is to deal with none types which are still placed in output
        #     row_cells = table.add_row().cells
        #     row_cells[0].text = milestone
        #     if milestone_date is None:
        #         row_cells[1].text = 'No date'
        #     else:
        #         row_cells[1].text = milestone_date.strftime("%d/%m/%Y")
        #     b_one_value = first_diff_data[project_name][milestone]
        #     row_cells[2].text = plus_minus_days(b_one_value)
        #     b_two_value = second_diff_data[project_name][milestone]
        #     row_cells[3].text = plus_minus_days(b_two_value)
        #     notes = p_current_milestones[project_name][milestone][milestone_date]
        #     row_cells[4].text = str(notes)
        #     paragraph = row_cells[4].paragraphs[0]
        #     run = paragraph.runs
        #     font = run[0].font
        #     font.size = Pt(8)  # font size = 8

    table.style = "Table Grid"

    # column widths
    column_widths = (Cm(6), Cm(2.6), Cm(2), Cm(2), Cm(8.95))
    set_col_widths(table, column_widths)
    # make_columns_bold([table.columns[0], table.columns[3]])  # make keys bold
    # make_text_red([table.columns[1], table.columns[4]])  # make 'not reported red'

    make_rows_bold(
        [table.rows[0]]
    )  # makes top of table bold. Found function on stack overflow.
    return doc


def compile_p_report(
    doc: Document,
    master,
    project_name: str,
) -> Document:
    wd_heading(doc, master, project_name)
    key_contacts(doc, master, project_name)
    dca_table(doc, master, project_name)
    # forward_look(doc, master, project_name)
    dca_narratives(doc, master, project_name)
    costs = CostData(master, group=[project_name], baseline=["standard"])
    costs.get_cost_profile()
    benefits = BenefitsData(master, group=[project_name], baseline=["standard"])
    milestones = MilestoneData(master, group=[project_name], baseline=["standard"])
    project_report_meta_data(doc, costs, milestones, benefits, project_name)
    change_word_doc_landscape(doc)
    cost_profile = cost_profile_graph(costs, master, show="No", group=milestones.group)
    put_matplotlib_fig_into_word(doc, cost_profile, transparent=False, size=8)
    total_profile = total_costs_benefits_bar_chart(
        costs, benefits, master, group=costs.group, show="No"
    )
    put_matplotlib_fig_into_word(doc, total_profile, transparent=False, size=8)
    #  handling of no milestones within filtered period.
    ab = master.abbreviations[project_name]["abb"]
    # requires refactor here.
    start_date = calculate_dates("minus 3 months").strftime("%d/%m/%Y")
    end_date = calculate_dates("2 years").strftime("%d/%m/%Y")
    try:
        milestones.filter_chart_info(dates=[start_date, end_date])
        milestones_chart = milestone_chart(
            milestones,
            master,
            blue_line="IPDC",
            title=ab + " schedule next two years",
            show="No",
        )
        put_matplotlib_fig_into_word(doc, milestones_chart, transparent=False, size=8)
    except ValueError:
        pass
    #     milestones = MilestoneData(master, group=[project_name], baseline=["standard"])
    #     milestones.filter_chart_info(dates=["1/9/2020", "31/12/2024"])
    #     milestones_chart = milestone_chart(
    #         milestones,
    #         master,
    #         blue_line="IPDC",
    #         title=ab + " schedule (2021 - 24)",
    #         show="No",
    #     )
    #     put_matplotlib_fig_into_word(doc, milestones_chart)
    print_out_project_milestones(doc, milestones, project_name)
    change_word_doc_portrait(doc)
    project_scope_text(doc, master, project_name)
    return doc


# TODO refactor all code below
# def grey_conditional_formatting(ws):
#     '''
#     function applies grey conditional formatting for 'Not Reporting'.
#     :param worksheet: ws
#     :return: cf of sheet
#     '''
#
#     grey_text = Font(color="f0f0f0")
#     grey_fill = PatternFill(bgColor="f0f0f0")
#     dxf = DifferentialStyle(font=grey_text, fill=grey_fill)
#     rule = Rule(type="containsText", operator="containsText", text="Not reporting", dxf=dxf)
#     rule.formula = ['NOT(ISERROR(SEARCH("Not reporting",A1)))']
#     ws.conditional_formatting.add('A1:X80', rule)
#
#     grey_text = Font(color="cfcfea")
#     grey_fill = PatternFill(bgColor="cfcfea")
#     dxf = DifferentialStyle(font=grey_text, fill=grey_fill)
#     rule = Rule(type="containsText", operator="containsText", text="Data not collected", dxf=dxf)
#     rule.formula = ['NOT(ISERROR(SEARCH("Data not collected",A1)))']
#     ws.conditional_formatting.add('A1:X80', rule)
#
#     return ws
#
#
def conditional_formatting(
    ws,
    list_columns,
    list_conditional_text,
    list_text_colours,
    list_background_colours,
    row_start,
    row_end,
):  # not working
    for column in list_columns:
        for i, txt in enumerate(list_conditional_text):
            text = list_text_colours[i]
            fill = list_background_colours[i]
            dxf = DifferentialStyle(font=text, fill=fill)
            rule = Rule(type="containsText", operator="containsText", text=txt, dxf=dxf)
            for_rule_formula = 'NOT(ISERROR(SEARCH("' + txt + '",' + column + "1)))"
            rule.formula = [for_rule_formula]
            ws.conditional_formatting.add(
                column + row_start + ":" + column + row_end, rule
            )

    return ws


#
#
# # data query stuff
# def return_data(master: Master,
#                 milestones: MilestoneData,
#                 project_group: List[str] or str,
#                 data_key_list: List[str] or str):
#     """Returns project values across multiple masters for specified keys of interest:
#     project_names_list: list of project names
#     data_key_list: list of data keys
#     """
#     wb = Workbook()
#
#     for i, key in enumerate(data_key_list):
#         '''worksheet is created for each project'''
#         try:
#             ws = wb.create_sheet(key[:29], i)  # creating worksheets
#             ws.title = key[:29]
#         except ValueError:
#             if "/" in key:
#                 newstr = key.replace("/", "")
#                 ws = wb.create_sheet(newstr[:29], i)  # creating worksheets
#                 ws.title = newstr[:29]  # title of worksheet
#
#         '''list project names, groups and stage in ws'''
#         for y, project_name in enumerate(project_group):
#             # get project group info
#             try:
#                 group = master[0].data[project_name]['DfT Group']
#             except KeyError:
#                 for m, master in enumerate(master):
#                     if project_name in master.projects:
#                         group = master[m].data[project_name]['DfT Group']
#
#             ws.cell(row=2 + y, column=1, value=group) # group info return
#             ws.cell(row=2 + y, column=2, value=project_name)  # project name returned
#
#             for x, master in enumerate(master):
#                 if project_name in master.projects:
#                     try:
#                         #standard keys
#                         if key in master[x].data[project_name].keys():
#                             value = master[x].data[project_name][key]
#                             ws.cell(row=2 + y, column=3 + x, value=value) # returns value
#
#                             if value is None:
#                                 ws.cell(row=2 + y, column=3 + x, value='md')
#
#                             try: # checks for change against last quarter
#                                 lst_value = master[x + 1].data[project_name][key]
#                                 if value != lst_value:
#                                     ws.cell(row=2 + y, column=3 + x).fill = SALMON_FILL
#                             except (KeyError, IndexError):
#                                 pass
#
#                         # milestone keys
#                         else:
#                             get_milestone_date(project_name)
#                             milestones = all_milestone_data_bulk([project_name], master[x])
#                             value = tuple(milestones[project_name][key])[0]
#                             ws.cell(row=2 + y, column=3 + x, value=value)
#                             ws.cell(row=2 + y, column=3 + x).number_format = 'dd/mm/yy'
#                             if value is None:
#                                 ws.cell(row=2 + y, column=3 + x, value='md')
#
#                             try:  # loop checks if value has changed since last quarter
#                                 old_milestones = all_milestone_data_bulk([project_name], master[x + 1])
#                                 lst_value = tuple(old_milestones[project_name][key])[0]
#                                 if value != lst_value:
#                                     ws.cell(row=2 + y, column=3 + x).fill = SALMON_FILL
#                             except (KeyError, IndexError):
#                                 pass
#
#                     except KeyError:
#                         if project_name in master.projects:
#                             #loop calculates if project was not reporting or data missing
#                             ws.cell(row=2 + y, column=3 + x, value='knc')
#                         else:
#                             ws.cell(row=2 + y, column=3 + x, value='pnr')
#
#                 else:
#                     ws.cell(row=2 + y, column=3 + x, value='pnr')
#
#         '''quarter tag information'''
#         ws.cell(row=1, column=1, value='Group')
#         ws.cell(row=1, column=2, value='Projects')
#         quarter_labels = get_quarter_stamp(master)
#         for l, label in enumerate(quarter_labels):
#             ws.cell(row=1, column=l + 3, value=label)
#
#         list_columns = list_column_ltrs[2:len(master) + 2]
#
#         if key in list_of_rag_keys:
#             conditional_formatting(ws, list_columns, rag_txt_list_full, rag_txt_colours, rag_fill_colours, '1', '80')
#
#         conditional_formatting(ws, list_columns, gen_txt_list, gen_txt_colours, gen_fill_colours, '1', '80')
#
#     return wb
#
# def return_baseline_data(project_name_list, data_key_list):
#     '''
#     returns values of interest across multiple ws for baseline values only.
#     project_name_list: list of project names
#     data_key_list: list of data keys containing values of interest.
#     '''
#     wb = Workbook()
#
#     for i, key in enumerate(data_key_list):
#         '''worksheet is created for each project'''
#         try:
#             ws = wb.create_sheet(key[:29], i)  # creating worksheets
#             ws.title = key[:29]
#         except ValueError:
#             if "/" in key:
#                 newstr = key.replace("/", "")
#                 ws = wb.create_sheet(newstr[:29], i)  # creating worksheets
#                 ws.title = newstr[:29]  # title of worksheet
#
#         key_type = get_key_type(key)
#         '''list project names, groups and stage in ws'''
#         for y, project_name in enumerate(project_name_list):
#
#             # get project group info
#             try:
#                 group = list_of_masters_all[0].data[project_name]['DfT Group']
#             except KeyError:
#                 for m, master in enumerate(list_of_masters_all):
#                     if project_name in master.projects:
#                         group = list_of_masters_all[m].data[project_name]['DfT Group']
#
#             ws.cell(row=2 + y, column=1, value=group) # group info
#             ws.cell(row=2 + y, column=2, value=project_name)  # project name returned
#
#             if key_type == 'benefits':
#                 bc_index = benefits_bl_index
#             else:
#                 bc_index = costs_bl_index
#
#             for x in range(0, len(bc_index[project_name])):
#                 index = bc_index[project_name][x]
#                 try: # standard keys
#                     value = list_of_masters_all[index].data[project_name][key]
#                     if value is None:
#                         ws.cell(row=2 + y, column=3 + x).value = 'md'
#                     else:
#                         ws.cell(row=2 + y, column=3 + x, value=value)
#                 except KeyError:
#                     try: # Milestones
#                         index = milestone_bl_index[project_name][x]
#                         milestones = all_milestone_data_bulk([project_name], list_of_masters_all[index])
#                         value = tuple(milestones[project_name][key])[0]
#                         if value is None:
#                             ws.cell(row=2 + y, column=3 + x).value = 'md'
#                         else:
#                             ws.cell(row=2 + y, column=3 + x).value = value
#                             ws.cell(row=2 + y, column=3 + x).number_format = 'dd/mm/yy'
#                     except KeyError: # exception catches both standard and milestone keys
#                         ws.cell(row=2 + y, column=3 + x).value = 'knc'
#                     except IndexError:
#                         pass
#                 except TypeError:
#                     ws.cell(row=2 + y, column=3 + x).value = 'pnr'
#
#         ws.cell(row=1, column=1, value='Group')
#         ws.cell(row=1, column=2, value='Project')
#         ws.cell(row=1, column=3, value='Latest')
#         ws.cell(row=1, column=4, value='Last quarter')
#         ws.cell(row=1, column=5, value='BL 1')
#         ws.cell(row=1, column=6, value='BL 2')
#         ws.cell(row=1, column=7, value='BL 3')
#         ws.cell(row=1, column=8, value='BL 4')
#         ws.cell(row=1, column=9, value='BL 5')
#
#         list_columns = list_column_ltrs[2:10] # hard coded so not ideal
#
#         if key in list_of_rag_keys:
#             conditional_formatting(ws, list_columns, rag_txt_list_full, rag_txt_colours, rag_fill_colours, '1', '80')
#
#         conditional_formatting(ws, list_columns, gen_txt_list, gen_txt_colours, gen_fill_colours, '1', '80')
#
#     return wb


def get_data_query_key_names(key_file: csv) -> List[str]:
    key_list = []
    with open(key_file) as csv_file:
        csv_reader = csv.reader(csv_file, delimiter=",")
        for row in csv_reader:
            key_list.append(row[0])
    return key_list[1:]


def data_query_into_wb_by_key(master, **kwargs) -> Workbook:
    """
    Returns data values for keys of interest. Keys placed on one page.
    Quarter data placed across different wbs.
    """

    wb = Workbook()
    iter_list = get_iter_list(kwargs, master)
    # for z, tp in enumerate(iter_list):
    for z, key in enumerate(kwargs["key"]):
        # i = master.quarter_list.index(tp)  # handling here. for wrong quarter string
        ws = wb.create_sheet(
            make_file_friendly(key)
        )  # creating worksheets. names restricted to 30 characters.
        ws.title = make_file_friendly(key)  # title of worksheet
        """list project names, groups and stage in ws"""
        for y, p in enumerate(list(master.project_information.keys())):
            if master.project_information[p]["Pipeline"] == "Yes":
                break
            else:
                abb = master.abbreviations[p]["abb"]
                ws.cell(row=2 + y, column=1).value = master.project_information[p][
                    "Group"
                ]
                ws.cell(row=2 + y, column=2).value = p
                ws.cell(row=2 + y, column=3).value = abb
                ws.cell(row=2 + y, column=4).value = master.project_information[p][
                    "GMPP"
                ]
                for x, tp in enumerate(iter_list):
                    p_data = get_correct_p_data(kwargs, master, "ipdc_costs", p, tp)
                    try:
                        p_data_last = get_correct_p_data(
                            kwargs, master, "ipdc_costs", p, iter_list[z + 1]
                        )
                    except IndexError:
                        p_data_last = None

                    ws.cell(row=1, column=5 + x, value=tp)
                    try:  # standard keys
                        value = convert_date(p_data[key])
                        if value is None:
                            ws.cell(row=2 + y, column=5 + x).value = "md"
                            ws.cell(row=2 + y, column=5 + x).fill = AMBER_FILL
                        else:
                            ws.cell(row=2 + y, column=5 + x, value=value)
                            if isinstance(value, datetime.datetime):
                                ws.cell(
                                    row=2 + y, column=5 + x, value=value
                                ).number_format = "dd/mm/yy"

                        try:  # checks for change against next master in loop
                            lst_value = convert_date(p_data_last[key])
                            if value != lst_value:
                                ws.cell(row=2 + y, column=5 + x).fill = SALMON_FILL
                        except (KeyError, UnboundLocalError, TypeError):
                            # KeyError is key not present in master.
                            # UnboundLocalError if there is no last_value.
                            # TypeError if project not in master. p_data_last becomes None.
                            pass
                    except TypeError:
                        pass
                    except KeyError:  # milestone keys
                        if "quarter" in kwargs:
                            milestones_one = MilestoneData(
                                master, quarter=[tp], group=[p]
                            )
                            try:
                                milestones_two = MilestoneData(
                                    master, quarter=[iter_list[z + 1]], group=[p]
                                )
                            except IndexError:
                                pass
                        if "baseline" in kwargs:
                            milestones_one = MilestoneData(
                                master, baseline=[tp], group=[p]
                            )
                            try:
                                milestones_two = MilestoneData(
                                    master, baseline=[iter_list[z + 1]], group=[p]
                                )
                            except IndexError:
                                pass
                        date = get_milestone_date(
                            abb, milestones_one.milestone_dict, tp, " " + key
                        )
                        if date is None:
                            ws.cell(row=2 + y, column=5 + x).value = "md"
                            ws.cell(row=2 + y, column=5 + x).fill = AMBER_FILL
                        else:
                            ws.cell(row=2 + y, column=5 + x).value = date
                            ws.cell(row=2 + y, column=5 + x).number_format = "dd/mm/yy"
                        try:  # checks for changes against next master in loop
                            lst_date = get_milestone_date(
                                abb,
                                milestones_two.milestone_dict,
                                iter_list[z + 1],
                                " " + key,
                            )
                            if date != lst_date:
                                ws.cell(row=2 + y, column=5 + x).fill = SALMON_FILL
                        except (KeyError, UnboundLocalError, TypeError, IndexError):
                            pass

        ws.cell(row=1, column=1).value = "Group"
        ws.cell(row=1, column=2).value = "Project Name"
        ws.cell(row=1, column=3).value = "Project Acronym"
        ws.cell(row=1, column=4).value = "GMPP"

    wb.remove(wb["Sheet"])
    return wb


def data_query_into_wb(master, **kwargs) -> Workbook:
    """
    Returns data values for keys of interest. Keys placed on one page.
    Quarter data placed across different wbs.
    """

    wb = Workbook()
    iter_list = get_iter_list(kwargs, master)
    for z, tp in enumerate(iter_list):
        # i = master.quarter_list.index(tp)  # handling here. for wrong quarter string
        ws = wb.create_sheet(
            make_file_friendly(tp)
        )  # creating worksheets. names restricted to 30 characters.
        ws.title = make_file_friendly(tp)  # title of worksheet
        group = get_group(master, tp, kwargs)

        """list project names, groups and stage in ws"""
        for y, p in enumerate(group):  # p is project name
            p_data = get_correct_p_data(kwargs, master, "ipdc_costs", p, tp)
            abb = master.abbreviations[p]["abb"]
            ws.cell(row=2 + y, column=1).value = master.project_information[p]["Group"]
            ws.cell(row=2 + y, column=2).value = p
            ws.cell(row=2 + y, column=3).value = abb
            ws.cell(row=2 + y, column=4).value = master.project_information[p]["GMPP"]
            try:
                p_data_last = get_correct_p_data(
                    kwargs, master, "ipdc_costs", p, iter_list[z + 1]
                )
            except IndexError:
                p_data_last = None
            for x, key in enumerate(kwargs["key"]):
                ws.cell(row=1, column=5 + x, value=key)
                try:  # standard keys
                    value = convert_date(p_data[key])
                    if value is None:
                        ws.cell(row=2 + y, column=5 + x).value = "md"
                        ws.cell(row=2 + y, column=5 + x).fill = AMBER_FILL
                    else:
                        ws.cell(row=2 + y, column=5 + x, value=value)
                        if isinstance(value, datetime.datetime):
                            ws.cell(
                                row=2 + y, column=5 + x, value=value
                            ).number_format = "dd/mm/yy"

                    try:  # checks for change against next master in loop
                        lst_value = convert_date(p_data_last[key])
                        if value != lst_value:
                            ws.cell(row=2 + y, column=5 + x).fill = SALMON_FILL
                    except (KeyError, UnboundLocalError, TypeError):
                        # KeyError is key not present in master.
                        # UnboundLocalError if there is no last_value.
                        # TypeError if project not in master. p_data_last becomes None.
                        pass
                except KeyError:  # milestone keys
                    if "quarter" in kwargs:
                        milestones_one = MilestoneData(master, quarter=[tp], group=[p])
                        try:
                            milestones_two = MilestoneData(
                                master, quarter=[iter_list[z + 1]], group=[p]
                            )
                        except IndexError:
                            pass
                    if "baseline" in kwargs:
                        milestones_one = MilestoneData(master, baseline=[tp], group=[p])
                        try:
                            milestones_two = MilestoneData(
                                master, baseline=[iter_list[z + 1]], group=[p]
                            )
                        except IndexError:
                            pass
                    date = get_milestone_date(
                        abb, milestones_one.milestone_dict, tp, key
                    )
                    if date is None:
                        ws.cell(row=2 + y, column=5 + x).value = "md"
                        ws.cell(row=2 + y, column=5 + x).fill = AMBER_FILL
                    else:
                        ws.cell(row=2 + y, column=5 + x).value = date
                        ws.cell(row=2 + y, column=5 + x).number_format = "dd/mm/yy"
                    try:  # checks for changes against next master in loop
                        lst_date = get_milestone_date(
                            abb,
                            milestones_two.milestone_dict,
                            iter_list[z + 1],
                            key,
                        )
                        if date != lst_date:
                            ws.cell(row=2 + y, column=5 + x).fill = SALMON_FILL
                    except (KeyError, UnboundLocalError, TypeError, IndexError):
                        pass

        ws.cell(row=1, column=1).value = "Group"
        ws.cell(row=1, column=2).value = "Project Name"
        ws.cell(row=1, column=3).value = "Project Acronym"
        ws.cell(row=1, column=4).value = "GMPP"

    wb.remove(wb["Sheet"])
    return wb


#
#
# # financial analysis_engine stuff
# def place_complex_comparision_excel(master_data_latest, master_data_last, master_data_baseline):
#     '''
#     Function that places all information structured via the get_wlc_costs and get_yearly_costs programmes into an
#     excel spreadsheet. It does some calculations on the level of change that has taken place.
#     This function places in data for a chart that shows changes in financial profile between latest, last and baseline
#     :param master_data_latest: data representing latest quarter information
#     :param master_data_last: data representing last quarter information.
#     :param master_data_baseline: data representing baseline quarter information
#     :return: excel workbook
#     '''
#     wb = Workbook()
#
#     for i, key in enumerate(list(master_data_latest.keys())):
#         ws = wb.create_sheet(key, i)  # creating worksheets
#         ws.title = key  # title of worksheet
#
#         data_latest = master_data_latest[key]
#         data_last = master_data_last[key]
#         data_baseline = master_data_baseline[key]
#
#         for i, project_name in enumerate(data_latest):
#             '''place project names into ws'''
#             ws.cell(row=i+2, column=1).value = project_name
#
#             '''loop for placing data into ws. highlight changes between quarters in red'''
#             latest_value = data_latest[project_name]
#             ws.cell(row=i + 2, column=2).value = latest_value
#
#             '''comparision data against last quarter'''
#             if project_name in data_last.keys():
#                 try:
#                     last_value = data_last[project_name]
#                     ws.cell(row=i + 2, column=3).value = last_value
#                     change = latest_value - last_value
#                     if last_value > 0:
#                         percent_change = (latest_value - last_value)/last_value
#                     else:
#                         percent_change = (latest_value - last_value)/(last_value + 1)
#                     ws.cell(row=i + 2, column=7).value = change
#                     ws.cell(row=i + 2, column=8).value = percent_change
#                     if change >= 100 or change <= -100:
#                         ws.cell(row=i + 2, column=7).font = red_text
#                     if percent_change >= 0.05 or percent_change <= -0.05:
#                         ws.cell(row=i + 2, column=8).font = red_text
#                 except TypeError:
#                     ws.cell(row=i + 2, column=3).value = 'check project data'
#             else:
#                 ws.cell(row=i + 2, column=3).value = 'None'
#
#             if project_name in data_baseline.keys():
#                 try:
#                     last_value = data_last[project_name]
#                     baseline_value = data_baseline[project_name]
#                     ws.cell(row=i + 2, column=4).value = baseline_value
#                     change = last_value - baseline_value
#                     if baseline_value > 0:
#                         percent_change = (last_value - baseline_value) / baseline_value
#                     else:
#                         percent_change = (last_value - baseline_value) / (baseline_value + 1)
#                     ws.cell(row=i + 2, column=5).value = change
#                     ws.cell(row=i + 2, column=6).value = percent_change
#                     if change >= 100 or change <= -100:
#                         ws.cell(row=i + 2, column=5).font = red_text
#                     if percent_change >= 0.05 or percent_change <= -0.05:
#                         ws.cell(row=i + 2, column=6).font = red_text
#                 except TypeError:
#                     ws.cell(row=i + 2, column=4).value = 'check project data'
#                 except KeyError:
#                     ws.cell(row=i + 2, column=4).value = 'not reporting'
#             else:
#                 ws.cell(row=i + 2, column=4).value = 'None'
#
#
#         # Note the ordering of data. Done in this manner so that data is displayed in graph in the correct way.
#         ws.cell(row=1, column=1).value = 'Project Name'
#         ws.cell(row=1, column=2).value = 'latest quarter (£m)'
#         ws.cell(row=1, column=3).value = 'last quarter (£m)'
#         ws.cell(row=1, column=4).value = 'baseline (£m)'
#         ws.cell(row=1, column=7).value = '£m change between latest and last quarter'
#         ws.cell(row=1, column=8).value = 'percentage change between latest and last quarter'
#         ws.cell(row=1, column=5).value = '£m change between last and baseline quarter'
#         ws.cell(row=1, column=6).value = 'percentage change between last and baseline quarter'
#
#     return wb
#
# def place_standard_comparision_excel(master_data_latest, master_data_baseline):
#     '''
#     Function that places all information structured via the get_wlc_costs and get_yearly_costs programmes into an
#     excel spreadsheet. It does some calculations on the level of change that has taken place.
#     This function places in data for a chart that shows changes in financial profile between latest and baseline.
#     :param master_data_latest: data representing latest quarter information
#     :param master_data_baseline: data representing baseline quarter information
#     :return: excel workbook
#     '''
#     wb = Workbook()
#
#     for i, key in enumerate(list(master_data_latest.keys())):
#         ws = wb.create_sheet(key, i)  # creating worksheets
#         ws.title = key  # title of worksheet
#
#         data_latest = master_data_latest[key]
#         data_baseline = master_data_baseline[key]
#
#         for i, project_name in enumerate(data_latest):
#             '''place project names into ws'''
#             ws.cell(row=i+2, column=1).value = project_name
#
#             '''loop for placing data into ws. highlight changes between quarters in red'''
#             latest_value = data_latest[project_name]
#             ws.cell(row=i + 2, column=2).value = latest_value
#
#             '''comparision data against last quarter'''
#             if project_name in data_baseline.keys():
#                 try:
#                     baseline_value = data_baseline[project_name]
#                     ws.cell(row=i + 2, column=3).value = baseline_value
#                     change = latest_value - baseline_value
#                     if baseline_value > 0:
#                         percent_change = (latest_value - baseline_value)/baseline_value
#                     else:
#                         percent_change = (latest_value - baseline_value)/(baseline_value + 1)
#                     ws.cell(row=i + 2, column=4).value = change
#                     ws.cell(row=i + 2, column=5).value = percent_change
#                     if change >= 100 or change <= -100:
#                         ws.cell(row=i + 2, column=4).font = red_text
#                     if percent_change >= 0.05 or percent_change <= -0.05:
#                         ws.cell(row=i + 2, column=5).font = red_text
#                 except TypeError:
#                     ws.cell(row=i + 2, column=3).value = 'check project data'
#             else:
#                 ws.cell(row=i + 2, column=3).value = 'None'
#
#
#         ws.cell(row=1, column=1).value = 'Project Name'
#         ws.cell(row=1, column=2).value = 'latest quarter (£m)'
#         ws.cell(row=1, column=3).value = 'baseline (£m)'
#         ws.cell(row=1, column=4).value = '£m change between latest and baseline'
#         ws.cell(row=1, column=5).value = 'percentage change between latest and baseline'
#
#     return wb
#
# def get_wlc(project_name_list, wlc_key, index):
#     '''
#     Function that gets projects wlc cost information and returns it in a python dictionary format.
#     :param project_name_list: list of project names
#     :param wlc_key: project whole life cost (wlc) keys
#     :param index: index value for which master to use from the q_master_data_list . 0 is for latest, 1 last and
#     2 baseline. The actual index list q_master_list is set at a global level in this programme.
#     :return: a dictionary structured 'wlc: 'project_name': total
#     '''
#     upper_dictionary = {}
#     lower_dictionary = {}
#     for project_name in project_name_list:
#         try:
#             project_data = list_of_masters_all[costs_bl_index[project_name][index]].data[project_name]
#             total = project_data[wlc_key]
#             lower_dictionary[project_name] = total
#         except TypeError:
#             lower_dictionary[project_name] = 0
#
#     upper_dictionary['wlc'] = lower_dictionary
#
#     return upper_dictionary
#
#
# '''getting financial wlc cost breakdown'''
# latest_wlc = get_wlc(list_of_masters_all[0].projects, wlc_key, 0)
# last_wlc = get_wlc(list_of_masters_all[0].projects, wlc_key, 1)
# baseline_wlc = get_wlc(list_of_masters_all[0].projects, wlc_key, 2)
#
# '''creating excel outputs'''
# output_one = place_complex_comparision_excel(latest_wlc, last_wlc, baseline_wlc)
# output_two = place_complex_comparision_excel(latest_cost_profiles, last_cost_profiles, baseline_1_cost_profiles)
# output_three = place_standard_comparision_excel(latest_wlc, baseline_wlc)
# output_four = place_standard_comparision_excel(latest_cost_profiles, baseline_1_cost_profiles)
#
# '''INSTRUCTIONS FOR RUNNING PROGRAMME'''
#
# '''Valid file paths for all the below need to be provided'''
#
# '''ONE. Provide file path to where to save complex wlc breakdown'''
# output_one.save(root_path/'output/comparing_wlc_complex_q2_2021.xlsx')
#
# '''TWO. Provide file path to where to save complex yearly cost profile breakdown'''
# output_two.save(root_path/'output/comparing_cost_profiles_complex_q2_2021.xlsx')
#
# '''THREE. Provide file path to where to save standard wlc breakdown'''
# output_three.save(root_path/'output/comparing_wlc_standard_q2_2021.xlsx')
#
# '''FOUR. Provide file path to where to save standard yearly cost profile breakdown'''
# output_four.save(root_path/'output/comparing_cost_profiles_standard_q2_2021.xlsx')
#
#
# # possible use in milestone analysis_engine
# PARLIAMENT = [
#     "Bill",
#     "bill",
#     "hybrid",
#     "Hybrid",
#     "reading",
#     "royal",
#     "Royal",
#     "assent",
#     "Assent",
#     "legislation",
#     "Legislation",
#     "Passed",
#     "NAO",
#     "nao",
#     "PAC",
#     "pac",
# ]
# CONSTRUCTION = [
#     "Start of Construction/build",
#     "Complete",
#     "complete",
#     "Tender",
#     "tender",
# ]
# OPERATIONS = [
#     "Full Operations",
#     "Start of Operation",
#     "operational",
#     "Operational",
#     "operations",
#     "Operations",
#     "operation",
#     "Operation",
# ]
# OTHER_GOV = ["TAP", "MPRG", "Cabinet Office", " CO ", "HMT"]
# CONSULTATIONS = [
#     "Consultation",
#     "consultation",
#     "Preferred",
#     "preferred",
#     "Route",
#     "route",
#     "Announcement",
#     "announcement",
#     "Statutory",
#     "statutory",
#     "PRA",
# ]
# PLANNING = [
#     "DCO",
#     "dco",
#     "Planning",
#     "planning",
#     "consent",
#     "Consent",
#     "Pre-PIN",
#     "Pre-OJEU",
#     "Initiation",
#     "initiation",
# ]
# IPDC = ["IPDC", "BICC"]
# HE_SPECIFIC = [
#     "Start of Construction/build",
#     "DCO",
#     "dco",
#     "PRA",
#     "Preferred",
#     "preferred",
#     "Route",
#     "route",
#     "Annoucement",
#     "announcement",
#     "submission",
#     "PVR" "Submission",
# ]

# def put_combined_data_into_wb(combined_data):
#     """
#     places combined_data object into excel wb. Data in wb
#     is milestone name, current data, movement from baseline
#     data and milestone notes.
#     """
#
#     wb = Workbook()
#     ws = wb.active
#
#     row_num = 2
#
#     for i, milestone in enumerate(combined_data.group_current.keys()):
#         ws.cell(row=row_num + i, column=2).value = milestone
#         try:
#             milestone_date = tuple(combined_data.group_current[milestone])[0]
#             ws.cell(row=row_num + i, column=3).value = milestone_date
#             ws.cell(row=row_num + i, column=3).number_format = 'dd/mm/yy'
#         except KeyError:
#             ws.cell(row=row_num + i, column=3).value = ''
#
#         try:
#             baseline_milestone_date = tuple(combined_data.group_baseline[milestone])[0]
#             time_delta = (milestone_date - baseline_milestone_date).days
#             ws.cell(row=row_num + i, column=4).value = time_delta
#         except (KeyError, TypeError):
#             ws.cell(row=row_num + i, column=4).value = ''
#
#         try:
#             ws.cell(row=row_num + i, column=5).value = combined_data.group_current[milestone][
#                 milestone_date]  # provides notes
#         except (IndexError, KeyError):
#             ws.cell(row=row_num + i, column=5).value = ''
#
#
#     #ws.cell(row=1, column=1).value = 'Project'
#     ws.cell(row=1, column=2).value = 'Milestone'
#     ws.cell(row=1, column=3).value = 'Date'
#     #ws.cell(row=1, column=4).value = '3/m change'
#     ws.cell(row=1, column=4).value = 'Movement from baseline'
#     # ws.cell(row=1, column=6).value = 'Baseline change (last)'
#     ws.cell(row=1, column=5).value = 'Notes'
#
#     return wb


# def rcf_data(master_dict, project_title, start_row, output_wb):
#     # output_wb = Workbook()
#     data = project_data_from_master(master_dict)
#     project_data = data[project_title]
#
#     cells_we_want_to_capture = ['Reporting period (GMPP - Snapshot Date)',
#                                 'Approval MM1',
#                                 'Approval MM1 Forecast / Actual',
#                                 'Approval MM3',
#                                 'Approval MM3 Forecast / Actual',
#                                 'Approval MM10',
#                                 'Approval MM10 Forecast / Actual',
#                                 'Project MM18',
#                                 'Project MM18 Forecast - Actual',
#                                 'Project MM19',
#                                 'Project MM19 Forecast - Actual',
#                                 'Project MM20',
#                                 'Project MM20 Forecast - Actual',
#                                 'Project MM21',
#                                 'Project MM21 Forecast - Actual']
#     output_list = []
#     for item in project_data.items():
#         if item[0] in cells_we_want_to_capture:
#             output_list.append(item)
#
#     output_list = list(enumerate(output_list, start=1))
#     print(output_list)
#
#     output_list2 = [output_list[2][1][1],
#                     output_list[4][1][1],
#                     output_list[6][1][1],
#                     output_list[8][1][1],
#                     output_list[10][1][1],
#                     output_list[12][1][1],
#                     output_list[14][1][1]]
#
#     SOBC = output_list2[0]
#     print('SOBC', SOBC)
#     OBC = output_list2[1]
#     print('OBC', OBC)
#     FBC = output_list2[2]
#     print('FBC', FBC)
#     start_project = output_list2[3]
#     print('Start of Project', start_project)
#     start_construction = output_list2[4]
#     print('Start of construction', start_construction)
#     start_ops = output_list2[5]
#     print('Start of Ops', start_ops)
#     end_project = output_list2[6]
#     print('End of project', end_project)
#
#     try:
#         time_delta1 = (SOBC - start_project).days
#     except TypeError:
#         time_delta1 = None
#     print(time_delta1)
#     try:
#         time_delta2 = (OBC - SOBC).days
#     except TypeError:
#         time_delta2 = None
#     print(time_delta2)
#     try:
#         time_delta3 = (FBC - OBC).days
#     except TypeError:
#         time_delta3 = None
#     print(time_delta3)
#     try:
#         time_delta4 = (start_construction - FBC).days
#     except TypeError:
#         time_delta4 = None
#     print(time_delta4)
#     try:
#         time_delta5 = (start_ops - start_construction).days
#     except TypeError:
#         time_delta5 = None
#     print(time_delta5)
#     try:
#         time_delta6 = (end_project - start_ops).days
#     except TypeError:
#         time_delta6 = None
#     print(time_delta6)
#
#     ws = output_wb.active
#
#     for x in output_list[:3]:
#         ws.cell(row=2, column=x[0] + 1, value=x[1][0])
#         ws.cell(row=start_row + 1, column=x[0] + 1, value=x[1][1])
#         ws.cell(row=start_row + 1, column=x[0] + 2, value=time_delta1)
#
#     for x in output_list[3:5]:
#         ws.cell(row=2, column=x[0] + 2, value=x[1][0])
#         ws.cell(row=start_row + 1, column=x[0] + 2, value=x[1][1])
#         ws.cell(row=start_row + 1, column=x[0] + 3, value=time_delta2)
#
#     for x in output_list[5:7]:
#         ws.cell(row=2, column=x[0] + 3, value=x[1][0])
#         ws.cell(row=start_row + 1, column=x[0] + 3, value=x[1][1])
#         ws.cell(row=start_row + 1, column=x[0] + 4, value=time_delta3)
#
#     for x in output_list[7:9]:
#         ws.cell(row=2, column=x[0] + 4, value=x[1][0])
#         ws.cell(row=start_row + 1, column=x[0] + 4, value=x[1][1])
#         # ws.cell(row=start_row+1, column=series_one[0]+5, value=time_delta3)
#
#     for x in output_list[9:11]:
#         ws.cell(row=2, column=x[0] + 5, value=x[1][0])
#         ws.cell(row=start_row + 1, column=x[0] + 5, value=x[1][1])
#         ws.cell(row=start_row + 1, column=x[0] + 6, value=time_delta4)
#
#     for x in output_list[11:13]:
#         ws.cell(row=2, column=x[0] + 6, value=x[1][0])
#         ws.cell(row=start_row + 1, column=x[0] + 6, value=x[1][1])
#         ws.cell(row=start_row + 1, column=x[0] + 7, value=time_delta5)
#
#     for x in output_list[13:15]:
#         ws.cell(row=2, column=x[0] + 7, value=x[1][0])
#         ws.cell(row=start_row + 1, column=x[0] + 7, value=x[1][1])
#         ws.cell(row=start_row + 1, column=x[0] + 8, value=time_delta6)
#
#     for x in output_list[1:3]:
#         ws.cell(row=start_row + 13, column=x[0] + 1, value=x[1][1])
#         # ws.cell(row=start_row+10, column=series_one[0]+2, value=series_one[1][1])
#         ws.cell(row=start_row + 13, column=5, value=time_delta1)
#         ws.cell(row=start_row + 13, column=6, value=1)
#
#     for x in output_list[3:5]:
#         ws.cell(row=start_row + 13 + len(master_list), column=x[0] - 1, value=x[1][1])
#         ws.cell(row=start_row + 13 + len(master_list), column=5, value=time_delta2)
#         ws.cell(row=start_row + 13 + len(master_list), column=6, value=2)
#
#     for x in output_list[5:7]:
#         ws.cell(row=start_row + 13 + (len(master_list) * 2), column=x[0] - 3, value=x[1][1])
#         ws.cell(row=start_row + 13 + (len(master_list) * 2), column=5, value=time_delta3)
#         ws.cell(row=start_row + 13 + (len(master_list) * 2), column=6, value=3)
#
#     for x in output_list[9:11]:
#         ws.cell(row=start_row + 13 + (len(master_list) * 3), column=x[0] - 7, value=x[1][1])
#         ws.cell(row=start_row + 13 + (len(master_list) * 3), column=5, value=time_delta4)
#         ws.cell(row=start_row + 13 + (len(master_list) * 3), column=6, value=4)
#
#     for x in output_list[11:13]:
#         ws.cell(row=start_row + 13 + (len(master_list) * 4), column=x[0] - 9, value=x[1][1])
#         ws.cell(row=start_row + 13 + (len(master_list) * 4), column=5, value=time_delta5)
#         ws.cell(row=start_row + 13 + (len(master_list) * 4), column=6, value=5)
#
#     for x in output_list[13:15]:
#         ws.cell(row=start_row + 13 + (len(master_list) * 5), column=x[0] - 11, value=x[1][1])
#         ws.cell(row=start_row + 13 + (len(master_list) * 5), column=5, value=time_delta6)
#         ws.cell(row=start_row + 13 + (len(master_list) * 5), column=6, value=6)
#
#     return output_wb
#
#
# def rcf_chart(data, p, output_wb):
#     # wb = load_workbook(workbook)
#     ws = output_wb.active
#     # approval_point = data[p]['BICC approval point']
#     chart = ScatterChart()
#     # chart.title = 'Time Delta Schedule \n Last BC agreed by BICC ' + str(approval_point)
#     chart.style = 18
#     chart.x_axis.title = 'Days'
#     chart.y_axis.title = 'Time Delta'
#     chart.height = 11  # default is 7.5
#     chart.width = 22  # default is 15
#
#     xvalues = Reference(ws, min_col=5, min_row=1 + (len(master_list) * 3), max_row=(len(master_list) * 4) - 2)
#     yvalues = Reference(ws, min_col=6, min_row=1 + (len(master_list) * 3), max_row=(len(master_list) * 4) - 2)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s2 = chart.series[0]
#     s2.marker.symbol = "diamond"
#     s2.marker.size = 10
#     s2.marker.graphicalProperties.solidFill = "dcc7aa"  # Marker filling grey
#     s2.marker.graphicalProperties.line.solidFill = "dcc7aa"  # Marker outline grey
#     s2.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=len(master_list) * 3, max_row=len(master_list) * 3)
#     yvalues = Reference(ws, min_col=6, min_row=len(master_list) * 3, max_row=len(master_list) * 3)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s1 = chart.series[1]
#     s1.marker.symbol = "diamond"
#     s1.marker.size = 10
#     s1.marker.graphicalProperties.solidFill = "f7c331"  # Marker filling yellow
#     s1.marker.graphicalProperties.line.solidFill = "f7c331"  # Marker outline yellow
#     s1.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=(len(master_list) * 4) - 1, max_row=(len(master_list) * 4) - 1)
#     yvalues = Reference(ws, min_col=6, min_row=(len(master_list) * 4) - 1, max_row=(len(master_list) * 4) - 1)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s3 = chart.series[2]
#     s3.marker.symbol = "diamond"
#     s3.marker.size = 10
#     s3.marker.graphicalProperties.solidFill = "f7882f"  # Marker filling orange
#     s3.marker.graphicalProperties.line.solidFill = "f7882f"  # Marker outline orange
#     s3.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=1 + (len(master_list) * 4), max_row=(len(master_list) * 5) - 2)
#     yvalues = Reference(ws, min_col=6, min_row=1 + (len(master_list) * 4), max_row=(len(master_list) * 5) - 2)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s4 = chart.series[3]
#     s4.marker.symbol = "diamond"
#     s4.marker.size = 10
#     s4.marker.graphicalProperties.solidFill = "dcc7aa"  # Marker filling grey
#     s4.marker.graphicalProperties.line.solidFill = "dcc7aa"  # Marker outline grey
#     s4.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=len(master_list) * 4, max_row=len(master_list) * 4)
#     yvalues = Reference(ws, min_col=6, min_row=len(master_list) * 4, max_row=len(master_list) * 4)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s5 = chart.series[4]
#     s5.marker.symbol = "diamond"
#     s5.marker.size = 10
#     s5.marker.graphicalProperties.solidFill = "f7c331"  # Marker filling yellow
#     s5.marker.graphicalProperties.line.solidFill = "f7c331"  # Marker outline yellow
#     s5.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=(len(master_list) * 5) - 1, max_row=(len(master_list) * 5) - 1)
#     yvalues = Reference(ws, min_col=6, min_row=(len(master_list) * 5) - 1, max_row=(len(master_list) * 5) - 1)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s6 = chart.series[5]
#     s6.marker.symbol = "diamond"
#     s6.marker.size = 10
#     s6.marker.graphicalProperties.solidFill = "f7882f"  # Marker filling orange
#     s6.marker.graphicalProperties.line.solidFill = "f7882f"  # Marker outline orange
#     s6.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=1 + (len(master_list) * 5), max_row=(len(master_list) * 6) - 2)
#     yvalues = Reference(ws, min_col=6, min_row=1 + (len(master_list) * 5), max_row=(len(master_list) * 6) - 2)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s7 = chart.series[6]
#     s7.marker.symbol = "diamond"
#     s7.marker.size = 10
#     s7.marker.graphicalProperties.solidFill = "dcc7aa"  # Marker filling grey
#     s7.marker.graphicalProperties.line.solidFill = "dcc7aa"  # Marker outline grey
#     s7.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=len(master_list) * 5, max_row=len(master_list) * 5)
#     yvalues = Reference(ws, min_col=6, min_row=len(master_list) * 5, max_row=len(master_list) * 5)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s8 = chart.series[7]
#     s8.marker.symbol = "diamond"
#     s8.marker.size = 10
#     s8.marker.graphicalProperties.solidFill = "f7c331"  # Marker filling yellow
#     s8.marker.graphicalProperties.line.solidFill = "f7c331"  # Marker outline yellow
#     s8.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=(len(master_list) * 6) - 1, max_row=(len(master_list) * 6) - 1)
#     yvalues = Reference(ws, min_col=6, min_row=(len(master_list) * 6) - 1, max_row=(len(master_list) * 6) - 1)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s9 = chart.series[8]
#     s9.marker.symbol = "diamond"
#     s9.marker.size = 10
#     s9.marker.graphicalProperties.solidFill = "f7882f"  # Marker filling orange
#     s9.marker.graphicalProperties.line.solidFill = "f7882f"  # Marker outline orange
#     s9.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=1 + (len(master_list) * 6), max_row=(len(master_list) * 7) - 2)
#     yvalues = Reference(ws, min_col=6, min_row=1 + (len(master_list) * 6), max_row=(len(master_list) * 7) - 2)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s10 = chart.series[9]
#     s10.marker.symbol = "diamond"
#     s10.marker.size = 10
#     s10.marker.graphicalProperties.solidFill = "dcc7aa"  # Marker filling grey
#     s10.marker.graphicalProperties.line.solidFill = "dcc7aa"  # Marker outline grey
#     s10.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=len(master_list) * 6, max_row=len(master_list) * 6)
#     yvalues = Reference(ws, min_col=6, min_row=len(master_list) * 6, max_row=len(master_list) * 6)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s11 = chart.series[10]
#     s11.marker.symbol = "diamond"
#     s11.marker.size = 10
#     s11.marker.graphicalProperties.solidFill = "f7c331"  # Marker filling yellow
#     s11.marker.graphicalProperties.line.solidFill = "f7c331"  # Marker outline yellow
#     s11.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=(len(master_list) * 7) - 1, max_row=(len(master_list) * 7) - 1)
#     yvalues = Reference(ws, min_col=6, min_row=(len(master_list) * 7) - 1, max_row=(len(master_list) * 7) - 1)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s12 = chart.series[11]
#     s12.marker.symbol = "diamond"
#     s12.marker.size = 10
#     s12.marker.graphicalProperties.solidFill = "f7882f"  # Marker filling orange
#     s12.marker.graphicalProperties.line.solidFill = "f7882f"  # Marker outline orange
#     s12.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=1 + (len(master_list) * 7), max_row=(len(master_list) * 8) - 2)
#     yvalues = Reference(ws, min_col=6, min_row=1 + (len(master_list) * 7), max_row=(len(master_list) * 8) - 2)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s13 = chart.series[12]
#     s13.marker.symbol = "diamond"
#     s13.marker.size = 10
#     s13.marker.graphicalProperties.solidFill = "dcc7aa"  # Marker filling grey
#     s13.marker.graphicalProperties.line.solidFill = "dcc7aa"  # Marker outline grey
#     s13.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=len(master_list) * 7, max_row=len(master_list) * 7)
#     yvalues = Reference(ws, min_col=6, min_row=len(master_list) * 7, max_row=len(master_list) * 7)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s14 = chart.series[13]
#     s14.marker.symbol = "diamond"
#     s14.marker.size = 10
#     s14.marker.graphicalProperties.solidFill = "f7c331"  # Marker filling yellow
#     s14.marker.graphicalProperties.line.solidFill = "f7c331"  # Marker outline yellow
#     s14.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=(len(master_list) * 8) - 1, max_row=(len(master_list) * 8) - 1)
#     yvalues = Reference(ws, min_col=6, min_row=(len(master_list) * 8) - 1, max_row=(len(master_list) * 8) - 1)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s15 = chart.series[14]
#     s15.marker.symbol = "diamond"
#     s15.marker.size = 10
#     s15.marker.graphicalProperties.solidFill = "f7882f"  # Marker filling orange
#     s15.marker.graphicalProperties.line.solidFill = "f7882f"  # Marker outline orange
#     s15.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=1 + (len(master_list) * 8), max_row=(len(master_list) * 9) - 2)
#     yvalues = Reference(ws, min_col=6, min_row=1 + (len(master_list) * 8), max_row=(len(master_list) * 9) - 2)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s16 = chart.series[15]
#     s16.marker.symbol = "diamond"
#     s16.marker.size = 10
#     s16.marker.graphicalProperties.solidFill = "dcc7aa"  # Marker filling grey
#     s16.marker.graphicalProperties.line.solidFill = "dcc7aa"  # Marker outline grey
#     s16.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=len(master_list) * 8, max_row=len(master_list) * 8)
#     yvalues = Reference(ws, min_col=6, min_row=len(master_list) * 8, max_row=len(master_list) * 8)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s17 = chart.series[16]
#     s17.marker.symbol = "diamond"
#     s17.marker.size = 10
#     s17.marker.graphicalProperties.solidFill = "f7c331"  # Marker filling yellow
#     s17.marker.graphicalProperties.line.solidFill = "f7c331"  # Marker outline yellow
#     s17.graphicalProperties.line.noFill = True
#
#     xvalues = Reference(ws, min_col=5, min_row=(len(master_list) * 9) - 1, max_row=(len(master_list) * 9) - 1)
#     yvalues = Reference(ws, min_col=6, min_row=(len(master_list) * 9) - 1, max_row=(len(master_list) * 9) - 1)
#     series = Series(values=yvalues, xvalues=xvalues, title=None)
#     chart.series.append(series)
#     s18 = chart.series[17]
#     s18.marker.symbol = "diamond"
#     s18.marker.size = 10
#     s18.marker.graphicalProperties.solidFill = "f7882f"  # Marker filling orange
#     s18.marker.graphicalProperties.line.solidFill = "f7882f"  # Marker outline orange
#     s18.graphicalProperties.line.noFill = True
#
#     ws.add_chart(chart, "I12")
#
#     return output_wb


class Pickle:
    def __init__(self, master, save_path: str):
        self.master = master
        self.path = save_path
        self.in_a_pickle()

    def in_a_pickle(self) -> None:
        with open(self.path + ".pickle", "wb") as handle:
            pickle.dump(self.master, handle, protocol=pickle.HIGHEST_PROTOCOL)


def open_pickle_file(path: str):
    with open(path, "rb") as handle:
        return pickle.load(handle)


# def json_date_converter(o):
#     if isinstance(o, datetime.date):
#         return o.__str__()


# class JsonData:
#     def __init__(self, master: Master, save_path: str):
#         self.master = master
#         self.path = save_path
#         self.put_into_json()
#
#     def put_into_json(self) -> None:
#         master_list = []
#         for m in self.master.master_data:
#             data = m.data
#             projects = m.projects
#             qrt = str(str(m.quarter))
#             d = {
#                 "data": data,
#                 "projects": projects,
#                 "quarter": qrt,
#             }
#             master_list.append(d)
#
#         json_dict = {
#             "abbreviations": self.master.abbreviations,
#             "bl_index": self.master.bl_index,
#             "bl_info": self.master.bl_info,
#             "current_projects": self.master.current_projects,
#             "current_quarter": str(self.master.current_quarter),
#             "dft_groups": self.master.dft_groups,
#             "full_names": self.master.full_names,
#             "kwargs": self.master.kwargs,
#             "master_data": master_list,
#             "pipeline_dict": self.master.pipeline_dict,
#             "pipeline_list": self.master.pipeline_list,
#             "project_group": self.master.project_group,
#             "project_information": self.master.project_information,
#             "project_stage": self.master.project_stage,
#             "quarter_list": self.master.quarter_list,
#         }
#         with open(self.path + ".json", "w") as write_file:
#             json.dump(json_dict, write_file, default=json_date_converter)


def open_json_file(path: str):
    with open(path, "r") as handle:
        return json.load(handle)


"""NOTE. these three lists need to have rag ratings in the same order"""
"""different colours are placed into a list"""
txt_colour_list = [ag_text, ar_text, red_text, green_text, amber_text]
fill_colour_list = [ag_fill, ar_fill, red_fill, green_fill, amber_fill]
"""list of how rag ratings are shown in document"""
rag_txt_list = ["A/G", "A/R", "R", "G", "A"]


def concatenate_dates(date: date):
    """
    function for converting dates into concatenated written time periods
    :param date: datetime.date
    :return: concatenated date
    """
    IPDC_DATE = parser.parse(
        get_ipdc_date(str(root_path) + "/core_data/ipdc_config.ini", "ipdc_date"),
        dayfirst=True,
    ).date()
    if date is not None:
        a = (date - IPDC_DATE).days
        year = 365
        month = 30
        fortnight = 14
        week = 7
        if a >= 365:
            yrs = int(a / year)
            holding_days_years = a % year
            months = int(holding_days_years / month)
            holding_days_months = a % month
            fortnights = int(holding_days_months / fortnight)
            weeks = int(holding_days_months / week)
        elif 0 <= a <= 365:
            yrs = 0
            months = int(a / month)
            holding_days_months = a % month
            fortnights = int(holding_days_months / fortnight)
            weeks = int(holding_days_months / week)
            # if 0 <= a <=60:
        elif a <= -365:
            yrs = int(a / year)
            holding_days = a % -year
            months = int(holding_days / month)
            holding_days_months = a % -month
            fortnights = int(holding_days_months / fortnight)
            weeks = int(holding_days_months / week)
        elif -365 <= a <= 0:
            yrs = 0
            months = int(a / month)
            holding_days_months = a % -month
            fortnights = int(holding_days_months / fortnight)
            weeks = int(holding_days_months / week)
            # if -60 <= a <= 0:
        else:
            print("something is wrong and needs checking")

        if yrs == 1:
            if months == 1:
                return "{} yr, {} mth".format(yrs, months)
            if months > 1:
                return "{} yr, {} mths".format(yrs, months)
            else:
                return "{} yr".format(yrs)
        elif yrs > 1:
            if months == 1:
                return "{} yrs, {} mth".format(yrs, months)
            if months > 1:
                return "{} yrs, {} mths".format(yrs, months)
            else:
                return "{} yrs".format(yrs)
        elif yrs == 0:
            if a == 0:
                return "Today"
            elif 1 <= a <= 6:
                return "This week"
            elif 7 <= a <= 13:
                return "Next week"
            elif -7 <= a <= -1:
                return "Last week"
            elif -14 <= a <= -8:
                return "-2 weeks"
            elif 14 <= a <= 20:
                return "2 weeks"
            elif 20 <= a <= 60:
                if IPDC_DATE.month == date.month:
                    return "Later this mth"
                elif (date.month - IPDC_DATE.month) == 1:
                    return "Next mth"
                else:
                    return "2 mths"
            elif -60 <= a <= -15:
                if IPDC_DATE.month == date.month:
                    return "Earlier this mth"
                elif (date.month - IPDC_DATE.month) == -1:
                    return "Last mth"
                else:
                    return "-2 mths"
            elif months == 12:
                return "1 yr"
            else:
                return "{} mths".format(months)

        elif yrs == -1:
            if months == -1:
                return "{} yr, {} mth".format(yrs, -(months))
            if months < -1:
                return "{} yr, {} mths".format(yrs, -(months))
            else:
                return "{} yr".format(yrs)
        elif yrs < -1:
            if months == -1:
                return "{} yrs, {} mth".format(yrs, -(months))
            if months < -1:
                return "{} yrs, {} mths".format(yrs, -(months))
            else:
                return "{} yrs".format(yrs)
    else:
        return "None"


class ResourceData:
    def __init__(self, master, **kwargs):
        self.master = master
        self.baseline_type = "ipdc_costs"
        self.kwargs = kwargs
        self.start_group = []
        self.group = []
        self.iter_list = []
        self.get_resource_totals()
        # self.ps_resource = 0
        # self.contractor_resource = 0
        # self.total_resource = 0

    def get_resource_totals(self) -> None:
        self.iter_list = get_iter_list(self.kwargs, self.master)
        for tp in self.iter_list:
            self.group = get_group(self.master, tp, self.kwargs)
            public_sector_resource = []
            c_resource = []
            t_resource = []
            fp_resource = []
            for project_name in self.group:
                p_data = get_correct_p_data(
                    self.kwargs, self.master, self.baseline_type, project_name, tp
                )
                if p_data is None:
                    break
                else:
                    ps = convert_none_types(p_data["DfTc Public Sector Employees"])
                    public_sector_resource.append(ps)
                    c = convert_none_types(p_data["DfTc External Contractors"])
                    c_resource.append(c)
                    t = convert_none_types(p_data["DfTc Project Team Total"])
                    t_resource.append(t)
                    fp = convert_none_types(p_data["DfTc Funded Posts"])
                    fp_resource.append(fp)

        self.ps_resource = sum(public_sector_resource)
        self.contractor_resource = sum(c_resource)
        self.total_resource = sum(t_resource)
        self.funded = sum(fp_resource)


## old and hashing out for now
# class DandelionChart:
#     def __init__(self, area, bubble_spacing=0):
#         """
#         Setup for bubble collapse.
#
#         @param area: array-like. Area of the bubbles.
#         @param bubble_spacing: float, default:0. Minimal spacing between bubbles after collapsing.
#
#         @note
#         If "area" is sorted, the results might look weird.
#         """
#         area = np.asarray(area)
#         r = np.sqrt(area / np.pi)
#
#         self.bubble_spacing = bubble_spacing
#         self.bubbles = np.ones((len(area), 4))
#         self.bubbles[:, 2] = r
#         self.bubbles[:, 3] = area
#         self.maxstep = 2 * self.bubbles[:, 2].max() + self.bubble_spacing
#         self.step_dist = self.maxstep / 2
#
#         # calculate initial grid layout for bubbles
#         length = np.ceil(np.sqrt(len(self.bubbles)))
#         grid = np.arange(length) * self.maxstep  # arrange might cause trouble
#         gx, gy = np.meshgrid(grid, grid)
#         self.bubbles[:, 0] = gx.flatten()[: len(self.bubbles)]
#         self.bubbles[:, 1] = gy.flatten()[: len(self.bubbles)]
#
#         self.com = self.center_of_mass()
#
#     def center_of_mass(self):
#         return np.average(self.bubbles[:, :2], axis=0, weights=self.bubbles[:, 3])
#
#     def center_distance(self, bubble, bubbles):
#         return np.hypot(bubble[0] - bubbles[:, 0], bubble[1] - bubbles[:, 1])
#
#     def outline_distance(self, bubble, bubbles):
#         center_distance = self.center_distance(bubble, bubbles)
#         return center_distance - bubble[2] - bubbles[:, 2] - self.bubble_spacing
#
#     def check_collisions(self, bubble, bubbles):
#         distance = self.outline_distance(bubble, bubbles)
#         return len(distance[distance < 0])
#
#     def collides_with(self, bubble, bubbles):
#         distance = self.outline_distance(bubble, bubbles)
#         idx_min = np.argmin(distance)
#         return idx_min if type(idx_min) == np.ndarray else [idx_min]
#
#     def collapse(self, n_iterations=50):
#         """
#         Move bubbles to the center of mass.
#
#         @param n_iterations: int, default: 50. Number of moves to perform.
#         @return:
#         """
#         for _i in range(n_iterations):
#             moves = 0
#             for i in range(len(self.bubbles)):
#                 rest_bub = np.delete(self.bubbles, i, 0)
#                 # try to move directly towards the center of mass
#                 # direction vector from bubble to the center of mass
#                 dir_vec = self.com - self.bubbles[i, :2]
#
#                 # shorten direction vector to have length of 1
#                 try:
#                     dir_vec = dir_vec / np.sqrt(dir_vec.dot(dir_vec))
#                 except (RuntimeWarning, RuntimeError):
#                     dir_vec = 1
#
#                 # calculate new bubble position
#                 new_point = self.bubbles[i, :2] + dir_vec * self.step_dist
#                 new_bubble = np.append(new_point, self.bubbles[i, 2:4])
#
#                 # check whether new bubble collides with other bubbles
#                 if not self.check_collisions(new_bubble, rest_bub):
#                     self.bubbles[i, :] = new_bubble
#                     self.com = self.center_of_mass()
#                     moves += 1
#                 else:
#                     # try to move around a bubble that you collide with
#                     # find colliding bubble
#                     for colliding in self.collides_with(new_bubble, rest_bub):
#                         # calculate direction vector
#                         dir_vec = rest_bub[colliding, :2] - self.bubbles[i, :2]
#                         dir_vec = dir_vec / np.sqrt(dir_vec.dot(dir_vec))
#                         # calculate orthogonal vector
#                         orth = np.array([dir_vec[1], -dir_vec[0]])
#                         # test which direction to go
#                         new_point1 = self.bubbles[i, :2] + orth * self.step_dist
#                         new_point2 = self.bubbles[i, :2] - orth * self.step_dist
#                         dist1 = self.center_distance(self.com, np.array([new_point1]))
#                         dist2 = self.center_distance(self.com, np.array([new_point2]))
#                         new_point = new_point1 if dist1 < dist2 else new_point2
#                         new_bubble = np.append(new_point, self.bubbles[i, 2:4])
#                         if not self.check_collisions(new_bubble, rest_bub):
#                             self.bubbles[i, :] = new_bubble
#                             self.com = self.center_of_mass()
#
#             if moves / len(self.bubbles) < 0.1:
#                 self.step_dist = self.step_dist / 2
#
#     def plot(self, ax, labels, colors):
#         """
#         Draw the bubble plot.
#
#         @param ax: matplotlib.axes.Axes
#         @param labels: list. labels of the bubbles.
#         @param colors: list. colour of the bubbles.
#         @return:
#         """
#         for i in range(len(self.bubbles)):
#             circ = plt.Circle(self.bubbles[i, :2], self.bubbles[i, 2], color=colors[i])
#             ax.add_patch(circ)
#             ax.text(
#                 *self.bubbles[i, :2],
#                 labels[i],
#                 horizontalalignment="center",
#                 verticalalignment="center",
#             )
#
#
# def run_dandelion_matplotlib_chart(dandelion: Dict[str, list], **kwargs) -> plt.figure:
#     bubble_chart = DandelionChart(area=dandelion["cost"], bubble_spacing=20)
#     bubble_chart.collapse()
#     fig, ax = plt.subplots(subplot_kw=dict(aspect="equal"))
#     bubble_chart.plot(ax, dandelion["projects"], dandelion["colour"])
#     ax.axis("off")
#     ax.relim()
#     ax.autoscale_view()
#     # ax.set_title(str(DandelionData.)
#     if "chart" in kwargs:
#         if kwargs["chart"]:
#             plt.show()
#     return fig
#


# def get_sp_data(
#     master: Master, g_list: List[str], quarter: List[str], **kwargs
# ) -> plt.figure:
#     sp_dict = {}  # stacked plot dict
#     # tp_index = master.quarter_list.index(quarter[0])
#     if kwargs["type"] == "comp":  # composition
#         for g in g_list:  # group list
#             # HERE. figure out how to remove g from sp_dict keys if it needs to be removed.
#             if "remove" in kwargs:
#                 costs = CostData(master, group=[g], quarter=quarter, remove=kwargs["remove"])
#                 if costs is None:
#                     pass
#             else:
#                 costs = CostData(master, group=[g], quarter=quarter)
#             try:
#                 sp_dict[master.abbreviations[g]['abb']] = costs.c_profiles[quarter[0]]["prof"]
#             except KeyError:
#                 sp_dict[g] = costs.c_profiles[quarter[0]]["prof"]
#     elif kwargs["type"] == "cat":  # category
#         costs = CostData(master, group=[g_list], quarter=quarter)
#         cat_list = ["cdel", "rdel", "ngov"]
#         for c in cat_list:
#             sp_dict[c] = costs.c_profiles[quarter[0]][c]
#
#     return sp_dict


def put_stackplot_data_into_wb(sp_data: Dict) -> workbook:
    wb = Workbook()
    ws = wb.active

    for x, g in enumerate(sp_data.keys()):
        ws.cell(row=1, column=2 + x).value = g
        for i, pv in enumerate(sp_data[g]):
            ws.cell(row=2 + i, column=1).value = YEAR_LIST[i]
            ws.cell(row=2 + i, column=2 + x).value = pv

    ws.cell(row=1, column=1).value = "Year"

    wb.save(root_path / "output/sp_data_all.xlsx")


def cost_stackplot_graph(sp_dict: Dict[str, float], master, **kwargs) -> plt.figure:
    sp_list = []  # stackplot list
    labels = list(sp_dict[str(master.current_quarter)].keys())
    for g in labels:
        sp_list.append(sp_dict[str(master.current_quarter)][g])
    y = np.vstack(sp_list)
    x = YEAR_LIST
    fig, ax = plt.subplots()
    fig.set_size_inches(18.5, 10.5)

    ## HERE
    title = get_chart_title(master, "costs stack plot", **kwargs)
    plt.suptitle(title, fontweight="bold", fontsize=20)
    ax.stackplot(x, y, labels=labels)

    # Chart styling
    plt.xticks(rotation=45, size=16)
    plt.yticks(size=16)
    ax.set_ylabel("Cost (£m)")
    ax.set_xlabel("Financial Year")
    xlab1 = ax.xaxis.get_label()
    xlab1.set_style("italic")
    xlab1.set_size(20)
    ylab1 = ax.yaxis.get_label()
    ylab1.set_style("italic")
    ylab1.set_size(20)
    ax.grid(color="grey", linestyle="-", linewidth=0.2)
    # ax.legend()
    ax.legend(loc="upper right", prop={"size": 16})
    # ax.legend(prop={"size": 12})

    if "chart" in kwargs:
        if kwargs["chart"]:
            plt.show()

    return fig
    # fig.savefig(root_path /"output/portfolio_cost_composition_cat.png")


# def make_a_dandelion_manual(wb: Union[str, bytes, os.PathLike]):
#     wb = load_workbook(wb, data_only=True)
#     ws = wb.active
#
#     d_list = []  # data list
#     for row_num in range(3, ws.max_row + 1):
#         x_axis = ws.cell(row=row_num, column=5).value
#         y_axis = ws.cell(row=row_num, column=6).value
#         b_size = ws.cell(row=row_num, column=13).value
#         colour = random.choice(list(COLOUR_DICT.values()))
#         d_list.append(((x_axis, y_axis), b_size / 10, colour, "-"))
#
#     # fig = plt.figure()
#     # d_list.insert(0, ((515, 450), 455, 'r'))
#     # , ((590, 422), 52, 'g')]
#     plt.figure(figsize=(20, 10))
#     for c in range(len(d_list)):
#         # circle = plt.Circle(d_list[c][0], radius=d_list[c][1], fc=d_list[c][2], linestyle='-')
#         circle = plt.Circle(
#             d_list[c][0], radius=d_list[c][1], linestyle="--", fill=False
#         )
#         plt.gca().add_patch(circle)
#     plt.axis("scaled")
#     plt.axis("off")
#
#     plt.show()
#
#     return plt


def get_horizontal_bar_chart_data(wb_path: TextIO) -> Dict:
    wb = load_workbook(wb_path)
    ws = wb.active

    category_names = []
    for x in range(2, ws.max_row + 1):
        cat = ws.cell(row=x, column=1).value
        if cat is not None:
            category_names.append(cat)

    results = {}
    for y in range(2, ws.max_column + 1):
        type = ws.cell(row=1, column=y).value
        results_list = []
        for x in range(2, ws.max_row + 1):
            r = ws.cell(row=x, column=y).value
            if r is not None:
                if isinstance(r, float) or r < 1:
                    results_list.append(int(r * 100))
                else:
                    results_list.append(r)
        if type is not None:
            results[type] = results_list

    def survey(results, category_names):
        """
        Parameters
        ----------
        results : dict
            A mapping from question labels to a list of answers per category.
            It is assumed all lists contain the same number of entries and that
            it matches the length of *category_names*.
        category_names : list of str
            The category labels.
        """
        labels = list(results.keys())
        data = np.array(list(results.values()))
        data_cum = data.cumsum(axis=1)
        category_colors = plt.get_cmap("RdYlGn")(np.linspace(0.15, 0.85, data.shape[1]))

        fig, ax = plt.subplots(facecolor=FACE_COLOUR)  # figsize=(9.2, 5)
        fig.set_size_inches(18.5, 10.5)
        ax.set_facecolor(FACE_COLOUR)
        ax.invert_yaxis()
        ax.xaxis.set_visible(False)
        ax.set_xlim(0, np.sum(data, axis=1).max())

        for i, (colname, color) in enumerate(zip(category_names, category_colors)):
            widths = data[:, i]
            starts = data_cum[:, i] - widths
            widths = np.ma.masked_equal(widths, 0)  # removes zeros
            rects = ax.barh(
                labels,
                widths,
                left=starts,
                height=0.5,
                label=colname,
                color=COLOUR_DICT[colname],
            )

            # r, g, b, _ = color
            # text_color = 'white'  # if r * g * b < 0.5 else 'darkgrey'
            ax.bar_label(rects, label_type="center")  # color=text_color

        ax.legend(
            ncol=len(category_names),
            bbox_to_anchor=(0.5, -0.1),
            loc="center",
            fontsize="medium",
            facecolor=FACE_COLOUR,
        )

        plt.axis("off")
        # plt.show()

        return fig

    fig = survey(results, category_names)

    return fig


def simple_horz_bar_chart(wb_path: TextIO) -> Dict:
    wb = load_workbook(wb_path)
    ws = wb.active

    type_names = []
    for y in range(2, ws.max_column + 1):
        cat = ws.cell(row=1, column=y).value
        if cat is not None:
            type_names.append(cat)

    # results = {}
    # for y in range(2, ws.max_column + 1):
    #     type = ws.cell(row=1, column=y).value
    #     results_list = []
    #     for x in range(2, ws.max_row + 1):
    #         r = ws.cell(row=x, column=y).value
    #         if r is not None:
    #             if isinstance(r, float) or r < 1:
    #                 results_list.append(int(r * 100))
    #             else:
    #                 results_list.append(r)
    #     if type is not None:
    #         results[type] = results_list
    #
    results = {}
    for x in range(2, ws.max_row + 1):
        cat = ws.cell(row=x, column=1).value
        results_list = []
        for y in range(2, ws.max_column + 1):
            r = ws.cell(row=x, column=y).value
            if r is not None:
                if isinstance(r, float) or r < 1:
                    results_list.append(int(r * 100))
                else:
                    results_list.append(r)
        if cat is not None:
            results[cat] = results_list

    fig, ax = plt.subplots(figsize=(5, 3))
    left_list = [0, 0]
    for r in results:
        result = np.ma.masked_equal(results[r], 0)  # removes zeros
        p = ax.barh(
            type_names,
            result,
            align="center",
            left=left_list,
            height=0.25,
            color=COLOUR_DICT[r],
            label=r,
        )
        left_list = [results[r][i] + left_list[i] for i in range(len(left_list))]
        ax.bar_label(p, label_type="center")
    # for (p, r) in zip(ax.patches, results):
    #     percentage = '{:.1f}%'.format(100 * p.get_width() / results[r])
    #     x = p.get_x() + p.get_width() + 0.02
    #     y = p.get_y() + p.get_height() / 2
    #     ax.annotate(percentage, (x, y))

    ax.set_yticks(type_names)
    # ax.set_xlabel('Percentage')
    # ax.set_title('Run Time by Job')
    # ax.grid(True)
    # ax.legend()
    # plt.tight_layout()
    # # plt.savefig('C:\\Data\\stackedbar.png')
    plt.show()

    return fig


def so_matplotlib():
    jobs = ["Q4", "Q3", "Q2"]
    waittimes = [3.3472222222222223, 4.752777777777778, 6.177777777777778]
    runtimes = [0.3472222222222222, 1.0027777777777778, 0.5111111111111111]

    fig = plt.figure()
    ax = fig.add_subplot(111)
    ax.barh(
        jobs, waittimes, align="center", height=0.25, color="#00ff00", label="wait time"
    )
    ax.barh(
        jobs,
        runtimes,
        align="center",
        height=0.25,
        left=waittimes,
        color="g",
        label="run time",
    )
    ax.set_yticks(jobs)
    ax.set_xlabel("Hour")
    ax.set_title("Run Time by Job")
    ax.grid(True)
    ax.legend()
    plt.tight_layout()
    plt.show()


# def radar_chart():
#     categories = ["Food Quality", "Food Variety", "Service Quality", "Ambiance", "Affordability"]
#     categories = [*categories, categories[0]]
#     restaurant_1 = [4,4,5,4,3]
#     restaurant_2 = [5,5,4,5,2]
#     restaurant_3 = [3,4,5,8,5]
#     restaurant_1 = [*restaurant_1, restaurant_1[0]]
#     restaurant_2 = [*restaurant_2, restaurant_2[0]]
#     restaurant_3 = [*restaurant_3, restaurant_3[0]]
#
#     label_loc = np.linspace(start=0, stop=2 * np.pi, num=len(restaurant_1))
#     # plt.figure(figsize=(8, 8))
#     # plt.subplot(polar=True)
#     # plt.plot(label_loc, restaurant_1, color='g', label="Rest 1")
#     # plt.plot(label_loc, restaurant_2, label="Rest 2")
#     # plt.plot(label_loc, restaurant_3, label="Rest 3")
#     # lines, labels = plt.thetagrids(np.degrees(label_loc), labels=categories)
#     # plt.legend()
#     # plt.show()
#     ax = plt.subplot(111, polar=true)
#     plt.xticks(label_loc, categories, color="grey", size=8)
#     ax.set


def radar_chart(sp_data_path: TextIO, master, **kwargs):
    from matplotlib.path import Path

    data = get_strategic_priorities_data(sp_data_path, master, **kwargs)

    def radar_factory(num_vars, frame="circle"):
        """Create a radar chart with `num_vars` axes.

        This function creates a RadarAxes projection and registers it.

        Parameters
        ----------
        num_vars : int
            Number of variables for radar chart.
        frame : {'circle' | 'polygon'}
            Shape of frame surrounding axes.

        """
        # calculate evenly-spaced axis angles
        theta = np.linspace(0, 2 * np.pi, num_vars, endpoint=False)

        class RadarAxes(PolarAxes):
            name = "radar"

            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                # rotate plot such that the first axis is at the top
                self.set_theta_zero_location("N")

            def fill(self, *args, closed=True, **kwargs):
                """Override fill so that line is closed by default"""
                return super().fill(closed=closed, *args, **kwargs)

            def plot(self, *args, **kwargs):
                """Override plot so that line is closed by default"""
                lines = super().plot(*args, **kwargs)
                for line in lines:
                    self._close_line(line)

            def _close_line(self, line):
                x, y = line.get_data()
                # FIXME: markers at x[0], y[0] get doubled-up
                if x[0] != x[-1]:
                    x = np.concatenate((x, [x[0]]))
                    y = np.concatenate((y, [y[0]]))
                    line.set_data(x, y)

            def set_varlabels(self, labels):
                self.set_thetagrids(np.degrees(theta), labels)

            def _gen_axes_patch(self):
                # The Axes patch must be centered at (0.5, 0.5) and of radius 0.5
                # in axes coordinates.
                if frame == "circle":
                    return Circle((0.5, 0.5), 0.5)
                elif frame == "polygon":
                    return RegularPolygon(
                        (0.5, 0.5), num_vars, radius=0.5, edgecolor="k"
                    )
                else:
                    raise ValueError("unknown value for 'frame': %s" % frame)

            def draw(self, renderer):
                """Draw. If frame is polygon, make gridlines polygon-shaped"""
                if frame == "polygon":
                    gridlines = self.yaxis.get_gridlines()
                    for gl in gridlines:
                        gl.get_path()._interpolation_steps = num_vars
                super().draw(renderer)

            def _gen_axes_spines(self):
                if frame == "circle":
                    return super()._gen_axes_spines()
                elif frame == "polygon":
                    # spine_type must be 'left'/'right'/'top'/'bottom'/'circle'.
                    spine = Spine(
                        axes=self,
                        spine_type="circle",
                        path=Path.unit_regular_polygon(num_vars),
                    )
                    # unit_regular_polygon gives a polygon of radius 1 centered at
                    # (0, 0) but we want a polygon of radius 0.5 centered at (0.5,
                    # 0.5) in axes coordinates.
                    spine.set_transform(
                        Affine2D().scale(0.5).translate(0.5, 0.5) + self.transAxes
                    )

                    return {"polar": spine}
                else:
                    raise ValueError("unknown value for 'frame': %s" % frame)

        register_projection(RadarAxes)
        return theta

    N = len(data[0])
    theta = radar_factory(N, frame="polygon")

    spoke_labels = data.pop(0)
    title, case_data = data[0]

    fig, ax = plt.subplots(figsize=(7, 6), subplot_kw=dict(projection="radar"))
    # fig.set_size_inches(18.5, 10.5)
    fig.subplots_adjust(top=0.85, bottom=0.05)

    # ax.set_rgrids([0, 1, 2, 3, 4, 5, 6, 7])
    ax.set_yticklabels([1, 2, 3, 4, 5])
    ax.set_title(title, fontweight="bold", position=(0.5, 1), fontsize=15)

    for d in case_data:
        line = ax.plot(theta, d[:-1], alpha=0.25, color=COLOUR_DICT[d[-1]])
        ax.fill(theta, d[:-1], alpha=0.25, facecolor=COLOUR_DICT[d[-1]])
    ax.set_varlabels(spoke_labels)

    if "chart" in kwargs:
        if kwargs["chart"]:
            plt.show()

    return fig


def get_strategic_priorities_data(
    wb_path: TextIO,
    master,
    **kwargs,
):
    sp_data = project_data_from_master(wb_path, 1, 2021)

    categories = [
        "Improving transport for the user",
        "Reduce environmental impacts",
        "Increase our global impact",
        "Grow an level up the economy",
        "Be an excellent department",
        # "VfM",
    ]

    sp_scoring = {
        None: 0,
        "adverse impact": 1,
        "zero impact": 2,
        "minor postive impact": 3,
        "moderate positive impact": 4,
        "major positive impact": 5,
    }

    vfm_scoring = {
        None: 0,
        "Very Poor": 0,
        "Poor": 1,
        "Low": 2,
        "Medium": 3,
        "High": 4,
        "Very High": 5,
        "Very High and Financially Positive": 5,
        "Economically Efficient Cost Saving": 5,
    }

    if "group" in kwargs:
        group_list = kwargs["group"]
    else:
        group_list = sp_data.data.keys()

    top_list = []
    for p in group_list:
        if p not in sp_data.data.keys():
            print(p)
            continue
        p_list = []
        for c in categories:  # no vfm
            # for c in categories[:-1]:  # if vfm included
            p_list.append(sp_scoring[sp_data.data[p][c]])
        try:
            # p_list.append(
            #     vfm_scoring[master.master_data[0].data[p]["VfM Category single entry"]]
            # )
            p_list.append(master.master_data[0]["data"][p]["Departmental DCA"])
        except KeyError:
            print(p)
            # p_list.append(0)  # if vfm included
            p_list.append("Amber")
        top_list.append(p_list)

    if "title" in kwargs:
        title = kwargs["title"]
    else:
        title = ""
    radar_data = [categories, (title, top_list)]

    return radar_data


# def get_map(wb, commas=False, gaps=False, flip=False):
#     ws = wb.active
#     output_dict = {}
#
#     for x in range(2, ws.max_row + 1):
#         ipa_key = ws.cell(row=x, column=2).value
#         if ipa_key in output_dict.keys():
#             pass
#         if ipa_key is None:
#             pass
#         else:
#             dft_key = ws.cell(row=x, column=1).value
#             ipa_key = ws.cell(row=x, column=2).value
#             if not commas:
#                 ipa_key = ipa_key.replace(',', '')
#             if not gaps:
#                 ipa_key = ipa_key.replace('  ', ' ')
#             if flip:
#                 output_dict[ipa_key] = dft_key
#             else:
#                 output_dict[dft_key] = ipa_key
#
#     return output_dict


# def semantic_ordering(wb):
#     # wb = load_workbook("/home/will/Downloads/semantic_v_nos.xlsx")
#     ws = wb.active
#
#     output_list = []
#     for x in range(2, ws.max_row + 1):
#         no = ws.cell(row=x, column=1).value
#         gmpp_key = ws.cell(row=x, column=2).value
#         dft_key = ws.cell(row=x, column=3).value
#         output_list.append((no, gmpp_key, dft_key))
#
#     output_list.sort(key=lambda s: list(map(int, s[0].split('.'))))
#
#     output_dict = {}
#     for x in output_list:
#         output_dict[x[1]] = x[2]
#
#     return output_dict


# def get_gmpp_keys():
#     wb_one = load_workbook("/home/will/Downloads/GMPP_DATA_Q4.xlsm")
#     wb_two = load_workbook("/home/will/Downloads/GMPP_DATA_Q1.xlsm")
#
#     wb_list = [wb_one, wb_two]
#     pure_key_list = []
#     output_list = []
#     for wb in wb_list:
#         ws = wb.active
#         for x in range(24, ws.max_row):  # row 24 is where gmpp data starts.
#             key_pure = ws.cell(row=x, column=6).value
#             key = key_pure.split(':')
#             if key_pure not in pure_key_list:
#                 if 'a' in key[0] or 'b' in key[0] or 'c' in key[0]:
#                     pass
#                 else:
#                     output_list.append(key)
#                     pure_key_list.append(key_pure)
#
#     output_list.sort(key=lambda s: list(map(int, s[0].split('.'))))
#
#     key_map = get_map(load_workbook("/home/will/Downloads/KEY_MAP_v8.xlsx"))
#
#     wb = Workbook()
#     ws = wb.active
#
#     for i, key in enumerate(output_list):
#         ws.cell(row=i + 1, column=1).value = key[0] + ":" + key[1]
#         try:
#             ws.cell(row=i + 1, column=2).value = key_map[key[0] + ":" + key[1]]
#         except KeyError:
#             ws.cell(row=i + 1, column=2).value = ""
#
#     wb.save("/home/will/Downloads/KEY_MAP_v9.xlsx")


def sort_gmpp_on_key_order(wb) -> List:
    """
    This function puts gmpp keys in semantic order. Was used when first
    working with gmpp_online data and the keys were in random order. Not
    currently required as order now handled via key_map.
    @return:
    """
    pure_key_list = []
    output_list = []
    errant_keys = [
        "6.03a: Has the project been using an evidenced range of  Project End Dates for planning and approval "
        "decisions?",
        "6.03b: Schedule Project End Date Range (FORECAST),  From",
        "6.03c: To",
        "7.02a: Is this the Projects first GMPP data return?",
    ]
    # for wb in wb_list:  # option to pass in list here rather than single wb
    ws = wb.active
    for x in range(24, ws.max_row):  # row 24 is where gmpp data starts.
        key_pure = ws.cell(row=x, column=6).value
        key = key_pure.split(":")
        if key_pure not in pure_key_list:
            if "a" in key[0] or "b" in key[0] or "c" in key[0]:
                if key[0] == "6.03a":
                    key[0] = "6.031"
                    if key_pure not in pure_key_list:
                        output_list.append(key)
                        pure_key_list.append(key_pure)
                if key[0] == "6.03b":
                    key[0] = "6.032"
                    if key_pure not in pure_key_list:
                        output_list.append(key)
                        pure_key_list.append(key_pure)
                if key[0] == "6.03c":
                    key[0] = "6.033"
                    if key_pure not in pure_key_list:
                        output_list.append(key)
                        pure_key_list.append(key_pure)
                if key[0] == "7.02a":
                    key[0] = "7.021"
                    if key_pure not in pure_key_list:
                        output_list.append(key)
                        pure_key_list.append(key_pure)
                if key[0] == "7.02b":
                    key[0] = "7.022"
                    if key_pure not in pure_key_list:
                        output_list.append(key)
                        pure_key_list.append(key_pure)
                else:
                    if key_pure not in errant_keys:
                        print(key_pure)
            elif "12" in key[0][:2]:
                pass
            else:
                output_list.append(key)
                pure_key_list.append(key_pure)

    output_list.sort(key=lambda s: list(map(int, s[0].split("."))))

    ipa_key_list = []
    for i, key in enumerate(output_list):
        if key[0] == "6.031":
            key[0] = "6.03a"
        if key[0] == "6.032":
            key[0] = "6.03b"
        if key[0] == "6.033":
            key[0] = "6.03c"
        if key[0] == "7.021":
            key[0] = "7.02a"
        if key[0] == "7.022":
            key[0] = "7.02b"
        ipa_key_list.append(key[0] + ":" + key[1])

    return ipa_key_list


# def get_gmpp_data(
#         file_name: str,
#         # file_name_two: str
# ):
#     from datetime import datetime
#     import xlrd
#
#     try:
#         wb = load_workbook(root_path / 'input/{}.xlsx'.format(file_name), data_only=True)
#     except: # bit of flexibility to help user with different file types
#         wb = load_workbook(root_path / 'input/{}.xlsm'.format(file_name), data_only=True)
#     ws = wb.active
#     ws_list = [ws]
#
#     key_map = get_map(load_workbook(root_path / "input/GMPP_INTEGRATION_KEY_MAP.xlsx"), commas=True, gaps=True)
#
#     initial_dict = {}
#     missing_keys = []
#     for ws in ws_list:
#         for x in range(24, ws.max_row + 1):
#             project_name = ws.cell(row=x, column=2).value
#             key = ws.cell(row=x, column=6).value
#             if key not in key_map.values():
#                 if key not in missing_keys:
#                         missing_keys.append(key)
#             s_value = ws.cell(row=x, column=7).value
#             n_value = ws.cell(row=x, column=8).value
#             if n_value != 0:
#                 s_value = n_value
#             if "Date" in key or "date" in key or "6.03c: To" in key:
#                 # s_value = datetime(*xlrd.xldate_as_tuple(n_value, 0))
#                 # print(project_name, key, n_value)
#                 if n_value > 20000:
#                     # if "7.02.10" in key:
#                     #     pass
#                     # else:
#                     s_value = datetime(*xlrd.xldate_as_tuple(n_value, 0))
#                     # else:
#                     #     s_value = n_value
#             if "Grade" in key:  # to make grade 6 consistent with dft data
#                 try:
#                     s_value = int(s_value)
#                 except ValueError:
#                     pass
#
#             if project_name in list(initial_dict.keys()):
#                 initial_dict[project_name][key] = s_value
#             else:
#                 initial_dict[project_name] = {key: s_value}
#
#     ## Use for checking to see if key map missing keys.
#     # for x in missing_keys:
#     #     print(x)
#
#     return initial_dict
#

# def place_gmpp_online_keys_into_dft_master_format(
#         initial_dict: Dict,
#         km_file_name: str,
#         ipdc_d_file_path,
#         project_list=False,
#     ):
#     wb = Workbook()
#     ws = wb.active
#
#     # get non-gmpp keys from team
#
#     key_map = get_map(load_workbook(
#         root_path / "input/{}.xlsx".format(km_file_name)
#     ), commas=True, gaps=True)
#     ipdc_data = project_data_from_master(root_path / "core_data/{}.xlsx".format(ipdc_d_file_path), 2, 2021)
#
#     a_proj_name = ipdc_data.projects[1]
#     # a_proj_name = 'East Coast Mainline Programme'
#     if project_list:
#         list_of_projects = project_list
#     else:
#         list_of_projects = list(initial_dict.keys())
#
#     for x, project in enumerate(list_of_projects):
#         ws.cell(row=1, column=3 + x).value = project
#         i = 0
#         for v in ipdc_data.data[a_proj_name].keys():
#             # for ipa_key, dft_key in key_map.items():
#             ws.cell(row=2 + i, column=1).value = v
#             try:
#                 ipa_key = key_map[v]
#                 ws.cell(row=2 + i, column=2).value = ipa_key
#                 # try:
#                 ipa_value = initial_dict[project][ipa_key]
#                 if isinstance(ipa_value, datetime.datetime):
#                     ipa_value = ipa_value.date()
#                     ws.cell(row=2 + i, column=3 + x, value=ipa_value).number_format = "dd/mm/yy"
#                 ws.cell(row=2 + i, column=3 + x).value = ipa_value
#                 # except KeyError:
#                 #     pass
#             except KeyError:
#                 pass
#             i += 1
#
#     ws.cell(row=1, column=1).value = "Project Name (DfT Keys)"
#     ws.cell(row=1, column=2).value = "Project Name (IPA Keys)"
#
#     wb.save(root_path / "output/gmpp_online_data_dft_master_format.xlsx")
#
#     ws.delete_cols(1, 1)
#
#     wb.save(root_path / "input/gmpp_online_data_temp.xlsx")
#

GMPP_M_DICT = {
    "Project - End Date": "Project End Date",
    "Project - Start Date": "Start of Project",
    "IPA Gate 0": "Gateway Review 0",
    "IPA Gate 1": "Gateway Review 1",
    "IPA Gate 2": "Gateway Review 2",
    "IPA Gate 3": "Gateway Review 3",
    "IPA Gate 4": "Gateway Review 4",
    "IPA Gate 5": "Gateway Review 5",
    "IPA Gate 0 (Final)": "Gateway Review 0 (Final)",
    "Approval - HMT OBC": "OBC - HMT Approval",
    "Approval - HMT SOC": "SOBC - HMT Approval",
    "Approval - HMT FBC": "FBC - HMT Approval",
    "Approval - HMT PBC": "Latest HMT Approved - Programme Business Case",
    "Yes": "Yes - Permanent",
    "DFT": "DfT",
}

RK_LIST = [
    "0.1.3: Project Name",
    "Variance",
    "Information Technology Name",
    "Digital Name",
    "Project Delivery Name",
    "Change Implementation Name",
    "Technical Name",
    "Industry Knowledge Name",
    "Finance Name",
    "Analysis Name",
    "Communications & Stakeholder Engagement Name",
    "Legal Commercial & Contract Management Name",
    "Approval Name",
    "DCA - SRO Narrative",
    "Project DCA - SRO Assessment Delivery confidence for whole project duration",
    "2.01.3: Project IQA - IPA Quarterly Assessment for whole project duration",
    "2.01.4: IQA - IPA Narrative",
]

IGNORE_LIST = [
    "Reporting period (GMPP - Snapshot Date)",
    "Pre 19-20 BL Income both Revenue and Capital",
    "Pre 19-20 BL Non-Gov",
    "Pre 19-20 CDEL BL one off new costs",
    "Pre 19-20 RDEL BL one off new costs",
    "Pre-profile Forecast Non-Gov",
    "Pre-profile CDEL Forecast one off new costs",
    "Pre-profile RDEL Forecast one off new costs",
    "Pre 19-20 Actual Income both Revenue and Capital",
    "Unprofiled Remainder BEN Baseline - Gov. Cashable",
    "Unprofiled Remainder BEN Baseline - Gov. Non-Cashable",
    "Unprofiled Remainder BEN Baseline - Economic (inc Private Partner)",
    "Unprofiled Remainder BEN Baseline - Disbenefit UK Economic",
    "Unprofiled Remainder BEN Forecast - Gov. Cashable",
    "Unprofiled Remainder BEN Forecast - Gov. Non-Cashable",
    "Unprofiled Remainder BEN Forecast - Economic (inc Private Partner)",
    "Unprofiled Remainder BEN Forecast - Disbenefit UK Economic",
    "Assurance MM1",
    "Assurance MM1 LoD",
    "Assurance MM1 Version No",
    "Assurance MM2",
    "Assurance MM2 LoD",
    "Assurance MM2 Version No",
    "Assurance MM3",
    "Assurance MM3 LoD",
    "Assurance MM3 Version No",
    "Assurance MM4",
    "Assurance MM4 LoD",
    "Assurance MM4 Version No",
    "Assurance MM5",
    "Assurance MM5 LoD",
    "Assurance MM5 Version No",
    "Assurance MM6",
    "Assurance MM6 LoD",
    "Assurance MM6 Version No",
    "Assurance MM7",
    "Assurance MM7 LoD",
    "Assurance MM7 Version No",
    "Project MM18",
    "Project MM18 DMC",
    "Project MM18 CP",
    "Project MM19",
    "Project MM19 DMC",
    "Project MM19 CP",
    "Project MM20",
    "Project MM20 DMC",
    "Project MM20 CP",
    "Classification",
    "GMPP - Main reason for joining GMPP",
    "If yes please select SRO from this list:",
    "GMPP - SRO ID",
    "If yes please select PD from this list:",
    "GMPP - PD ID",
]


# def print_gmpp_data(gmpp_dict: Dict):
#     wb = Workbook()
#     ws = wb.active
#
#     key_map = get_map(load_workbook("/home/will/Downloads/KEY_MAP_v8.xlsx"))
#     # key_map = semantic_ordering(load_workbook("/home/will/Downloads/semantic_v_nos.xlsx"))
#
#     for x, project in enumerate(list(gmpp_dict.keys())):
#         ws.cell(row=1, column=3 + x).value = project
#         for i, k in enumerate(gmpp_dict[project]):  # k is key
#             if x == 0:
#                 ws.cell(row=2 + i, column=1).value = k
#                 try:
#                     ws.cell(row=2 + i, column=2).value = key_map[k]
#                 except KeyError:
#                     ws.cell(row=2 + i, column=2).value = ''
#             try:
#                 ws.cell(row=2 + i, column=3 + x).value = gmpp_dict[project][k]
#             except KeyError:
#                 pass
#
#     ws.cell(row=1, column=1).value = "Project Name"
#
#     wb.save("/home/will/Downloads/GMPP_DATA_DFT_FORMAT_v7.xlsx")
#

# def change_gmpp_keys_order(ipdc_dict):
#     key_map = get_map(load_workbook("/home/will/Downloads/KEY_MAP_v8.xlsx"))
#
#     output_dict = {}
#     for key in ipdc_dict['Crossrail Programme']:
#         for x, value in key_map.items():
#             if key == value:
#                 output_dict[key] = x
#
#     wb = Workbook()
#     ws = wb.active
#     for i, key in enumerate(ipdc_dict['Crossrail Programme']):
#         ws.cell(row=i + 2, column=1).value = key
#         try:
#             ws.cell(row=i + 2, column=2).value = output_dict[key]
#         except KeyError:
#             print(key)
#             ws.cell(row=i + 2, column=2).value = None
#
#     wb.save("/home/will/Downloads/GMPP_KEY_MAP_DFT_ORDER_v2.xlsx")
#
#     return output_dict
#

# def put_n02_into_master():
#     # q1 = get_project_info_data("/home/will/Documents/ipdc/core_data/master_1_2021.xlsx")
#     no2 = get_project_info_data("/home/will/Downloads/GMPP_DATA_DFT_FORMAT_NO2_FINAL.xlsx")
#
#     wb = load_workbook("/home/will/Documents/ipdc/core_data/master_1_2021.xlsx")
#     ws = wb.active
#
#     for x in range(2, ws.max_column + 1):
#         if ws.cell(row=1, column=x).value == "NO2 Reduction":
#             for z in range(2, ws.max_row + 1):
#                 key = ws.cell(row=z, column=1).value
#                 try:
#                     value = no2["NO2 Reduction"][key]
#                     # ws.cell(row=z, column=x).value = value
#                     if isinstance(value, datetime.datetime):
#                         value = value.date()
#                         ws.cell(row=z, column=x, value=value).number_format = "dd/mm/yy"
#                     else:
#                         ws.cell(row=z, column=x).value = value
#                 except KeyError:
#                     pass
#
#     wb.save("/home/will/Documents/ipdc/core_data/master_1_2021.xlsx")
#


def get_risk_data():
    wb = load_workbook("/home/will/Downloads/BasicRiskReport.xlsx")
    ws = wb.active

    output_dict = {}
    for x in range(2, ws.max_row + 1):
        code = ws.cell(row=x, column=2).value
        output_dict[code] = {
            "name": ws.cell(row=x, column=3).value,
            "likelihood": ws.cell(row=x, column=4).value,
            "impact": ws.cell(row=x, column=5).value,
            "exposure": ws.cell(row=x, column=6).value,
            "date": ws.cell(row=x, column=12).value,
            "planned": ws.cell(row=x, column=13).value,
            "implemented": ws.cell(row=x, column=15).value,
            "cause": ws.cell(row=x, column=17).value,
            "event": ws.cell(row=x, column=18).value,
            "effect": ws.cell(row=x, column=19).value,
        }

    return output_dict


def print_risk_data(risk_data: Dict):
    wb = Workbook()
    ws = wb.active

    for x, code in enumerate(list(risk_data.keys())):
        ws.cell(row=x + 2, column=1).value = code
        ws.cell(row=x + 2, column=2).value = risk_data[code]["name"]
        ws.cell(row=x + 2, column=3).value = risk_data[code]["likelihood"]
        ws.cell(row=x + 2, column=4).value = risk_data[code]["impact"]
        ws.cell(row=x + 2, column=5).value = risk_data[code]["exposure"]
        ws.cell(row=x + 2, column=6).value = risk_data[code]["date"]
        ws.cell(row=x + 2, column=6).number_format = "dd/mm/yy"
        ws.cell(row=x + 2, column=7).value = risk_data[code]["planned"]
        ws.cell(row=x + 2, column=8).value = risk_data[code]["implemented"]
        ws.cell(row=x + 2, column=9).value = risk_data[code]["cause"]
        ws.cell(row=x + 2, column=10).value = risk_data[code]["event"]
        ws.cell(row=x + 2, column=11).value = risk_data[code]["effect"]

    ws.cell(row=1, column=1).value = "code"
    ws.cell(row=1, column=2).value = "name"
    ws.cell(row=1, column=3).value = "likelihood"
    ws.cell(row=1, column=4).value = "impact"
    ws.cell(row=1, column=5).value = "exposure"
    ws.cell(row=1, column=6).value = "date"
    ws.cell(row=1, column=7).value = "planned"
    ws.cell(row=1, column=8).value = "implemented"
    ws.cell(row=1, column=9).value = "cause"
    ws.cell(row=1, column=10).value = "event"
    ws.cell(row=1, column=11).value = "effect"

    wb.save("/home/will/Downloads/exco_risks_HSRG.xlsx")


def calculate_dates(date_str):
    if date_str == "today":
        return datetime.date.today()
    if date_str == "minus 3 months":
        return datetime.date.today() + relativedelta(months=-3)
    if date_str == "minus 6 months":
        return datetime.date.today() + relativedelta(months=-6)
    if date_str == "2 years":
        return datetime.date.today() + relativedelta(months=24)


def doughut(dca_data: DcaData, **kwargs):
    fig, ax = plt.subplots(figsize=(6, 3), subplot_kw=dict(aspect="equal"))

    number = []
    colours = []
    (q,) = kwargs["quarter"]
    for c in ["Green", "Amber/Green", "Amber", "Amber/Red", "Red", None]:
        count = dca_data.dca_count[q][kwargs["conf_type"]][c][kwargs["weighting"]]
        if kwargs["weighting"] == "ct":
            count = round(count * 100, 2)
        if count != 0:
            number.append(count)
            colours.append(COLOUR_DICT[c])

    wedges, texts = ax.pie(
        number, colors=colours, wedgeprops=dict(width=0.5), startangle=-40
    )

    bbox_props = dict(boxstyle="square,pad=0.3", fc="w", ec="k", lw=0.72)
    kw = dict(arrowprops=dict(arrowstyle="-"), bbox=bbox_props, zorder=0, va="center")

    for i, p in enumerate(wedges):
        ang = (p.theta2 - p.theta1) / 2.0 + p.theta1
        y = np.sin(np.deg2rad(ang))
        x = np.cos(np.deg2rad(ang))
        horizontalalignment = {-1: "right", 1: "left"}[int(np.sign(x))]
        connectionstyle = "angle,angleA=0,angleB={}".format(ang)
        kw["arrowprops"].update({"connectionstyle": connectionstyle})
        ax.annotate(
            number[i],
            xy=(x, y),
            xytext=(1.35 * np.sign(x), 1.4 * y),
            horizontalalignment=horizontalalignment,
            **kw,
        )

    ax.set_title(kwargs["conf_type"] + " confidence, simple count: a donut")

    plt.show()
    return fig


def amend_project_information():
    key_list = [
        "16 - 17 RDEL Baseline - One off new costs - investment in change",
        "16 - 17 RDEL Baseline - Recurring new costs - investment in change",
        "16 - 17 RDEL Baseline - Recurring old costs",
        "16 - 17 RDEL Baseline - Whole Life Cost breakdown",
        "16 - 17 CDEL Baseline - One off new costs - investment in change",
        "16 - 17 CDEL Baseline - Recurring new costs - investment in change",
        "16 - 17 CDEL Baseline - Recurring old costs",
        "16 - 17 Baseline - Non - Gov both Revenue and Capital",
        "16 - 17 CDEL Baseline - Whole Life Cost breakdown",
        "16 - 17 Baseline - Income(£m) both Revenue and Capital",
    ]

    master = get_project_info_data(root_path / "core_data/master_4_2016.xlsx")

    wb = load_workbook(root_path / "core_data/project_info.xlsx")
    ws = wb.active

    for x in range(2, ws.max_column + 1):
        project_name = ws.cell(row=1, column=x).value
        if project_name in list(master.keys()):
            for y in range(2, ws.max_row + 1):
                key = ws.cell(row=y, column=1).value
                if key in key_list:
                    print("yes")
                    ws.cell(row=y, column=x).value = master[project_name][key]

    wb.save(root_path / "core_data/project_info_amended.xlsx")
