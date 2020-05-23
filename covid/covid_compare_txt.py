'''
quick bit of code to compare reported narratives for covid-19 reporting. output is word document which
flags changes to text in red

requirements:
1) covid master wbs to go into analysis_engine core data directory
2) output file goes into output directory

follow instructions below.
'''


from datamaps.api import project_data_from_master
from analysis.engine_functions import compare_text_newandold
from analysis.data import root_path
from docx import Document

def printing():

     doc = Document()

     for project_name in covid_master_list[0].projects:
         #project_name = latest_master.data[project]['Project']
         heading = str(project_name)
         intro = doc.add_heading(str(heading), 0)
         intro.alignment = 1
         intro.bold = True

         new = doc.add_paragraph()
         new.add_run('')


         latest_txt = covid_master_list[0].data[project_name]['Project Narrative']
         if latest_txt is None:
             latest_txt = covid_master_list[0].data[project_name]['Group Narrative']

         try:
             last_txt = covid_master_list[1].data[project_name]['Project Narrative']
             if last_txt is None:
                 last_txt = covid_master_list[1].data[project_name]['Group Narrative']

             try:
                compare_text_newandold(latest_txt, last_txt, doc)
             except AttributeError:
                 pass

             if last_txt is None:
                 sentence = 'Project did not provide text last time. Latest text provided, but not compared' \
                            ' against anything.'
                 p = doc.add_paragraph()
                 runner = p.add_run(sentence)
                 runner.bold = True
                 doc.add_paragraph(latest_txt)
                 doc.add_paragraph()

             if latest_txt is None:
                 sentence = 'Project did not provide text this time. Text provided last time given here'
                 p = doc.add_paragraph()
                 runner = p.add_run(sentence)
                 runner.bold = True
                 doc.add_paragraph(latest_txt)
                 doc.add_paragraph()

             new_2 = doc.add_paragraph()
             new_2.add_run('')

         except KeyError:

             sentence = 'First time project is reporting or the project has changed name in master'
             p = doc.add_paragraph()
             runner = p.add_run(sentence)
             runner.bold = True
             doc.add_paragraph(latest_txt)
             doc.add_paragraph()


     return doc

'''RUNNING PROGRAMME INSTRUCTIONS'''

'''Ensure the latest master is being uploaded and placed in the masters list'''

master_29_05 = project_data_from_master(root_path/'core_data/covid_19/master_290520.xlsx', 1, 2020)
master_13_05 = project_data_from_master(root_path/'core_data/covid_19/master_130520.xlsx', 1, 2020)
master_01_05 = project_data_from_master(root_path/'core_data/covid_19/master_010520.xlsx', 1, 2020)

covid_master_list = [master_29_05,
                     master_13_05,
                     master_01_05]

'''The name of the output document can be changed via changing the file path here, if desired. The standard
name output name is covid_compare_text. Note make sure .docx is at end'''
run = printing()
run.save(root_path/'output/covid_19/covid_compare_text.docx')